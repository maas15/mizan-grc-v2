"""Synthetic board-ready defect sections mirroring live export (34)/(61)."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Dict, Tuple

FIXTURE_DIR = Path(__file__).resolve().parent
DOCX_FIXTURE = FIXTURE_DIR / 'cyber_strategy_34.docx'
PDF_FIXTURE = FIXTURE_DIR / 'cyber_strategy_61.pdf'


def content_quality_defect_sections() -> Dict[str, str]:
    """Sections with shallow pillars, weak roadmap, bad KPI/KRI, generic risks."""
    pillars = (
        '## الركائز الاستراتيجية\n\n'
        '### حوكمة ونموذج التشغيل\n'
        'تشغيل المركز وخطط الاستجابة المحددةفي المنظمة.\n'
        '| المبادرة | الوصف | المخرج المتوقع | المالك |\n'
        '|---|---|---|---|\n'
        '| سياسات الحوكمة | اعتماد سياسات | منصة حوكمة معتمدة | — |\n'
        '| لجنة الحوكمة | تشغيل اللجنة | لجنة حوكمة فعّالة | — |\n'
        '| RACI | توزيع الأدوار | مصفوفة RACI معتمدة | CISO |\n'
        '### الحماية والكشف والاستجابة\n'
        '| المبادرة | الوصف | المخرج المتوقع | المالك |\n'
        '|---|---|---|---|\n'
        '| SOC | تشغيل المركز | مركز SOC تشغيلي | — |\n'
        '| SIEM | قواعد المراقبة | تغطية SIEM للأصول الحرجة | مدير SOC |\n'
        '| CSIRT | خطط الاستجابة | فريق CSIRT جاهز | — |\n'
        '### الهوية وحماية البيانات\n'
        '| المبادرة | الوصف | المخرج المتوقع | المالك |\n'
        '|---|---|---|---|\n'
        '| IAM | إدارة الهوية | تغطية MFA للحسابات الحرجة | مدير الهوية |\n'
        '| DLP | بال منصات DLP | منصة DLP مفعّلة | — |\n'
        '### المرونة واستمرارية الأعمال\n'
        '| المبادرة | الوصف | المخرج المتوقع | المالك |\n'
        '|---|---|---|---|\n'
        '| BCP | استمرارية | خطط استمرارية معتمدة | مدير BCP |\n'
        '| DR | تعافي | خطة DR مختبرة | مدير BCP |\n'
        '| Backup | نسخ | خطة نسخ احتياطي معتمدة | IT |\n'
    )
    roadmap_preview = (
        '## خارطة الطريق التنفيذية\n\n'
        '| المرحلة | الفترة | المبادرة | المسؤول | المخرج | الإطار |\n'
        '|---|---|---|---|---|---|\n'
    )
    for i in range(1, 13):
        roadmap_preview += (
            f'| M{i//4+1} | Q{i} | مبادرة {i} | المسؤول أمن السيبرانيe Lead | '
            f'مخرج {i} | ECC |\n'
        )
    roadmap_docx = (
        '## خارطة الطريق التنفيذية\n\n'
        '| المرحلة | الفترة | المبادرة | المسؤول | المخرج | الإطار |\n'
        '|---|---|---|---|---|---|\n'
    )
    for i in range(1, 4):
        roadmap_docx += (
            f'| M1 | Q{i} | مبادرة مختصرة {i} | Owner | هيكل | ECC |\n'
        )
    kpis = (
        '## مؤشرات الأداء الرئيسية\n\n'
        '| # | وصف المؤشر | القيمة المستهدفة | صيغة الاحتساب | مصدر | تواتر |\n'
        '|---|---|---|---|---|---|\n'
        '| 1 | نسبة محاولات الدخول الفاشلة الشاذة | KPI 100% | '
        '(المنجز المقيس ÷ الهدف التشغيلي) × 100 | SIEM | شهري |\n'
        '| 2 | عدد حوادث تسرب البيانات الحرجة | KPI ≥ 95% | '
        '(عدد العناصر المطابقة / إجمالي العناصر) × 100 | DLP | شهري |\n'
        '| 3 | درجة مخاطر الأطراف الثالثة | KRI 100% | '
        '(المنجز / المخطط) × 100 | GRC | ربع سنوي |\n'
        '| 4 | متوسط زمن الكشف MTTD | ≤ 60 دقيقة | زمن الكشف | SOC | شهري |\n'
        '| 5 | متوسط زمن الاستجابة MTTR | ≤ 4 ساعات | زمن الاستجابة | CSIRT | شهري |\n'
    )
    confidence = (
        '## تقييم الثقة والمخاطر\n\n'
        'سجل المخاطر\n'
        '| # | الخطر | الاحتمالية | التأثير | خطة المعالجة | المالك |\n'
        '|---|---|---|---|---|---|\n'
        '| 1 | ضعف الحوكمة والسياسات | متوسط | عالٍ | '
        'ضوابط تقنية وإجراءات تشغيلية ومراقبة مستمرة | CISO |\n'
        '| 2 | ضعف IAM/MFA | عالٍ | عالٍ | '
        'ضوابط تقنية وإجراءات تشغيلية ومراقبة مستمرة | CISO |\n'
        '| 3 | غياب SOC/SIEM | عالٍ | عالٍ | '
        'ضوابط تقنية وإجراءات تشغيلية ومراقبة مستمرة | CISO |\n'
        '| 4 | هجمات التصيد والفدية | متوسط | عالٍ | '
        'ضوابط تقنية وإجراءات تشغيلية ومراقبة مستمرة | CSIRT |\n'
        '| 5 | مخاطر الأطراف الثالثة | متوسط | متوسط | '
        'ضوابط تقنية وإجراءات تشغيلية ومراقبة مستمرة | المشتريات |\n'
        '| 6 | ثغرات غير معالجة | عالٍ | عالٍ | '
        'ضوابط تقنية وإجراءات تشغيلية ومراقبة مستمرة | IT |\n'
    )
    traceability = (
        '## مصفوفة تتبع الأطر المرجعية\n\n'
        '| الإطار | مجال القدرة | الفجوة المرتبطة | المبادرة | المؤشر | الخطر |\n'
        '|---|---|---|---|---|---|\n'
        '| NCA DCC | تصنيف البيانات | مراجعة حسابات متميزة IAM/PAM | '
        'IAM | KPI | خطر |\n'
        '| NCA DCC | حماية البيانات | DLP فقط بدون تشفير | DLP | KPI | خطر |\n'
        '| NCA DCC | التشفير | ضعف عام | تشفير | KPI | خطر |\n'
        '| NCA DCC | DLP | — | DLP | KPI | خطر |\n'
        '| NCA ECC | الاستجابة للحوادث | عدم وجود مركز عمليات أمنية SOC/SIEM | '
        'SOC | MTTR | خطر |\n'
    )
    return {
        'vision': '## الرؤية والأهداف\n\nأهداف استراتيجية للأمن السيبراني.\n',
        'pillars': pillars,
        'roadmap': roadmap_docx,
        'roadmap_preview': roadmap_preview,
        'kpis': kpis,
        'confidence': confidence,
        'traceability': traceability,
        'environment': 'ال ال معتمد في المنظمة لل معالجة الحوادث.\n',
    }


def _sections_to_docx_bytes(sections: Dict[str, str]) -> bytes:
    from docx import Document

    doc = Document()
    for key in (
            'vision', 'pillars', 'roadmap', 'kpis',
            'confidence', 'traceability', 'environment'):
        value = sections.get(key) or ''
        if key == 'roadmap_preview':
            continue
        doc.add_paragraph(f'[{key}]')
        for line in str(value).splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def ensure_content_quality_fixtures() -> Tuple[Path, Path]:
    """Write cyber_strategy_34.docx and cyber_strategy_61.pdf if missing."""
    sections = content_quality_defect_sections()
    # Raw DOCX preserves intentional preview/docx roadmap drift (6 vs 12 rows).
    DOCX_FIXTURE.write_bytes(_sections_to_docx_bytes(sections))
    if not PDF_FIXTURE.is_file():
        try:
            import importlib.util
            import os
            import sys

            root = Path(__file__).resolve().parents[3]
            spec = importlib.util.spec_from_file_location(
                'app', root / 'app.py')
            app = importlib.util.module_from_spec(spec)
            sys.path.insert(0, str(root))
            os.environ.setdefault('REL2_SKIP_EXPORT_EVIDENCE', '1')
            spec.loader.exec_module(app)
            backend = app._rel2_backend_callables()
            build_pdf = backend.get('build_pdf_bytes')
            build_docx = backend.get('build_docx_bytes')
            if build_docx and build_pdf:
                md = app._prcy65_rebuild_content_from_sections(sections, None)
                pdf_bytes = build_pdf(
                    md, 'strategy', 'ar', domain='cyber',
                    selected_frameworks=['NCA ECC', 'NCA DCC'])
                PDF_FIXTURE.write_bytes(pdf_bytes)
            else:
                PDF_FIXTURE.write_bytes(_sections_to_docx_bytes(sections))
        except Exception:
            PDF_FIXTURE.write_bytes(_sections_to_docx_bytes(sections))
    return DOCX_FIXTURE, PDF_FIXTURE
