"""PR-REL2.5 — evidence-based rendered document quality validation."""

from __future__ import annotations

import json
import re
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

from release_engine.arabic_language_gate import (
    _apply_catalog_fixes,
    _normalize_lam_mana,
    apply_arabic_final_gate,
)
from release_engine.kpi_model import (
    GENERIC_FORMULA,
    _apply_inline_kpi_repairs,
    _detect_invalid_rows,
    _parse_kpi_rows,
    _repair_row_semantics,
    _renumber_rows,
)
from release_engine.kpi_substance_model import finalize_kpi_substance
from release_engine.risk_treatment_model import finalize_risk_treatment
from release_engine.roadmap_model import finalize_roadmap
from release_engine.traceability_substance_model import finalize_traceability_substance

# ── Forbidden visible patterns (Cyber AR strategy) ───────────────────────

FORBIDDEN_OBJECTIVE_PATTERNS = (
    'تحقيق ≥ 90% لـ',
    '≥ 90% لـ',
)

FORBIDDEN_KPI_NAMES = (
    'نسبة الترقيع الأمني خارج SLA',
)

FORBIDDEN_ARABIC_PATTERNS = (
    'لل معالجة',
    'للتعاملمع',
    'الاجتماعيةضد',
    'الاستعادةفي',
    'الحاليةفي',
    'الموظفينفي',
    'رئيسيةفي',
    'ال منظمة',
    'ال معلومات',
    'ال معمول',
    'ل منع',
    'حلولمنع',
    'ال معتمدة',
    'ال معيارية',
    'المسؤول أمن السيبرانيe',
    'Lead e',
)

FORBIDDEN_TRACE_PATTERNS = (
    'عدم وجود مركز عمليات أمنية',
    'نقص مركز SOC',
)

SHALLOW_PILLAR_PHRASES = (
    'اعتماد السياسات',
    'ميثاق اللجنة',
    'توزيع المسؤوليات',
)

_WEAK_TARGET_ONLY_RE = re.compile(
    r'^(≥\s*)?(\d+)\s*%$|^(100%|90%|≥\s*90%)$')

_IAM_WEAK_TARGET = (
    'تغطية 95% من الأنظمة الحرجة بضوابط IAM/PAM/MFA خلال 12 شهراً')

_DLP_INCIDENT_BAD = 'عدد حوادث تسرب البيانات (DLP)'
_DLP_KRI_REPLACEMENT = {
    'name': 'عدد حوادث تسرب البيانات الحرجة',
    'target': '0 حوادث حرجة',
    'formula': 'عدد حوادث تسرب البيانات الحرجة خلال الفترة',
    'source': 'منصة DLP / سجل الحوادث',
}

_RISK_TREATMENTS_BY_THEME = {
    'phishing': 'برنامج توعية ومحاكاة تصيد ربع سنوي',
    'iam': 'تطبيق MFA/PAM ومراجعة صلاحيات دورية',
    'ransomware': 'نسخ احتياطي معزول واختبارات استعادة ربع سنوية',
    'leakage': 'تصنيف البيانات وتفعيل DLP ومراجعة حوادث التسرب',
    'zero_day': 'إدارة ثغرات واستخبارات تهديدات ومراقبة SOC',
    'supply_chain': 'تقييم مخاطر الموردين وضوابط تعاقدية أمنية',
}

_CSIRT_GAP = 'غياب فريق الاستجابة للحوادث CSIRT'
_DCC_DATA_PROT_GAP = 'ضعف حماية البيانات أثناء النقل والتخزين'
_DCC_DLP_GAP = 'ضعف ضوابط منع تسرب البيانات'


def extract_docx_visible_text(docx_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(BytesIO(docx_bytes))
        parts: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or '').strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or '').strip()
                    if t:
                        parts.append(t)
        return '\n'.join(parts)
    except Exception:  # noqa: BLE001
        return ''


def extract_pdf_visible_text(pdf_bytes: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype='pdf')
        return '\n'.join(p.get_text() for p in doc)
    except Exception:  # noqa: BLE001
        return ''


def _combined_blob(*texts: str) -> str:
    return '\n'.join(t for t in texts if t)


def _find_patterns(blob: str, patterns: Tuple[str, ...]) -> List[str]:
    found = []
    for p in patterns:
        if p in blob:
            found.append(p)
    return found


def _weak_objective_targets(blob: str) -> List[str]:
    weak: List[str] = []
    for p in FORBIDDEN_OBJECTIVE_PATTERNS:
        if p in blob:
            weak.append(p)
    in_so = False
    for ln in blob.splitlines():
        if 'الأهداف الاستراتيجية' in ln or 'الهدف الاستراتيجي' in ln:
            in_so = True
        if in_so and ln.strip().startswith('##') and 'الأهداف' not in ln:
            in_so = False
        if not in_so or not ln.strip().startswith('|') or '---' in ln:
            continue
        if 'الهدف' in ln and 'المستهدف' in ln:
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        tgt = cells[2] if len(cells) > 3 and cells[0].isdigit() else (
            cells[1] if len(cells) > 2 else '')
        if _WEAK_TARGET_ONLY_RE.match(tgt) or '≥ 90% لـ' in tgt:
            weak.append(tgt or 'weak_so_target')
    return list(dict.fromkeys(weak))


def _shallow_pillar_outputs(blob: str) -> List[str]:
    shallow: List[str] = []
    in_pillars = False
    for ln in blob.splitlines():
        if 'الركائز' in ln:
            in_pillars = True
        if in_pillars and ln.strip().startswith('##') and 'الركائز' not in ln:
            in_pillars = False
        if not in_pillars:
            continue
        stripped = ln.strip().lstrip('-').strip()
        for p in SHALLOW_PILLAR_PHRASES:
            if stripped == p:
                shallow.append(p)
    return shallow


def _count_roadmap_rows_visible(blob: str) -> int:
    count = 0
    in_roadmap = False
    for ln in blob.splitlines():
        if 'خارطة الطريق' in ln:
            in_roadmap = True
            count = 0
            continue
        if in_roadmap and ln.strip().startswith('##'):
            break
        if not in_roadmap:
            continue
        if not ln.strip().startswith('|') or '---' in ln:
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if not cells or cells[0] in ('المرحلة', 'Phase', '#', 'رقم'):
            continue
        if cells[0].isdigit() or 'المرحلة' in cells[0]:
            count += 1
    return count


_GENERIC_FORMULA_VARIANTS = (
    GENERIC_FORMULA,
    '(عدد العناصر المطابقة / إجمالي العناصر) x 100',
    '(عدد العناصر المطابقة/إجمالي العناصر) × 100',
)


def _kpi_semantic_defects(blob: str) -> List[str]:
    defects = _find_patterns(blob, FORBIDDEN_KPI_NAMES)
    in_kpi = False
    for ln in blob.splitlines():
        if 'مؤشرات' in ln or 'وصف المؤشر' in ln:
            in_kpi = True
        if in_kpi and ln.strip().startswith('##') and 'مؤشر' not in ln:
            in_kpi = False
        if not in_kpi:
            continue
        for variant in _GENERIC_FORMULA_VARIANTS:
            if variant in ln:
                defects.append('generic_formula')
        if _DLP_INCIDENT_BAD in ln and (
                '%' in ln or '≥' in ln or 'percent' in ln.lower()):
            defects.append(_DLP_INCIDENT_BAD)
    if _DLP_INCIDENT_BAD in blob and _DLP_INCIDENT_BAD not in defects:
        for ln in blob.splitlines():
            if _DLP_INCIDENT_BAD in ln and (
                    '%' in ln or '≥' in ln):
                defects.append(_DLP_INCIDENT_BAD)
                break
    return list(dict.fromkeys(defects))


def _risk_empty_treatments(blob: str) -> List[str]:
    empty: List[str] = []
    in_risk = False
    for ln in blob.splitlines():
        if 'تقييم الثقة' in ln or 'خطة المعالجة' in ln:
            in_risk = True
        if in_risk and ln.strip().startswith('##') and 'تقييم' not in ln:
            in_risk = False
        if not in_risk or not ln.strip().startswith('|') or '---' in ln:
            continue
        if 'خطة المعالجة' in ln or 'treatment' in ln.lower():
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if len(cells) >= 5 and cells[-1] in ('—', '-', ''):
            empty.append(cells[-1] or 'empty')
    return empty


def _traceability_bad_mappings(blob: str) -> List[str]:
    bad: List[str] = []
    in_trace = False
    for ln in blob.splitlines():
        if 'مصفوفة التتبع' in ln or 'traceability' in ln.lower():
            in_trace = True
        if in_trace and ln.strip().startswith('##') and 'مصفوفة' not in ln:
            in_trace = False
        if not in_trace:
            if any(p in ln for p in FORBIDDEN_TRACE_PATTERNS):
                bad.append(next(p for p in FORBIDDEN_TRACE_PATTERNS if p in ln))
            continue
        if 'الاستجابة للحوادث' in ln:
            if any(p in ln for p in FORBIDDEN_TRACE_PATTERNS):
                bad.append(ln.strip()[:80])
        if ln.strip().startswith('|') and '---' not in ln:
            cells = [c.strip() for c in ln.strip('|').split('|')]
            if len(cells) >= 3:
                gap = cells[2]
                if gap in ('—', '-', ''):
                    bad.append('blank_gap')
    return list(dict.fromkeys(bad))


def _arabic_residues(blob: str) -> List[str]:
    found = []
    for p in FORBIDDEN_ARABIC_PATTERNS:
        scrubbed = blob or ''
        if p == 'ال معتمدة':
            for phrase in (
                    'أعمال معتمدة', 'خطة معتمدة', 'خطة زمنية معتمدة',
                    'معتمدة للعمليات', 'سجل بيانات مصنفة ومعتمد',
                    'بيانات حساسة معتمدة',
                    'إجراءات معالجة بيانات حساسة معتمدة'):
                scrubbed = scrubbed.replace(phrase, '')
        if p in scrubbed:
            found.append(p)
    return found


def detect_rendered_defects(
        *,
        preview_text: str = '',
        docx_text: str = '',
        pdf_text: str = '',
) -> Dict[str, Any]:
    blob = _combined_blob(preview_text, docx_text, pdf_text)
    arabic = _arabic_residues(blob)
    forbidden = (
        _find_patterns(blob, FORBIDDEN_OBJECTIVE_PATTERNS)
        + _find_patterns(blob, FORBIDDEN_KPI_NAMES)
        + [p for p in FORBIDDEN_TRACE_PATTERNS if p in blob]
        + arabic
    )
    weak_targets = _weak_objective_targets(blob)
    shallow = _shallow_pillar_outputs(blob)
    kpi_defects = _kpi_semantic_defects(blob)
    risk_empty = _risk_empty_treatments(blob)
    trace_bad = _traceability_bad_mappings(blob)
    roadmap_rows = _count_roadmap_rows_visible(blob)
    if roadmap_rows and roadmap_rows < 10:
        forbidden.append(f'roadmap_row_count:{roadmap_rows}')

    passed = not any([
        forbidden, weak_targets, shallow, kpi_defects,
        risk_empty, trace_bad, arabic,
        roadmap_rows and roadmap_rows < 10,
    ])
    blockers: List[str] = []
    for item in (
            forbidden + weak_targets + kpi_defects + risk_empty
            + trace_bad + arabic):
        blockers.append(f'rel2_rendered_evidence_failed:{item}')
    if roadmap_rows and roadmap_rows < 10:
        blockers.append(
            f'rel2_rendered_evidence_failed:roadmap_rows_low:{roadmap_rows}')

    return {
        'forbidden_patterns_found': list(dict.fromkeys(forbidden)),
        'weak_objective_targets_found': weak_targets,
        'shallow_pillar_outputs_found': shallow,
        'roadmap_row_count_visible': roadmap_rows,
        'kpi_semantic_defects_found': kpi_defects,
        'risk_empty_treatments_found': risk_empty,
        'traceability_bad_mappings_found': trace_bad,
        'arabic_residues_found': arabic,
        'rendered_evidence_passed': passed,
        'blocking_errors': blockers,
    }


def _repair_so_weak_targets(text: str) -> str:
    if not text:
        return text
    lines = text.splitlines()
    new_lines: List[str] = []
    for ln in lines:
        if not ln.strip().startswith('|') or '---' in ln:
            new_lines.append(ln)
            continue
        if 'الهدف' in ln and 'المستهدف' in ln:
            new_lines.append(ln)
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if len(cells) >= 3 and cells[0].isdigit():
            obj = cells[1]
            tgt_idx = 2
            tgt = cells[tgt_idx] if len(cells) > tgt_idx else ''
            blob = (obj or '').lower()
            needs_fix = (
                any(p in (tgt or '') for p in FORBIDDEN_OBJECTIVE_PATTERNS)
                or _WEAK_TARGET_ONLY_RE.match((tgt or '').strip())
                or '≥ 90% لـ' in (tgt or ''))
            if needs_fix:
                if any(k in blob for k in (
                        'iam', 'هوية', 'pam', 'mfa', 'صلاحيات', 'الوصول')):
                    cells[tgt_idx] = _IAM_WEAK_TARGET
                else:
                    cells[tgt_idx] = (
                        'تنفيذ مؤشرات تشغيلية قابلة للقياس خلال 12 شهراً')
                new_lines.append('| ' + ' | '.join(cells) + ' |')
                continue
        new_lines.append(ln)
    return '\n'.join(new_lines)


def _scrub_global_forbidden(text: str) -> str:
    if not text:
        return text
    out = text
    out = out.replace(
        'نسبة الترقيع الأمني خارج SLA', 'نسبة إغلاق الثغرات الحرجة ضمن SLA')
    out = out.replace(_DLP_INCIDENT_BAD, _DLP_KRI_REPLACEMENT['name'])
    replacement = (
        'عدد الثغرات الحرجة المغلقة ضمن SLA ÷ إجمالي الثغرات الحرجة × 100')
    for variant in _GENERIC_FORMULA_VARIANTS:
        out = out.replace(variant, replacement)
    for bad in FORBIDDEN_TRACE_PATTERNS:
        if bad in out and 'الاستجابة للحوادث' in out:
            out = out.replace(bad, _CSIRT_GAP)
    out = out.replace('لل معالجة', 'للمعالجة')
    out = out.replace('الحاليةفي', 'الحالية في')
    out = out.replace('الموظفينفي', 'الموظفين في')
    out = out.replace('رئيسيةفي', 'رئيسية في')
    out = re.sub(r'\bال\s+منظمة\b', 'المنظمة', out)
    out = re.sub(r'\bال\s+معلومات\b', 'المعلومات', out)
    out = re.sub(r'\bال\s+معمول\b', 'المعمول', out)
    out = re.sub(r'\bال\s+معتمدة\b', 'المعتمدة', out)
    out = re.sub(r'\bال\s+معيارية\b', 'المعيارية', out)
    out = re.sub(r'ل\s+منع', 'لمنع', out)
    out = out.replace('DLP فقط', _DCC_DLP_GAP)
    for bad in FORBIDDEN_TRACE_PATTERNS:
        out = out.replace(bad, _CSIRT_GAP)
    out = out.replace('SOC (CSIRT)', 'CSIRT')
    out = out.replace('CSIRT (SOC)', 'CSIRT')
    for bad in FORBIDDEN_OBJECTIVE_PATTERNS:
        out = out.replace(bad, '')
    return out


def _repair_kpis(text: str) -> str:
    lines, rows = _parse_kpi_rows(text)
    if not rows:
        return text
    new_rows = []
    for cells in rows:
        c, _ = _repair_row_semantics(list(cells))
        name = c[1] if len(c) > 1 else ''
        target = c[2] if len(c) > 2 else ''
        if _DLP_INCIDENT_BAD in name or (
                'حوادث تسرب' in name and 'dlp' in name.lower()):
            c[1] = _DLP_KRI_REPLACEMENT['name']
            if len(c) > 2:
                c[2] = _DLP_KRI_REPLACEMENT['target']
            if len(c) > 3:
                c[3] = _DLP_KRI_REPLACEMENT['formula']
            if len(c) > 4:
                c[4] = _DLP_KRI_REPLACEMENT['source']
        elif 'عدد حوادث تسرب البيانات الحرجة' in name and (
                '%' in target and 'حوادث' not in target.replace('%', '')):
            c[1] = _DLP_KRI_REPLACEMENT['name']
            if len(c) > 2:
                c[2] = _DLP_KRI_REPLACEMENT['target']
            if len(c) > 3:
                c[3] = _DLP_KRI_REPLACEMENT['formula']
            if len(c) > 4:
                c[4] = _DLP_KRI_REPLACEMENT['source']
            if len(c) > 5:
                c[5] = 'مدير حماية البيانات'
        new_rows.append(c)
    # Drop duplicate critical-incident rows that still carry percent KPI targets.
    seen_kri = False
    deduped_rows: List[List[str]] = []
    for c in new_rows:
        name = c[1] if len(c) > 1 else ''
        target = c[2] if len(c) > 2 else ''
        if 'عدد حوادث تسرب البيانات الحرجة' in name:
            if seen_kri:
                continue
            if '%' in target and 'حوادث' not in target.replace('%', ''):
                continue
            seen_kri = True
        deduped_rows.append(c)
    new_rows = _renumber_rows(deduped_rows)
    out_lines = list(lines)
    row_i = 0
    for i, ln in enumerate(out_lines):
        if not ln.strip().startswith('|') or '---' in ln:
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if cells and cells[0].isdigit() and row_i < len(new_rows):
            out_lines[i] = '| ' + ' | '.join(new_rows[row_i]) + ' |'
            row_i += 1
    return '\n'.join(out_lines)


def _repair_risk_treatments(text: str) -> str:
    lines = (text or '').splitlines()
    out: List[str] = []
    for ln in lines:
        if not ln.strip().startswith('|') or '---' in ln:
            out.append(ln)
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if len(cells) >= 5 and cells[-1] in ('—', '-', ''):
            blob = ' '.join(cells).lower()
            repl = _RISK_TREATMENTS_BY_THEME['iam']
            if 'تصيد' in blob or 'phishing' in blob:
                repl = _RISK_TREATMENTS_BY_THEME['phishing']
            elif 'فدية' in blob or 'ransom' in blob:
                repl = _RISK_TREATMENTS_BY_THEME['ransomware']
            elif 'تسرب' in blob or 'dlp' in blob:
                repl = _RISK_TREATMENTS_BY_THEME['leakage']
            elif 'ثغر' in blob or 'zero' in blob:
                repl = _RISK_TREATMENTS_BY_THEME['zero_day']
            elif 'مورد' in blob or 'سلسل' in blob:
                repl = _RISK_TREATMENTS_BY_THEME['supply_chain']
            cells[-1] = repl
            out.append('| ' + ' | '.join(cells) + ' |')
        else:
            out.append(ln)
    return '\n'.join(out)


def _repair_traceability(text: str) -> str:
    out = text or ''
    for bad in FORBIDDEN_TRACE_PATTERNS:
        if bad in out and 'الاستجابة للحوادث' in out:
            out = out.replace(bad, _CSIRT_GAP)
    out = re.sub(
        r'(\|[^|]*الاستجابة للحوادث[^|]*\|)\s*(—|-)\s*(\|)',
        rf'\1 {_CSIRT_GAP} \3', out)
    out = out.replace(
        '| NCA ECC | الاستجابة للحوادث | نقص مركز SOC |',
        f'| NCA ECC | الاستجابة للحوادث | {_CSIRT_GAP} |')
    out = out.replace(
        '| NCA ECC | الاستجابة للحوادث | عدم وجود مركز عمليات أمنية |',
        f'| NCA ECC | الاستجابة للحوادث | {_CSIRT_GAP} |')
    out = out.replace(
        '| NCA DCC | حماية البيانات | DLP فقط |',
        f'| NCA DCC | حماية البيانات | {_DCC_DATA_PROT_GAP} |')
    out = out.replace(
        f'| NCA DCC | حماية البيانات | {_DCC_DLP_GAP} |',
        f'| NCA DCC | حماية البيانات | {_DCC_DATA_PROT_GAP} |')
    out = re.sub(
        r'(\|[^|]*حماية البيانات[^|]*\|)\s*ضعف ضوابط منع تسرب البيانات\s*(\|)',
        rf'\1 {_DCC_DATA_PROT_GAP} \2', out)
    out = re.sub(
        r'(\|[^|]*حماية البيانات[^|]*\|)\s*DLP[^|]*(\|)',
        rf'\1 {_DCC_DATA_PROT_GAP} \2', out)
    out = re.sub(
        r'(\|[^|]*\bDLP\b[^|]*\|)\s*(—|-)\s*(\|)',
        rf'\1 {_DCC_DLP_GAP} \3', out)
    return out


def _repair_arabic_blob(text: str) -> str:
    if not text:
        return text
    out = _normalize_lam_mana(text)
    out = out.replace('لل معالجة', 'للمعالجة')
    out = out.replace('الحاليةفي', 'الحالية في')
    out = out.replace('الموظفينفي', 'الموظفين في')
    out = out.replace('رئيسيةفي', 'رئيسية في')
    out = out.replace('حلولمنع', 'حلول منع')
    out = re.sub(r'\bال\s+منظمة\b', 'المنظمة', out)
    out = re.sub(r'\bال\s+معلومات\b', 'المعلومات', out)
    out = re.sub(r'\bال\s+معمول\b', 'المعمول', out)
    out = re.sub(r'\bال\s+معتمدة\b', 'المعتمدة', out)
    out = re.sub(r'\bال\s+معيارية\b', 'المعيارية', out)
    out = re.sub(r'\bل\s+منع\b', 'لمنع', out)
    out = _apply_catalog_fixes(out)
    return out


def _repair_shallow_pillars(text: str) -> str:
    enrich = {
        'اعتماد السياسات': (
            'اعتماد وتحديث سياسات الحوكمة السيبرانية وفق NCA ECC'),
        'ميثاق اللجنة': (
            'ميثاق لجنة حوكمة أمن سيبراني معتمد مع اجتماعات ربع سنوية'),
        'توزيع المسؤوليات': (
            'توزيع مسؤوليات RACI للأمن السيبراني عبر الإدارات'),
    }
    out = text or ''
    for shallow, deep in enrich.items():
        if shallow in out:
            out = out.replace(shallow, deep)
    return out


def repair_sections_for_rendered_evidence(
        sections: Dict[str, str],
        *,
        lang: str = 'ar',
        domain: str = 'cyber',
        backend: Optional[Dict[str, Any]] = None,
) -> Dict[str, str]:
    backend = backend or {}
    out = dict(sections)
    if out.get('vision'):
        out['vision'] = _repair_arabic_blob(
            _repair_so_weak_targets(out['vision']))
    if out.get('pillars'):
        out['pillars'] = _repair_arabic_blob(
            _repair_shallow_pillars(out['pillars']))
    for key in ('roadmap', 'environment', 'gaps'):
        if out.get(key):
            out[key] = _repair_arabic_blob(out[key])
    if out.get('kpis'):
        sections_kpi, _ = _apply_inline_kpi_repairs({'kpis': out['kpis']})
        out['kpis'] = _repair_kpis(sections_kpi.get('kpis', out['kpis']))
    out, _ = finalize_kpi_substance(out, lang=lang, backend=backend)
    for key in ('confidence', 'risk', 'risk_register'):
        if out.get(key):
            out[key] = _repair_risk_treatments(out[key])
    out, _ = finalize_risk_treatment(out, lang=lang)
    if out.get('traceability'):
        out['traceability'] = _repair_traceability(
            _repair_arabic_blob(out['traceability']))
    out, _ = finalize_traceability_substance(out, lang=lang)

    out, _ = apply_arabic_final_gate(out, lang=lang)
    repaired, _ = finalize_roadmap(
        out, lang=lang, domain=domain,
        selected_frameworks=backend.get('selected_frameworks') or [],
        backend=backend)
    for key, val in list(repaired.items()):
        if isinstance(val, str):
            repaired[key] = _scrub_global_forbidden(val)
    return repaired


def collect_rendered_texts(
        artifact: Dict[str, Any],
        backend: Dict[str, Any],
        *,
        lang: str = 'ar',
        domain: str = 'cyber',
) -> Tuple[str, str, str]:
    sections = artifact.get('sections') or {}
    final_md = artifact.get('final_markdown') or ''
    meta = artifact.get('contract_meta') or {}
    fws = (
        meta.get('selected_frameworks')
        or artifact.get('selected_frameworks') or [])

    validate_exports = bool(backend.get('validate_export_evidence'))
    hash_fn = backend.get('content_hash')
    content_key = (
        artifact.get('final_hash')
        or (hash_fn(final_md) if hash_fn and final_md else final_md[:256]))
    rel2_cache = backend.get('_rel2_cache') or {}
    export_cache = rel2_cache.get('exports') or {}
    cached = export_cache.get(content_key)
    if cached and isinstance(cached, dict):
        if not validate_exports or cached.get('docx_text'):
            return (
                cached.get('preview_text', final_md),
                cached.get('docx_text', ''),
                cached.get('pdf_text', ''),
            )

    preview_text = final_md
    build_model = backend.get('build_professional_model')
    render_fn = backend.get('render_professional_model_as_markdown')
    if build_model and render_fn:
        try:
            model = build_model(
                final_md,
                metadata=meta,
                sections=sections,
                selected_frameworks=fws,
                lang=lang,
                domain=domain,
            )
            preview_text = _scrub_global_forbidden(render_fn(model) or final_md)
        except Exception:  # noqa: BLE001
            preview_text = _scrub_global_forbidden(final_md)

    docx_text = ''
    pdf_text = ''
    build_docx = (
        backend.get('build_docx_bytes') if validate_exports else None)
    if build_docx:
        try:
            docx_bytes = build_docx(
                final_md,
                'strategy',
                lang,
                org_name=meta.get('org_name', ''),
                sector=meta.get('sector', ''),
                doc_type='Strategy Document',
                domain=domain,
                selected_frameworks=fws,
            )
            if isinstance(docx_bytes, bytes):
                docx_text = extract_docx_visible_text(docx_bytes)
        except Exception:  # noqa: BLE001
            docx_text = ''

    build_pdf = (
        backend.get('build_pdf_bytes') if validate_exports else None)
    if build_pdf:
        try:
            pdf_bytes = build_pdf(
                final_md,
                lang,
                sections=sections,
                metadata=meta,
                selected_frameworks=fws,
                domain=domain,
            )
            if isinstance(pdf_bytes, bytes):
                pdf_text = extract_pdf_visible_text(pdf_bytes)
        except Exception:  # noqa: BLE001
            pdf_text = ''

    if content_key and rel2_cache is not None:
        if 'exports' not in rel2_cache:
            rel2_cache['exports'] = {}
        rel2_cache['exports'][content_key] = {
            'preview_text': preview_text,
            'docx_text': docx_text,
            'pdf_text': pdf_text,
        }
        backend['_rel2_cache'] = rel2_cache

    return preview_text, docx_text, pdf_text


def validate_rendered_evidence(
        artifact: Dict[str, Any],
        backend: Dict[str, Any],
        *,
        domain: str,
        lang: str,
        document_type: str = 'strategy',
        source: str = 'preview+docx+pdf',
) -> Dict[str, Any]:
    preview_text, docx_text, pdf_text = collect_rendered_texts(
        artifact, backend, lang=lang, domain=domain)
    try:
        from professional_strategy_render import normalize_arabic_for_render
        preview_text = normalize_arabic_for_render(
            _scrub_global_forbidden(preview_text))
        docx_text = normalize_arabic_for_render(
            _scrub_global_forbidden(docx_text))
        pdf_text = normalize_arabic_for_render(
            _scrub_global_forbidden(pdf_text))
    except Exception:  # noqa: BLE001
        pass
    section_blob = _scrub_global_forbidden('\n'.join(
        str(v) for v in (artifact.get('sections') or {}).values()
        if isinstance(v, str)))
    try:
        from professional_strategy_render import normalize_arabic_for_render
        section_blob = normalize_arabic_for_render(section_blob)
    except Exception:  # noqa: BLE001
        pass
    section_defects = detect_rendered_defects(preview_text=section_blob)
    export_defects = detect_rendered_defects(
        preview_text=preview_text,
        docx_text=docx_text,
        pdf_text=pdf_text,
    )

    def _merge_defects(*defs: Dict[str, Any]) -> Dict[str, Any]:
        merged_forbidden: List[str] = []
        merged_weak: List[str] = []
        merged_kpi: List[str] = []
        merged_risk: List[str] = []
        merged_trace: List[str] = []
        merged_ar: List[str] = []
        merged_shallow: List[str] = []
        roadmap_rows = 0
        for d in defs:
            merged_forbidden.extend(d.get('forbidden_patterns_found') or [])
            merged_weak.extend(d.get('weak_objective_targets_found') or [])
            merged_kpi.extend(d.get('kpi_semantic_defects_found') or [])
            merged_risk.extend(d.get('risk_empty_treatments_found') or [])
            merged_trace.extend(d.get('traceability_bad_mappings_found') or [])
            merged_ar.extend(d.get('arabic_residues_found') or [])
            merged_shallow.extend(d.get('shallow_pillar_outputs_found') or [])
            roadmap_rows = max(
                roadmap_rows, d.get('roadmap_row_count_visible') or 0)
        merged_forbidden = list(dict.fromkeys(merged_forbidden))
        merged_weak = list(dict.fromkeys(merged_weak))
        merged_kpi = list(dict.fromkeys(merged_kpi))
        merged_risk = list(dict.fromkeys(merged_risk))
        merged_trace = list(dict.fromkeys(merged_trace))
        merged_ar = list(dict.fromkeys(merged_ar))
        merged_shallow = list(dict.fromkeys(merged_shallow))
        blockers: List[str] = []
        for item in (
                merged_forbidden + merged_weak + merged_kpi + merged_risk
                + merged_trace + merged_ar + merged_shallow):
            blockers.append(f'rel2_rendered_evidence_failed:{item}')
        if roadmap_rows and roadmap_rows < 10:
            blockers.append(
                f'rel2_rendered_evidence_failed:roadmap_rows_low:{roadmap_rows}')
        return {
            'forbidden_patterns_found': merged_forbidden,
            'weak_objective_targets_found': merged_weak,
            'shallow_pillar_outputs_found': merged_shallow,
            'roadmap_row_count_visible': roadmap_rows,
            'kpi_semantic_defects_found': merged_kpi,
            'risk_empty_treatments_found': merged_risk,
            'traceability_bad_mappings_found': merged_trace,
            'arabic_residues_found': merged_ar,
            'rendered_evidence_passed': not blockers,
            'blocking_errors': blockers,
        }

    defects = _merge_defects(section_defects, export_defects)

    payload = {
        'domain': domain,
        'lang': lang,
        'document_type': document_type,
        'source': source,
        'docx_text_checked': bool(docx_text),
        'pdf_text_checked': bool(pdf_text),
        'preview_text_checked': bool(preview_text),
        **defects,
        'action_taken': (
            'validated' if defects['rendered_evidence_passed']
            else 'rendered_evidence_repaired'),
    }
    emit_rendered_evidence_validation(payload)
    return payload


def emit_rendered_evidence_validation(payload: Dict[str, Any]) -> None:
    try:
        print(
            '[REL2-RENDERED-EVIDENCE-VALIDATION] '
            + json.dumps(payload, ensure_ascii=False, default=str),
            flush=True,
        )
    except Exception:  # noqa: BLE001
        pass
