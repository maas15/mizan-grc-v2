"""
Mizan GRC - Enterprise Governance, Risk & Compliance Platform
Flask Application - Full RTL Support
Created by: Eng. Mohammad Abbas Alsaadon
"""

import os
import json
import sqlite3
import hashlib
import secrets
import re
import html
import uuid
from contextlib import closing
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timedelta
from functools import wraps
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, flash, abort, Response, g

from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Stable secret key — persists across restarts to keep sessions alive
def _get_stable_secret_key():
    """Get or create a stable secret key that survives server restarts."""
    env_key = os.getenv('SECRET_KEY')
    if env_key:
        return env_key
    key_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.secret_key')
    try:
        if os.path.exists(key_file):
            with open(key_file, 'r') as f:
                key = f.read().strip()
                if len(key) >= 32:
                    return key
        # Generate and persist a new key
        key = secrets.token_hex(32)
        with open(key_file, 'w') as f:
            f.write(key)
        os.chmod(key_file, 0o600)
        return key
    except Exception:
        return secrets.token_hex(32)  # Final fallback

app.secret_key = _get_stable_secret_key()
app.permanent_session_lifetime = timedelta(days=7)

# Security configurations
app.config['SESSION_COOKIE_SECURE'] = os.getenv('FLASK_ENV') == 'production'
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Thread pool for background AI tasks (prevents OOM from unbounded thread spawning)
ai_executor = ThreadPoolExecutor(max_workers=5, thread_name_prefix='mizan-ai')

# CSRF Protection
def generate_csrf_token():
    if 'csrf_token' not in session:
        session['csrf_token'] = secrets.token_hex(32)
    return session['csrf_token']

def csrf_token():
    return generate_csrf_token()

# Make csrf_token available in all templates
app.jinja_env.globals['csrf_token'] = csrf_token

# ============================================================================
# BACKGROUND TASK SYSTEM for long-running AI operations
# ============================================================================
# Render has a 30s request timeout + gunicorn uses multiple workers.
# We store task state in SQLite so any worker can read/write it.

def create_background_task(task_id, user_id, domain):
    """Create a new pending task in the database."""
    try:
        with closing(get_db_direct()) as conn:
            conn.execute(
                'INSERT INTO background_tasks (task_id, user_id, status, callback_domain) VALUES (?, ?, ?, ?)',
                (task_id, user_id, 'pending', domain)
            )
            conn.commit()
    except Exception as e:
        print(f"Task create error: {e}", flush=True)

def complete_background_task(task_id, result):
    """Mark task as done with result."""
    try:
        with closing(get_db_direct()) as conn:
            conn.execute(
                'UPDATE background_tasks SET status = ?, result = ? WHERE task_id = ?',
                ('done', result, task_id)
            )
            conn.commit()
    except Exception as e:
        print(f"Task complete error: {e}", flush=True)

def fail_background_task(task_id, error):
    """Mark task as failed."""
    try:
        with closing(get_db_direct()) as conn:
            conn.execute(
                'UPDATE background_tasks SET status = ?, error = ? WHERE task_id = ?',
                ('error', error, task_id)
            )
            conn.commit()
    except Exception as e:
        print(f"Task fail error: {e}", flush=True)

def get_background_task(task_id):
    """Get task status from database."""
    try:
        with closing(get_db_direct()) as conn:
            task = conn.execute(
                'SELECT task_id, user_id, status, result, error, callback_domain FROM background_tasks WHERE task_id = ?',
                (task_id,)
            ).fetchone()
            return task
    except Exception as e:
        print(f"Task get error: {e}", flush=True)
        return None

def delete_background_task(task_id):
    """Remove completed task."""
    try:
        with closing(get_db_direct()) as conn:
            conn.execute('DELETE FROM background_tasks WHERE task_id = ?', (task_id,))
            conn.commit()
    except Exception:
        pass

def run_ai_task(task_id, prompt, lang, content_type=None):
    """Run AI generation in a background thread, store result in DB."""
    try:
        result = generate_ai_content(prompt, lang, content_type=content_type)
        complete_background_task(task_id, result)
        print(f"✅ Background task {task_id[:8]} completed ({len(result)} chars)", flush=True)
    except Exception as e:
        fail_background_task(task_id, str(e))
        print(f"❌ Background task {task_id[:8]} failed: {e}", flush=True)

@app.before_request
def csrf_protect():
    """Validate CSRF token for state-changing POST/PUT/DELETE requests."""
    if request.method in ("POST", "PUT", "DELETE"):
        # Read-only POST endpoints (downloads, exports) — exempt from CSRF
        # These don't change server state; they just return generated files.
        # Login/register are pre-authentication — no session to protect.
        csrf_exempt_prefixes = (
            '/api/generate-docx',
            '/api/generate-pdf',
            '/api/generate-excel',
            '/api/export-all',
            '/api/task-status/',
            '/login',
            '/register',
        )
        if request.path.startswith(csrf_exempt_prefixes):
            return
        
        token = session.get('csrf_token', None)
        
        # If session has no token yet (first request after login/new session),
        # generate one and skip validation this once — the NEXT request will validate.
        if not token:
            generate_csrf_token()
            return
        
        # Accept token from either header (AJAX) or form field (traditional form)
        req_token = request.headers.get('X-CSRFToken', '') or request.form.get('csrf_token', '')
        
        if not req_token or req_token != token:
            print(f"CSRF REJECT: method={request.method} path={request.path} "
                  f"has_header={'yes' if request.headers.get('X-CSRFToken') else 'no'} "
                  f"has_form={'yes' if request.form.get('csrf_token') else 'no'} "
                  f"content_type={request.content_type}", flush=True)
            abort(403)

@app.after_request
def add_security_headers(response):
    """Add security headers to all responses."""
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    if os.getenv('FLASK_ENV') == 'production':
        response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    return response

# Input sanitization
def sanitize_input(text, max_length=1000):
    """Sanitize user input to prevent XSS and injection attacks."""
    if not text:
        return ''
    # Convert to string and limit length
    text = str(text)[:max_length]
    # HTML escape
    text = html.escape(text)
    # Remove potential SQL injection patterns
    dangerous_patterns = ['--', ';--', '/*', '*/', 'xp_', 'UNION', 'SELECT', 'DROP', 'DELETE', 'INSERT', 'UPDATE']
    for pattern in dangerous_patterns:
        text = re.sub(re.escape(pattern), '', text, flags=re.IGNORECASE)
    return text.strip()

def validate_username(username):
    """Validate username format."""
    if not username or not re.match(r'^[a-zA-Z0-9_]{3,50}$', username):
        return False
    return True

def validate_email(email):
    """Validate email format."""
    if not email:
        return True  # Email is optional
    return bool(re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email))

def validate_password(password):
    """Validate password strength."""
    if len(password) < 8:
        return False, 'Password must be at least 8 characters'
    if not re.search(r'[A-Z]', password):
        return False, 'Password must contain uppercase letter'
    if not re.search(r'[a-z]', password):
        return False, 'Password must contain lowercase letter'
    if not re.search(r'[0-9]', password):
        return False, 'Password must contain a number'
    return True, ''

# Simple rate limiting (in-memory, resets on restart)
rate_limit_store = {}

def check_rate_limit(key, max_requests=10, window_seconds=60):
    """Check if request is within rate limit."""
    now = datetime.now()
    if key not in rate_limit_store:
        rate_limit_store[key] = []
    
    # Remove old entries
    rate_limit_store[key] = [t for t in rate_limit_store[key] if (now - t).seconds < window_seconds]
    
    if len(rate_limit_store[key]) >= max_requests:
        return False
    
    rate_limit_store[key].append(now)
    return True

# Configuration
class Config:
    APP_NAME = "Mizan"
    APP_VERSION = "3.0.0"
    APP_TAGLINE = "Governance • Risk • Compliance"
    CREATOR_NAME = "Eng. Mohammad Abbas Alsaadon"
    CREATOR_NAME_AR = "المهندس: محمد بن عباس السعدون"
    CREATOR_TITLE = "Consultant/Expert"
    COPYRIGHT_YEAR = "2026"
    DB_PATH = "mizan.db"
    OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', '')
    ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY', '')
    GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY', '')
    GROQ_API_KEY = os.getenv('GROQ_API_KEY', '')  # Groq (Llama 3.1, Mistral, etc.)
    AI_PROVIDER = os.getenv('AI_PROVIDER', 'auto')  # auto, openai, anthropic, google, groq
    AI_MODEL = os.getenv('AI_MODEL', '')  # Override model name
    # Email settings (using SMTP)
    SMTP_SERVER = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
    SMTP_PORT = int(os.getenv('SMTP_PORT', '587'))
    SMTP_USERNAME = os.getenv('SMTP_USERNAME', '')
    SMTP_PASSWORD = os.getenv('SMTP_PASSWORD', '')
    SMTP_FROM_EMAIL = os.getenv('SMTP_FROM_EMAIL', 'noreply@mizan.app')
    SMTP_USE_TLS = os.getenv('SMTP_USE_TLS', 'true').lower() == 'true'
    # Production domain — used for absolute URLs in emails, shared links, etc.
    PRODUCTION_DOMAIN = os.getenv('PRODUCTION_DOMAIN', 'https://mizan.app')
    # Password reset token expiry (minutes)
    PASSWORD_RESET_EXPIRY = int(os.getenv('PASSWORD_RESET_EXPIRY', '30'))

config = Config()

def send_otp_email(to_email, otp_code, doc_title, shared_by):
    """Send OTP code via email."""
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    
    if not config.SMTP_USERNAME or not config.SMTP_PASSWORD:
        print("SMTP not configured, OTP:", otp_code, flush=True)
        return False, "Email service not configured"
    
    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = f'Your Mizan Document Access Code - {otp_code}'
        msg['From'] = config.SMTP_FROM_EMAIL
        msg['To'] = to_email
        
        # Plain text version
        text = f"""
Your Mizan Document Access Code

You have been granted access to a document shared by {shared_by}.

Document: {doc_title}

Your access code is: {otp_code}

This code expires in 10 minutes.

If you did not request this, please ignore this email.

---
Mizan GRC Platform
        """
        
        # HTML version
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ font-family: Arial, sans-serif; background: #f5f5f5; padding: 20px; }}
        .container {{ max-width: 500px; margin: 0 auto; background: white; border-radius: 10px; padding: 30px; box-shadow: 0 4px 15px rgba(0,0,0,0.1); }}
        .header {{ text-align: center; margin-bottom: 30px; }}
        .logo {{ font-size: 28px; font-weight: bold; background: linear-gradient(135deg, #667eea, #764ba2); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }}
        .otp-box {{ background: linear-gradient(135deg, #667eea, #764ba2); color: white; font-size: 32px; letter-spacing: 8px; text-align: center; padding: 20px; border-radius: 10px; margin: 20px 0; }}
        .info {{ color: #666; margin: 15px 0; }}
        .footer {{ text-align: center; color: #999; font-size: 12px; margin-top: 30px; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="logo">MIZAN</div>
            <p>Secure Document Access</p>
        </div>
        <p>You have been granted access to a document shared by <strong>{shared_by}</strong>.</p>
        <p class="info"><strong>Document:</strong> {doc_title}</p>
        <p>Your access code is:</p>
        <div class="otp-box">{otp_code}</div>
        <p class="info">⏱️ This code expires in 10 minutes.</p>
        <p class="info">If you did not request this, please ignore this email.</p>
        <div class="footer">
            <p>Mizan GRC Platform</p>
        </div>
    </div>
</body>
</html>
        """
        
        part1 = MIMEText(text, 'plain')
        part2 = MIMEText(html, 'html')
        msg.attach(part1)
        msg.attach(part2)
        
        with smtplib.SMTP(config.SMTP_SERVER, config.SMTP_PORT) as server:
            server.starttls()
            server.login(config.SMTP_USERNAME, config.SMTP_PASSWORD)
            server.sendmail(config.SMTP_FROM_EMAIL, to_email, msg.as_string())
        
        return True, "Email sent"
    except Exception as e:
        print(f"Email error: {e}", flush=True)
        return False, str(e)

# ============================================================================
# DATABASE
# ============================================================================

def get_db():
    """Get database connection (request-scoped, auto-closed at request end)."""
    if '_database' not in g:
        g._database = sqlite3.connect(config.DB_PATH)
        g._database.row_factory = sqlite3.Row
    return g._database

@app.teardown_appcontext
def close_db(exception):
    """Auto-close DB connection at end of every request."""
    db = g.pop('_database', None)
    if db is not None:
        db.close()

def get_db_direct():
    """Get a standalone DB connection for background threads (no request context).
    Caller MUST use: with closing(get_db_direct()) as conn: ..."""
    conn = sqlite3.connect(config.DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def hash_password(password):
    """Hash password with salt."""
    salt = secrets.token_hex(16)
    hash_obj = hashlib.pbkdf2_hmac('sha256', password.encode(), salt.encode(), 100000)
    return f"{salt}${hash_obj.hex()}"

def verify_password(password, stored_hash):
    """Verify password against stored hash."""
    try:
        salt, hash_value = stored_hash.split('$')
        hash_obj = hashlib.pbkdf2_hmac('sha256', password.encode(), salt.encode(), 100000)
        return hash_obj.hex() == hash_value
    except:
        return False

def init_db():
    """Initialize database tables."""
    conn = get_db_direct()
    cursor = conn.cursor()
    
    # Users table with role and email
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE,
            password_hash TEXT NOT NULL,
            role TEXT DEFAULT 'user',
            is_active INTEGER DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            last_login TIMESTAMP
        )
    ''')
    
    # Add columns if they don't exist (for existing databases)
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN email TEXT UNIQUE')
    except:
        pass
    
    # AI preference columns
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN ai_provider_generate TEXT DEFAULT "auto"')
    except:
        pass
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN ai_provider_review TEXT DEFAULT "auto"')
    except:
        pass
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN role TEXT DEFAULT "user"')
    except:
        pass
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN is_active INTEGER DEFAULT 1')
    except:
        pass
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN last_login TIMESTAMP')
    except:
        pass
    
    # Strategies table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS strategies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            domain TEXT,
            org_name TEXT,
            sector TEXT,
            content TEXT,
            language TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Policies table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS policies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            domain TEXT,
            policy_name TEXT,
            framework TEXT,
            content TEXT,
            language TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Audits table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS audits (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            domain TEXT,
            framework TEXT,
            scope TEXT,
            content TEXT,
            language TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Risks table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS risks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            domain TEXT,
            asset_name TEXT,
            threat TEXT,
            risk_level TEXT,
            analysis TEXT,
            language TEXT DEFAULT 'en',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Add language column to risks if not exists (for existing databases)
    try:
        cursor.execute('ALTER TABLE risks ADD COLUMN language TEXT DEFAULT "en"')
    except:
        pass
    
    # Shared documents table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS shared_documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            share_id TEXT UNIQUE NOT NULL,
            user_id INTEGER,
            doc_type TEXT NOT NULL,
            doc_id INTEGER NOT NULL,
            title TEXT,
            domain TEXT,
            content TEXT,
            language TEXT DEFAULT 'en',
            view_count INTEGER DEFAULT 0,
            is_active INTEGER DEFAULT 1,
            requires_otp INTEGER DEFAULT 0,
            recipient_email TEXT,
            otp_code TEXT,
            otp_verified INTEGER DEFAULT 0,
            otp_expires_at TIMESTAMP,
            expires_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Industry benchmarks table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS benchmarks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            sector TEXT NOT NULL,
            sector_ar TEXT,
            compliance_score_avg REAL DEFAULT 60,
            maturity_level_avg REAL DEFAULT 2.5,
            risk_coverage_avg REAL DEFAULT 50,
            policy_count_avg INTEGER DEFAULT 5,
            audit_count_avg INTEGER DEFAULT 3,
            risk_assessment_avg INTEGER DEFAULT 4,
            source TEXT,
            source_year INTEGER DEFAULT 2024,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Insert default benchmark data if empty
    existing = cursor.execute('SELECT COUNT(*) FROM benchmarks').fetchone()[0]
    if existing == 0:
        default_benchmarks = [
            ('Government', 'حكومي', 62, 2.5, 55, 6, 3, 4, 'NCA Annual Report', 2024),
            ('Banking/Finance', 'بنوك/مالي', 78, 3.8, 72, 12, 6, 8, 'SAMA CSF Report', 2024),
            ('Healthcare', 'رعاية صحية', 58, 2.3, 48, 5, 2, 3, 'MOH Compliance Survey', 2024),
            ('Energy', 'طاقة', 70, 3.2, 65, 8, 4, 6, 'SEC Industry Report', 2024),
            ('Telecom', 'اتصالات', 72, 3.4, 68, 9, 5, 7, 'CITC Benchmark Study', 2024),
            ('Retail', 'تجزئة', 52, 2.0, 42, 4, 2, 3, 'Industry Average Estimate', 2024),
            ('Manufacturing', 'تصنيع', 55, 2.2, 45, 4, 2, 3, 'Industry Average Estimate', 2024),
        ]
        cursor.executemany('''
            INSERT INTO benchmarks (sector, sector_ar, compliance_score_avg, maturity_level_avg, 
                                   risk_coverage_avg, policy_count_avg, audit_count_avg, 
                                   risk_assessment_avg, source, source_year)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', default_benchmarks)
    
    # Background tasks table for async AI operations
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS background_tasks (
            task_id TEXT PRIMARY KEY,
            user_id INTEGER,
            status TEXT DEFAULT 'pending',
            result TEXT,
            error TEXT,
            callback_domain TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Clean up old tasks on startup
    cursor.execute("DELETE FROM background_tasks WHERE created_at < datetime('now', '-1 hour')")
    
    # ── KPIs table for Silent Pipe integration ──
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS kpis (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            domain TEXT DEFAULT 'General',
            name TEXT NOT NULL,
            description TEXT,
            current_value TEXT DEFAULT 'TBD',
            target_value TEXT,
            timeframe TEXT,
            data_source TEXT,
            status TEXT DEFAULT 'pending',
            source_type TEXT DEFAULT 'manual',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    ''')
    
    # ERM Risk Register table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS risk_register (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            name TEXT NOT NULL,
            description TEXT,
            category TEXT DEFAULT 'Operational',
            likelihood INTEGER DEFAULT 3,
            impact INTEGER DEFAULT 3,
            owner TEXT,
            treatment TEXT DEFAULT 'Mitigate',
            treatment_plan TEXT,
            status TEXT DEFAULT 'Open',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    ''')
    
    # Awareness module scores
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS awareness_scores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            domain TEXT NOT NULL,
            module_id TEXT NOT NULL,
            score INTEGER DEFAULT 0,
            total INTEGER DEFAULT 0,
            passed INTEGER DEFAULT 0,
            language TEXT DEFAULT 'en',
            completed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    ''')
    
    # Project management tasks
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS project_tasks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            domain TEXT DEFAULT 'General',
            title TEXT NOT NULL,
            description TEXT,
            status TEXT DEFAULT 'todo',
            priority TEXT DEFAULT 'medium',
            owner TEXT,
            due_date TEXT,
            category TEXT DEFAULT 'implementation',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    ''')
    
    # Compliance score history for trend tracking
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS compliance_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            score REAL DEFAULT 0,
            maturity_avg REAL DEFAULT 0,
            strategies INTEGER DEFAULT 0,
            policies INTEGER DEFAULT 0,
            audits INTEGER DEFAULT 0,
            risks INTEGER DEFAULT 0,
            domains_covered INTEGER DEFAULT 0,
            recorded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    ''')
    
    # ── NEW: Pillar 1 — Audit Logs ──
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS audit_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            action TEXT NOT NULL,
            metadata TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    ''')
    
    # ── NEW: Pillar 3 — GRC Frameworks (replaces hardcoded dict) ──
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS grc_frameworks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            domain TEXT NOT NULL,
            region TEXT NOT NULL,
            name TEXT NOT NULL,
            description TEXT,
            is_active INTEGER DEFAULT 1,
            sort_order INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(domain, region, name)
        )
    ''')
    
    # ── NEW: Pillar 4 — Form Drafts (Save Progress) ──
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS form_drafts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            domain TEXT NOT NULL,
            form_type TEXT NOT NULL,
            draft_data TEXT NOT NULL,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id),
            UNIQUE(user_id, domain, form_type)
        )
    ''')
    
    # ── NEW COLUMNS: Pillar 2 — Risk-to-Task linkage ──
    try:
        cursor.execute('ALTER TABLE project_tasks ADD COLUMN mitigation_for_risk_id INTEGER')
    except:
        pass
    try:
        cursor.execute('ALTER TABLE project_tasks ADD COLUMN source_type TEXT DEFAULT "manual"')
    except:
        pass
    
    # ── NEW COLUMNS: risk_register — Domain + source tracking ──
    try:
        cursor.execute('ALTER TABLE risk_register ADD COLUMN domain TEXT DEFAULT "General"')
    except:
        pass
    try:
        cursor.execute('ALTER TABLE risk_register ADD COLUMN source_type TEXT DEFAULT "manual"')
    except:
        pass
    
    # ── NEW COLUMNS: Pillar 3 — Token usage tracking ──
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN token_usage INTEGER DEFAULT 0')
    except:
        pass
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN token_limit INTEGER DEFAULT 500000')
    except:
        pass
    # Update existing users who still have old 50000 limit
    try:
        cursor.execute('UPDATE users SET token_limit = 500000 WHERE token_limit = 50000 OR token_limit IS NULL')
    except:
        pass
    
    # ── NEW COLUMNS: Pillar 4 — Branded export logo ──
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN logo_path TEXT')
    except:
        pass
    
    # NOTE: Framework seeding deferred to seed_frameworks() — called after DOMAIN_FRAMEWORKS is defined
    
    # Create default admin user if not exists
    admin_pw = os.getenv('ADMIN_PASSWORD')
    if not admin_pw:
        raise RuntimeError(
            "FATAL: ADMIN_PASSWORD environment variable is not set. "
            "Set it in Render Dashboard → Environment Variables before deploying."
        )
    admin_hash = hash_password(admin_pw)
    try:
        cursor.execute('''
            INSERT OR IGNORE INTO users (username, email, password_hash, role) 
            VALUES (?, ?, ?, ?)
        ''', ('admin', 'admin@mizan.local', admin_hash, 'admin'))
    except:
        pass
    
    conn.commit()
    conn.close()
# Initialize database on startup
init_db()

# ============================================================================
# AUTHENTICATION
# ============================================================================

def login_required(f):
    """Decorator to require login. Returns JSON 401 for API calls, redirect for pages."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            # For API endpoints, return JSON error instead of redirect
            if request.path.startswith('/api/') or request.is_json or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': 'Session expired. Please login again.', 'session_expired': True}), 401
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    """Decorator to require admin role."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        if session.get('role') != 'admin':
            flash('Access denied. Admin privileges required.')
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function

# Registration limit
MAX_USERS = 20000

# Usage limits per user PER DOMAIN (for free tier)
USAGE_LIMITS = {
    'strategies': 1,
    'policies': 2,
    'audits': 2,
    'risks': 2
}

def get_user_usage_by_domain(user_id, domain):
    """Get current usage counts for a user in a specific domain."""
    conn = get_db()
    usage = {
        'strategies': conn.execute('SELECT COUNT(*) FROM strategies WHERE user_id = ? AND domain = ?', (user_id, domain)).fetchone()[0],
        'policies': conn.execute('SELECT COUNT(*) FROM policies WHERE user_id = ? AND domain = ?', (user_id, domain)).fetchone()[0],
        'audits': conn.execute('SELECT COUNT(*) FROM audits WHERE user_id = ? AND domain = ?', (user_id, domain)).fetchone()[0],
        'risks': conn.execute('SELECT COUNT(*) FROM risks WHERE user_id = ? AND domain = ?', (user_id, domain)).fetchone()[0]
    }
    return usage

def check_usage_limit(user_id, doc_type, domain):
    """Check if user has reached usage limit for a document type in a domain."""
    usage = get_user_usage_by_domain(user_id, domain)
    current = usage.get(doc_type, 0)
    limit = USAGE_LIMITS.get(doc_type, 1)
    return current < limit, current, limit

def get_remaining_usage(user_id, domain):
    """Get remaining usage for all document types in a specific domain."""
    usage = get_user_usage_by_domain(user_id, domain)
    remaining = {}
    for doc_type, limit in USAGE_LIMITS.items():
        remaining[doc_type] = {
            'used': usage.get(doc_type, 0),
            'limit': limit,
            'remaining': max(0, limit - usage.get(doc_type, 0))
        }
    return remaining

# ============================================================================
# ENTERPRISE HELPERS — Audit Logging, Token Tracking, Risk-Task Linkage
# ============================================================================

def log_action(user_id, action, metadata=None):
    """Write an entry to the audit_logs table (fire-and-forget)."""
    try:
        meta_str = json.dumps(metadata, default=str) if metadata else None
        conn = get_db()
        conn.execute(
            'INSERT INTO audit_logs (user_id, action, metadata) VALUES (?, ?, ?)',
            (user_id, action, meta_str)
        )
        conn.commit()
    except Exception as e:
        print(f"Audit log error (non-fatal): {e}", flush=True)

TOKEN_LIMIT_DEFAULT = 500000

def check_token_quota(user_id):
    """Return (allowed, used, limit). Blocks generation if over limit."""
    try:
        conn = get_db()
        row = conn.execute(
            'SELECT token_usage, token_limit FROM users WHERE id = ?', (user_id,)
        ).fetchone()
        if row:
            used = row['token_usage'] or 0
            limit = row['token_limit'] or TOKEN_LIMIT_DEFAULT
            return used < limit, used, limit
    except Exception:
        pass
    return True, 0, TOKEN_LIMIT_DEFAULT

def increment_token_usage(user_id, tokens):
    """Add tokens to the user's running total."""
    try:
        conn = get_db()
        conn.execute(
            'UPDATE users SET token_usage = COALESCE(token_usage, 0) + ? WHERE id = ?',
            (tokens, user_id)
        )
        conn.commit()
    except Exception as e:
        print(f"Token tracking error (non-fatal): {e}", flush=True)

def auto_create_mitigation_task(user_id, risk_id, risk_name, risk_score, domain='General'):
    """Pillar 2: Auto-generate a task for Critical/High risks (score >= 12)."""
    if risk_score < 12:
        return None
    priority = 'critical' if risk_score >= 20 else 'high'
    title = f"Mitigate Risk: {risk_name[:80]}"
    desc = f"Auto-generated mitigation task for risk (score {risk_score}). Review and assign an action plan."
    try:
        conn = get_db()
        conn.execute(
            '''INSERT INTO project_tasks 
               (user_id, domain, title, description, status, priority, category, mitigation_for_risk_id)
               VALUES (?, ?, ?, ?, 'todo', ?, 'risk_mitigation', ?)''',
            (user_id, domain, title, desc, priority, risk_id)
        )
        conn.commit()
        task_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        log_action(user_id, 'auto_task_from_risk', {'risk_id': risk_id, 'task_id': task_id, 'score': risk_score})
        return task_id
    except Exception as e:
        print(f"Auto-task error (non-fatal): {e}", flush=True)
        return None

def pipe_risk_assessment_to_register(content, user_id, domain='General', asset='', threat='', category='Operational'):
    """Auto-extract risk from AI risk assessment output → risk_register.
    
    Parses the Risk Assessment Summary table for risk level, score, category,
    and creates an entry in the risk_register with proper domain tagging.
    Returns the inserted risk_register id, or None.
    """
    import re
    
    # Extract risk level
    level_match = re.search(
        r'(?:Risk Level|مستوى الخطر)\s*\|\s*(Critical|High|Medium|Low|حرج|عالي|متوسط|منخفض)',
        content, re.IGNORECASE
    )
    # Extract risk score
    score_match = re.search(r'(?:Risk Score|درجة الخطر)\s*\|\s*(\d+)\s*/\s*10', content)
    
    if not level_match and not score_match:
        return None
    
    # Map text levels to likelihood/impact scores
    level_text = level_match.group(1).strip().lower() if level_match else 'medium'
    score_val = int(score_match.group(1)) if score_match else 5
    
    # Map to 1-5 scores
    if score_val >= 9 or level_text in ('critical', 'حرج'):
        likelihood, impact = 5, 5
    elif score_val >= 7 or level_text in ('high', 'عالي'):
        likelihood, impact = 4, 4
    elif score_val >= 4 or level_text in ('medium', 'متوسط'):
        likelihood, impact = 3, 3
    else:
        likelihood, impact = 2, 2
    
    # Build risk name
    risk_name = threat[:100] if threat else 'Risk from assessment'
    if asset:
        risk_name = f"{threat[:80]} → {asset[:40]}" if threat else f"Risk to {asset[:80]}"
    
    # Extract category from AI output
    cat_match = re.search(r'(?:Risk Category|فئة الخطر)\s*\|\s*([^|\n]+)', content)
    if cat_match:
        category = cat_match.group(1).strip()[:50]
    
    try:
        conn = get_db()
        conn.execute('''
            INSERT INTO risk_register 
            (user_id, name, description, category, likelihood, impact,
             treatment, treatment_plan, status, domain, source_type)
            VALUES (?, ?, ?, ?, ?, ?, 'Mitigate', '', 'Open', ?, 'ai_risk_assessment')
        ''', (user_id, risk_name[:200],
              f"Auto-extracted from {domain} risk assessment. Score: {score_val}/10",
              category, likelihood, impact, domain))
        conn.commit()
        reg_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        log_action(user_id, 'pipe_risk_to_register', {
            'domain': domain, 'risk_register_id': reg_id,
            'score': score_val, 'level': level_text
        })
        print(f"✅ PIPE: Risk assessment → Risk Register #{reg_id} ({domain})", flush=True)
        return reg_id
    except Exception as e:
        print(f"Pipe risk-to-register error (non-fatal): {e}", flush=True)
        return None


def pipe_initiatives_to_tasks(content, user_id, domain='General'):
    """Parse Strategic Pillars section for initiatives → project_tasks (category=strategic_initiative)."""
    import re
    ids = []
    # Match both EN (• Initiative:) and AR (• مبادرة:) formats, and also plain bullet lines under pillar headers
    pillar_name = None
    for line in content.split('\n'):
        # Detect pillar header
        pm = re.match(r'###\s*(?:Pillar|الركيزة)\s*\d+\s*[:：]\s*(.+)', line, re.IGNORECASE)
        if pm:
            pillar_name = pm.group(1).strip().rstrip(')')
            continue
        # Match initiative bullet lines
        im = re.match(r'\s*[•\-]\s*(?:Initiative|مبادرة)\s*[\d]*\s*[:\-—]?\s*(.+)', line, re.IGNORECASE)
        if not im:
            im = re.match(r'\s*[•\-]\s*(?:Initiative|مبادرة)\s*(.+)', line, re.IGNORECASE)
        if im:
            text = im.group(1).strip().strip('[]')
            if len(text) < 5 or text.lower().startswith('[') or 'concrete action' in text.lower():
                continue
            title = f"{pillar_name}: {text[:120]}" if pillar_name else text[:150]
            try:
                conn = get_db()
                conn.execute('''INSERT INTO project_tasks 
                    (user_id, domain, title, description, status, priority, category, source_type)
                    VALUES (?, ?, ?, ?, 'todo', 'high', 'strategic_initiative', 'ai_strategy')''',
                    (user_id, domain, title, text[:500]))
                conn.commit()
                ids.append(conn.execute('SELECT last_insert_rowid()').fetchone()[0])
            except Exception:
                pass
    if ids:
        log_action(user_id, 'pipe_initiatives_to_tasks', {'domain': domain, 'count': len(ids)})
        print(f"✅ PIPE: {len(ids)} initiative(s) → Tasks ({domain})", flush=True)
    return ids


def pipe_gaps_to_tasks(content, user_id, domain='General'):
    """Parse Gap Analysis table → project_tasks (category=gap_mitigation)."""
    import re
    ids = []
    # Match table rows: | # | Gap | Description | Priority | Status |
    for m in re.finditer(r'\|\s*\d+\s*\|\s*([^|]{5,80}?)\s*\|\s*([^|]{5,}?)\s*\|\s*(Critical|High|Medium|Low|حرج|عالي|متوسط|منخفض)\s*\|', content, re.IGNORECASE):
        gap_name = m.group(1).strip()
        description = m.group(2).strip()
        priority_raw = m.group(3).strip().lower()
        if gap_name.lower() in ('gap', 'الفجوة', '---', ''):
            continue
        priority_map = {'critical': 'critical', 'حرج': 'critical', 'high': 'high', 'عالي': 'high',
                        'medium': 'medium', 'متوسط': 'medium', 'low': 'low', 'منخفض': 'low'}
        priority = priority_map.get(priority_raw, 'medium')
        try:
            conn = get_db()
            conn.execute('''INSERT INTO project_tasks 
                (user_id, domain, title, description, status, priority, category, source_type)
                VALUES (?, ?, ?, ?, 'todo', ?, 'gap_mitigation', 'ai_strategy')''',
                (user_id, domain, gap_name[:200], description[:500], priority))
            conn.commit()
            ids.append(conn.execute('SELECT last_insert_rowid()').fetchone()[0])
        except Exception:
            pass
    if ids:
        log_action(user_id, 'pipe_gaps_to_tasks', {'domain': domain, 'count': len(ids)})
        print(f"✅ PIPE: {len(ids)} gap(s) → Tasks ({domain})", flush=True)
    return ids


def pipe_roadmap_to_tasks(content, user_id, domain='General'):
    """Parse Implementation Roadmap tables → project_tasks (category=implementation)."""
    import re
    ids = []
    # Match table rows: | # | Activity | Owner | Timeline | Deliverable |
    for m in re.finditer(r'\|\s*\d+\s*\|\s*([^|]{5,100}?)\s*\|\s*([^|]{2,60}?)\s*\|\s*([^|]{2,60}?)\s*\|\s*([^|]{2,}?)\s*\|', content):
        activity = m.group(1).strip()
        owner = m.group(2).strip()
        timeline = m.group(3).strip()
        deliverable = m.group(4).strip()
        # Skip header rows and placeholders
        if activity.lower() in ('activity', 'النشاط', '---', '') or 'activity' in activity.lower() and activity.startswith('['):
            continue
        if '---' in activity or activity.startswith('|'):
            continue
        # Determine priority from timeline
        priority = 'high' if re.search(r'Month\s*1|الشهر\s*1|شهر\s*1', timeline) else 'medium'
        title = activity[:200]
        desc = f"Owner: {owner}\nTimeline: {timeline}\nDeliverable: {deliverable}"
        try:
            conn = get_db()
            conn.execute('''INSERT INTO project_tasks 
                (user_id, domain, title, description, status, priority, owner, category, source_type)
                VALUES (?, ?, ?, ?, 'todo', ?, ?, 'implementation', 'ai_strategy')''',
                (user_id, domain, title, desc[:500], priority, owner[:100]))
            conn.commit()
            ids.append(conn.execute('SELECT last_insert_rowid()').fetchone()[0])
        except Exception:
            pass
    if ids:
        log_action(user_id, 'pipe_roadmap_to_tasks', {'domain': domain, 'count': len(ids)})
        print(f"✅ PIPE: {len(ids)} roadmap item(s) → Tasks ({domain})", flush=True)
    return ids


def pipe_kpis_to_tasks(content, user_id, domain='General'):
    """Parse KPI table → project_tasks (category=kpi)."""
    import re
    ids = []
    # Match KPI rows: | # | KPI | Description | Current | Target | Justification | Timeframe | Source |
    for m in re.finditer(r'\|\s*\d+\s*\|\s*([^|]{5,}?)\s*\|\s*([^|]{5,}?)\s*\|\s*([^|]*?)\s*\|\s*([^|]*?)\s*\|\s*([^|]*?)\s*\|\s*([^|]*?)\s*\|\s*([^|]*?)\s*\|', content):
        kpi_name = m.group(1).strip()
        description = m.group(2).strip()
        target = m.group(4).strip()
        timeframe = m.group(6).strip()
        # Skip headers and placeholders
        if kpi_name.lower() in ('kpi', 'المؤشر', '---', '') or kpi_name.startswith('['):
            continue
        if 'kpi name' in kpi_name.lower() or 'اسم المؤشر' in kpi_name:
            continue
        title = kpi_name[:200]
        desc = f"{description}\nTarget: {target}\nTimeframe: {timeframe}"
        try:
            conn = get_db()
            conn.execute('''INSERT INTO project_tasks 
                (user_id, domain, title, description, status, priority, category, source_type)
                VALUES (?, ?, ?, ?, 'todo', 'medium', 'kpi', 'ai_strategy')''',
                (user_id, domain, title, desc[:500]))
            conn.commit()
            ids.append(conn.execute('SELECT last_insert_rowid()').fetchone()[0])
        except Exception:
            pass
    if ids:
        log_action(user_id, 'pipe_kpis_to_tasks', {'domain': domain, 'count': len(ids)})
        print(f"✅ PIPE: {len(ids)} KPI(s) → Tasks ({domain})", flush=True)
    return ids


# ============================================================================
# TRANSLATIONS
# ============================================================================

TRANSLATIONS = {
    "en": {
        "app_name": "Mizan",
        "tagline": "Governance • Risk • Compliance",
        "login": "Sign In",
        "register": "Register",
        "logout": "Logout",
        "username": "Username",
        "password": "Password",
        "welcome": "Welcome to Enterprise GRC Platform",
        "welcome_sub": "Secure • Compliant • Intelligent",
        "disclaimer": "AI-generated content. Verify with professionals. This app does not save or store your data.",
        "no_sensitive": "Do not enter sensitive or confidential data",
        "usage_limit": "Usage Limit",
        "usage_reached": "You have reached your limit for this feature",
        "remaining": "remaining",
        "domains": ["Cyber Security", "Data Management", "Artificial Intelligence", "Digital Transformation", "Global Standards", "Enterprise Risk Management"],
        "tabs": ["Strategy", "Policy Lab", "Audit", "Risk Radar"],
        "org_name": "Organization Name",
        "sector": "Sector",
        "sectors": ["Government", "Banking/Finance", "Healthcare", "Energy", "Telecom", "Retail", "Manufacturing"],
        "size": "Organization Size",
        "sizes": ["Small (<100)", "Medium (100-1000)", "Large (1000+)"],
        "budget": "Budget Range",
        "budgets": ["< 1M SAR", "1M-5M SAR", "5M-20M SAR", "20M+ SAR"],
        "frameworks": "Regulatory Frameworks",
        "horizon": "Strategic Horizon (Months)",
        "generate": "Generate Strategy",
        "generating": "Generating...",
        "current_state": "Current State Assessment",
        "technologies": "Current Technologies",
        "challenges": "Key Challenges",
        "strategy_sections": {
            "vision": "Executive Vision & Strategic Objectives",
            "gaps": "Current State Assessment (Gap Analysis)",
            "pillars": "Strategic Pillars & Initiatives",
            "roadmap": "Implementation Roadmap",
            "kpis": "Measuring Success (KPIs & KRIs)",
            "confidence": "Confidence Score"
        },
        "policy_name": "Policy Title",
        "policy_framework": "Framework/Standard",
        "generate_policy": "Generate Policy",
        "risk_category": "Risk Category",
        "risk_scenario": "Risk Scenario",
        "asset_name": "Asset Name",
        "analyze_risk": "Analyze Risk",
        "download": "Download",
        "created_by": "Created by",
        "settings": "Settings",
        "clear_history": "Clear History",
        "ai_connected": "AI Core Connected",
        "ai_disconnected": "Simulation Mode",
        "logged_in_as": "Logged in as",
        "my_documents": "My Documents",
        "document_history": "Document History",
        "no_documents": "No documents yet. Start generating!",
        "view_document": "View",
        "delete_document": "Delete",
        "document_type": "Type",
        "document_date": "Date",
        "all_types": "All Types",
        "all_domains": "All Domains",
        "strategies": "Strategies",
        "policies": "Policies",
        "audits": "Audits",
        "risks": "Risk Analyses",
        "filter_by": "Filter by",
        "confirm_delete": "Are you sure you want to delete this document?",
        "share": "Share",
        "share_document": "Share Document",
        "share_link": "Share Link",
        "copy_link": "Copy Link",
        "link_copied": "Link copied to clipboard!",
        "shared_by": "Shared by",
        "views": "Views",
        "share_expires": "Expires",
        "never_expires": "Never",
        "stop_sharing": "Stop Sharing",
        "my_shared": "My Shared Links",
        "no_shared": "No shared documents yet",
        "shared_document": "Shared Document",
        "document_not_found": "Document not found or link expired",
        "secure_share": "Secure Share (OTP)",
        "public_share": "Public Share",
        "recipient_email": "Recipient Email",
        "send_otp": "Send Access Code",
        "otp_sent": "Access code sent to",
        "enter_otp": "Enter Access Code",
        "verify_otp": "Verify",
        "otp_invalid": "Invalid or expired code",
        "otp_verified": "Verified! Loading document...",
        "otp_required": "This document requires verification",
        "otp_email_subject": "Your Mizan Document Access Code",
        "resend_otp": "Resend Code",
        "analytics": "Analytics",
        "insights": "Insights & Analytics",
        "compliance_score": "Compliance Score",
        "risk_heatmap": "Risk Heatmap",
        "maturity_radar": "Maturity Radar",
        "benchmark": "Benchmark Comparison",
        "your_score": "Your Score",
        "industry_avg": "Industry Average",
        "gap": "Gap",
        "above_avg": "Above Average",
        "below_avg": "Below Average",
        "likelihood": "Likelihood",
        "impact": "Impact",
        "low": "Low",
        "medium": "Medium",
        "high": "High",
        "critical": "Critical",
        "governance": "Governance",
        "risk_mgmt": "Risk Management",
        "compliance": "Compliance",
        "technology": "Technology",
        "process": "Process",
        "maturity_level": "Maturity Level",
        "initial": "Initial",
        "developing": "Developing",
        "defined": "Defined",
        "managed": "Managed",
        "optimized": "Optimized",
        "no_data": "No data yet. Generate documents to see analytics.",
        "select_sector": "Select your sector to compare",
        "benchmark_source": "Source",
        "manage_benchmarks": "Manage Benchmarks",
        "templates": "Templates",
        "use_template": "Use Template",
        "select_template": "Select a Template",
        "template_description": "Description",
        "apply_template": "Apply Template",
        "no_templates": "No templates available",
        "bilingual": "Bilingual",
        "generate_both": "Generate Both (EN & AR)",
        "chat_document": "Chat with Document",
        "ask_question": "Ask a question about this document...",
        "send": "Send",
        "review_policy": "Review Policy",
        "review_type": "Review Type",
        "comprehensive_review": "Comprehensive Review",
        "compliance_review": "Compliance Check",
        "gap_analysis": "Gap Analysis",
        "start_review": "Start Review",
        "gap_remediation": "Gap Remediation",
        "enter_gaps": "Enter identified gaps (one per line)",
        "generate_plan": "Generate Remediation Plan",
        "reviewing": "Reviewing...",
        "chatting": "Processing...",
        "profile": "Profile",
        "my_profile": "My Profile",
        "change_password": "Change Password",
        "current_password": "Current Password",
        "new_password": "New Password",
        "confirm_password": "Confirm New Password",
        "update_password": "Update Password",
        "password_updated": "Password updated successfully",
        "password_mismatch": "Passwords do not match",
        "wrong_password": "Current password is incorrect",
        "account_info": "Account Information",
        "member_since": "Member Since",
        "last_active": "Last Active",
        "total_documents": "Total Documents",
        "usage_summary": "Usage Summary",
        "used": "Used",
        "of": "of",
        "export_all": "Export All",
        "export_all_docs": "Export All Documents",
        "export_desc": "Download all your documents as a ZIP file",
        "exporting": "Exporting...",
        "no_docs_export": "No documents to export",
        "recent_documents": "Recent Documents",
        "quick_actions": "Quick Actions",
        "new_strategy": "New Strategy",
        "new_policy": "New Policy",
        "new_audit": "New Audit",
        "new_risk": "New Risk Analysis",
        "view_all": "View All",
        "usage_overview": "Usage Overview",
        "domains_active": "Active Domains",
        "landing_title": "Enterprise GRC Platform",
        "landing_subtitle": "AI-Powered Governance, Risk & Compliance Management",
        "landing_cta": "Get Started Free",
        "landing_login": "Already have an account?",
        "feature_ai": "AI-Powered Generation",
        "feature_ai_desc": "Generate strategies, policies, audits, and risk assessments using advanced AI",
        "feature_domains": "6 GRC Domains",
        "feature_domains_desc": "Cyber Security, Data Management, AI, Digital Transformation & Global Standards",
        "feature_bilingual": "Bilingual Support",
        "feature_bilingual_desc": "Full Arabic and English support with RTL layout and professional exports",
        "feature_export": "Professional Exports",
        "feature_export_desc": "Export to PDF, Word, and Excel with proper formatting and branding",
        "feature_analytics": "Analytics Dashboard",
        "feature_analytics_desc": "Track compliance scores, risk heatmaps, and maturity assessments",
        "feature_templates": "150+ Templates",
        "feature_templates_desc": "Ready-to-use policy, audit, and risk templates across all domains",
        "trusted_by": "Trusted by GRC professionals across the Kingdom",
        "how_it_works": "How It Works",
        "step1_title": "Select Your Domain",
        "step1_desc": "Choose from 6 specialized GRC domains",
        "step2_title": "Configure & Generate",
        "step2_desc": "Set parameters and let AI generate professional documents",
        "step3_title": "Review & Export",
        "step3_desc": "Review, customize, and export in multiple formats"
    },
    "ar": {
        "app_name": "ميزان",
        "tagline": "الحوكمة • المخاطر • الامتثال",
        "login": "تسجيل الدخول",
        "register": "إنشاء حساب",
        "logout": "تسجيل الخروج",
        "username": "اسم المستخدم",
        "password": "كلمة المرور",
        "welcome": "مرحباً بك في منصة حوكمة المؤسسات",
        "welcome_sub": "آمن • ممتثل • ذكي",
        "disclaimer": "محتوى مُنشأ بالذكاء الاصطناعي. يُرجى التحقق مع المختصين. هذا التطبيق لا يحفظ أو يخزن بياناتك.",
        "no_sensitive": "لا تدخل بيانات حساسة أو سرية",
        "usage_limit": "حد الاستخدام",
        "usage_reached": "لقد وصلت إلى الحد الأقصى لهذه الميزة",
        "remaining": "متبقي",
        "domains": ["الأمن السيبراني", "إدارة البيانات", "الذكاء الاصطناعي", "التحول الرقمي", "المعايير العالمية", "إدارة المخاطر المؤسسية"],
        "tabs": ["الاستراتيجية", "معمل السياسات", "التدقيق", "رادار المخاطر"],
        "org_name": "اسم المنظمة",
        "sector": "القطاع",
        "sectors": ["حكومي", "بنوك/مالي", "رعاية صحية", "طاقة", "اتصالات", "تجزئة", "تصنيع"],
        "size": "حجم المنظمة",
        "sizes": ["صغيرة (أقل من 100)", "متوسطة (100-1000)", "كبيرة (أكثر من 1000)"],
        "budget": "نطاق الميزانية",
        "budgets": ["< 1 مليون ريال", "1-5 مليون ريال", "5-20 مليون ريال", "20+ مليون ريال"],
        "frameworks": "الأطر التنظيمية",
        "horizon": "الأفق الاستراتيجي (أشهر)",
        "generate": "إنشاء الاستراتيجية",
        "generating": "جاري الإنشاء...",
        "current_state": "تقييم الوضع الحالي",
        "technologies": "التقنيات الحالية",
        "challenges": "التحديات الرئيسية",
        "strategy_sections": {
            "vision": "الرؤية التنفيذية والأهداف الاستراتيجية",
            "gaps": "تقييم الوضع الراهن (تحليل الفجوات)",
            "pillars": "الركائز الاستراتيجية والمبادرات",
            "roadmap": "خارطة طريق التنفيذ",
            "kpis": "قياس النجاح (مؤشرات الأداء والمخاطر)",
            "confidence": "درجة الثقة والتحقق"
        },
        "policy_name": "عنوان السياسة",
        "policy_framework": "الإطار/المعيار",
        "generate_policy": "إنشاء السياسة",
        "risk_category": "فئة المخاطر",
        "risk_scenario": "سيناريو الخطر",
        "asset_name": "اسم الأصل",
        "analyze_risk": "تحليل الخطر",
        "download": "تحميل",
        "created_by": "تم الإنشاء بواسطة",
        "settings": "الإعدادات",
        "clear_history": "مسح السجل",
        "ai_connected": "الذكاء الاصطناعي متصل",
        "ai_disconnected": "وضع المحاكاة",
        "logged_in_as": "مسجل الدخول كـ",
        "my_documents": "مستنداتي",
        "document_history": "سجل المستندات",
        "no_documents": "لا توجد مستندات بعد. ابدأ بالإنشاء!",
        "view_document": "عرض",
        "delete_document": "حذف",
        "document_type": "النوع",
        "document_date": "التاريخ",
        "all_types": "جميع الأنواع",
        "all_domains": "جميع المجالات",
        "strategies": "الاستراتيجيات",
        "policies": "السياسات",
        "audits": "التدقيقات",
        "risks": "تحليلات المخاطر",
        "filter_by": "تصفية حسب",
        "confirm_delete": "هل أنت متأكد من حذف هذا المستند؟",
        "share": "مشاركة",
        "share_document": "مشاركة المستند",
        "share_link": "رابط المشاركة",
        "copy_link": "نسخ الرابط",
        "link_copied": "تم نسخ الرابط!",
        "shared_by": "تمت المشاركة بواسطة",
        "views": "المشاهدات",
        "share_expires": "ينتهي",
        "never_expires": "لا ينتهي",
        "stop_sharing": "إيقاف المشاركة",
        "my_shared": "روابطي المشتركة",
        "no_shared": "لا توجد مستندات مشتركة بعد",
        "shared_document": "مستند مشترك",
        "document_not_found": "المستند غير موجود أو انتهت صلاحية الرابط",
        "secure_share": "مشاركة آمنة (رمز التحقق)",
        "public_share": "مشاركة عامة",
        "recipient_email": "بريد المستلم",
        "send_otp": "إرسال رمز الوصول",
        "otp_sent": "تم إرسال رمز الوصول إلى",
        "enter_otp": "أدخل رمز الوصول",
        "verify_otp": "تحقق",
        "otp_invalid": "رمز غير صالح أو منتهي الصلاحية",
        "otp_verified": "تم التحقق! جاري تحميل المستند...",
        "otp_required": "هذا المستند يتطلب التحقق",
        "otp_email_subject": "رمز الوصول لمستند ميزان",
        "resend_otp": "إعادة إرسال الرمز",
        "analytics": "التحليلات",
        "insights": "الرؤى والتحليلات",
        "compliance_score": "درجة الامتثال",
        "risk_heatmap": "خريطة المخاطر الحرارية",
        "maturity_radar": "رادار النضج",
        "benchmark": "مقارنة معيارية",
        "your_score": "درجتك",
        "industry_avg": "متوسط القطاع",
        "gap": "الفجوة",
        "above_avg": "أعلى من المتوسط",
        "below_avg": "أقل من المتوسط",
        "likelihood": "الاحتمالية",
        "impact": "الأثر",
        "low": "منخفض",
        "medium": "متوسط",
        "high": "عالي",
        "critical": "حرج",
        "governance": "الحوكمة",
        "risk_mgmt": "إدارة المخاطر",
        "compliance": "الامتثال",
        "technology": "التقنية",
        "process": "العمليات",
        "maturity_level": "مستوى النضج",
        "initial": "أولي",
        "developing": "تطويري",
        "defined": "محدد",
        "managed": "مُدار",
        "optimized": "محسّن",
        "no_data": "لا توجد بيانات بعد. أنشئ مستندات لعرض التحليلات.",
        "select_sector": "اختر قطاعك للمقارنة",
        "benchmark_source": "المصدر",
        "manage_benchmarks": "إدارة المعايير",
        "templates": "القوالب",
        "use_template": "استخدم قالب",
        "select_template": "اختر قالباً",
        "template_description": "الوصف",
        "apply_template": "تطبيق القالب",
        "no_templates": "لا توجد قوالب متاحة",
        "bilingual": "ثنائي اللغة",
        "generate_both": "إنشاء بالعربية والإنجليزية",
        "chat_document": "محادثة مع الوثيقة",
        "ask_question": "اطرح سؤالاً عن هذه الوثيقة...",
        "send": "إرسال",
        "review_policy": "مراجعة السياسة",
        "review_type": "نوع المراجعة",
        "comprehensive_review": "مراجعة شاملة",
        "compliance_review": "فحص الامتثال",
        "gap_analysis": "تحليل الفجوات",
        "start_review": "بدء المراجعة",
        "gap_remediation": "معالجة الفجوات",
        "enter_gaps": "أدخل الفجوات المحددة (واحدة في كل سطر)",
        "generate_plan": "إنشاء خطة المعالجة",
        "reviewing": "جارٍ المراجعة...",
        "chatting": "جارٍ المعالجة...",
        "profile": "الملف الشخصي",
        "my_profile": "ملفي الشخصي",
        "change_password": "تغيير كلمة المرور",
        "current_password": "كلمة المرور الحالية",
        "new_password": "كلمة المرور الجديدة",
        "confirm_password": "تأكيد كلمة المرور الجديدة",
        "update_password": "تحديث كلمة المرور",
        "password_updated": "تم تحديث كلمة المرور بنجاح",
        "password_mismatch": "كلمتا المرور غير متطابقتين",
        "wrong_password": "كلمة المرور الحالية غير صحيحة",
        "account_info": "معلومات الحساب",
        "member_since": "عضو منذ",
        "last_active": "آخر نشاط",
        "total_documents": "إجمالي المستندات",
        "usage_summary": "ملخص الاستخدام",
        "used": "مستخدم",
        "of": "من",
        "export_all": "تصدير الكل",
        "export_all_docs": "تصدير جميع المستندات",
        "export_desc": "تحميل جميع مستنداتك كملف ZIP",
        "exporting": "جارٍ التصدير...",
        "no_docs_export": "لا توجد مستندات للتصدير",
        "recent_documents": "المستندات الأخيرة",
        "quick_actions": "إجراءات سريعة",
        "new_strategy": "استراتيجية جديدة",
        "new_policy": "سياسة جديدة",
        "new_audit": "تدقيق جديد",
        "new_risk": "تحليل مخاطر جديد",
        "view_all": "عرض الكل",
        "usage_overview": "نظرة عامة على الاستخدام",
        "domains_active": "المجالات النشطة",
        "landing_title": "منصة الحوكمة والمخاطر والامتثال",
        "landing_subtitle": "إدارة الحوكمة والمخاطر والامتثال المدعومة بالذكاء الاصطناعي",
        "landing_cta": "ابدأ مجاناً",
        "landing_login": "لديك حساب بالفعل؟",
        "feature_ai": "إنشاء بالذكاء الاصطناعي",
        "feature_ai_desc": "إنشاء استراتيجيات وسياسات وتقارير تدقيق وتقييم مخاطر باستخدام الذكاء الاصطناعي المتقدم",
        "feature_domains": "6 مجالات للحوكمة",
        "feature_domains_desc": "الأمن السيبراني، إدارة البيانات، الذكاء الاصطناعي، التحول الرقمي والمعايير العالمية",
        "feature_bilingual": "دعم ثنائي اللغة",
        "feature_bilingual_desc": "دعم كامل للعربية والإنجليزية مع تنسيق RTL وتصدير احترافي",
        "feature_export": "تصدير احترافي",
        "feature_export_desc": "تصدير إلى PDF و Word و Excel بتنسيق ومظهر احترافي",
        "feature_analytics": "لوحة تحليلات",
        "feature_analytics_desc": "تتبع درجات الامتثال وخرائط المخاطر الحرارية وتقييمات النضج",
        "feature_templates": "أكثر من 150 قالب",
        "feature_templates_desc": "قوالب جاهزة للسياسات والتدقيق والمخاطر في جميع المجالات",
        "trusted_by": "موثوق من متخصصي الحوكمة في المملكة",
        "how_it_works": "كيف تعمل المنصة",
        "step1_title": "اختر المجال",
        "step1_desc": "اختر من 6 مجالات متخصصة في الحوكمة",
        "step2_title": "حدد المعايير وأنشئ",
        "step2_desc": "حدد المتطلبات واترك الذكاء الاصطناعي ينشئ المستندات الاحترافية",
        "step3_title": "راجع وصدّر",
        "step3_desc": "راجع المستندات وخصصها وصدّرها بعدة تنسيقات"
    }
}

def get_text(lang='en'):
    """Get translations for language."""
    return TRANSLATIONS.get(lang, TRANSLATIONS['en'])

# ============================================================================
# DOMAIN DATA
# ============================================================================

DOMAIN_FRAMEWORKS = {
    "cyber": {
        "KSA": [
            "NCA ECC (Essential Cybersecurity Controls)",
            "NCA CSCC (Critical Systems Cybersecurity Controls)", 
            "NCA DCC (Data Cybersecurity Controls)",
            "NCA OTCC (Operational Technology Cybersecurity Controls)",
            "NCA TCC (Telework Cybersecurity Controls)",
            "NCA OSMACC (Social Media Cybersecurity Controls)",
            "NCA CCC (Cloud Cybersecurity Controls)",
            "NCA NCS (National Cryptographic Standards)",
            "NCA CGIoT (Cybersecurity Guidelines for IoT)",
            "SAMA CSF (Cybersecurity Framework)",
            "SAMA BCM (Business Continuity Management)",
            "CMA Cybersecurity Guidelines",
            "CITC Cybersecurity Regulations"
        ],
        "GCC": [
            "UAE NESA (National Electronic Security Authority)",
            "UAE ADSIC (Abu Dhabi Information Security)",
            "Qatar NCSA (National Cyber Security Agency)",
            "Bahrain NCSRC",
            "Oman ITA Cybersecurity",
            "Kuwait NCSC"
        ],
        "EU": [
            "NIS2 Directive (Network and Information Security)",
            "DORA (Digital Operational Resilience Act)",
            "Cyber Resilience Act",
            "ENISA Guidelines",
            "UK NCSC Cyber Essentials",
            "UK NCSC Cyber Essentials Plus",
            "German BSI IT-Grundschutz"
        ],
        "US": [
            "NIST CSF 2.0 (Cybersecurity Framework)",
            "NIST SP 800-53 (Security Controls)",
            "NIST SP 800-171 (CUI Protection)",
            "CMMC (Cybersecurity Maturity Model)",
            "CISA Guidelines",
            "FedRAMP",
            "FISMA",
            "NERC CIP (Critical Infrastructure)"
        ],
        "International": [
            "ISO 27001:2022 (ISMS)",
            "ISO 27002:2022 (Security Controls)",
            "ISO 27017 (Cloud Security)",
            "ISO 27018 (Cloud Privacy)",
            "ISO 27701 (Privacy Information)",
            "CIS Controls v8",
            "COBIT 2019 (Security)",
            "CSA CCM (Cloud Controls Matrix)",
            "PCI DSS v4.0"
        ]
    },
    "data": {
        "KSA": [
            "PDPL (Personal Data Protection Law)",
            "NCA DCC (Data Cybersecurity Controls)",
            "NDMO Data Governance Framework",
            "SDAIA Data Classification Policy",
            "SAMA Data Management Requirements",
            "CITC Data Regulations"
        ],
        "GCC": [
            "UAE PDPL (Personal Data Protection Law)",
            "UAE DIFC Data Protection Law",
            "UAE ADGM Data Protection Regulations",
            "Qatar DPL (Data Protection Law)",
            "Bahrain PDPL",
            "Oman Data Protection Law",
            "Kuwait Data Protection Draft"
        ],
        "EU": [
            "GDPR (General Data Protection Regulation)",
            "Data Governance Act (DGA)",
            "Data Act",
            "ePrivacy Directive",
            "UK Data Protection Act 2018",
            "UK GDPR"
        ],
        "US": [
            "CCPA/CPRA (California Privacy)",
            "HIPAA (Health Data)",
            "GLBA (Financial Data)",
            "FERPA (Education Data)",
            "COPPA (Children's Data)",
            "State Privacy Laws (Virginia, Colorado, etc.)"
        ],
        "International": [
            "ISO 27701 (Privacy Information Management)",
            "ISO 38505 (Data Governance)",
            "DAMA DMBOK (Data Management)",
            "DCAM (Data Capability Assessment)",
            "COBIT 2019 (Data Governance)"
        ]
    },
    "ai": {
        "KSA": [
            "SDAIA AI Ethics Principles",
            "SDAIA AI Governance Framework",
            "NCA AI Cybersecurity Guidelines",
            "MCIT AI Strategy"
        ],
        "GCC": [
            "UAE National AI Strategy",
            "UAE AI Ethics Guidelines",
            "Qatar National AI Strategy",
            "Bahrain AI Strategy"
        ],
        "EU": [
            "EU AI Act",
            "EU AI Liability Directive",
            "ALTAI (Assessment List for Trustworthy AI)",
            "UK AI Regulation Framework",
            "UK ICO AI Guidance"
        ],
        "US": [
            "NIST AI RMF (Risk Management Framework)",
            "Executive Order on AI Safety",
            "FTC AI Guidelines",
            "EEOC AI Hiring Guidelines",
            "State AI Regulations"
        ],
        "International": [
            "ISO 42001 (AI Management System)",
            "ISO 23894 (AI Risk Management)",
            "IEEE 7000 (Ethical AI)",
            "OECD AI Principles",
            "UNESCO AI Ethics"
        ]
    },
    "dt": {
        "KSA": [
            "DGA Digital Government Policy",
            "DGA Digital Transformation Standards",
            "MCIT National Digital Strategy",
            "Vision 2030 Digital Programs",
            "NCA Cloud Security Controls"
        ],
        "GCC": [
            "UAE Digital Government Strategy",
            "UAE Smart Government Initiative",
            "Qatar Digital Agenda",
            "Bahrain eGovernment",
            "Oman Digital Strategy"
        ],
        "EU": [
            "EU Digital Compass 2030",
            "Digital Services Act (DSA)",
            "Digital Markets Act (DMA)",
            "eIDAS 2.0",
            "UK Digital Strategy"
        ],
        "US": [
            "Federal Digital Strategy",
            "21st Century IDEA Act",
            "Cloud Smart Strategy",
            "Zero Trust Architecture (EO 14028)"
        ],
        "International": [
            "COBIT 2019",
            "TOGAF (Enterprise Architecture)",
            "ITIL 4 (Service Management)",
            "SAFe (Scaled Agile)",
            "ISO 38500 (IT Governance)",
            "ISO 20000 (Service Management)"
        ]
    },
    "global": {
        "Quality & Management": [
            "ISO 9001:2015 (Quality Management)",
            "ISO 14001:2015 (Environmental)",
            "ISO 45001:2018 (Occupational Health)",
            "ISO 22000 (Food Safety)",
            "ISO 50001 (Energy Management)"
        ],
        "Business Continuity": [
            "ISO 22301:2019 (BCMS)",
            "ISO 22313 (BC Guidance)",
            "ISO 22317 (BIA)",
            "BS 11200 (Crisis Management)",
            "ASIS BCM Standard"
        ],
        "Audit & Assurance": [
            "ISO 19011 (Auditing Guidelines)",
            "ISAE 3402 (Service Organization Controls)",
            "SOC 1/SOC 2/SOC 3",
            "SSAE 18"
        ],
        "Governance": [
            "ISO 37000 (Governance of Organizations)",
            "ISO 37001 (Anti-Bribery)",
            "ISO 37002 (Whistleblowing)",
            "COSO Internal Control Framework",
            "King IV (Corporate Governance)"
        ],
        "Sector-Specific": [
            "PCI DSS v4.0 (Payment Card)",
            "SWIFT CSCF (Financial Messaging)",
            "SOX (Sarbanes-Oxley)",
            "Basel III/IV (Banking)",
            "Solvency II (Insurance)"
        ]
    },
    "erm": {
        "International Standards": [
            "ISO 31000:2018 (Risk Management)",
            "ISO 31010:2019 (Risk Assessment Techniques)",
            "ISO 31022 (Legal Risk)",
            "ISO 31030 (Travel Risk)",
            "ISO 31050 (Emerging Risks)",
            "COSO ERM Framework (2017)",
            "IRM Risk Management Standard"
        ],
        "Financial Risk": [
            "Basel III/IV (Banking Risk)",
            "Solvency II (Insurance Risk)",
            "IFRS 9 (Financial Instruments)",
            "SAMA Risk Management Guidelines",
            "CMA Risk Management Requirements"
        ],
        "Operational Risk": [
            "NIST SP 800-30 (Risk Assessment)",
            "NIST RMF (SP 800-37)",
            "FAIR (Factor Analysis of Information Risk)",
            "OCTAVE (Operationally Critical Threat Assessment)"
        ],
        "Governance": [
            "King IV (Corporate Governance)",
            "FERMA Risk Management Standard",
            "AS/NZS 4360 (Australia/NZ Standard)",
            "AIRMIC/ALARM/IRM Risk Management Guide"
        ],
        "Regional": [
            "NCA Risk Management Requirements (KSA)",
            "SAMA Operational Risk Guidelines (KSA)",
            "CMA Corporate Governance Code (KSA)",
            "UAE Corporate Governance Code",
            "Qatar Corporate Governance Code"
        ]
    }
}

# Legacy flat list for backward compatibility
DOMAIN_FRAMEWORKS_FLAT = {
    "cyber": [fw for region in DOMAIN_FRAMEWORKS["cyber"].values() for fw in region],
    "data": [fw for region in DOMAIN_FRAMEWORKS["data"].values() for fw in region],
    "ai": [fw for region in DOMAIN_FRAMEWORKS["ai"].values() for fw in region],
    "dt": [fw for region in DOMAIN_FRAMEWORKS["dt"].values() for fw in region],
    "global": [fw for region in DOMAIN_FRAMEWORKS["global"].values() for fw in region],
    "erm": [fw for region in DOMAIN_FRAMEWORKS["erm"].values() for fw in region]
}



# ============================================================================
# FRAMEWORK DETAILS - Metadata for dynamic procedure & awareness generation
# Covers every regulation, framework, standard, and act in DOMAIN_FRAMEWORKS


# ============================================================================
# DYNAMIC BUILDERS - Generate procedures & awareness for ALL frameworks
# ============================================================================

def build_framework_awareness(domain_code, lang='en'):
    """Generate awareness training modules for EVERY framework in a domain."""
    frameworks_dict = DOMAIN_FRAMEWORKS.get(domain_code, {})
    modules = []
    counter = 0
    
    # Icon mapping by domain/type
    domain_icons = {
        'cyber': 'fas fa-shield-alt', 'data': 'fas fa-database',
        'ai': 'fas fa-robot', 'dt': 'fas fa-digital-tachograph',
        'global': 'fas fa-globe', 'erm': 'fas fa-chart-line'
    }
    region_icons = {
        'KSA': 'fas fa-flag', 'GCC': 'fas fa-mosque', 'EU': 'fas fa-landmark',
        'US': 'fas fa-flag-usa', 'International': 'fas fa-globe',
        'Quality & Management': 'fas fa-certificate', 'Business Continuity': 'fas fa-sync',
        'Audit & Assurance': 'fas fa-search', 'Governance': 'fas fa-university',
        'Sector-Specific': 'fas fa-industry', 'International Standards': 'fas fa-globe',
        'Financial Risk': 'fas fa-money-bill-wave', 'Operational Risk': 'fas fa-cogs',
        'Regional': 'fas fa-map-marked-alt'
    }
    
    for region_or_cat, fw_list in frameworks_dict.items():
        for fw_name in fw_list:
            counter += 1
            short = fw_name.split('(')[0].strip() if '(' in fw_name else fw_name
            desc_part = fw_name.split('(')[1].rstrip(')') if '(' in fw_name else fw_name
            fw_id = f"fw_{domain_code}_{counter}"
            icon = region_icons.get(region_or_cat, domain_icons.get(domain_code, 'fas fa-book'))
            
            # Check if we have detailed quiz data in FRAMEWORK_DETAILS
            details = FRAMEWORK_DETAILS.get(fw_name)
            
            if details and lang == 'en':
                quiz_data = [{"q": q["q"], "options": q["options"], "answer": q["answer"]} for q in details["quiz"]]
                learning_points = [
                    f"{details['short']} is a {details['type']} issued by {details['regulator']} ({details['region']})",
                    f"Scope: {details['controls']}",
                    f"Applicability: {details['applicability']}",
                ] + [f"Key area: {area}" for area in details.get('key_areas', [])[:3]]
                scenario_title = f"{details['short']} Compliance Gap"
                scenario_text = f"During a compliance review, your organization discovers it has not fully implemented the mandatory requirements of {details['short']}. The regulator ({details['regulator']}) has indicated that enforcement actions may follow non-compliance."
                scenario_action = f"Immediately conduct a gap assessment against {details['short']} requirements. Prioritize critical gaps, assign owners, create remediation timeline, and establish regular reporting to management. Engage with {details['regulator']} proactively about your compliance roadmap."
            elif details and lang == 'ar':
                quiz_data = [{"q": q["q_ar"], "options": q["options_ar"], "answer": q["answer"]} for q in details["quiz"]]
                learning_points = [
                    f"{details['short']} هو {details['type']} صادر من {details['regulator']} ({details['region']})",
                    f"النطاق: {details['domains_ar']}",
                    f"قابلية التطبيق: {details['applicability_ar']}",
                ] + [f"مجال رئيسي: {area}" for area in details.get('key_areas_ar', [])[:3]]
                scenario_title = f"فجوة امتثال {details['short']}"
                scenario_text = f"خلال مراجعة الامتثال، اكتشفت منظمتك أنها لم تطبق بالكامل المتطلبات الإلزامية لـ {details['short']}. أشارت الجهة التنظيمية ({details['regulator']}) إلى أن إجراءات إنفاذ قد تتبع عدم الامتثال."
                scenario_action = f"أجرِ فوراً تقييم فجوات مقابل متطلبات {details['short']}. رتّب الفجوات الحرجة حسب الأولوية وعيّن مسؤولين وأنشئ جدولاً زمنياً للمعالجة وأنشئ إبلاغاً منتظماً للإدارة. تواصل مع {details['regulator']} بشكل استباقي حول خارطة طريق الامتثال."
            else:
                # Generic module for frameworks without detailed data
                if lang == 'en':
                    quiz_data = [
                        {"q": f"What is {short}?", "options": [f"A {desc_part} framework/standard", "A programming language", "A hardware device", "An operating system"], "answer": 0},
                        {"q": f"What is the primary purpose of {short}?", "options": ["Entertainment", f"Establishing controls and requirements for {domain_code.replace('_',' ')} governance", "Social networking", "Marketing"], "answer": 1},
                        {"q": f"Who should be aware of {short} requirements?", "options": ["Only IT staff", "Only management", "All relevant personnel in scope", "Only external auditors"], "answer": 2},
                        {"q": f"How often should {short} compliance be reviewed?", "options": ["Never", "Only once during implementation", "Periodically as defined by the standard", "Every 10 years"], "answer": 2},
                        {"q": f"What happens if {short} requirements are not met?", "options": ["Nothing", "Potential regulatory penalties, fines, and reputational damage", "Free consulting", "Automatic extension"], "answer": 1}
                    ]
                    learning_points = [
                        f"{fw_name} provides a structured framework for {domain_code.replace('_',' ')} governance and compliance",
                        f"Region/Category: {region_or_cat} - applicable to organizations within its jurisdiction",
                        f"Key focus: Establishing mandatory controls, procedures, and monitoring requirements",
                        f"Organizations must assess their compliance status and implement required controls",
                        f"Regular audits and reviews are essential to maintain ongoing compliance"
                    ]
                    scenario_title = f"{short} Implementation Challenge"
                    scenario_text = f"Your organization needs to comply with {fw_name}. An initial assessment reveals gaps in several control areas. Management is asking for a prioritized implementation plan."
                    scenario_action = f"Conduct a detailed gap assessment against {short} requirements. Map current controls to required controls. Prioritize based on risk and regulatory deadlines. Create an implementation roadmap with clear milestones, owners, and budgets. Report progress regularly to stakeholders."
                else:
                    quiz_data = [
                        {"q": f"ما هو {short}؟", "options": [f"إطار/معيار لـ {desc_part}", "لغة برمجة", "جهاز أجهزة", "نظام تشغيل"], "answer": 0},
                        {"q": f"ما الغرض الأساسي من {short}؟", "options": ["الترفيه", f"وضع ضوابط ومتطلبات لحوكمة المجال", "التواصل الاجتماعي", "التسويق"], "answer": 1},
                        {"q": f"من يجب أن يكون على دراية بمتطلبات {short}؟", "options": ["فريق تقنية المعلومات فقط", "الإدارة فقط", "جميع الموظفين المعنيين ضمن النطاق", "المدققون الخارجيون فقط"], "answer": 2},
                        {"q": f"كم مرة يجب مراجعة الامتثال لـ {short}؟", "options": ["أبداً", "مرة واحدة فقط أثناء التطبيق", "دورياً كما يحدده المعيار", "كل 10 سنوات"], "answer": 2},
                        {"q": f"ماذا يحدث إذا لم تُستوفَ متطلبات {short}؟", "options": ["لا شيء", "عقوبات تنظيمية وغرامات وأضرار سمعية محتملة", "استشارة مجانية", "تمديد تلقائي"], "answer": 1}
                    ]
                    learning_points = [
                        f"{fw_name} يوفر إطاراً منظماً للحوكمة والامتثال",
                        f"المنطقة/الفئة: {region_or_cat} - ينطبق على المنظمات ضمن نطاقه",
                        f"التركيز الرئيسي: وضع ضوابط وإجراءات ومتطلبات مراقبة إلزامية",
                        f"يجب على المنظمات تقييم حالة امتثالها وتطبيق الضوابط المطلوبة",
                        f"التدقيقات والمراجعات المنتظمة ضرورية للحفاظ على الامتثال المستمر"
                    ]
                    scenario_title = f"تحدي تطبيق {short}"
                    scenario_text = f"تحتاج منظمتك للامتثال لـ {fw_name}. كشف تقييم أولي عن فجوات في عدة مجالات ضوابط. الإدارة تطلب خطة تطبيق مرتبة حسب الأولوية."
                    scenario_action = f"أجرِ تقييم فجوات تفصيلي مقابل متطلبات {short}. حدد الضوابط الحالية مقابل المطلوبة. رتّب حسب المخاطر والمواعيد التنظيمية. أنشئ خارطة طريق بمعالم ومسؤولين وميزانيات واضحة. أبلغ أصحاب المصلحة بالتقدم بانتظام."
            
            module = {
                "id": fw_id,
                "title": fw_name if lang == 'en' else fw_name,
                "icon": icon,
                "level": "Intermediate" if lang == 'en' else "متوسط",
                "duration": "15 min" if lang == 'en' else "15 دقيقة",
                "description": (f"Comprehensive training on {fw_name} - understanding requirements, mandatory controls, compliance obligations, and implementation best practices."
                               if lang == 'en' else
                               f"تدريب شامل على {fw_name} - فهم المتطلبات والضوابط الإلزامية والتزامات الامتثال وأفضل ممارسات التطبيق."),
                "learning_points": learning_points,
                "scenario": {
                    "title": scenario_title,
                    "text": scenario_text,
                    "correct_action": scenario_action
                },
                "quiz": quiz_data
            }
            modules.append(module)
    
    return modules


# ── Merge framework awareness into AWARENESS_MODULES ──
def merge_framework_awareness():
    """Add dynamically generated framework awareness to AWARENESS_MODULES."""
    for dc in ['cyber', 'data', 'ai', 'dt', 'global', 'erm']:
        if dc not in AWARENESS_MODULES:
            AWARENESS_MODULES[dc] = {}
        
        for lang in ['en', 'ar']:
            generated = build_framework_awareness(dc, lang)
            if lang not in AWARENESS_MODULES[dc]:
                AWARENESS_MODULES[dc][lang] = []
            
            # Get existing module IDs to avoid duplicates
            existing_ids = {m['id'] for m in AWARENESS_MODULES[dc][lang]}
            for mod in generated:
                if mod['id'] not in existing_ids:
                    AWARENESS_MODULES[dc][lang].append(mod)


FRAMEWORK_DETAILS = {
    # ── CYBER - KSA ──
    "NCA ECC (Essential Cybersecurity Controls)": {
        "short": "NCA ECC", "type": "regulation", "regulator": "NCA", "region": "KSA",
        "controls": "114 controls across 5 domains", "applicability": "All government entities and critical national infrastructure",
        "domains_ar": "5 مجالات و114 ضابطة", "applicability_ar": "جميع الجهات الحكومية والبنية التحتية الحرجة",
        "key_areas": ["Cybersecurity Governance", "Cybersecurity Defense", "Cybersecurity Resilience", "Third-Party & Cloud", "ICS/OT Security"],
        "key_areas_ar": ["حوكمة الأمن السيبراني", "الدفاع السيبراني", "صمود الأمن السيبراني", "الأطراف الثالثة والسحابة", "أمن أنظمة التحكم الصناعي"],
        "quiz": [
            {"q": "How many main domains are in NCA ECC?", "q_ar": "كم عدد المجالات الرئيسية في NCA ECC؟", "options": ["3","5","7","10"], "options_ar": ["3","5","7","10"], "answer": 1},
            {"q": "Who must comply with NCA ECC?", "q_ar": "من يجب أن يمتثل لـ NCA ECC؟", "options": ["Only banks","Only tech companies","All government entities & critical infrastructure","Only large enterprises"], "options_ar": ["البنوك فقط","شركات التقنية فقط","جميع الجهات الحكومية والبنية التحتية الحرجة","المنشآت الكبيرة فقط"], "answer": 2},
            {"q": "What does the Cybersecurity Defense domain cover?", "q_ar": "ماذا يغطي مجال الدفاع السيبراني؟", "options": ["Employee training","Threat detection, monitoring & incident response","Physical security only","Budget planning"], "options_ar": ["تدريب الموظفين","كشف التهديدات والمراقبة والاستجابة للحوادث","الأمن المادي فقط","تخطيط الميزانية"], "answer": 1},
            {"q": "Which organization enforces NCA ECC?", "q_ar": "أي جهة تنفذ ضوابط NCA ECC؟", "options": ["SAMA","SDAIA","National Cybersecurity Authority (NCA)","CITC"], "options_ar": ["ساما","سدايا","الهيئة الوطنية للأمن السيبراني","هيئة الاتصالات"], "answer": 2},
            {"q": "Is NCA ECC compliance voluntary?", "q_ar": "هل الامتثال لـ NCA ECC طوعي؟", "options": ["Yes, fully voluntary","Mandatory for government, voluntary for private","Mandatory for all targeted entities","Only for ISO-certified organizations"], "options_ar": ["نعم، طوعي بالكامل","إلزامي للحكومة، طوعي للقطاع الخاص","إلزامي لجميع الجهات المستهدفة","فقط للمنظمات الحاصلة على ISO"], "answer": 2}
        ]
    },
    "NCA CSCC (Critical Systems Cybersecurity Controls)": {
        "short": "NCA CSCC", "type": "regulation", "regulator": "NCA", "region": "KSA",
        "controls": "Enhanced controls for critical national systems", "applicability": "Entities operating national critical systems and infrastructure",
        "domains_ar": "ضوابط معززة للأنظمة الحساسة الوطنية", "applicability_ar": "الجهات المشغلة للأنظمة والبنية التحتية الحرجة الوطنية",
        "key_areas": ["Critical System Identification", "Enhanced Access Control", "24/7 SOC Monitoring", "Advanced Threat Detection", "Incident Response for Critical Systems"],
        "key_areas_ar": ["تحديد الأنظمة الحساسة", "تحكم وصول معزز", "مراقبة مركز عمليات أمنية على مدار الساعة", "كشف متقدم للتهديدات", "استجابة للحوادث في الأنظمة الحساسة"],
        "quiz": [
            {"q": "What is NCA CSCC designed to protect?", "q_ar": "ما الذي صُمم NCA CSCC لحمايته؟", "options": ["Small businesses","National critical systems and infrastructure","Only financial institutions","Social media accounts"], "options_ar": ["المنشآت الصغيرة","الأنظمة والبنية التحتية الحرجة الوطنية","المؤسسات المالية فقط","حسابات التواصل الاجتماعي"], "answer": 1},
            {"q": "What monitoring level does CSCC require?", "q_ar": "ما مستوى المراقبة الذي يتطلبه CSCC؟", "options": ["Weekly reviews","Business hours only","24/7 continuous monitoring","Monthly assessments"], "options_ar": ["مراجعات أسبوعية","ساعات العمل فقط","مراقبة مستمرة على مدار الساعة","تقييمات شهرية"], "answer": 2},
            {"q": "How does CSCC relate to ECC?", "q_ar": "ما علاقة CSCC بـ ECC؟", "options": ["CSCC replaces ECC","CSCC adds enhanced controls on top of ECC","They are unrelated","ECC is optional if CSCC is implemented"], "options_ar": ["CSCC يحل محل ECC","CSCC يضيف ضوابط معززة فوق ECC","غير مرتبطين","ECC اختياري إذا طُبق CSCC"], "answer": 1},
            {"q": "What type of access control does CSCC emphasize?", "q_ar": "ما نوع التحكم بالوصول الذي يؤكد عليه CSCC؟", "options": ["Basic password protection","Privileged access management with MFA","No access restrictions","Username-only access"], "options_ar": ["حماية أساسية بكلمات المرور","إدارة الوصول المميز مع MFA","لا قيود وصول","وصول باسم المستخدم فقط"], "answer": 1},
            {"q": "What must entities do when a critical system incident occurs?", "q_ar": "ما يجب على الجهات فعله عند وقوع حادث في نظام حساس؟", "options": ["Wait for next business day","Immediately activate incident response and notify NCA","Handle internally without reporting","Send an email report within a month"], "options_ar": ["الانتظار ليوم العمل التالي","تفعيل الاستجابة للحوادث فوراً وإخطار الهيئة","المعالجة داخلياً دون إبلاغ","إرسال تقرير بالبريد خلال شهر"], "answer": 1}
        ]
    },
    "NCA DCC (Data Cybersecurity Controls)": {
        "short": "NCA DCC", "type": "regulation", "regulator": "NCA", "region": "KSA",
        "controls": "Data classification, protection, and transfer controls", "applicability": "All entities subject to NCA regulations handling sensitive data",
        "domains_ar": "ضوابط تصنيف وحماية ونقل البيانات", "applicability_ar": "جميع الجهات الخاضعة لأنظمة الهيئة التي تتعامل مع بيانات حساسة",
        "key_areas": ["Data Classification", "Data Protection", "Data Transfer Controls", "Data Retention", "Data Disposal"],
        "key_areas_ar": ["تصنيف البيانات", "حماية البيانات", "ضوابط نقل البيانات", "الاحتفاظ بالبيانات", "التخلص من البيانات"],
        "quiz": [
            {"q": "What is the primary focus of NCA DCC?", "q_ar": "ما التركيز الأساسي لـ NCA DCC؟", "options": ["Network security","Data cybersecurity controls","Physical security","Employee management"], "options_ar": ["أمن الشبكات","ضوابط الأمن السيبراني للبيانات","الأمن المادي","إدارة الموظفين"], "answer": 1},
            {"q": "What must be done before data is shared externally?", "q_ar": "ما يجب فعله قبل مشاركة البيانات خارجياً؟", "options": ["Nothing special","Classification and protection assessment per DCC","Only get verbal approval","Just encrypt it"], "options_ar": ["لا شيء خاص","تقييم التصنيف والحماية وفق DCC","الحصول على موافقة شفهية فقط","تشفيرها فقط"], "answer": 1},
            {"q": "How does NCA DCC relate to PDPL?", "q_ar": "ما علاقة NCA DCC بنظام حماية البيانات الشخصية؟", "options": ["DCC replaces PDPL","DCC provides technical controls to support PDPL compliance","They conflict with each other","PDPL is only for private sector"], "options_ar": ["DCC يحل محل النظام","DCC يوفر ضوابط تقنية لدعم الامتثال للنظام","يتعارضان مع بعضهما","النظام للقطاع الخاص فقط"], "answer": 1},
            {"q": "What is required for data at rest per DCC?", "q_ar": "ما المطلوب للبيانات المخزنة وفق DCC؟", "options": ["No requirements","Encryption based on classification level","Only access logs","Physical locks only"], "options_ar": ["لا متطلبات","تشفير حسب مستوى التصنيف","سجلات الوصول فقط","أقفال مادية فقط"], "answer": 1},
            {"q": "What controls does DCC require for data disposal?", "q_ar": "ما الضوابط التي يتطلبها DCC للتخلص من البيانات؟", "options": ["Simple deletion","Secure destruction with documented verification","Just empty the recycle bin","No specific requirements"], "options_ar": ["حذف بسيط","إتلاف آمن مع توثيق التحقق","تفريغ سلة المحذوفات فقط","لا متطلبات محددة"], "answer": 1}
        ]
    },
    "NCA OTCC (Operational Technology Cybersecurity Controls)": {
        "short": "NCA OTCC", "type": "regulation", "regulator": "NCA", "region": "KSA",
        "controls": "OT/ICS-specific cybersecurity controls", "applicability": "Entities operating operational technology and industrial control systems",
        "domains_ar": "ضوابط الأمن السيبراني الخاصة بالتقنية التشغيلية", "applicability_ar": "الجهات المشغلة للتقنية التشغيلية وأنظمة التحكم الصناعي",
        "key_areas": ["OT Asset Management", "OT Network Security", "OT Access Control", "OT Patch Management", "OT Incident Response"],
        "key_areas_ar": ["إدارة أصول التقنية التشغيلية", "أمن شبكات التقنية التشغيلية", "التحكم بالوصول للتقنية التشغيلية", "إدارة تحديثات التقنية التشغيلية", "الاستجابة لحوادث التقنية التشغيلية"],
        "quiz": [
            {"q": "What does OT stand for?", "q_ar": "ماذا تعني OT؟", "options": ["Online Technology","Operational Technology","Outsourced Tasks","Optimal Testing"], "options_ar": ["التقنية عبر الإنترنت","التقنية التشغيلية","المهام الخارجية","الاختبار الأمثل"], "answer": 1},
            {"q": "Why is OT security different from IT security?", "q_ar": "لماذا يختلف أمن التقنية التشغيلية عن أمن تقنية المعلومات؟", "options": ["It's exactly the same","OT systems prioritize availability and safety; patching can disrupt operations","OT is less important","OT doesn't need security"], "options_ar": ["متماثلان تماماً","أنظمة OT تعطي الأولوية للتوفر والسلامة والتحديثات قد تعطل العمليات","OT أقل أهمية","OT لا يحتاج أمناً"], "answer": 1},
            {"q": "What is the first step in OTCC compliance?", "q_ar": "ما الخطوة الأولى في الامتثال لـ OTCC؟", "options": ["Install firewalls","Complete OT asset inventory and classification","Buy new equipment","Disconnect all OT systems"], "options_ar": ["تثبيت جدران حماية","إكمال جرد وتصنيف أصول التقنية التشغيلية","شراء معدات جديدة","فصل جميع أنظمة OT"], "answer": 1},
            {"q": "How should IT and OT networks be connected?", "q_ar": "كيف يجب ربط شبكات تقنية المعلومات والتقنية التشغيلية؟", "options": ["Directly without restrictions","Through segmented, controlled connections with DMZ","They should never connect","Through public internet"], "options_ar": ["مباشرة بدون قيود","عبر اتصالات مقسمة ومتحكم بها مع DMZ","لا يجب أن تتصل أبداً","عبر الإنترنت العام"], "answer": 1},
            {"q": "How is patching handled differently in OT environments?", "q_ar": "كيف تُدار التحديثات بشكل مختلف في بيئات التقنية التشغيلية؟", "options": ["Auto-update everything immediately","Test in staging, schedule during maintenance windows, verify system stability","Never patch OT systems","Apply patches during peak operations"], "options_ar": ["تحديث تلقائي فوري لكل شيء","اختبار في بيئة تجريبية وجدولة خلال فترات الصيانة والتحقق من استقرار النظام","لا تحدّث أنظمة OT أبداً","تطبيق التحديثات أثناء ذروة العمليات"], "answer": 1}
        ]
    },
    "NCA TCC (Telework Cybersecurity Controls)": {
        "short": "NCA TCC", "type": "regulation", "regulator": "NCA", "region": "KSA",
        "controls": "Remote work security requirements", "applicability": "All entities enabling telework/remote work arrangements",
        "domains_ar": "متطلبات أمن العمل عن بعد", "applicability_ar": "جميع الجهات التي تمكّن ترتيبات العمل عن بعد",
        "key_areas": ["Remote Access Security", "Endpoint Protection", "VPN Configuration", "Data Loss Prevention", "Home Network Security"],
        "key_areas_ar": ["أمن الوصول عن بعد", "حماية الأجهزة", "تكوين VPN", "منع فقدان البيانات", "أمن الشبكة المنزلية"],
        "quiz": [
            {"q": "What is the primary focus of NCA TCC?", "q_ar": "ما التركيز الأساسي لـ NCA TCC؟", "options": ["Office security","Cybersecurity controls for telework/remote work","Travel security","Telecom security"], "options_ar": ["أمن المكتب","ضوابط الأمن السيبراني للعمل عن بعد","أمن السفر","أمن الاتصالات"], "answer": 1},
            {"q": "What VPN requirement does TCC specify?", "q_ar": "ما متطلب VPN الذي يحدده TCC؟", "options": ["VPN is optional","Encrypted VPN with MFA for all remote connections","Only for managers","Free VPN services are acceptable"], "options_ar": ["VPN اختياري","VPN مشفر مع MFA لجميع الاتصالات عن بعد","للمديرين فقط","خدمات VPN المجانية مقبولة"], "answer": 1},
            {"q": "What endpoint controls does TCC require?", "q_ar": "ما ضوابط الأجهزة التي يتطلبها TCC؟", "options": ["None","Updated antivirus, disk encryption, and remote wipe capability","Only password protection","Just a screensaver"], "options_ar": ["لا شيء","مضاد فيروسات محدّث وتشفير القرص وقدرة المسح عن بعد","حماية بكلمة المرور فقط","شاشة توقف فقط"], "answer": 1},
            {"q": "Can employees use personal devices for remote work per TCC?", "q_ar": "هل يمكن للموظفين استخدام أجهزتهم الشخصية للعمل عن بعد وفق TCC؟", "options": ["Yes, without restrictions","Only with approved BYOD policy and security controls","Never allowed","Only on weekends"], "options_ar": ["نعم بدون قيود","فقط مع سياسة BYOD معتمدة وضوابط أمنية","غير مسموح أبداً","فقط في عطلات نهاية الأسبوع"], "answer": 1},
            {"q": "What data protection measure is required for remote workers?", "q_ar": "ما إجراء حماية البيانات المطلوب للعاملين عن بعد؟", "options": ["No special measures","DLP tools and restrictions on data transfers to personal storage","Just verbal agreements","Only for classified data"], "options_ar": ["لا إجراءات خاصة","أدوات DLP وقيود على نقل البيانات للتخزين الشخصي","اتفاقيات شفهية فقط","للبيانات المصنفة فقط"], "answer": 1}
        ]
    },
    "NCA OSMACC (Social Media Cybersecurity Controls)": {
        "short": "NCA OSMACC", "type": "regulation", "regulator": "NCA", "region": "KSA",
        "controls": "Social media account security controls", "applicability": "Government entities and organizations with official social media presence",
        "domains_ar": "ضوابط أمن حسابات التواصل الاجتماعي", "applicability_ar": "الجهات الحكومية والمنظمات ذات التواجد الرسمي على وسائل التواصل الاجتماعي",
        "key_areas": ["Account Registration & Governance", "Access Control & MFA", "Content Publishing Controls", "Incident Response", "Account Recovery"],
        "key_areas_ar": ["تسجيل الحسابات والحوكمة", "التحكم بالوصول والمصادقة المتعددة", "ضوابط نشر المحتوى", "الاستجابة للحوادث", "استعادة الحسابات"],
        "quiz": [
            {"q": "What does NCA OSMACC regulate?", "q_ar": "ماذا ينظم NCA OSMACC؟", "options": ["Email security","Social media cybersecurity for official accounts","Website design","Mobile app development"], "options_ar": ["أمن البريد الإلكتروني","الأمن السيبراني لحسابات التواصل الاجتماعي الرسمية","تصميم المواقع","تطوير تطبيقات الجوال"], "answer": 1},
            {"q": "What authentication is required for official social media accounts?", "q_ar": "ما المصادقة المطلوبة لحسابات التواصل الاجتماعي الرسمية؟", "options": ["Simple password","Multi-factor authentication with access logging","No authentication needed","Shared team password"], "options_ar": ["كلمة مرور بسيطة","مصادقة متعددة العوامل مع تسجيل الوصول","لا حاجة للمصادقة","كلمة مرور مشتركة للفريق"], "answer": 1},
            {"q": "Who should have access to official social media accounts?", "q_ar": "من يجب أن يكون لديه وصول لحسابات التواصل الاجتماعي الرسمية؟", "options": ["Any employee","Only authorized personnel with documented roles","The entire marketing team","External agencies without oversight"], "options_ar": ["أي موظف","الموظفون المصرح لهم فقط بأدوار موثقة","فريق التسويق بالكامل","وكالات خارجية بدون إشراف"], "answer": 1},
            {"q": "What should happen if a social media account is compromised?", "q_ar": "ماذا يجب فعله إذا تم اختراق حساب تواصل اجتماعي؟", "options": ["Just change the password","Activate incident response plan, secure account, notify NCA","Delete the account","Post an apology and continue"], "options_ar": ["تغيير كلمة المرور فقط","تفعيل خطة الاستجابة للحوادث وتأمين الحساب وإخطار الهيئة","حذف الحساب","نشر اعتذار والاستمرار"], "answer": 1},
            {"q": "What content publishing control does OSMACC require?", "q_ar": "ما ضوابط نشر المحتوى التي يتطلبها OSMACC؟", "options": ["No controls needed","Content approval workflow, scheduled publishing, and activity logging","Only spell-check","Just manager notification"], "options_ar": ["لا ضوابط مطلوبة","سير عمل اعتماد المحتوى والنشر المجدول وتسجيل الأنشطة","التدقيق الإملائي فقط","إخطار المدير فقط"], "answer": 1}
        ]
    },
    "NCA CCC (Cloud Cybersecurity Controls)": {
        "short": "NCA CCC", "type": "regulation", "regulator": "NCA", "region": "KSA",
        "controls": "Cloud adoption and security controls", "applicability": "Entities adopting cloud services for government or sensitive data",
        "domains_ar": "ضوابط تبني وأمن الحوسبة السحابية", "applicability_ar": "الجهات التي تتبنى خدمات سحابية للبيانات الحكومية أو الحساسة",
        "key_areas": ["Cloud Provider Assessment", "Cloud Data Residency", "Cloud Access Control", "Cloud Configuration Security", "Cloud Monitoring"],
        "key_areas_ar": ["تقييم مزود السحابة", "إقامة البيانات السحابية", "التحكم بالوصول السحابي", "أمن تكوين السحابة", "مراقبة السحابة"],
        "quiz": [
            {"q": "What must be assessed before adopting a cloud service per CCC?", "q_ar": "ما يجب تقييمه قبل تبني خدمة سحابية وفق CCC؟", "options": ["Only the price","Cloud provider security, data residency, and compliance capabilities","Nothing special","Just speed and uptime"], "options_ar": ["السعر فقط","أمن مزود السحابة وإقامة البيانات وقدرات الامتثال","لا شيء خاص","السرعة والتوفر فقط"], "answer": 1},
            {"q": "Where must government data reside per NCA CCC?", "q_ar": "أين يجب أن تُقيم البيانات الحكومية وفق NCA CCC؟", "options": ["Anywhere globally","Within KSA borders unless specifically authorized","Only in the US","In the nearest data center"], "options_ar": ["في أي مكان عالمياً","داخل حدود المملكة ما لم يُصرح تحديداً","في الولايات المتحدة فقط","في أقرب مركز بيانات"], "answer": 1},
            {"q": "What access control does CCC require for cloud environments?", "q_ar": "ما التحكم بالوصول الذي يتطلبه CCC لبيئات السحابة؟", "options": ["Basic username/password","Identity and access management with MFA, least privilege, and activity monitoring","Shared admin accounts","No specific requirements"], "options_ar": ["اسم مستخدم/كلمة مرور أساسية","إدارة الهوية والوصول مع MFA وأقل الصلاحيات ومراقبة الأنشطة","حسابات مدير مشتركة","لا متطلبات محددة"], "answer": 1},
            {"q": "What shared responsibility model does CCC reference?", "q_ar": "ما نموذج المسؤولية المشتركة الذي يشير إليه CCC؟", "options": ["Provider is fully responsible","Customer is fully responsible","Shared responsibility between provider and customer","Nobody is responsible"], "options_ar": ["المزود مسؤول بالكامل","العميل مسؤول بالكامل","مسؤولية مشتركة بين المزود والعميل","لا أحد مسؤول"], "answer": 2},
            {"q": "How should cloud security configurations be managed?", "q_ar": "كيف يجب إدارة تكوينات أمن السحابة؟", "options": ["Set once and forget","Continuous monitoring, baseline enforcement, and drift detection","Only during initial setup","No management needed"], "options_ar": ["ضبط مرة واحدة والنسيان","مراقبة مستمرة وإنفاذ خط الأساس وكشف الانحراف","أثناء الإعداد الأولي فقط","لا حاجة للإدارة"], "answer": 1}
        ]
    },
    "NCA NCS (National Cryptographic Standards)": {
        "short": "NCA NCS", "type": "standard", "regulator": "NCA", "region": "KSA",
        "controls": "Cryptographic standards and key management", "applicability": "All entities using cryptographic technologies for data protection",
        "domains_ar": "معايير التشفير وإدارة المفاتيح", "applicability_ar": "جميع الجهات التي تستخدم تقنيات التشفير لحماية البيانات",
        "key_areas": ["Algorithm Selection", "Key Generation", "Key Distribution", "Key Storage & Rotation", "Certificate Management"],
        "key_areas_ar": ["اختيار الخوارزميات", "إنشاء المفاتيح", "توزيع المفاتيح", "تخزين وتدوير المفاتيح", "إدارة الشهادات"],
        "quiz": [
            {"q": "What does NCA NCS primarily regulate?", "q_ar": "ما الذي ينظمه NCA NCS بشكل أساسي؟", "options": ["Network cables","National cryptographic standards and encryption requirements","Social media","Employee training"], "options_ar": ["كابلات الشبكة","معايير التشفير الوطنية ومتطلبات التشفير","وسائل التواصل الاجتماعي","تدريب الموظفين"], "answer": 1},
            {"q": "Which encryption algorithms are approved per NCS?", "q_ar": "أي خوارزميات تشفير معتمدة وفق NCS؟", "options": ["Any algorithm","Only NCA-approved algorithms meeting national security standards","Only proprietary algorithms","No restrictions"], "options_ar": ["أي خوارزمية","فقط الخوارزميات المعتمدة من الهيئة التي تلبي معايير الأمن الوطني","الخوارزميات المملوكة فقط","لا قيود"], "answer": 1},
            {"q": "How often should encryption keys be rotated per NCS?", "q_ar": "كم مرة يجب تدوير مفاتيح التشفير وفق NCS؟", "options": ["Never","Based on key type and classification - annually or more frequently","Only when compromised","Every 10 years"], "options_ar": ["أبداً","حسب نوع المفتاح والتصنيف - سنوياً أو أكثر تكراراً","فقط عند الاختراق","كل 10 سنوات"], "answer": 1},
            {"q": "What must be done with expired cryptographic keys?", "q_ar": "ما يجب فعله بمفاتيح التشفير المنتهية؟", "options": ["Keep them active","Securely destroy with documented verification per NCS","Delete from one system","Archive indefinitely"], "options_ar": ["إبقاؤها نشطة","إتلافها بشكل آمن مع توثيق التحقق وفق NCS","حذفها من نظام واحد","أرشفتها إلى أجل غير مسمى"], "answer": 1},
            {"q": "What certificate management does NCS require?", "q_ar": "ما إدارة الشهادات التي يتطلبها NCS؟", "options": ["No requirements","Full lifecycle management: issuance, renewal, revocation, and monitoring","Self-signed certificates only","Let certificates auto-expire"], "options_ar": ["لا متطلبات","إدارة دورة حياة كاملة: الإصدار والتجديد والإلغاء والمراقبة","شهادات موقعة ذاتياً فقط","ترك الشهادات تنتهي تلقائياً"], "answer": 1}
        ]
    },
    "NCA CGIoT (Cybersecurity Guidelines for IoT)": {
        "short": "NCA CGIoT", "type": "guideline", "regulator": "NCA", "region": "KSA",
        "controls": "IoT device security guidelines", "applicability": "Entities deploying IoT devices in government or sensitive environments",
        "domains_ar": "إرشادات أمن أجهزة إنترنت الأشياء", "applicability_ar": "الجهات التي تنشر أجهزة إنترنت الأشياء في البيئات الحكومية أو الحساسة",
        "key_areas": ["IoT Device Inventory", "Firmware Security", "Network Isolation", "Authentication", "Monitoring & Updates"],
        "key_areas_ar": ["جرد أجهزة IoT", "أمن البرامج الثابتة", "عزل الشبكة", "المصادقة", "المراقبة والتحديثات"],
        "quiz": [
            {"q": "What is the first step in IoT security per CGIoT?", "q_ar": "ما الخطوة الأولى في أمن IoT وفق CGIoT؟", "options": ["Buy expensive devices","Discover and inventory all IoT devices on the network","Ignore IoT devices","Only secure cameras"], "options_ar": ["شراء أجهزة مكلفة","اكتشاف وجرد جميع أجهزة IoT على الشبكة","تجاهل أجهزة IoT","تأمين الكاميرات فقط"], "answer": 1},
            {"q": "How should IoT devices be connected to the network?", "q_ar": "كيف يجب ربط أجهزة IoT بالشبكة؟", "options": ["On the main corporate network","On an isolated, segmented network with controlled access","Through public WiFi","No network requirements"], "options_ar": ["على الشبكة المؤسسية الرئيسية","على شبكة معزولة ومقسمة مع وصول متحكم","عبر WiFi العامة","لا متطلبات شبكية"], "answer": 1},
            {"q": "What firmware practice does CGIoT require?", "q_ar": "ما ممارسة البرامج الثابتة التي يتطلبها CGIoT؟", "options": ["Never update firmware","Verify firmware integrity and apply security updates regularly","Only use default firmware","Update only when convenient"], "options_ar": ["عدم تحديث البرامج الثابتة أبداً","التحقق من سلامة البرامج الثابتة وتطبيق تحديثات الأمان بانتظام","استخدام البرامج الثابتة الافتراضية فقط","التحديث فقط عند الملاءمة"], "answer": 1},
            {"q": "What must be changed on IoT devices before deployment?", "q_ar": "ما يجب تغييره في أجهزة IoT قبل النشر؟", "options": ["Nothing","Default credentials, disable unnecessary services, and configure security settings","Only the device name","The device color"], "options_ar": ["لا شيء","بيانات الاعتماد الافتراضية وتعطيل الخدمات غير الضرورية وتكوين إعدادات الأمان","اسم الجهاز فقط","لون الجهاز"], "answer": 1},
            {"q": "How should decommissioned IoT devices be handled?", "q_ar": "كيف يجب التعامل مع أجهزة IoT المُوقفة؟", "options": ["Just throw them away","Securely wipe data, revoke credentials, and remove from network inventory","Give them to employees","Sell them online"], "options_ar": ["رميها فقط","مسح البيانات بأمان وإلغاء بيانات الاعتماد وإزالتها من جرد الشبكة","إعطاؤها للموظفين","بيعها عبر الإنترنت"], "answer": 1}
        ]
    },
    "SAMA CSF (Cybersecurity Framework)": {
        "short": "SAMA CSF", "type": "framework", "regulator": "SAMA", "region": "KSA",
        "controls": "4 domains covering governance, risk, operations, third-party", "applicability": "All SAMA-regulated financial institutions (banks, insurance, finance companies)",
        "domains_ar": "4 مجالات تغطي الحوكمة والمخاطر والعمليات والأطراف الثالثة", "applicability_ar": "جميع المؤسسات المالية المنظمة من ساما (البنوك، التأمين، شركات التمويل)",
        "key_areas": ["Cyber Security Leadership & Governance", "Cyber Security Risk Management", "Cyber Security Operations & Technology", "Third Party Cyber Security"],
        "key_areas_ar": ["قيادة وحوكمة الأمن السيبراني", "إدارة مخاطر الأمن السيبراني", "عمليات وتقنية الأمن السيبراني", "أمن الأطراف الثالثة السيبراني"],
        "quiz": [
            {"q": "Which institutions must comply with SAMA CSF?", "q_ar": "أي مؤسسات يجب أن تمتثل لإطار ساما؟", "options": ["All Saudi companies","Only international banks","All SAMA-regulated financial institutions","Only fintech startups"], "options_ar": ["جميع الشركات السعودية","البنوك الدولية فقط","جميع المؤسسات المالية المنظمة من ساما","شركات التقنية المالية الناشئة فقط"], "answer": 2},
            {"q": "How many domains does SAMA CSF have?", "q_ar": "كم عدد مجالات إطار ساما؟", "options": ["3","4","5","7"], "options_ar": ["3","4","5","7"], "answer": 1},
            {"q": "What does SAMA require annually?", "q_ar": "ماذا تتطلب ساما سنوياً؟", "options": ["Nothing","Self-assessment submission and periodic examinations","Only a meeting","External audit only"], "options_ar": ["لا شيء","تقديم التقييم الذاتي والفحوصات الدورية","اجتماع فقط","تدقيق خارجي فقط"], "answer": 1},
            {"q": "What international frameworks does SAMA CSF align with?", "q_ar": "مع أي أطر دولية يتوافق إطار ساما؟", "options": ["PCI DSS only","ISO 27001 and NIST","COBIT only","No international alignment"], "options_ar": ["PCI DSS فقط","ISO 27001 و NIST","COBIT فقط","لا توافق دولي"], "answer": 1},
            {"q": "What is included in SAMA CSF Domain 4?", "q_ar": "ماذا يتضمن المجال 4 من إطار ساما؟", "options": ["Employee training","Physical security","Third-party cybersecurity management","Budget planning"], "options_ar": ["تدريب الموظفين","الأمن المادي","إدارة الأمن السيبراني للأطراف الثالثة","تخطيط الميزانية"], "answer": 2}
        ]
    },
    "SAMA BCM (Business Continuity Management)": {
        "short": "SAMA BCM", "type": "framework", "regulator": "SAMA", "region": "KSA",
        "controls": "Business continuity and disaster recovery requirements", "applicability": "All SAMA-regulated financial institutions",
        "domains_ar": "متطلبات استمرارية الأعمال والتعافي من الكوارث", "applicability_ar": "جميع المؤسسات المالية المنظمة من ساما",
        "key_areas": ["Business Impact Analysis", "BCM Strategy", "BCM Plans & Procedures", "Testing & Exercises", "Crisis Management"],
        "key_areas_ar": ["تحليل أثر الأعمال", "استراتيجية استمرارية الأعمال", "خطط وإجراءات الاستمرارية", "الاختبار والتمارين", "إدارة الأزمات"],
        "quiz": [
            {"q": "What is BIA in the context of SAMA BCM?", "q_ar": "ما هو BIA في سياق استمرارية أعمال ساما؟", "options": ["Business Intelligence Analysis","Business Impact Analysis","Basic IT Assessment","Budget Impact Assessment"], "options_ar": ["تحليل ذكاء الأعمال","تحليل أثر الأعمال","تقييم أساسي لتقنية المعلومات","تقييم أثر الميزانية"], "answer": 1},
            {"q": "How often must BCM exercises be conducted?", "q_ar": "كم مرة يجب إجراء تمارين استمرارية الأعمال؟", "options": ["Once every 5 years","At least annually with various exercise types","Only after an actual incident","Never required"], "options_ar": ["مرة كل 5 سنوات","سنوياً على الأقل مع أنواع مختلفة من التمارين","فقط بعد حادث فعلي","غير مطلوب أبداً"], "answer": 1},
            {"q": "What must BIA identify?", "q_ar": "ما يجب أن يحدده تحليل أثر الأعمال؟", "options": ["Only IT systems","Critical processes, RTOs, RPOs, and dependencies","Only employee roles","Marketing strategies"], "options_ar": ["أنظمة تقنية المعلومات فقط","العمليات الحرجة وRTO وRPO والاعتماديات","أدوار الموظفين فقط","استراتيجيات التسويق"], "answer": 1},
            {"q": "What is RTO?", "q_ar": "ما هو RTO؟", "options": ["Return To Office","Recovery Time Objective - maximum acceptable downtime","Real-Time Operations","Risk Transfer Obligation"], "options_ar": ["العودة للمكتب","هدف وقت الاستعادة - أقصى وقت توقف مقبول","العمليات الفورية","التزام نقل المخاطر"], "answer": 1},
            {"q": "What does SAMA require during a crisis?", "q_ar": "ماذا تتطلب ساما أثناء الأزمة؟", "options": ["Wait for instructions","Activate crisis management plan, escalate, and communicate per BCM procedures","Close the institution","Only notify customers"], "options_ar": ["انتظار التعليمات","تفعيل خطة إدارة الأزمات والتصعيد والتواصل وفق إجراءات BCM","إغلاق المؤسسة","إخطار العملاء فقط"], "answer": 1}
        ]
    },
    "CMA Cybersecurity Guidelines": {
        "short": "CMA Cybersecurity", "type": "guideline", "regulator": "CMA", "region": "KSA",
        "controls": "Cybersecurity requirements for capital market participants", "applicability": "CMA-regulated entities including brokers, asset managers, and exchanges",
        "domains_ar": "متطلبات الأمن السيبراني للمشاركين في سوق رأس المال", "applicability_ar": "الجهات المنظمة من هيئة السوق المالية بما في ذلك الوسطاء ومديرو الأصول والبورصات",
        "key_areas": ["Trading System Security", "Market Data Protection", "Investor Data Privacy", "Operational Resilience", "Compliance Reporting"],
        "key_areas_ar": ["أمن أنظمة التداول", "حماية بيانات السوق", "خصوصية بيانات المستثمرين", "المرونة التشغيلية", "إبلاغ الامتثال"],
        "quiz": [
            {"q": "Who does CMA Cybersecurity Guidelines apply to?", "q_ar": "على من تنطبق إرشادات الأمن السيبراني لهيئة السوق المالية؟", "options": ["All companies","CMA-regulated capital market participants","Only banks","Only technology firms"], "options_ar": ["جميع الشركات","المشاركون في سوق رأس المال المنظمون من الهيئة","البنوك فقط","شركات التقنية فقط"], "answer": 1},
            {"q": "What trading system control is required?", "q_ar": "ما ضوابط أنظمة التداول المطلوبة؟", "options": ["No controls","Protection of trading systems against unauthorized access, manipulation, and disruption","Only virus scanning","Speed optimization only"], "options_ar": ["لا ضوابط","حماية أنظمة التداول من الوصول غير المصرح والتلاعب والتعطيل","فحص الفيروسات فقط","تحسين السرعة فقط"], "answer": 1},
            {"q": "What investor data protection does CMA require?", "q_ar": "ما حماية بيانات المستثمرين التي تتطلبها الهيئة؟", "options": ["None","Encryption, access controls, and privacy protection for all investor personal and financial data","Only name protection","Public disclosure required"], "options_ar": ["لا شيء","التشفير والتحكم بالوصول وحماية الخصوصية لجميع بيانات المستثمرين الشخصية والمالية","حماية الاسم فقط","الكشف العام مطلوب"], "answer": 1},
            {"q": "What operational resilience does CMA mandate?", "q_ar": "ما المرونة التشغيلية التي تفرضها الهيئة؟", "options": ["No requirement","Business continuity plans for critical trading and settlement operations","Only backup power","Office relocation plan only"], "options_ar": ["لا متطلب","خطط استمرارية أعمال لعمليات التداول والتسوية الحرجة","طاقة احتياطية فقط","خطة نقل المكتب فقط"], "answer": 1},
            {"q": "How must cybersecurity incidents be reported to CMA?", "q_ar": "كيف يجب الإبلاغ عن حوادث الأمن السيبراني للهيئة؟", "options": ["No reporting required","Immediately per CMA incident reporting guidelines with full impact assessment","Only if requested","Annual summary only"], "options_ar": ["لا إبلاغ مطلوب","فوراً وفق إرشادات الإبلاغ عن الحوادث مع تقييم أثر كامل","فقط إذا طُلب","ملخص سنوي فقط"], "answer": 1}
        ]
    },
    "CITC Cybersecurity Regulations": {
        "short": "CITC Cybersecurity", "type": "regulation", "regulator": "CITC", "region": "KSA",
        "controls": "Telecom and ICT sector cybersecurity requirements", "applicability": "Telecommunications service providers, ISPs, and ICT companies regulated by CITC",
        "domains_ar": "متطلبات الأمن السيبراني لقطاع الاتصالات وتقنية المعلومات", "applicability_ar": "مزودو خدمات الاتصالات ومزودو خدمات الإنترنت وشركات تقنية المعلومات المنظمة من هيئة الاتصالات",
        "key_areas": ["Telecom Infrastructure Security", "Customer Data Protection", "Network Monitoring", "Incident Reporting to CITC", "Service Availability"],
        "key_areas_ar": ["أمن البنية التحتية للاتصالات", "حماية بيانات العملاء", "مراقبة الشبكة", "الإبلاغ عن الحوادث لهيئة الاتصالات", "توفر الخدمة"],
        "quiz": [
            {"q": "Which sector does CITC cybersecurity regulations primarily target?", "q_ar": "أي قطاع تستهدف لوائح الأمن السيبراني لهيئة الاتصالات بشكل أساسي؟", "options": ["Healthcare","Telecommunications and ICT sector","Education","Manufacturing"], "options_ar": ["الرعاية الصحية","قطاع الاتصالات وتقنية المعلومات","التعليم","التصنيع"], "answer": 1},
            {"q": "What infrastructure must telecom providers protect?", "q_ar": "ما البنية التحتية التي يجب على مزودي الاتصالات حمايتها؟", "options": ["Only offices","Core network infrastructure, subscriber data, and service platforms","Only customer-facing websites","Only billing systems"], "options_ar": ["المكاتب فقط","البنية التحتية للشبكة الأساسية وبيانات المشتركين ومنصات الخدمة","مواقع العملاء فقط","أنظمة الفوترة فقط"], "answer": 1},
            {"q": "What monitoring does CITC require?", "q_ar": "ما المراقبة التي تتطلبها هيئة الاتصالات؟", "options": ["No monitoring","Continuous network monitoring with anomaly detection and threat intelligence","Monthly checks only","Only during business hours"], "options_ar": ["لا مراقبة","مراقبة مستمرة للشبكة مع كشف الشذوذ واستخبارات التهديدات","فحوصات شهرية فقط","أثناء ساعات العمل فقط"], "answer": 1},
            {"q": "How must telecom incidents be reported?", "q_ar": "كيف يجب الإبلاغ عن حوادث الاتصالات؟", "options": ["No reporting needed","Promptly to CITC with impact assessment and remediation plan","Only if the media reports it","Annual summary only"], "options_ar": ["لا إبلاغ مطلوب","فوراً لهيئة الاتصالات مع تقييم الأثر وخطة المعالجة","فقط إذا نشرتها وسائل الإعلام","ملخص سنوي فقط"], "answer": 1},
            {"q": "What service availability requirement exists?", "q_ar": "ما متطلب توفر الخدمة الموجود؟", "options": ["No requirement","Minimum uptime SLAs and redundancy for critical telecom services","Best effort only","60% uptime is acceptable"], "options_ar": ["لا متطلب","اتفاقيات مستوى خدمة للحد الأدنى من التوفر والتكرار للخدمات الحرجة","أفضل جهد فقط","توفر 60% مقبول"], "answer": 1}
        ]
    },
}

# ── Seed grc_frameworks from DOMAIN_FRAMEWORKS (deferred from init_db) ──
def seed_frameworks():
    """Populate grc_frameworks table from DOMAIN_FRAMEWORKS dict if empty."""
    try:
        conn = sqlite3.connect(config.DB_PATH)
        conn.row_factory = sqlite3.Row
        existing_fw = conn.execute('SELECT COUNT(*) FROM grc_frameworks').fetchone()[0]
        if existing_fw == 0:
            fw_rows = []
            sort_idx = 0
            for domain_key, regions in DOMAIN_FRAMEWORKS.items():
                for region, frameworks in regions.items():
                    for fw_name in frameworks:
                        fw_rows.append((domain_key, region, fw_name, '', 1, sort_idx))
                        sort_idx += 1
            conn.executemany(
                'INSERT OR IGNORE INTO grc_frameworks (domain, region, name, description, is_active, sort_order) VALUES (?,?,?,?,?,?)',
                fw_rows
            )
            conn.commit()
            print(f"✅ Seeded {len(fw_rows)} frameworks into grc_frameworks table", flush=True)
        conn.close()
    except Exception as e:
        print(f"Framework seeding (non-fatal): {e}", flush=True)

seed_frameworks()

# Arabic translations for framework names
FRAMEWORK_AR_NAMES = {
    "NCA ECC (Essential Cybersecurity Controls)": "الضوابط الأساسية للأمن السيبراني (NCA ECC)",
    "NCA CSCC (Critical Systems Cybersecurity Controls)": "ضوابط الأمن السيبراني للأنظمة الحساسة (NCA CSCC)",
    "NCA DCC (Data Cybersecurity Controls)": "ضوابط الأمن السيبراني للبيانات (NCA DCC)",
    "NCA OTCC (Operational Technology Cybersecurity Controls)": "ضوابط الأمن السيبراني للتقنيات التشغيلية (NCA OTCC)",
    "NCA TCC (Telework Cybersecurity Controls)": "ضوابط الأمن السيبراني للعمل عن بعد (NCA TCC)",
    "NCA OSMACC (Social Media Cybersecurity Controls)": "ضوابط الأمن السيبراني لوسائل التواصل الاجتماعي (NCA OSMACC)",
    "NCA CCC (Cloud Cybersecurity Controls)": "ضوابط الأمن السيبراني السحابية (NCA CCC)",
    "NCA NCS (National Cryptographic Standards)": "المعايير الوطنية للتشفير (NCA NCS)",
    "NCA CGIoT (Cybersecurity Guidelines for IoT)": "إرشادات الأمن السيبراني لإنترنت الأشياء (NCA CGIoT)",
    "SAMA CSF (Cybersecurity Framework)": "إطار الأمن السيبراني لمؤسسة النقد (SAMA CSF)",
    "SAMA BCM (Business Continuity Management)": "إدارة استمرارية الأعمال لمؤسسة النقد (SAMA BCM)",
    "PDPL (Personal Data Protection Law)": "نظام حماية البيانات الشخصية (PDPL)",
    "NDMO Data Governance Framework": "إطار حوكمة البيانات - مكتب إدارة البيانات الوطنية (NDMO)",
    "SDAIA AI Ethics Framework": "إطار أخلاقيات الذكاء الاصطناعي - سدايا (SDAIA)",
}

def translate_framework_ar(fw_name):
    """Translate framework name to Arabic if available."""
    return FRAMEWORK_AR_NAMES.get(fw_name, fw_name)

def translate_frameworks_list_ar(frameworks_list_str):
    """Translate a comma-separated frameworks list to Arabic."""
    frameworks = [f.strip() for f in frameworks_list_str.split(',')]
    translated = [translate_framework_ar(f) for f in frameworks]
    return '، '.join(translated)
DOMAIN_TECHNOLOGIES = {
    "cyber": {
        "en": {
            "Security Operations": ["SIEM", "SOAR", "SOC", "Threat Intelligence Platform", "Log Management"],
            "Endpoint Security": ["EDR/XDR", "Antivirus/Anti-malware", "Mobile Device Management (MDM)", "Endpoint DLP"],
            "Network Security": ["Next-Gen Firewall", "IDS/IPS", "Network Access Control (NAC)", "Web Proxy", "DNS Security"],
            "Identity & Access": ["IAM", "PAM", "MFA/2FA", "SSO", "Directory Services (AD/LDAP)"],
            "Data Protection": ["DLP", "Encryption (at-rest/in-transit)", "Backup & Recovery", "Data Classification"],
            "Application Security": ["WAF", "SAST/DAST", "API Security", "Code Review Tools"],
            "Cloud Security": ["CASB", "CSPM", "CWPP", "Cloud IAM"],
            "GRC Tools": ["Vulnerability Scanner", "Penetration Testing", "Compliance Management", "Risk Register"]
        },
        "ar": {
            "العمليات الأمنية": ["SIEM", "SOAR", "مركز العمليات الأمنية", "منصة استخبارات التهديدات", "إدارة السجلات"],
            "أمن النقاط الطرفية": ["EDR/XDR", "مكافحة الفيروسات", "إدارة الأجهزة المحمولة", "DLP للنقاط الطرفية"],
            "أمن الشبكات": ["جدار الحماية المتقدم", "IDS/IPS", "التحكم بالوصول للشبكة", "بروكسي الويب", "أمن DNS"],
            "الهوية والوصول": ["IAM", "PAM", "المصادقة متعددة العوامل", "SSO", "خدمات الدليل"],
            "حماية البيانات": ["DLP", "التشفير", "النسخ الاحتياطي والاستعادة", "تصنيف البيانات"],
            "أمن التطبيقات": ["WAF", "SAST/DAST", "أمن API", "أدوات مراجعة الكود"],
            "أمن السحابة": ["CASB", "CSPM", "CWPP", "IAM السحابي"],
            "أدوات الحوكمة": ["ماسح الثغرات", "اختبار الاختراق", "إدارة الامتثال", "سجل المخاطر"]
        }
    },
    "data": {
        "en": {
            "Data Governance": ["Data Catalog", "Metadata Management", "Data Lineage", "Business Glossary"],
            "Data Quality": ["Data Profiling", "Data Cleansing", "Master Data Management (MDM)", "Data Validation"],
            "Data Security": ["Data Masking", "Tokenization", "Database Encryption", "Access Controls"],
            "Data Privacy": ["Consent Management", "Privacy Impact Assessment", "Data Subject Rights Management", "Cookie Management"],
            "Data Integration": ["ETL/ELT Tools", "Data Virtualization", "API Management", "Data Replication"],
            "Data Analytics": ["BI Platform", "Data Warehouse", "Data Lake", "Reporting Tools"],
            "Data Lifecycle": ["Archiving Solutions", "Retention Management", "Secure Disposal", "Backup Systems"]
        },
        "ar": {
            "حوكمة البيانات": ["كتالوج البيانات", "إدارة البيانات الوصفية", "تتبع مسار البيانات", "قاموس الأعمال"],
            "جودة البيانات": ["تحليل البيانات", "تنظيف البيانات", "إدارة البيانات الرئيسية", "التحقق من البيانات"],
            "أمن البيانات": ["إخفاء البيانات", "الترميز", "تشفير قواعد البيانات", "ضوابط الوصول"],
            "خصوصية البيانات": ["إدارة الموافقات", "تقييم أثر الخصوصية", "إدارة حقوق أصحاب البيانات", "إدارة الكوكيز"],
            "تكامل البيانات": ["أدوات ETL/ELT", "المحاكاة الافتراضية للبيانات", "إدارة API", "نسخ البيانات"],
            "تحليل البيانات": ["منصة ذكاء الأعمال", "مستودع البيانات", "بحيرة البيانات", "أدوات التقارير"],
            "دورة حياة البيانات": ["حلول الأرشفة", "إدارة الاحتفاظ", "الإتلاف الآمن", "أنظمة النسخ الاحتياطي"]
        }
    },
    "ai": {
        "en": {
            "ML Infrastructure": ["ML Platform", "Model Registry", "Feature Store", "Experiment Tracking"],
            "Data for AI": ["Data Labeling", "Training Data Management", "Synthetic Data Generation", "Data Versioning"],
            "Model Development": ["AutoML", "Notebook Environment", "Model Training Pipeline", "Hyperparameter Tuning"],
            "Model Operations": ["Model Deployment", "Model Monitoring", "A/B Testing", "Model Versioning"],
            "AI Governance": ["Model Documentation", "Bias Detection", "Explainability Tools", "Audit Trail"],
            "AI Security": ["Adversarial Testing", "Model Encryption", "Secure Inference", "Access Controls"]
        },
        "ar": {
            "بنية تعلم الآلة": ["منصة ML", "سجل النماذج", "مخزن الميزات", "تتبع التجارب"],
            "بيانات الذكاء الاصطناعي": ["تصنيف البيانات", "إدارة بيانات التدريب", "توليد البيانات الاصطناعية", "إصدارات البيانات"],
            "تطوير النماذج": ["AutoML", "بيئة Notebook", "خط أنابيب التدريب", "ضبط المعاملات"],
            "عمليات النماذج": ["نشر النماذج", "مراقبة النماذج", "اختبار A/B", "إصدارات النماذج"],
            "حوكمة الذكاء الاصطناعي": ["توثيق النماذج", "كشف التحيز", "أدوات التفسير", "سجل المراجعة"],
            "أمن الذكاء الاصطناعي": ["الاختبار العدائي", "تشفير النماذج", "الاستدلال الآمن", "ضوابط الوصول"]
        }
    },
    "dt": {
        "en": {
            "Digital Platforms": ["Enterprise Portal", "Mobile Apps", "Customer Experience Platform", "Digital Workplace"],
            "Integration": ["ESB/Integration Platform", "API Gateway", "iPaaS", "Microservices"],
            "Process Automation": ["RPA", "BPM", "Workflow Engine", "Low-Code Platform"],
            "Cloud Services": ["IaaS", "PaaS", "SaaS", "Hybrid Cloud"],
            "Analytics & Insights": ["Big Data Platform", "Real-time Analytics", "Predictive Analytics", "Dashboard/KPI Tools"],
            "Collaboration": ["Unified Communications", "Document Management", "Project Management", "Knowledge Management"],
            "Customer Engagement": ["CRM", "Marketing Automation", "Chatbots", "Omnichannel Platform"]
        },
        "ar": {
            "المنصات الرقمية": ["البوابة المؤسسية", "تطبيقات الجوال", "منصة تجربة العميل", "مكان العمل الرقمي"],
            "التكامل": ["منصة التكامل", "بوابة API", "iPaaS", "الخدمات المصغرة"],
            "أتمتة العمليات": ["RPA", "إدارة العمليات", "محرك سير العمل", "منصة Low-Code"],
            "الخدمات السحابية": ["IaaS", "PaaS", "SaaS", "السحابة الهجينة"],
            "التحليلات": ["منصة البيانات الضخمة", "التحليلات الفورية", "التحليلات التنبؤية", "لوحات المؤشرات"],
            "التعاون": ["الاتصالات الموحدة", "إدارة الوثائق", "إدارة المشاريع", "إدارة المعرفة"],
            "تفاعل العملاء": ["CRM", "أتمتة التسويق", "روبوتات المحادثة", "منصة القنوات المتعددة"]
        }
    },
    "global": {
        "en": {
            "Quality Management": ["QMS Software", "Document Control", "CAPA Management", "Audit Management"],
            "Risk Management": ["ERM Platform", "Risk Register", "Risk Assessment Tools", "Incident Management"],
            "Compliance": ["Compliance Management", "Policy Management", "Training Management", "Certification Tracking"],
            "Business Continuity": ["BCP Platform", "DR Solutions", "Crisis Management", "Emergency Notification"],
            "Information Security": ["ISMS Platform", "Asset Management", "Vulnerability Management", "Security Awareness"]
        },
        "ar": {
            "إدارة الجودة": ["برنامج QMS", "التحكم بالوثائق", "إدارة CAPA", "إدارة التدقيق"],
            "إدارة المخاطر": ["منصة ERM", "سجل المخاطر", "أدوات تقييم المخاطر", "إدارة الحوادث"],
            "الامتثال": ["إدارة الامتثال", "إدارة السياسات", "إدارة التدريب", "تتبع الشهادات"],
            "استمرارية الأعمال": ["منصة BCP", "حلول DR", "إدارة الأزمات", "الإشعارات الطارئة"],
            "أمن المعلومات": ["منصة ISMS", "إدارة الأصول", "إدارة الثغرات", "التوعية الأمنية"]
        }
    },
    "erm": {
        "en": {
            "Risk Identification": ["Risk Register Software", "Risk Taxonomy", "Scenario Analysis Tools", "Risk Surveys/Questionnaires", "Emerging Risk Radar"],
            "Risk Assessment": ["Quantitative Risk Analysis", "Monte Carlo Simulation", "Bow-Tie Analysis", "FMEA Tools", "Bayesian Networks"],
            "Risk Monitoring": ["KRI Dashboards", "Risk Heat Maps", "Real-time Risk Monitoring", "Early Warning Systems", "Risk Aggregation"],
            "Risk Appetite & Tolerance": ["Risk Appetite Framework", "Tolerance Thresholds", "Risk Capacity Models", "Board Reporting"],
            "Business Continuity": ["BCP/DRP Platform", "Crisis Management System", "Emergency Communication", "Business Impact Analysis"],
            "Compliance & Controls": ["Control Self-Assessment", "Internal Audit Software", "Policy Management", "Regulatory Change Management"],
            "Insurance & Transfer": ["Insurance Management Platform", "Claims Management", "Risk Transfer Analysis", "Captive Management"],
            "GRC Integration": ["Enterprise GRC Platform", "Integrated Assurance", "Three Lines Model", "Risk Culture Assessment"]
        },
        "ar": {
            "تحديد المخاطر": ["برنامج سجل المخاطر", "تصنيف المخاطر", "أدوات تحليل السيناريوهات", "استبيانات المخاطر", "رادار المخاطر الناشئة"],
            "تقييم المخاطر": ["التحليل الكمي للمخاطر", "محاكاة مونت كارلو", "تحليل ربطة العنق", "أدوات FMEA", "الشبكات البايزية"],
            "مراقبة المخاطر": ["لوحات مؤشرات المخاطر", "خرائط المخاطر الحرارية", "المراقبة الفورية للمخاطر", "أنظمة الإنذار المبكر", "تجميع المخاطر"],
            "شهية المخاطر والتحمل": ["إطار شهية المخاطر", "عتبات التحمل", "نماذج القدرة على المخاطر", "تقارير مجلس الإدارة"],
            "استمرارية الأعمال": ["منصة BCP/DRP", "نظام إدارة الأزمات", "الاتصالات الطارئة", "تحليل تأثير الأعمال"],
            "الامتثال والضوابط": ["التقييم الذاتي للضوابط", "برنامج التدقيق الداخلي", "إدارة السياسات", "إدارة التغييرات التنظيمية"],
            "التأمين والنقل": ["منصة إدارة التأمين", "إدارة المطالبات", "تحليل نقل المخاطر", "إدارة شركات التأمين الأسيرة"],
            "تكامل الحوكمة": ["منصة GRC المتكاملة", "التأكيد المتكامل", "نموذج الخطوط الثلاثة", "تقييم ثقافة المخاطر"]
        }
    }
}

# Enhanced Risk Categories with scenarios
RISK_CATEGORIES = {
    "cyber": {
        "en": {
            "Access Control": ["Unauthorized access to systems", "Privilege escalation", "Credential theft", "Insider threat", "Session hijacking"],
            "Network Security": ["Network intrusion", "DDoS attack", "Man-in-the-middle attack", "DNS poisoning", "Lateral movement"],
            "Data Protection": ["Data breach", "Data leakage", "Unauthorized data access", "Data corruption", "Ransomware encryption"],
            "Endpoint Security": ["Malware infection", "Zero-day exploit", "USB-based attack", "Remote access trojan", "Cryptomining"],
            "Application Security": ["SQL injection", "XSS attack", "API abuse", "Broken authentication", "Insecure deserialization"],
            "Cloud Security": ["Cloud misconfiguration", "Data exposure in cloud", "Account hijacking", "Insecure APIs", "Shadow IT"],
            "Social Engineering": ["Phishing attack", "Spear phishing", "Business email compromise", "Vishing", "Pretexting"],
            "Third Party Risk": ["Vendor breach", "Supply chain attack", "Third-party data exposure", "Service provider failure"],
            "Operational Technology": ["SCADA attack", "ICS compromise", "Physical-cyber attack", "OT network breach"],
            "Incident Response": ["Delayed detection", "Inadequate response", "Evidence loss", "Communication failure"]
        },
        "ar": {
            "التحكم بالوصول": ["وصول غير مصرح للأنظمة", "تصعيد الصلاحيات", "سرقة بيانات الاعتماد", "التهديد الداخلي", "اختطاف الجلسة"],
            "أمن الشبكات": ["اختراق الشبكة", "هجوم DDoS", "هجوم الوسيط", "تسميم DNS", "الحركة الجانبية"],
            "حماية البيانات": ["خرق البيانات", "تسرب البيانات", "وصول غير مصرح للبيانات", "تلف البيانات", "تشفير الفدية"],
            "أمن النقاط الطرفية": ["إصابة بالبرمجيات الخبيثة", "استغلال الثغرات الصفرية (Zero-Day)", "هجوم USB", "حصان طروادة", "التعدين الخبيث"],
            "أمن التطبيقات": ["حقن SQL", "هجوم XSS", "إساءة استخدام API", "مصادقة معطلة", "إلغاء تسلسل غير آمن"],
            "أمن السحابة": ["سوء تكوين السحابة", "كشف البيانات السحابية", "اختطاف الحساب", "APIs غير آمنة", "Shadow IT"],
            "الهندسة الاجتماعية": ["هجوم التصيد", "التصيد الموجه", "اختراق البريد التجاري", "التصيد الصوتي", "الذريعة"],
            "مخاطر الأطراف الثالثة": ["اختراق المورد", "هجوم سلسلة التوريد", "كشف بيانات الطرف الثالث", "فشل مزود الخدمة"],
            "التقنيات التشغيلية": ["هجوم SCADA", "اختراق ICS", "هجوم فيزيائي-سيبراني", "اختراق شبكة OT"],
            "الاستجابة للحوادث": ["تأخر الكشف", "استجابة غير كافية", "فقدان الأدلة", "فشل الاتصال"]
        }
    },
    "data": {
        "en": {
            "Data Quality": ["Incomplete data", "Duplicate records", "Data inconsistency", "Outdated information", "Invalid data formats"],
            "Data Privacy": ["Personal data exposure", "Consent violation", "Cross-border transfer issues", "Right to erasure failure", "Purpose limitation breach"],
            "Data Governance": ["Undefined data ownership", "Missing data lineage", "Inconsistent data definitions", "Policy non-compliance", "Metadata gaps"],
            "Data Security": ["Unauthorized data access", "Database breach", "Encryption failure", "Backup exposure", "Insider data theft"],
            "Data Lifecycle": ["Retention policy violation", "Improper disposal", "Archive corruption", "Recovery failure", "Storage overflow"],
            "Data Integration": ["ETL failure", "Data sync issues", "API data exposure", "Migration errors", "Real-time feed disruption"],
            "Regulatory Compliance": ["PDPL violation", "GDPR non-compliance", "Audit findings", "Reporting failures", "Documentation gaps"]
        },
        "ar": {
            "جودة البيانات": ["بيانات ناقصة", "سجلات مكررة", "عدم اتساق البيانات", "معلومات قديمة", "تنسيقات بيانات غير صالحة"],
            "خصوصية البيانات": ["كشف البيانات الشخصية", "انتهاك الموافقة", "مشاكل النقل عبر الحدود", "فشل حق المحو", "انتهاك تحديد الغرض"],
            "حوكمة البيانات": ["ملكية بيانات غير محددة", "مسار بيانات مفقود", "تعريفات بيانات غير متسقة", "عدم الامتثال للسياسة", "فجوات البيانات الوصفية"],
            "أمن البيانات": ["وصول غير مصرح للبيانات", "اختراق قاعدة البيانات", "فشل التشفير", "كشف النسخ الاحتياطية", "سرقة البيانات الداخلية"],
            "دورة حياة البيانات": ["انتهاك سياسة الاحتفاظ", "التخلص غير السليم", "تلف الأرشيف", "فشل الاستعادة", "تجاوز التخزين"],
            "تكامل البيانات": ["فشل ETL", "مشاكل مزامنة البيانات", "كشف بيانات API", "أخطاء الترحيل", "انقطاع التغذية الفورية"],
            "الامتثال التنظيمي": ["انتهاك PDPL", "عدم الامتثال لـ GDPR", "نتائج التدقيق", "فشل التقارير", "فجوات التوثيق"]
        }
    },
    "ai": {
        "en": {
            "Model Bias": ["Demographic bias", "Selection bias", "Measurement bias", "Algorithmic discrimination", "Feedback loop bias"],
            "Data Quality for AI": ["Training data poisoning", "Label errors", "Data drift", "Insufficient training data", "Unrepresentative samples"],
            "Model Performance": ["Model degradation", "Concept drift", "Overfitting", "Underfitting", "Poor generalization"],
            "Explainability": ["Black box decisions", "Lack of interpretability", "Unexplainable outcomes", "Audit trail gaps", "Stakeholder confusion"],
            "AI Security": ["Adversarial attacks", "Model theft", "Model inversion", "Data extraction attacks", "Prompt injection"],
            "Ethical Concerns": ["Autonomous harm", "Privacy invasion", "Unfair treatment", "Manipulation", "Job displacement"],
            "Operational Risks": ["Model failure in production", "Integration issues", "Scalability problems", "Resource constraints", "Dependency failures"],
            "Compliance": ["Regulatory non-compliance", "Documentation gaps", "Consent issues", "Cross-border AI use", "Audit failures"]
        },
        "ar": {
            "تحيز النموذج": ["تحيز ديموغرافي", "تحيز الاختيار", "تحيز القياس", "تمييز خوارزمي", "تحيز حلقة التغذية"],
            "جودة بيانات الذكاء الاصطناعي": ["تسميم بيانات التدريب", "أخطاء التصنيف", "انحراف البيانات", "بيانات تدريب غير كافية", "عينات غير ممثلة"],
            "أداء النموذج": ["تدهور النموذج", "انحراف المفهوم", "الإفراط في التخصيص", "نقص التخصيص", "ضعف التعميم"],
            "قابلية التفسير": ["قرارات الصندوق الأسود", "نقص قابلية التفسير", "نتائج غير قابلة للتفسير", "فجوات سجل المراجعة", "إرباك أصحاب المصلحة"],
            "أمن الذكاء الاصطناعي": ["الهجمات العدائية", "سرقة النموذج", "عكس النموذج", "هجمات استخراج البيانات", "حقن الأوامر"],
            "المخاوف الأخلاقية": ["الضرر المستقل", "انتهاك الخصوصية", "المعاملة غير العادلة", "التلاعب", "استبدال الوظائف"],
            "المخاطر التشغيلية": ["فشل النموذج في الإنتاج", "مشاكل التكامل", "مشاكل قابلية التوسع", "قيود الموارد", "فشل التبعيات"],
            "الامتثال": ["عدم الامتثال التنظيمي", "فجوات التوثيق", "مشاكل الموافقة", "استخدام AI عبر الحدود", "فشل التدقيق"]
        }
    },
    "dt": {
        "en": {
            "Change Management": ["Resistance to change", "Inadequate training", "Cultural barriers", "Communication gaps", "Leadership misalignment"],
            "Technology Integration": ["System incompatibility", "Data migration failures", "API integration issues", "Legacy system constraints", "Vendor lock-in"],
            "Digital Skills": ["Skills shortage", "Knowledge gaps", "Training inadequacy", "Talent retention", "Digital literacy"],
            "Process Disruption": ["Workflow disruption", "Process automation failure", "Business continuity impact", "Operational inefficiency", "Service degradation"],
            "Customer Experience": ["Poor digital experience", "Channel inconsistency", "Accessibility issues", "Response time delays", "Customer data exposure"],
            "Strategic Alignment": ["Misaligned objectives", "ROI uncertainty", "Scope creep", "Priority conflicts", "Resource constraints"],
            "Vendor & Cloud": ["Vendor dependency", "Cloud service outage", "Contract issues", "Cost overruns", "SLA breaches"],
            "Security in Transformation": ["Security gaps during migration", "New attack surfaces", "Access control issues", "Data exposure", "Compliance gaps"]
        },
        "ar": {
            "إدارة التغيير": ["مقاومة التغيير", "تدريب غير كافٍ", "حواجز ثقافية", "فجوات الاتصال", "عدم توافق القيادة"],
            "تكامل التقنية": ["عدم توافق الأنظمة", "فشل ترحيل البيانات", "مشاكل تكامل API", "قيود الأنظمة القديمة", "الارتباط بالمورد"],
            "المهارات الرقمية": ["نقص المهارات", "فجوات المعرفة", "قصور التدريب", "الاحتفاظ بالمواهب", "الثقافة الرقمية"],
            "تعطيل العمليات": ["تعطيل سير العمل", "فشل أتمتة العمليات", "تأثير استمرارية الأعمال", "عدم كفاءة التشغيل", "تدهور الخدمة"],
            "تجربة العميل": ["تجربة رقمية سيئة", "عدم اتساق القنوات", "مشاكل إمكانية الوصول", "تأخر وقت الاستجابة", "كشف بيانات العميل"],
            "التوافق الاستراتيجي": ["أهداف غير متوافقة", "عدم يقين العائد", "زحف النطاق", "تعارض الأولويات", "قيود الموارد"],
            "المورد والسحابة": ["الاعتماد على المورد", "انقطاع الخدمة السحابية", "مشاكل العقود", "تجاوز التكاليف", "انتهاكات SLA"],
            "الأمن في التحول": ["فجوات أمنية أثناء الترحيل", "أسطح هجوم جديدة", "مشاكل التحكم بالوصول", "كشف البيانات", "فجوات الامتثال"]
        }
    },
    "global": {
        "en": {
            "Strategic Risk": ["Market changes", "Competitive disruption", "Regulatory changes", "Technology obsolescence", "Geopolitical factors"],
            "Operational Risk": ["Process failures", "System outages", "Human errors", "Resource constraints", "Supplier disruptions"],
            "Financial Risk": ["Budget overruns", "Cost escalation", "Revenue impact", "Investment loss", "Currency fluctuation"],
            "Compliance Risk": ["Regulatory violations", "Audit findings", "Certification loss", "Legal penalties", "Reporting failures"],
            "Reputational Risk": ["Brand damage", "Customer trust loss", "Media exposure", "Stakeholder concerns", "Social media crisis"],
            "Business Continuity": ["Natural disasters", "Pandemic impact", "Infrastructure failure", "Key person dependency", "Supply chain disruption"],
            "Information Security": ["Data breaches", "Cyber attacks", "Insider threats", "Third-party risks", "Physical security"],
            "Quality Risk": ["Product defects", "Service failures", "Customer complaints", "Non-conformance", "Continuous improvement gaps"]
        },
        "ar": {
            "المخاطر الاستراتيجية": ["تغيرات السوق", "التعطيل التنافسي", "التغييرات التنظيمية", "تقادم التقنية", "العوامل الجيوسياسية"],
            "المخاطر التشغيلية": ["فشل العمليات", "انقطاع الأنظمة", "الأخطاء البشرية", "قيود الموارد", "اضطرابات الموردين"],
            "المخاطر المالية": ["تجاوز الميزانية", "تصاعد التكاليف", "تأثير الإيرادات", "خسارة الاستثمار", "تقلب العملة"],
            "مخاطر الامتثال": ["انتهاكات تنظيمية", "نتائج التدقيق", "فقدان الشهادات", "العقوبات القانونية", "فشل التقارير"],
            "مخاطر السمعة": ["ضرر العلامة التجارية", "فقدان ثقة العميل", "التعرض الإعلامي", "مخاوف أصحاب المصلحة", "أزمة وسائل التواصل"],
            "استمرارية الأعمال": ["الكوارث الطبيعية", "تأثير الجائحة", "فشل البنية التحتية", "الاعتماد على أشخاص رئيسيين", "تعطل سلسلة التوريد"],
            "أمن المعلومات": ["خروقات البيانات", "الهجمات السيبرانية", "التهديدات الداخلية", "مخاطر الأطراف الثالثة", "الأمن المادي"],
            "مخاطر الجودة": ["عيوب المنتج", "فشل الخدمة", "شكاوى العملاء", "عدم المطابقة", "فجوات التحسين المستمر"]
        }
    },
    "erm": {
        "en": {
            "Strategic Risk": ["Market disruption", "Competitive pressure", "Strategic misalignment", "Mergers & acquisitions risk", "Innovation failure", "Stakeholder expectation gap"],
            "Operational Risk": ["Process failure", "System downtime", "Supply chain disruption", "Human error", "Capacity constraints", "Quality failures"],
            "Financial Risk": ["Liquidity risk", "Credit risk", "Market risk", "Currency risk", "Interest rate risk", "Capital adequacy"],
            "Compliance & Legal Risk": ["Regulatory breach", "Contract dispute", "Litigation exposure", "License/permit risk", "Anti-corruption violation", "Sanctions risk"],
            "Reputational Risk": ["Brand damage", "Public trust erosion", "Social media crisis", "Ethical misconduct", "ESG criticism", "Customer complaint escalation"],
            "Geopolitical Risk": ["Political instability", "Trade restrictions", "Sanctions", "Cross-border regulatory conflict", "Nationalization risk", "Regional conflict"],
            "Environmental & Climate Risk": ["Physical climate risk", "Transition risk", "Regulatory carbon risk", "Resource scarcity", "Natural disaster", "Environmental liability"],
            "People & Culture Risk": ["Key person dependency", "Talent attrition", "Skills shortage", "Workplace safety", "Labor dispute", "Culture misalignment"],
            "Technology & Cyber Risk": ["Cyber attack", "Technology obsolescence", "Digital transformation failure", "Data breach", "IT infrastructure failure", "Third-party tech risk"],
            "Emerging & Systemic Risk": ["Pandemic risk", "AI disruption", "Black swan events", "Systemic market collapse", "Interconnected risk cascades", "Regulatory paradigm shift"]
        },
        "ar": {
            "المخاطر الاستراتيجية": ["تعطيل السوق", "الضغط التنافسي", "عدم التوافق الاستراتيجي", "مخاطر الاندماج والاستحواذ", "فشل الابتكار", "فجوة توقعات أصحاب المصلحة"],
            "المخاطر التشغيلية": ["فشل العمليات", "توقف الأنظمة", "تعطل سلسلة التوريد", "الخطأ البشري", "قيود السعة", "فشل الجودة"],
            "المخاطر المالية": ["مخاطر السيولة", "مخاطر الائتمان", "مخاطر السوق", "مخاطر العملة", "مخاطر سعر الفائدة", "كفاية رأس المال"],
            "مخاطر الامتثال والقانون": ["خرق تنظيمي", "نزاع تعاقدي", "التعرض للتقاضي", "مخاطر التراخيص", "انتهاك مكافحة الفساد", "مخاطر العقوبات"],
            "مخاطر السمعة": ["ضرر العلامة التجارية", "تآكل الثقة العامة", "أزمة وسائل التواصل", "سوء السلوك الأخلاقي", "انتقاد ESG", "تصعيد شكاوى العملاء"],
            "المخاطر الجيوسياسية": ["عدم الاستقرار السياسي", "القيود التجارية", "العقوبات", "تعارض تنظيمي عابر للحدود", "مخاطر التأميم", "الصراع الإقليمي"],
            "المخاطر البيئية والمناخية": ["مخاطر المناخ المادية", "مخاطر التحول", "مخاطر الكربون التنظيمية", "شح الموارد", "الكوارث الطبيعية", "المسؤولية البيئية"],
            "مخاطر الأفراد والثقافة": ["الاعتماد على أشخاص رئيسيين", "استنزاف المواهب", "نقص المهارات", "سلامة بيئة العمل", "النزاعات العمالية", "عدم توافق الثقافة"],
            "مخاطر التقنية والسيبرانية": ["الهجمات السيبرانية", "تقادم التقنية", "فشل التحول الرقمي", "خرق البيانات", "فشل البنية التحتية لتقنية المعلومات", "مخاطر تقنية الأطراف الثالثة"],
            "المخاطر الناشئة والنظامية": ["مخاطر الجائحة", "اضطراب الذكاء الاصطناعي", "أحداث البجعة السوداء", "انهيار السوق النظامي", "تتابع المخاطر المترابطة", "تحول النموذج التنظيمي"]
        }
    }
}

DOMAIN_CODES = {
    "Cyber Security": "cyber",
    "Data Management": "data", 
    "Artificial Intelligence": "ai",
    "Digital Transformation": "dt",
    "Global Standards": "global",
    "Enterprise Risk Management": "erm",
    "الأمن السيبراني": "cyber",
    "إدارة البيانات": "data",
    "الذكاء الاصطناعي": "ai",
    "التحول الرقمي": "dt",
    "المعايير العالمية": "global",
    "إدارة المخاطر المؤسسية": "erm"
}

# ============================================================================
# AWARENESS MODULES - Domain Specific Training & Quizzes
# ============================================================================

AWARENESS_MODULES = {
    "cyber": {
        "en": [
            {
                "id": "cyber_phishing",
                "title": "Phishing & Social Engineering",
                "icon": "fas fa-fish",
                "level": "Beginner",
                "duration": "15 min",
                "description": "Learn to identify phishing emails, social engineering tactics, and protect yourself from common attack vectors.",
                "learning_points": [
                    "How to spot phishing emails by checking sender addresses, URLs, and urgent language",
                    "Common social engineering tactics: pretexting, baiting, tailgating, and quid pro quo",
                    "Safe practices: never click suspicious links, verify requests through official channels",
                    "Reporting procedures when you suspect a phishing attempt"
                ],
                "scenario": {
                    "title": "The Urgent Invoice",
                    "text": "You receive an email from 'finance-dept@yourcompany.co' (note: your company domain is yourcompany.com) saying: 'URGENT: Overdue invoice attached. Pay within 24 hours to avoid service interruption.' The email has a ZIP attachment and asks you to enable macros.",
                    "correct_action": "Do not open the attachment. Report it to IT security. The domain is suspicious (.co vs .com), the urgency is a red flag, and ZIP files with macros are a common malware delivery method."
                },
                "quiz": [
                    {"q": "Which of the following is a strong indicator of a phishing email?", "options": ["The email is from a colleague", "The sender domain has a slight misspelling", "The email contains your name", "The email is about a meeting"], "answer": 1},
                    {"q": "An unknown person calls claiming to be IT support and asks for your password. What should you do?", "options": ["Provide the password since IT needs it", "Hang up and call the official IT helpdesk number", "Give only your username", "Ask them to call back later"], "answer": 1},
                    {"q": "What is 'tailgating' in cybersecurity?", "options": ["Following too closely while driving", "Following an authorized person through a secured door", "Sending multiple spam emails", "Monitoring network traffic"], "answer": 1},
                    {"q": "You receive an email with an attachment from your CEO asking you to urgently wire money. What's the best first step?", "options": ["Wire the money immediately", "Reply to the email asking for confirmation", "Call the CEO using a known phone number to verify", "Forward it to all employees as a warning"], "answer": 2},
                    {"q": "Which file type is most commonly used to deliver malware?", "options": [".txt files", ".jpg images", ".exe or macro-enabled documents", ".pdf files only"], "answer": 2}
                ]
            },
            {
                "id": "cyber_passwords",
                "title": "Password Security & MFA",
                "icon": "fas fa-key",
                "level": "Beginner",
                "duration": "10 min",
                "description": "Best practices for creating strong passwords, managing credentials, and using multi-factor authentication.",
                "learning_points": [
                    "Strong passwords: 12+ characters, mix of upper/lower case, numbers, and symbols",
                    "Never reuse passwords across different accounts or systems",
                    "Use a password manager to generate and store unique passwords",
                    "Enable MFA (Multi-Factor Authentication) on all critical accounts",
                    "Understand authentication factors: something you know, have, and are"
                ],
                "scenario": {
                    "title": "The Shared Password",
                    "text": "Your colleague asks to borrow your login credentials to access a shared system while their account is being reset. They say it's just for today and the deadline is tight.",
                    "correct_action": "Never share your credentials. Suggest they contact IT for a temporary account or expedited reset. Sharing passwords violates security policy and makes you accountable for their actions."
                },
                "quiz": [
                    {"q": "What is the recommended minimum password length?", "options": ["6 characters", "8 characters", "12 characters", "4 characters"], "answer": 2},
                    {"q": "Which is the strongest password?", "options": ["Password123", "MyDog'sName", "j7$kL9!mQ2#xP4", "12345678"], "answer": 2},
                    {"q": "What does MFA stand for?", "options": ["Multiple File Access", "Multi-Factor Authentication", "Main Firewall Application", "Managed File Audit"], "answer": 1},
                    {"q": "Which is NOT a valid MFA factor?", "options": ["Fingerprint scan", "SMS code", "Security question", "Your favorite color"], "answer": 3},
                    {"q": "A colleague asks for your password to finish an urgent task. What should you do?", "options": ["Share it just this once", "Give them a hint", "Refuse and suggest contacting IT", "Write it on a sticky note for them"], "answer": 2}
                ]
            },
            {
                "id": "cyber_data_protection",
                "title": "Data Classification & Protection",
                "icon": "fas fa-shield-alt",
                "level": "Intermediate",
                "duration": "20 min",
                "description": "Understanding data classification levels, handling sensitive information, and compliance with data protection regulations.",
                "learning_points": [
                    "Data classification levels: Public, Internal, Confidential, Restricted/Top Secret",
                    "Each classification level requires specific handling, storage, and transmission controls",
                    "Encryption requirements for data at rest and in transit",
                    "Clean desk policy and secure disposal of sensitive documents",
                    "Compliance with NCA ECC, PDPL, and sector-specific regulations"
                ],
                "scenario": {
                    "title": "The USB Drive",
                    "text": "You find a USB drive in the parking lot with a label 'Salary Report Q4'. You're curious about it and consider plugging it into your work computer to find the owner.",
                    "correct_action": "Never plug unknown USB devices into any computer. Report it to security or IT. This is a known attack vector called 'USB baiting' where malware is deliberately planted on drives."
                },
                "quiz": [
                    {"q": "Which data classification typically requires encryption?", "options": ["Public", "Internal only", "Confidential and Restricted", "None of them"], "answer": 2},
                    {"q": "What should you do with printed confidential documents you no longer need?", "options": ["Put them in the regular trash", "Shred them using a cross-cut shredder", "Leave them on your desk", "Recycle them normally"], "answer": 1},
                    {"q": "You need to send a confidential file to a partner. What is the safest method?", "options": ["Personal email attachment", "Encrypted file via approved secure channel", "WhatsApp message", "Unencrypted USB drive"], "answer": 1},
                    {"q": "What does PDPL stand for in Saudi regulations?", "options": ["Private Data Protocol Law", "Personal Data Protection Law", "Public Domain Privacy License", "Primary Data Processing Limit"], "answer": 1},
                    {"q": "An employee accidentally emails confidential data to the wrong person. What should happen first?", "options": ["Ignore it", "Notify the security team immediately", "Delete the sent email", "Ask the recipient to delete it"], "answer": 1}
                ]
            },
            {
                "id": "cyber_incident",
                "title": "Incident Response Basics",
                "icon": "fas fa-ambulance",
                "level": "Intermediate",
                "duration": "15 min",
                "description": "How to recognize, report, and respond to cybersecurity incidents effectively.",
                "learning_points": [
                    "Signs of a security incident: unusual system behavior, unauthorized access alerts, data anomalies",
                    "Incident response steps: Identify, Contain, Eradicate, Recover, Lessons Learned",
                    "Your role: report immediately, preserve evidence, follow established procedures",
                    "Do NOT try to investigate or fix the issue yourself — let the security team handle it",
                    "Documentation is critical: record what you observed, when, and what actions you took"
                ],
                "scenario": {
                    "title": "The Ransomware Alert",
                    "text": "Your computer screen suddenly displays a message saying all your files are encrypted and demands payment in Bitcoin. Some files on the shared drive also appear affected.",
                    "correct_action": "Disconnect from the network immediately (unplug Ethernet or disable WiFi). Do NOT pay the ransom. Call the IT security team right away. Do not restart the computer as this may destroy forensic evidence."
                },
                "quiz": [
                    {"q": "What is the first thing to do when you suspect a cybersecurity incident?", "options": ["Try to fix it yourself", "Report it to the security team immediately", "Turn off your computer", "Wait to see if it resolves"], "answer": 1},
                    {"q": "If your computer is infected with ransomware, should you pay the ransom?", "options": ["Yes, to get files back quickly", "No, it funds criminals and there's no guarantee of recovery", "Only if the amount is small", "Yes, if the company approves"], "answer": 1},
                    {"q": "What is the correct order of incident response?", "options": ["Fix, Report, Forget", "Identify, Contain, Eradicate, Recover, Lessons Learned", "Panic, Restart, Hope", "Eradicate, Identify, Recover"], "answer": 1},
                    {"q": "Why should you preserve evidence during an incident?", "options": ["For social media posts", "For forensic analysis and legal proceedings", "It's not important", "To show your manager"], "answer": 1},
                    {"q": "During an active incident, who should communicate with the media?", "options": ["Any employee", "Only the designated spokesperson/PR team", "The IT team", "The person who found the incident"], "answer": 1}
                ]
            }
        ],
        "ar": [
            {
                "id": "cyber_phishing",
                "title": "التصيد الاحتيالي والهندسة الاجتماعية",
                "icon": "fas fa-fish",
                "level": "مبتدئ",
                "duration": "15 دقيقة",
                "description": "تعلم كيفية التعرف على رسائل التصيد الاحتيالي وأساليب الهندسة الاجتماعية وحماية نفسك من أشهر طرق الهجوم.",
                "learning_points": [
                    "كيفية اكتشاف رسائل التصيد من خلال التحقق من عنوان المرسل والروابط واللغة العاجلة",
                    "أساليب الهندسة الاجتماعية الشائعة: الذرائع، الإغراء، التتبع، والمقايضة",
                    "الممارسات الآمنة: لا تنقر على الروابط المشبوهة، تحقق من الطلبات عبر القنوات الرسمية",
                    "إجراءات الإبلاغ عند الاشتباه في محاولة تصيد احتيالي"
                ],
                "scenario": {
                    "title": "الفاتورة العاجلة",
                    "text": "تلقيت بريداً إلكترونياً من 'finance-dept@yourcompany.co' (ملاحظة: نطاق شركتك هو yourcompany.com) يقول: 'عاجل: فاتورة متأخرة مرفقة. ادفع خلال 24 ساعة لتجنب انقطاع الخدمة.' يحتوي البريد على مرفق ZIP ويطلب تفعيل وحدات الماكرو.",
                    "correct_action": "لا تفتح المرفق. أبلغ فريق أمن تقنية المعلومات. النطاق مشبوه (.co بدلاً من .com)، والاستعجال علامة تحذيرية، وملفات ZIP مع وحدات الماكرو وسيلة شائعة لنشر البرمجيات الخبيثة."
                },
                "quiz": [
                    {"q": "أي مما يلي مؤشر قوي على بريد تصيد احتيالي؟", "options": ["البريد من زميل", "نطاق المرسل به خطأ إملائي بسيط", "البريد يحتوي على اسمك", "البريد عن اجتماع"], "answer": 1},
                    {"q": "اتصل بك شخص مجهول يدعي أنه من الدعم الفني ويطلب كلمة المرور. ماذا تفعل؟", "options": ["أعطيه الكلمة لأن الدعم يحتاجها", "أغلق الخط واتصل بالرقم الرسمي للدعم", "أعطيه اسم المستخدم فقط", "اطلب منه معاودة الاتصال لاحقاً"], "answer": 1},
                    {"q": "ما هو 'التتبع' (Tailgating) في الأمن السيبراني؟", "options": ["القيادة خلف شخص عن قرب", "الدخول خلف شخص مصرح له عبر باب مؤمّن", "إرسال رسائل مزعجة متعددة", "مراقبة حركة الشبكة"], "answer": 1},
                    {"q": "تلقيت بريداً من الرئيس التنفيذي يطلب تحويل أموال عاجل. ما أفضل خطوة أولى؟", "options": ["حوّل المبلغ فوراً", "رد على البريد لتأكيد", "اتصل بالرئيس على رقم معروف للتحقق", "أعد توجيهه للجميع كتحذير"], "answer": 2},
                    {"q": "أي نوع ملفات يُستخدم أكثر لنشر البرمجيات الخبيثة؟", "options": ["ملفات .txt", "صور .jpg", "ملفات .exe أو مستندات بوحدات ماكرو", "ملفات .pdf فقط"], "answer": 2}
                ]
            },
            {
                "id": "cyber_passwords",
                "title": "أمان كلمات المرور والمصادقة المتعددة",
                "icon": "fas fa-key",
                "level": "مبتدئ",
                "duration": "10 دقائق",
                "description": "أفضل الممارسات لإنشاء كلمات مرور قوية وإدارة بيانات الاعتماد واستخدام المصادقة متعددة العوامل.",
                "learning_points": [
                    "كلمات المرور القوية: 12 حرفاً على الأقل، مزيج من الأحرف الكبيرة والصغيرة والأرقام والرموز",
                    "لا تعيد استخدام كلمات المرور عبر حسابات أو أنظمة مختلفة",
                    "استخدم مدير كلمات المرور لإنشاء وتخزين كلمات مرور فريدة",
                    "فعّل المصادقة متعددة العوامل (MFA) على جميع الحسابات المهمة",
                    "عوامل المصادقة: شيء تعرفه، شيء تملكه، شيء أنت عليه"
                ],
                "scenario": {
                    "title": "كلمة المرور المشتركة",
                    "text": "طلب زميلك استعارة بيانات تسجيل الدخول الخاصة بك للوصول إلى نظام مشترك بينما يتم إعادة تعيين حسابه. يقول إنها ليوم واحد فقط والموعد النهائي قريب.",
                    "correct_action": "لا تشارك بيانات الاعتماد أبداً. اقترح عليه الاتصال بتقنية المعلومات للحصول على حساب مؤقت. مشاركة كلمات المرور تنتهك سياسة الأمان وتجعلك مسؤولاً عن تصرفاته."
                },
                "quiz": [
                    {"q": "ما الحد الأدنى الموصى به لطول كلمة المرور؟", "options": ["6 أحرف", "8 أحرف", "12 حرفاً", "4 أحرف"], "answer": 2},
                    {"q": "أي كلمة مرور هي الأقوى؟", "options": ["Password123", "اسم_كلبي", "j7$kL9!mQ2#xP4", "12345678"], "answer": 2},
                    {"q": "ماذا تعني MFA؟", "options": ["الوصول المتعدد للملفات", "المصادقة متعددة العوامل", "تطبيق الجدار الناري الرئيسي", "تدقيق الملفات المُدارة"], "answer": 1},
                    {"q": "أي مما يلي ليس عامل مصادقة صالحاً؟", "options": ["بصمة الإصبع", "رمز SMS", "سؤال الأمان", "لونك المفضل"], "answer": 3},
                    {"q": "طلب زميل كلمة مرورك لإنهاء مهمة عاجلة. ماذا تفعل؟", "options": ["شاركها هذه المرة فقط", "أعطه تلميحاً", "ارفض واقترح الاتصال بتقنية المعلومات", "اكتبها على ورقة لاصقة له"], "answer": 2}
                ]
            },
            {
                "id": "cyber_data_protection",
                "title": "تصنيف البيانات وحمايتها",
                "icon": "fas fa-shield-alt",
                "level": "متوسط",
                "duration": "20 دقيقة",
                "description": "فهم مستويات تصنيف البيانات والتعامل مع المعلومات الحساسة والامتثال لأنظمة حماية البيانات.",
                "learning_points": [
                    "مستويات تصنيف البيانات: عامة، داخلية، سرية، مقيدة/سرية للغاية",
                    "كل مستوى تصنيف يتطلب ضوابط محددة للتعامل والتخزين والنقل",
                    "متطلبات التشفير للبيانات المخزنة والمنقولة",
                    "سياسة المكتب النظيف والتخلص الآمن من المستندات الحساسة",
                    "الامتثال لـ NCA ECC و PDPL واللوائح القطاعية"
                ],
                "scenario": {
                    "title": "ذاكرة USB",
                    "text": "وجدت ذاكرة USB في موقف السيارات عليها ملصق 'تقرير الرواتب Q4'. تشعر بالفضول وتفكر في توصيلها بجهاز العمل لمعرفة صاحبها.",
                    "correct_action": "لا تقم أبداً بتوصيل أجهزة USB مجهولة بأي حاسوب. أبلغ فريق الأمن أو تقنية المعلومات. هذا أسلوب هجوم معروف يُسمى 'إغراء USB' حيث تُزرع البرمجيات الخبيثة عمداً."
                },
                "quiz": [
                    {"q": "أي تصنيف بيانات يتطلب عادةً التشفير؟", "options": ["عامة", "داخلية فقط", "سرية ومقيدة", "لا شيء منها"], "answer": 2},
                    {"q": "ماذا تفعل بالمستندات السرية المطبوعة التي لم تعد تحتاجها؟", "options": ["ضعها في سلة المهملات العادية", "مزقها باستخدام آلة تمزيق متقاطعة", "اتركها على مكتبك", "أعد تدويرها بشكل عادي"], "answer": 1},
                    {"q": "تحتاج لإرسال ملف سري لشريك. ما هي الطريقة الأكثر أماناً؟", "options": ["مرفق بريد شخصي", "ملف مشفر عبر قناة آمنة معتمدة", "رسالة واتساب", "ذاكرة USB غير مشفرة"], "answer": 1},
                    {"q": "ماذا يعني PDPL في اللوائح السعودية؟", "options": ["قانون بروتوكول البيانات الخاصة", "نظام حماية البيانات الشخصية", "رخصة خصوصية النطاق العام", "حد معالجة البيانات الأساسي"], "answer": 1},
                    {"q": "أرسل موظف بيانات سرية بالخطأ لشخص غير مخول. ما أول إجراء؟", "options": ["تجاهل الأمر", "إبلاغ فريق الأمن فوراً", "حذف البريد المرسل", "طلب حذفها من المستلم"], "answer": 1}
                ]
            },
            {
                "id": "cyber_incident",
                "title": "أساسيات الاستجابة للحوادث",
                "icon": "fas fa-ambulance",
                "level": "متوسط",
                "duration": "15 دقيقة",
                "description": "كيفية التعرف على حوادث الأمن السيبراني والإبلاغ عنها والاستجابة لها بفعالية.",
                "learning_points": [
                    "علامات الحادث الأمني: سلوك غير عادي للنظام، تنبيهات وصول غير مصرح به، شذوذ في البيانات",
                    "خطوات الاستجابة للحوادث: التحديد، الاحتواء، الإزالة، الاستعادة، الدروس المستفادة",
                    "دورك: أبلغ فوراً، حافظ على الأدلة، اتبع الإجراءات المعتمدة",
                    "لا تحاول التحقيق أو الإصلاح بنفسك — دع فريق الأمن يتولى الأمر",
                    "التوثيق مهم: سجل ما لاحظته ومتى وما الإجراءات التي اتخذتها"
                ],
                "scenario": {
                    "title": "تنبيه فدية",
                    "text": "ظهرت فجأة رسالة على شاشة حاسوبك تقول إن جميع ملفاتك مشفرة وتطالب بالدفع بالبيتكوين. بعض الملفات على المجلد المشترك تبدو متأثرة أيضاً.",
                    "correct_action": "افصل عن الشبكة فوراً (افصل كابل الإيثرنت أو أوقف الواي فاي). لا تدفع الفدية. اتصل بفريق أمن تقنية المعلومات فوراً. لا تعد تشغيل الحاسوب لأن ذلك قد يدمر الأدلة الجنائية."
                },
                "quiz": [
                    {"q": "ما أول شيء تفعله عند الاشتباه بحادث أمن سيبراني؟", "options": ["حاول إصلاحه بنفسك", "أبلغ فريق الأمن فوراً", "أطفئ حاسوبك", "انتظر لترى إن حُل"], "answer": 1},
                    {"q": "إذا أصيب حاسوبك ببرنامج فدية، هل يجب دفع الفدية؟", "options": ["نعم لاستعادة الملفات بسرعة", "لا، فهذا يموّل المجرمين ولا ضمان للاستعادة", "فقط إذا كان المبلغ صغيراً", "نعم إذا وافقت الشركة"], "answer": 1},
                    {"q": "ما الترتيب الصحيح للاستجابة للحوادث؟", "options": ["إصلاح، إبلاغ، نسيان", "تحديد، احتواء، إزالة، استعادة، دروس مستفادة", "ذعر، إعادة تشغيل، أمل", "إزالة، تحديد، استعادة"], "answer": 1},
                    {"q": "لماذا يجب الحفاظ على الأدلة أثناء الحادث؟", "options": ["للنشر على وسائل التواصل", "للتحليل الجنائي والإجراءات القانونية", "ليست مهمة", "لإظهارها للمدير"], "answer": 1},
                    {"q": "أثناء حادث نشط، من يجب أن يتحدث مع الإعلام؟", "options": ["أي موظف", "المتحدث الرسمي / فريق العلاقات العامة فقط", "فريق تقنية المعلومات", "الشخص الذي اكتشف الحادث"], "answer": 1}
                ]
            }
        ]
    },
    "data": {
        "en": [
            {
                "id": "data_governance",
                "title": "Data Governance Fundamentals",
                "icon": "fas fa-database",
                "level": "Beginner",
                "duration": "15 min",
                "description": "Core principles of data governance, roles, and responsibilities for managing organizational data assets.",
                "learning_points": [
                    "Data governance ensures data quality, security, availability, and usability across the organization",
                    "Key roles: Data Owner, Data Steward, Data Custodian, and Data Consumer",
                    "Data lifecycle: Creation, Storage, Usage, Sharing, Archiving, and Destruction",
                    "Establishing data standards and policies for consistent management"
                ],
                "scenario": {"title": "The Duplicate Records", "text": "Your team discovers that the same customer exists in three different databases with slightly different names and addresses. Orders are being split across records, causing billing errors.", "correct_action": "Report the data quality issue to the Data Steward. Implement a master data management approach with a single source of truth. Establish data entry standards to prevent future duplicates."},
                "quiz": [
                    {"q": "Who is typically responsible for defining data quality standards?", "options": ["IT department only", "Data Steward", "Any employee", "External auditors"], "answer": 1},
                    {"q": "What is a 'single source of truth' in data management?", "options": ["A single database server", "One authoritative data source everyone references", "A backup copy of data", "The CEO's report"], "answer": 1},
                    {"q": "Which is NOT part of the data lifecycle?", "options": ["Creation", "Storage", "Monetization", "Destruction"], "answer": 2},
                    {"q": "What is a Data Custodian responsible for?", "options": ["Business decisions about data", "Technical management and storage of data", "Creating data policies", "Using data for reports"], "answer": 1},
                    {"q": "Why is data governance important?", "options": ["Only for compliance", "To increase IT budgets", "To ensure data quality, security, and proper use", "To slow down operations"], "answer": 2}
                ]
            },
            {
                "id": "data_quality",
                "title": "Data Quality & Integrity",
                "icon": "fas fa-check-double",
                "level": "Intermediate",
                "duration": "15 min",
                "description": "Ensuring data accuracy, completeness, consistency, and timeliness across systems.",
                "learning_points": ["Data quality dimensions: Accuracy, Completeness, Consistency, Timeliness, Validity, Uniqueness", "Impact of poor data quality: wrong decisions, compliance failures, customer dissatisfaction", "Data validation techniques: input validation, cross-referencing, automated checks", "Establishing data quality metrics and monitoring dashboards"],
                "scenario": {"title": "The Missing Records", "text": "A regulatory audit reveals that 15% of transaction records are missing required fields. The compliance team is asking for an explanation and a remediation plan.", "correct_action": "Conduct a root cause analysis to identify where the data quality breaks. Implement mandatory field validation at the data entry point. Create a remediation plan with timeline to fix existing records and prevent recurrence."},
                "quiz": [
                    {"q": "Which is NOT a dimension of data quality?", "options": ["Accuracy", "Completeness", "Popularity", "Timeliness"], "answer": 2},
                    {"q": "What causes the biggest impact from poor data quality?", "options": ["Slower computers", "Incorrect business decisions", "Bigger databases", "More storage costs"], "answer": 1},
                    {"q": "What is data validation?", "options": ["Deleting old data", "Checking data meets defined rules before accepting it", "Backing up data", "Sharing data externally"], "answer": 1},
                    {"q": "How often should data quality be measured?", "options": ["Once a year", "Only during audits", "Continuously through automated monitoring", "Never"], "answer": 2},
                    {"q": "What is a data quality metric?", "options": ["The size of a database", "A measurable indicator of data quality", "A type of database software", "An encryption method"], "answer": 1}
                ]
            },
            {
                "id": "data_privacy",
                "title": "Data Privacy & PDPL Compliance",
                "icon": "fas fa-user-lock",
                "level": "Intermediate",
                "duration": "20 min",
                "description": "Understanding personal data protection requirements under Saudi PDPL and international regulations.",
                "learning_points": ["PDPL principles: lawful processing, purpose limitation, data minimization, accuracy, storage limitation", "Individual rights: access, correction, deletion, portability, objection to processing", "Requirements for cross-border data transfers", "Data breach notification obligations and timelines"],
                "scenario": {"title": "The Marketing Database", "text": "Marketing wants to use customer data collected for service delivery to send promotional emails. They argue it will increase revenue and the data is already available.", "correct_action": "This likely violates purpose limitation. The data was collected for service delivery, not marketing. Marketing needs to obtain separate consent for promotional use or verify that the privacy policy covers marketing as a compatible purpose."},
                "quiz": [
                    {"q": "What is 'purpose limitation' in PDPL?", "options": ["Limiting data storage", "Using data only for the purpose it was collected for", "Limiting database size", "Restricting employee access"], "answer": 1},
                    {"q": "Under PDPL, what right do individuals have regarding their data?", "options": ["No rights", "Right to access and correct their data", "Only the right to delete", "Only government access"], "answer": 1},
                    {"q": "When must a data breach be reported under PDPL?", "options": ["Within 72 hours", "Within 30 days", "Only if data is lost", "Reporting is optional"], "answer": 0},
                    {"q": "What does 'data minimization' mean?", "options": ["Deleting all data", "Collecting only what is necessary for the stated purpose", "Using the smallest database", "Minimizing IT costs"], "answer": 1},
                    {"q": "Can personal data be transferred outside Saudi Arabia?", "options": ["Always freely", "Never under any circumstances", "Only with adequate protections and authorization", "Only to GCC countries"], "answer": 2}
                ]
            }
        ],
        "ar": [
            {
                "id": "data_governance", "title": "أساسيات حوكمة البيانات", "icon": "fas fa-database", "level": "مبتدئ", "duration": "15 دقيقة",
                "description": "المبادئ الأساسية لحوكمة البيانات والأدوار والمسؤوليات لإدارة أصول البيانات المؤسسية.",
                "learning_points": ["حوكمة البيانات تضمن جودة وأمن وتوفر وسهولة استخدام البيانات عبر المنظمة", "الأدوار الرئيسية: مالك البيانات، أمين البيانات، حارس البيانات، مستخدم البيانات", "دورة حياة البيانات: الإنشاء، التخزين، الاستخدام، المشاركة، الأرشفة، الإتلاف", "وضع معايير وسياسات البيانات للإدارة المتسقة"],
                "scenario": {"title": "السجلات المكررة", "text": "اكتشف فريقك أن نفس العميل موجود في ثلاث قواعد بيانات مختلفة بأسماء وعناوين مختلفة قليلاً. الطلبات مقسمة عبر السجلات مما يسبب أخطاء في الفواتير.", "correct_action": "أبلغ أمين البيانات بمشكلة الجودة. طبّق نهج إدارة البيانات الرئيسية بمصدر واحد للحقيقة. ضع معايير إدخال البيانات لمنع التكرار مستقبلاً."},
                "quiz": [
                    {"q": "من المسؤول عادةً عن تحديد معايير جودة البيانات؟", "options": ["قسم تقنية المعلومات فقط", "أمين البيانات", "أي موظف", "المدققون الخارجيون"], "answer": 1},
                    {"q": "ما هو 'المصدر الواحد للحقيقة'؟", "options": ["خادم قاعدة بيانات واحد", "مصدر بيانات واحد موثوق يرجع إليه الجميع", "نسخة احتياطية", "تقرير الرئيس التنفيذي"], "answer": 1},
                    {"q": "أي مما يلي ليس جزءاً من دورة حياة البيانات؟", "options": ["الإنشاء", "التخزين", "التسييل", "الإتلاف"], "answer": 2},
                    {"q": "ما مسؤولية حارس البيانات؟", "options": ["القرارات التجارية حول البيانات", "الإدارة التقنية وتخزين البيانات", "إنشاء سياسات البيانات", "استخدام البيانات للتقارير"], "answer": 1},
                    {"q": "لماذا حوكمة البيانات مهمة؟", "options": ["للامتثال فقط", "لزيادة ميزانية تقنية المعلومات", "لضمان جودة البيانات وأمنها واستخدامها السليم", "لإبطاء العمليات"], "answer": 2}
                ]
            },
            {
                "id": "data_quality", "title": "جودة البيانات وسلامتها", "icon": "fas fa-check-double", "level": "متوسط", "duration": "15 دقيقة",
                "description": "ضمان دقة البيانات واكتمالها واتساقها وتوقيتها عبر الأنظمة.",
                "learning_points": ["أبعاد جودة البيانات: الدقة، الاكتمال، الاتساق، التوقيت، الصلاحية، التفرد", "أثر ضعف الجودة: قرارات خاطئة، فشل في الامتثال، عدم رضا العملاء", "تقنيات التحقق: التحقق من المدخلات، المراجعة التبادلية، الفحوصات الآلية", "وضع مقاييس جودة البيانات ولوحات المراقبة"],
                "scenario": {"title": "السجلات المفقودة", "text": "كشف تدقيق تنظيمي أن 15% من سجلات المعاملات تفتقد حقولاً مطلوبة. فريق الامتثال يطلب تفسيراً وخطة معالجة.", "correct_action": "أجرِ تحليل السبب الجذري لتحديد مكان خلل جودة البيانات. طبّق التحقق الإلزامي من الحقول عند نقطة إدخال البيانات. أنشئ خطة معالجة بجدول زمني."},
                "quiz": [
                    {"q": "أي مما يلي ليس بُعداً لجودة البيانات؟", "options": ["الدقة", "الاكتمال", "الشعبية", "التوقيت"], "answer": 2},
                    {"q": "ما أكبر أثر لضعف جودة البيانات؟", "options": ["بطء الحواسيب", "قرارات تجارية خاطئة", "قواعد بيانات أكبر", "تكاليف تخزين أعلى"], "answer": 1},
                    {"q": "ما هو التحقق من البيانات؟", "options": ["حذف البيانات القديمة", "فحص البيانات وفق قواعد محددة قبل قبولها", "نسخ البيانات احتياطياً", "مشاركة البيانات خارجياً"], "answer": 1},
                    {"q": "كم مرة يجب قياس جودة البيانات؟", "options": ["مرة سنوياً", "أثناء التدقيق فقط", "باستمرار عبر المراقبة الآلية", "أبداً"], "answer": 2},
                    {"q": "ما هو مقياس جودة البيانات؟", "options": ["حجم قاعدة البيانات", "مؤشر قابل للقياس لجودة البيانات", "نوع من برامج قواعد البيانات", "طريقة تشفير"], "answer": 1}
                ]
            },
            {
                "id": "data_privacy", "title": "خصوصية البيانات والامتثال لنظام PDPL", "icon": "fas fa-user-lock", "level": "متوسط", "duration": "20 دقيقة",
                "description": "فهم متطلبات حماية البيانات الشخصية وفق نظام PDPL السعودي واللوائح الدولية.",
                "learning_points": ["مبادئ PDPL: المعالجة المشروعة، تحديد الغرض، تقليل البيانات، الدقة، تحديد التخزين", "حقوق الأفراد: الوصول، التصحيح، الحذف، النقل، الاعتراض على المعالجة", "متطلبات نقل البيانات عبر الحدود", "التزامات إخطار خرق البيانات والجداول الزمنية"],
                "scenario": {"title": "قاعدة بيانات التسويق", "text": "يريد التسويق استخدام بيانات العملاء المجمعة لتقديم الخدمات لإرسال رسائل ترويجية. يقولون إنها ستزيد الإيرادات والبيانات متاحة بالفعل.", "correct_action": "هذا على الأرجح ينتهك تحديد الغرض. جُمعت البيانات لتقديم الخدمات لا للتسويق. يحتاج التسويق للحصول على موافقة منفصلة أو التحقق من أن سياسة الخصوصية تغطي التسويق كغرض متوافق."},
                "quiz": [
                    {"q": "ما هو 'تحديد الغرض' في PDPL؟", "options": ["تحديد مساحة التخزين", "استخدام البيانات فقط للغرض الذي جُمعت من أجله", "تحديد حجم القاعدة", "تقييد وصول الموظفين"], "answer": 1},
                    {"q": "بموجب PDPL، ما الحق الذي يملكه الأفراد بشأن بياناتهم؟", "options": ["لا حقوق", "حق الوصول والتصحيح", "حق الحذف فقط", "وصول الحكومة فقط"], "answer": 1},
                    {"q": "متى يجب الإبلاغ عن خرق البيانات وفق PDPL؟", "options": ["خلال 72 ساعة", "خلال 30 يوماً", "فقط إذا فُقدت البيانات", "الإبلاغ اختياري"], "answer": 0},
                    {"q": "ماذا يعني 'تقليل البيانات'؟", "options": ["حذف جميع البيانات", "جمع الضروري فقط للغرض المحدد", "استخدام أصغر قاعدة بيانات", "تقليل تكاليف تقنية المعلومات"], "answer": 1},
                    {"q": "هل يمكن نقل البيانات الشخصية خارج السعودية؟", "options": ["دائماً بحرية", "أبداً تحت أي ظرف", "فقط مع حماية كافية وتفويض", "فقط لدول الخليج"], "answer": 2}
                ]
            }
        ]
    },
    "ai": {
        "en": [
            {
                "id": "ai_ethics", "title": "AI Ethics & Responsible Use", "icon": "fas fa-balance-scale", "level": "Beginner", "duration": "15 min",
                "description": "Understanding ethical principles for AI development, deployment, and governance.",
                "learning_points": ["Core AI ethics principles: fairness, transparency, accountability, privacy, safety", "Bias in AI: how it occurs in training data, model design, and deployment", "Explainability: users have the right to understand AI decisions that affect them", "Human oversight: AI should augment, not replace, human judgment in critical decisions"],
                "scenario": {"title": "The Biased Hiring Tool", "text": "Your company deployed an AI tool to screen resumes. After 3 months, data shows it consistently ranks male candidates higher. The tool was trained on historical hiring data from the past 10 years.", "correct_action": "Pause the tool immediately. The historical data likely reflects past gender bias. Conduct a fairness audit, retrain with balanced data, implement ongoing bias monitoring, and ensure human review of all AI-assisted hiring decisions."},
                "quiz": [
                    {"q": "What is AI bias?", "options": ["A type of AI software", "Systematic favoritism or discrimination in AI outputs", "Normal AI behavior", "An AI programming language"], "answer": 1},
                    {"q": "Why is AI explainability important?", "options": ["Only for developers", "So users understand how AI decisions are made", "It's not important", "Only for marketing purposes"], "answer": 1},
                    {"q": "What should happen when AI makes a high-risk decision?", "options": ["Accept it automatically", "A human should review it", "Ignore any concerns", "Speed up the process"], "answer": 1},
                    {"q": "How does bias typically enter AI systems?", "options": ["Through biased training data", "Through electricity", "AI cannot be biased", "Only through intentional programming"], "answer": 0},
                    {"q": "What is 'human-in-the-loop' in AI governance?", "options": ["Humans training AI to exercise", "Requiring human review of AI decisions", "A programming framework", "An error in AI"], "answer": 1}
                ]
            },
            {
                "id": "ai_data_handling", "title": "AI Data Handling & Privacy", "icon": "fas fa-brain", "level": "Intermediate", "duration": "15 min",
                "description": "Protecting sensitive data when using AI tools, chatbots, and automated systems.",
                "learning_points": ["Never input confidential, personal, or classified data into public AI tools", "Understand where AI tools store and process your data — cloud vs. on-premise", "AI-generated content may contain hallucinations — always verify critical information", "Organizational policies for approved AI tools and use cases"],
                "scenario": {"title": "The AI Shortcut", "text": "A team member copies a confidential contract into a public AI chatbot to summarize it, saving hours of work. They share the AI summary in a team meeting.", "correct_action": "This is a data breach. The confidential contract was sent to a third-party AI service. Report the incident, assess what data was exposed, and remind the team that confidential data must never be input into public AI tools."},
                "quiz": [
                    {"q": "Is it safe to paste confidential data into public AI tools?", "options": ["Yes, AI tools are always secure", "No, the data may be stored and used for training", "Only short texts", "Only if you delete the chat after"], "answer": 1},
                    {"q": "What is an AI 'hallucination'?", "options": ["AI having dreams", "AI generating false or fabricated information confidently", "A visual AI effect", "Normal AI behavior"], "answer": 1},
                    {"q": "Before using an AI tool at work, you should:", "options": ["Just use whatever is available", "Check if it's approved by your organization's policy", "Ask a friend", "Google it"], "answer": 1},
                    {"q": "Where is data typically processed by cloud AI services?", "options": ["Only on your device", "On external servers, potentially in another country", "Nowhere — it disappears", "Only on your company's servers"], "answer": 1},
                    {"q": "A colleague uses AI to draft an email with client financial data. What's the concern?", "options": ["No concern", "Client data exposure to third-party AI provider", "Email formatting issues", "Spelling errors"], "answer": 1}
                ]
            }
        ],
        "ar": [
            {
                "id": "ai_ethics", "title": "أخلاقيات الذكاء الاصطناعي والاستخدام المسؤول", "icon": "fas fa-balance-scale", "level": "مبتدئ", "duration": "15 دقيقة",
                "description": "فهم المبادئ الأخلاقية لتطوير الذكاء الاصطناعي ونشره وحوكمته.",
                "learning_points": ["المبادئ الأساسية لأخلاقيات الذكاء الاصطناعي: العدالة، الشفافية، المساءلة، الخصوصية، السلامة", "التحيز في الذكاء الاصطناعي: كيف يحدث في بيانات التدريب وتصميم النماذج والنشر", "قابلية التفسير: للمستخدمين الحق في فهم قرارات الذكاء الاصطناعي التي تؤثر عليهم", "الإشراف البشري: الذكاء الاصطناعي يجب أن يعزز الحكم البشري لا أن يحل محله"],
                "scenario": {"title": "أداة التوظيف المتحيزة", "text": "نشرت شركتك أداة ذكاء اصطناعي لفرز السير الذاتية. بعد 3 أشهر، تُظهر البيانات أنها تصنف المرشحين الذكور أعلى باستمرار. دُرّبت الأداة على بيانات توظيف تاريخية من آخر 10 سنوات.", "correct_action": "أوقف الأداة فوراً. البيانات التاريخية تعكس على الأرجح تحيزاً سابقاً. أجرِ تدقيق عدالة، أعد التدريب ببيانات متوازنة، طبّق مراقبة مستمرة للتحيز، وتأكد من المراجعة البشرية."},
                "quiz": [
                    {"q": "ما هو تحيز الذكاء الاصطناعي؟", "options": ["نوع من البرمجيات", "تفضيل أو تمييز منهجي في مخرجات الذكاء الاصطناعي", "سلوك طبيعي", "لغة برمجة"], "answer": 1},
                    {"q": "لماذا قابلية تفسير الذكاء الاصطناعي مهمة؟", "options": ["للمطورين فقط", "ليفهم المستخدمون كيف تُتخذ قرارات الذكاء الاصطناعي", "ليست مهمة", "لأغراض تسويقية فقط"], "answer": 1},
                    {"q": "ماذا يجب أن يحدث عندما يتخذ الذكاء الاصطناعي قراراً عالي المخاطر؟", "options": ["قبوله تلقائياً", "يجب أن يراجعه إنسان", "تجاهل المخاوف", "تسريع العملية"], "answer": 1},
                    {"q": "كيف يدخل التحيز عادةً إلى أنظمة الذكاء الاصطناعي؟", "options": ["عبر بيانات تدريب متحيزة", "عبر الكهرباء", "الذكاء الاصطناعي لا يمكن أن يكون متحيزاً", "فقط عبر البرمجة المتعمدة"], "answer": 0},
                    {"q": "ما هو 'الإنسان في الحلقة'؟", "options": ["تدريب البشر على التمارين", "اشتراط المراجعة البشرية لقرارات الذكاء الاصطناعي", "إطار برمجي", "خطأ في الذكاء الاصطناعي"], "answer": 1}
                ]
            },
            {
                "id": "ai_data_handling", "title": "التعامل مع بيانات الذكاء الاصطناعي والخصوصية", "icon": "fas fa-brain", "level": "متوسط", "duration": "15 دقيقة",
                "description": "حماية البيانات الحساسة عند استخدام أدوات الذكاء الاصطناعي والمحادثات الآلية والأنظمة المؤتمتة.",
                "learning_points": ["لا تدخل أبداً بيانات سرية أو شخصية أو مصنفة في أدوات الذكاء الاصطناعي العامة", "افهم أين تخزن وتعالج أدوات الذكاء الاصطناعي بياناتك — سحابية أم محلية", "المحتوى المُنتج بالذكاء الاصطناعي قد يحتوي على هلوسات — تحقق دائماً من المعلومات المهمة", "السياسات التنظيمية لأدوات الذكاء الاصطناعي المعتمدة وحالات الاستخدام"],
                "scenario": {"title": "اختصار الذكاء الاصطناعي", "text": "نسخ أحد أعضاء الفريق عقداً سرياً في روبوت محادثة ذكاء اصطناعي عام لتلخيصه، موفراً ساعات من العمل. شارك ملخص الذكاء الاصطناعي في اجتماع الفريق.", "correct_action": "هذا خرق بيانات. أُرسل العقد السري لخدمة ذكاء اصطناعي تابعة لطرف ثالث. أبلغ عن الحادث، قيّم البيانات المكشوفة، وذكّر الفريق بعدم إدخال البيانات السرية في أدوات الذكاء الاصطناعي العامة."},
                "quiz": [
                    {"q": "هل من الآمن لصق البيانات السرية في أدوات الذكاء الاصطناعي العامة؟", "options": ["نعم، أدوات الذكاء الاصطناعي آمنة دائماً", "لا، قد تُخزن البيانات وتُستخدم للتدريب", "النصوص القصيرة فقط", "فقط إذا حذفت المحادثة بعدها"], "answer": 1},
                    {"q": "ما هي 'هلوسة' الذكاء الاصطناعي؟", "options": ["أحلام الذكاء الاصطناعي", "إنتاج معلومات خاطئة أو ملفقة بثقة", "تأثير بصري", "سلوك طبيعي"], "answer": 1},
                    {"q": "قبل استخدام أداة ذكاء اصطناعي في العمل، يجب:", "options": ["استخدام أي أداة متاحة", "التحقق من اعتمادها في سياسة منظمتك", "سؤال صديق", "البحث في جوجل"], "answer": 1},
                    {"q": "أين تُعالج البيانات عادةً بواسطة خدمات الذكاء الاصطناعي السحابية؟", "options": ["على جهازك فقط", "على خوادم خارجية، ربما في دولة أخرى", "لا مكان — تختفي", "على خوادم شركتك فقط"], "answer": 1},
                    {"q": "استخدم زميل الذكاء الاصطناعي لصياغة بريد يتضمن بيانات مالية للعملاء. ما المشكلة؟", "options": ["لا مشكلة", "تعرض بيانات العميل لمزود ذكاء اصطناعي خارجي", "مشاكل تنسيق البريد", "أخطاء إملائية"], "answer": 1}
                ]
            }
        ]
    },
    "dt": {
        "en": [
            {
                "id": "dt_change", "title": "Digital Change Management", "icon": "fas fa-sync-alt", "level": "Beginner", "duration": "15 min",
                "description": "Managing organizational change during digital transformation initiatives effectively.",
                "learning_points": ["Change management frameworks: ADKAR, Kotter's 8 Steps", "Stakeholder engagement and communication strategies", "Overcoming resistance to digital change", "Measuring adoption rates and user satisfaction"],
                "scenario": {"title": "The Rejected System", "text": "After 6 months and significant investment, a new ERP system has only 30% adoption. Employees complain it's harder than the old system and continue using spreadsheets.", "correct_action": "Conduct user feedback sessions. Identify specific pain points. Provide role-based training. Identify champions in each department. Gradually phase out old systems while ensuring the new system addresses real workflow needs."},
                "quiz": [
                    {"q": "What does ADKAR stand for?", "options": ["A software framework", "Awareness, Desire, Knowledge, Ability, Reinforcement", "A type of database", "Automated Digital Knowledge And Review"], "answer": 1},
                    {"q": "What is the most common reason digital transformation projects fail?", "options": ["Bad technology", "Insufficient budget", "Poor change management and resistance", "Lack of internet"], "answer": 2},
                    {"q": "What is a 'change champion'?", "options": ["A person who resists change", "An employee who advocates for and supports change", "A project manager", "An external consultant"], "answer": 1},
                    {"q": "How should you handle employee resistance to new systems?", "options": ["Force adoption through mandates only", "Listen, address concerns, provide training and support", "Ignore complaints", "Threaten consequences"], "answer": 1},
                    {"q": "What metric best measures digital transformation success?", "options": ["Money spent", "User adoption rate and satisfaction", "Number of meetings held", "Lines of code written"], "answer": 1}
                ]
            },
            {
                "id": "dt_security", "title": "Security in Digital Transformation", "icon": "fas fa-lock", "level": "Intermediate", "duration": "15 min",
                "description": "Ensuring security-by-design in digital transformation initiatives and cloud migration.",
                "learning_points": ["Security-by-design: integrate security from the start, not as an afterthought", "Cloud security shared responsibility model", "API security and integration risks", "Digital identity management in modern architectures"],
                "scenario": {"title": "The Rush to Cloud", "text": "Management wants to migrate all systems to the cloud within 3 months to cut costs. The security team hasn't been consulted yet.", "correct_action": "Advocate for including security from the start. Conduct a cloud security assessment. Define a shared responsibility model. Prioritize migration of less sensitive systems first. Ensure proper access controls, encryption, and compliance before migrating critical data."},
                "quiz": [
                    {"q": "What is 'security by design'?", "options": ["Adding security after deployment", "Integrating security from the earliest design phase", "A security software brand", "Security only for designers"], "answer": 1},
                    {"q": "In cloud computing, who is responsible for data security?", "options": ["Only the cloud provider", "Only the customer", "Shared responsibility between both", "No one"], "answer": 2},
                    {"q": "What should happen before migrating to the cloud?", "options": ["Nothing special", "Cloud security assessment and risk analysis", "Just move everything at once", "Wait for competitors to go first"], "answer": 1},
                    {"q": "What is API security?", "options": ["A type of antivirus", "Protecting interfaces that connect different systems", "An encryption standard", "A cloud provider"], "answer": 1},
                    {"q": "Why should less sensitive systems be migrated to cloud first?", "options": ["They're smaller", "To learn and refine security controls before migrating critical data", "They're cheaper", "No specific reason"], "answer": 1}
                ]
            }
        ],
        "ar": [
            {
                "id": "dt_change", "title": "إدارة التغيير الرقمي", "icon": "fas fa-sync-alt", "level": "مبتدئ", "duration": "15 دقيقة",
                "description": "إدارة التغيير المؤسسي خلال مبادرات التحول الرقمي بفعالية.",
                "learning_points": ["أطر إدارة التغيير: ADKAR، خطوات كوتر الثمانية", "إشراك أصحاب المصلحة واستراتيجيات التواصل", "التغلب على مقاومة التغيير الرقمي", "قياس معدلات التبني ورضا المستخدمين"],
                "scenario": {"title": "النظام المرفوض", "text": "بعد 6 أشهر واستثمار كبير، نظام ERP الجديد بنسبة تبنٍّ 30% فقط. يشتكي الموظفون أنه أصعب من النظام القديم ويستمرون باستخدام جداول البيانات.", "correct_action": "أجرِ جلسات ملاحظات مع المستخدمين. حدد نقاط الألم المحددة. قدم تدريباً حسب الأدوار. حدد سفراء في كل قسم. أوقف الأنظمة القديمة تدريجياً."},
                "quiz": [
                    {"q": "ماذا تعني ADKAR؟", "options": ["إطار برمجي", "الوعي، الرغبة، المعرفة، القدرة، التعزيز", "نوع قاعدة بيانات", "مراجعة المعرفة الرقمية الآلية"], "answer": 1},
                    {"q": "ما أشهر سبب لفشل مشاريع التحول الرقمي؟", "options": ["تقنية سيئة", "ميزانية غير كافية", "ضعف إدارة التغيير والمقاومة", "عدم وجود إنترنت"], "answer": 2},
                    {"q": "ما هو 'سفير التغيير'؟", "options": ["شخص يقاوم التغيير", "موظف يدافع عن التغيير ويدعمه", "مدير مشروع", "مستشار خارجي"], "answer": 1},
                    {"q": "كيف تتعامل مع مقاومة الموظفين للأنظمة الجديدة؟", "options": ["إجبار التبني بالأوامر فقط", "الاستماع ومعالجة المخاوف وتقديم التدريب والدعم", "تجاهل الشكاوى", "التهديد بالعواقب"], "answer": 1},
                    {"q": "ما أفضل مقياس لنجاح التحول الرقمي؟", "options": ["المال المصروف", "معدل التبني ورضا المستخدمين", "عدد الاجتماعات", "سطور الكود المكتوبة"], "answer": 1}
                ]
            },
            {
                "id": "dt_security", "title": "الأمان في التحول الرقمي", "icon": "fas fa-lock", "level": "متوسط", "duration": "15 دقيقة",
                "description": "ضمان الأمان بالتصميم في مبادرات التحول الرقمي والترحيل السحابي.",
                "learning_points": ["الأمان بالتصميم: دمج الأمان من البداية وليس كإضافة لاحقة", "نموذج المسؤولية المشتركة للأمان السحابي", "أمان واجهات البرمجة (API) ومخاطر التكامل", "إدارة الهوية الرقمية في البنى الحديثة"],
                "scenario": {"title": "التسرع للسحابة", "text": "تريد الإدارة ترحيل جميع الأنظمة للسحابة خلال 3 أشهر لتقليل التكاليف. لم يُستشر فريق الأمن بعد.", "correct_action": "ادعُ لإشراك الأمن من البداية. أجرِ تقييم أمان سحابي. حدد نموذج المسؤولية المشتركة. رحّل الأنظمة الأقل حساسية أولاً. تأكد من ضوابط الوصول والتشفير والامتثال قبل ترحيل البيانات الحرجة."},
                "quiz": [
                    {"q": "ما هو 'الأمان بالتصميم'؟", "options": ["إضافة الأمان بعد النشر", "دمج الأمان من أولى مراحل التصميم", "علامة تجارية لبرنامج أمني", "أمان للمصممين فقط"], "answer": 1},
                    {"q": "في الحوسبة السحابية، من المسؤول عن أمان البيانات؟", "options": ["مزود السحابة فقط", "العميل فقط", "مسؤولية مشتركة بين الاثنين", "لا أحد"], "answer": 2},
                    {"q": "ماذا يجب أن يحدث قبل الترحيل للسحابة؟", "options": ["لا شيء خاص", "تقييم أمان سحابي وتحليل مخاطر", "نقل كل شيء دفعة واحدة", "انتظار المنافسين"], "answer": 1},
                    {"q": "ما هو أمان API؟", "options": ["نوع مضاد فيروسات", "حماية الواجهات التي تربط الأنظمة المختلفة", "معيار تشفير", "مزود سحابي"], "answer": 1},
                    {"q": "لماذا تُرحّل الأنظمة الأقل حساسية أولاً؟", "options": ["لأنها أصغر", "للتعلم وتحسين ضوابط الأمان قبل ترحيل البيانات الحرجة", "لأنها أرخص", "لا سبب محدد"], "answer": 1}
                ]
            }
        ]
    },
    "global": {
        "en": [
            {
                "id": "global_iso27001", "title": "ISO 27001 Essentials", "icon": "fas fa-certificate", "level": "Beginner", "duration": "15 min",
                "description": "Core concepts of ISO 27001 information security management system (ISMS).",
                "learning_points": ["ISO 27001 provides a systematic approach to managing sensitive information", "Key components: Context, Leadership, Planning, Support, Operation, Evaluation, Improvement", "Risk-based approach: identify, assess, and treat information security risks", "Annex A controls: 93 controls organized in 4 themes"],
                "scenario": {"title": "The Certification Audit", "text": "Your organization is preparing for ISO 27001 certification. The internal audit found that 40% of employees don't know the information security policy exists.", "correct_action": "This is a major finding. Address it by: conducting mandatory security awareness training, making the policy accessible on the intranet, requiring acknowledgment from all employees, and establishing ongoing communication about security responsibilities."},
                "quiz": [
                    {"q": "What does ISO 27001 primarily address?", "options": ["Quality management", "Information security management", "Environmental management", "Financial management"], "answer": 1},
                    {"q": "What is the PDCA cycle in ISO 27001?", "options": ["A type of encryption", "Plan-Do-Check-Act for continuous improvement", "A network protocol", "A risk assessment tool"], "answer": 1},
                    {"q": "How many controls are in ISO 27001:2022 Annex A?", "options": ["114", "93", "50", "200"], "answer": 1},
                    {"q": "What is the Statement of Applicability (SoA)?", "options": ["A legal document", "A document listing which controls apply and justification", "An audit report", "A user manual"], "answer": 1},
                    {"q": "Is ISO 27001 certification mandatory?", "options": ["Yes, for all companies", "No, it's voluntary but often required by clients", "Only for government", "Only in the EU"], "answer": 1}
                ]
            },
            {
                "id": "global_compliance", "title": "Regulatory Compliance Basics", "icon": "fas fa-gavel", "level": "Intermediate", "duration": "15 min",
                "description": "Understanding compliance requirements, frameworks, and your role in maintaining organizational compliance.",
                "learning_points": ["Compliance means adhering to laws, regulations, standards, and internal policies", "Key Saudi regulations: NCA ECC, PDPL, SAMA CSF, CITC regulations", "International frameworks: ISO 27001, NIST CSF, COBIT, PCI DSS", "Everyone's role: compliance is not just the compliance team's job"],
                "scenario": {"title": "The Compliance Gap", "text": "During a regulatory review, NCA finds that your organization hasn't implemented 30% of required ECC controls. The deadline for compliance passed 6 months ago.", "correct_action": "Immediately develop a remediation plan with clear timelines. Prioritize critical controls first. Assign owners for each gap. Report progress to management weekly. Engage with NCA proactively about your remediation timeline."},
                "quiz": [
                    {"q": "What is NCA ECC?", "options": ["A type of encryption", "Saudi National Cybersecurity Authority Essential Controls", "An email protocol", "A cloud provider"], "answer": 1},
                    {"q": "Whose responsibility is compliance?", "options": ["Only the compliance team", "Only management", "Everyone in the organization", "Only IT department"], "answer": 2},
                    {"q": "What happens if an organization fails to comply with NCA ECC?", "options": ["Nothing", "Possible fines, sanctions, and reputational damage", "Free consulting", "Automatic extension"], "answer": 1},
                    {"q": "What is a compliance audit?", "options": ["A financial review", "A systematic evaluation of adherence to regulations", "A technology upgrade", "A marketing survey"], "answer": 1},
                    {"q": "What should you do if you discover a compliance violation?", "options": ["Ignore it", "Report it through proper channels", "Fix it secretly", "Wait for an audit to find it"], "answer": 1}
                ]
            }
        ],
        "ar": [
            {
                "id": "global_iso27001", "title": "أساسيات ISO 27001", "icon": "fas fa-certificate", "level": "مبتدئ", "duration": "15 دقيقة",
                "description": "المفاهيم الأساسية لنظام إدارة أمن المعلومات ISO 27001.",
                "learning_points": ["يوفر ISO 27001 نهجاً منظماً لإدارة المعلومات الحساسة", "المكونات: السياق، القيادة، التخطيط، الدعم، التشغيل، التقييم، التحسين", "النهج القائم على المخاطر: تحديد وتقييم ومعالجة مخاطر أمن المعلومات", "ضوابط الملحق أ: 93 ضابطة منظمة في 4 محاور"],
                "scenario": {"title": "تدقيق الشهادة", "text": "تستعد منظمتك للحصول على شهادة ISO 27001. وجد التدقيق الداخلي أن 40% من الموظفين لا يعلمون بوجود سياسة أمن المعلومات.", "correct_action": "هذه ملاحظة جوهرية. عالجها بإجراء تدريب توعوي إلزامي، وإتاحة السياسة على الإنترانت، وطلب إقرار من جميع الموظفين، وإنشاء تواصل مستمر حول مسؤوليات الأمن."},
                "quiz": [
                    {"q": "ما الذي يعالجه ISO 27001 بشكل أساسي؟", "options": ["إدارة الجودة", "إدارة أمن المعلومات", "الإدارة البيئية", "الإدارة المالية"], "answer": 1},
                    {"q": "ما هي دورة PDCA في ISO 27001؟", "options": ["نوع تشفير", "خطط-نفذ-تحقق-صحح للتحسين المستمر", "بروتوكول شبكة", "أداة تقييم مخاطر"], "answer": 1},
                    {"q": "كم عدد الضوابط في ملحق أ ISO 27001:2022؟", "options": ["114", "93", "50", "200"], "answer": 1},
                    {"q": "ما هو بيان قابلية التطبيق (SoA)؟", "options": ["وثيقة قانونية", "وثيقة تسرد الضوابط المطبقة ومبرراتها", "تقرير تدقيق", "دليل مستخدم"], "answer": 1},
                    {"q": "هل شهادة ISO 27001 إلزامية؟", "options": ["نعم لجميع الشركات", "لا، طوعية لكن العملاء غالباً يطلبونها", "للحكومة فقط", "في الاتحاد الأوروبي فقط"], "answer": 1}
                ]
            },
            {
                "id": "global_compliance", "title": "أساسيات الامتثال التنظيمي", "icon": "fas fa-gavel", "level": "متوسط", "duration": "15 دقيقة",
                "description": "فهم متطلبات الامتثال والأطر التنظيمية ودورك في الحفاظ على امتثال المنظمة.",
                "learning_points": ["الامتثال يعني الالتزام بالقوانين واللوائح والمعايير والسياسات الداخلية", "اللوائح السعودية الرئيسية: NCA ECC، PDPL، SAMA CSF، لوائح CITC", "الأطر الدولية: ISO 27001، NIST CSF، COBIT، PCI DSS", "دور الجميع: الامتثال ليس مسؤولية فريق الامتثال فقط"],
                "scenario": {"title": "فجوة الامتثال", "text": "أثناء مراجعة تنظيمية، وجدت الهيئة الوطنية للأمن السيبراني أن منظمتك لم تطبق 30% من ضوابط ECC المطلوبة. انتهى الموعد النهائي قبل 6 أشهر.", "correct_action": "طوّر فوراً خطة معالجة بجداول زمنية واضحة. أعطِ الأولوية للضوابط الحرجة. عيّن مسؤولين لكل فجوة. أبلغ الإدارة أسبوعياً. تواصل مع الهيئة بشكل استباقي."},
                "quiz": [
                    {"q": "ما هو NCA ECC؟", "options": ["نوع تشفير", "الضوابط الأساسية للأمن السيبراني للهيئة الوطنية السعودية", "بروتوكول بريد إلكتروني", "مزود سحابي"], "answer": 1},
                    {"q": "مسؤولية الامتثال على عاتق من؟", "options": ["فريق الامتثال فقط", "الإدارة فقط", "الجميع في المنظمة", "قسم تقنية المعلومات فقط"], "answer": 2},
                    {"q": "ماذا يحدث إذا فشلت المنظمة في الامتثال لـ NCA ECC؟", "options": ["لا شيء", "غرامات وعقوبات محتملة وضرر سمعي", "استشارة مجانية", "تمديد تلقائي"], "answer": 1},
                    {"q": "ما هو تدقيق الامتثال؟", "options": ["مراجعة مالية", "تقييم منهجي للالتزام باللوائح", "ترقية تقنية", "استطلاع تسويقي"], "answer": 1},
                    {"q": "ماذا تفعل إذا اكتشفت مخالفة امتثال؟", "options": ["تجاهلها", "أبلغ عنها عبر القنوات المناسبة", "أصلحها سراً", "انتظر التدقيق ليجدها"], "answer": 1}
                ]
            }
        ]
    },
    "erm": {
        "en": [
            {
                "id": "erm_fundamentals", "title": "Risk Management Fundamentals", "icon": "fas fa-exclamation-circle", "level": "Beginner", "duration": "15 min",
                "description": "Core concepts of enterprise risk management, risk identification, and assessment methods.",
                "learning_points": ["Risk = Likelihood × Impact — understanding how risks are measured", "Risk categories: Strategic, Operational, Financial, Compliance, Reputational", "Risk appetite vs risk tolerance: how much risk the organization accepts", "The risk management cycle: Identify, Assess, Treat, Monitor, Report"],
                "scenario": {"title": "The Untracked Risk", "text": "A department head decides not to report a vendor's repeated service failures because they're 'managing it internally'. Three months later, the vendor goes bankrupt, causing a major service disruption.", "correct_action": "All significant risks must be reported and tracked in the risk register, regardless of who is managing them. The risk register provides visibility to leadership. Implement vendor risk monitoring with early warning indicators."},
                "quiz": [
                    {"q": "How is risk typically calculated?", "options": ["Cost × Time", "Likelihood × Impact", "Revenue × Profit", "Employees × Departments"], "answer": 1},
                    {"q": "What is 'risk appetite'?", "options": ["The amount of risk an organization is willing to accept", "A risk management software", "How hungry the risk manager is", "The maximum number of risks"], "answer": 0},
                    {"q": "Which is NOT a risk treatment option?", "options": ["Avoid", "Mitigate", "Transfer", "Ignore completely"], "answer": 3},
                    {"q": "Why must risks be reported to a central register?", "options": ["Bureaucracy", "Visibility and informed decision-making at leadership level", "To create more paperwork", "It's optional"], "answer": 1},
                    {"q": "What framework focuses on enterprise risk management?", "options": ["ISO 9001", "COSO ERM", "PCI DSS", "HTML"], "answer": 1}
                ]
            },
            {
                "id": "erm_culture", "title": "Building a Risk-Aware Culture", "icon": "fas fa-users", "level": "Intermediate", "duration": "15 min",
                "description": "Creating an organizational culture where everyone understands and actively manages risk.",
                "learning_points": ["Risk culture starts at the top — leadership must model risk-aware behavior", "Every employee is a risk manager in their area of work", "Psychological safety: employees must feel safe reporting risks without blame", "Regular risk communication through meetings, dashboards, and training"],
                "scenario": {"title": "The Silent Risk", "text": "An employee notices a critical system has been running without backups for 2 weeks due to a misconfiguration. They hesitate to report it because last time someone reported a problem, they were blamed for it.", "correct_action": "Create a blame-free reporting culture. Establish anonymous reporting channels. Reward risk identification. The employee should be thanked, not blamed. Fix the backup issue immediately and investigate how the monitoring gap occurred."},
                "quiz": [
                    {"q": "Where does risk culture start?", "options": ["IT department", "At the top — leadership", "New employees", "External consultants"], "answer": 1},
                    {"q": "Why is psychological safety important for risk management?", "options": ["It's not important", "Employees need to feel safe reporting risks without fear of blame", "Only for HR purposes", "To reduce stress levels"], "answer": 1},
                    {"q": "Who should be responsible for managing risk?", "options": ["Only the risk management team", "Only senior management", "Everyone in the organization within their role", "External auditors"], "answer": 2},
                    {"q": "How should an organization respond when someone reports a risk?", "options": ["Blame the reporter", "Thank and investigate the risk", "Ignore it", "Punish the department"], "answer": 1},
                    {"q": "What is the best way to communicate risk across an organization?", "options": ["Annual emails only", "Regular updates through multiple channels: meetings, dashboards, training", "Only during crises", "Never — it creates panic"], "answer": 1}
                ]
            }
        ],
        "ar": [
            {
                "id": "erm_fundamentals", "title": "أساسيات إدارة المخاطر", "icon": "fas fa-exclamation-circle", "level": "مبتدئ", "duration": "15 دقيقة",
                "description": "المفاهيم الأساسية لإدارة المخاطر المؤسسية وأساليب التحديد والتقييم.",
                "learning_points": ["المخاطر = الاحتمالية × الأثر — فهم كيفية قياس المخاطر", "فئات المخاطر: استراتيجية، تشغيلية، مالية، امتثال، سمعة", "شهية المخاطر مقابل تحمل المخاطر: مقدار المخاطر التي تقبلها المنظمة", "دورة إدارة المخاطر: التحديد، التقييم، المعالجة، المراقبة، الإبلاغ"],
                "scenario": {"title": "المخاطر غير المتتبعة", "text": "قرر مدير إدارة عدم الإبلاغ عن إخفاقات المورد المتكررة لأنه 'يديرها داخلياً'. بعد 3 أشهر، أفلس المورد مسبباً انقطاعاً كبيراً في الخدمة.", "correct_action": "يجب الإبلاغ عن جميع المخاطر الجوهرية وتتبعها في سجل المخاطر بغض النظر عمن يديرها. السجل يوفر رؤية للقيادة. طبّق مراقبة مخاطر الموردين مع مؤشرات إنذار مبكر."},
                "quiz": [
                    {"q": "كيف تُحسب المخاطر عادةً؟", "options": ["التكلفة × الوقت", "الاحتمالية × الأثر", "الإيرادات × الربح", "الموظفون × الأقسام"], "answer": 1},
                    {"q": "ما هي 'شهية المخاطر'؟", "options": ["مقدار المخاطر التي تقبلها المنظمة", "برنامج إدارة مخاطر", "مدى جوع مدير المخاطر", "الحد الأقصى لعدد المخاطر"], "answer": 0},
                    {"q": "أي مما يلي ليس خيار معالجة مخاطر؟", "options": ["التجنب", "التخفيف", "النقل", "التجاهل الكامل"], "answer": 3},
                    {"q": "لماذا يجب الإبلاغ عن المخاطر في سجل مركزي؟", "options": ["بيروقراطية", "الرؤية واتخاذ القرار المستنير على مستوى القيادة", "لإنشاء مزيد من الأوراق", "اختياري"], "answer": 1},
                    {"q": "أي إطار يركز على إدارة المخاطر المؤسسية؟", "options": ["ISO 9001", "COSO ERM", "PCI DSS", "HTML"], "answer": 1}
                ]
            },
            {
                "id": "erm_culture", "title": "بناء ثقافة واعية بالمخاطر", "icon": "fas fa-users", "level": "متوسط", "duration": "15 دقيقة",
                "description": "إنشاء ثقافة مؤسسية يفهم فيها الجميع المخاطر ويديرونها بنشاط.",
                "learning_points": ["ثقافة المخاطر تبدأ من القمة — القيادة يجب أن تكون قدوة", "كل موظف هو مدير مخاطر في مجال عمله", "الأمان النفسي: يجب أن يشعر الموظفون بالأمان عند الإبلاغ دون لوم", "التواصل المنتظم حول المخاطر عبر الاجتماعات ولوحات المعلومات والتدريب"],
                "scenario": {"title": "المخاطر الصامتة", "text": "لاحظ موظف أن نظاماً حرجاً يعمل بدون نسخ احتياطية منذ أسبوعين بسبب خطأ في الإعدادات. يتردد في الإبلاغ لأنه في المرة الأخيرة التي أبلغ فيها شخص عن مشكلة، تم لومه.", "correct_action": "أنشئ ثقافة إبلاغ بدون لوم. أوجد قنوات إبلاغ مجهولة. كافئ تحديد المخاطر. يجب شكر الموظف لا لومه. أصلح مشكلة النسخ الاحتياطي فوراً وحقق في كيفية حدوث فجوة المراقبة."},
                "quiz": [
                    {"q": "من أين تبدأ ثقافة المخاطر؟", "options": ["قسم تقنية المعلومات", "من القمة — القيادة", "الموظفين الجدد", "المستشارين الخارجيين"], "answer": 1},
                    {"q": "لماذا الأمان النفسي مهم لإدارة المخاطر؟", "options": ["ليس مهماً", "يحتاج الموظفون للشعور بالأمان عند الإبلاغ دون خوف من اللوم", "لأغراض الموارد البشرية فقط", "لتقليل مستويات التوتر"], "answer": 1},
                    {"q": "من يجب أن يكون مسؤولاً عن إدارة المخاطر؟", "options": ["فريق إدارة المخاطر فقط", "الإدارة العليا فقط", "الجميع في المنظمة ضمن دورهم", "المدققون الخارجيون"], "answer": 2},
                    {"q": "كيف يجب أن تستجيب المنظمة عند إبلاغ شخص عن خطر؟", "options": ["لوم المُبلِّغ", "شكره والتحقيق في الخطر", "تجاهله", "معاقبة القسم"], "answer": 1},
                    {"q": "ما أفضل طريقة للتواصل حول المخاطر عبر المنظمة؟", "options": ["رسائل بريد سنوية فقط", "تحديثات منتظمة عبر قنوات متعددة: اجتماعات، لوحات، تدريب", "أثناء الأزمات فقط", "أبداً — يسبب ذعراً"], "answer": 1}
                ]
            }
        ]
    },
    # GRC Professional Modules - Frameworks, Regulations & Standards Training
    "grc_professional": {
        "en": [
            {
                "id": "grc_nca_frameworks", "title": "NCA Cybersecurity Frameworks (KSA)", "icon": "fas fa-flag", "level": "Intermediate", "duration": "25 min",
                "description": "Comprehensive guide to Saudi National Cybersecurity Authority frameworks including ECC, CSCC, DCC, and other controls.",
                "learning_points": [
                    "NCA ECC (Essential Cybersecurity Controls): 114 controls in 5 domains - mandatory for all government entities",
                    "NCA CSCC (Critical Systems): Enhanced controls for national critical infrastructure",
                    "NCA DCC (Data Cybersecurity): Data protection controls aligned with PDPL",
                    "NCA CCC (Cloud): Security requirements for cloud service adoption",
                    "NCA TCC/OTCC: Telework and Operational Technology specific controls"
                ],
                "scenario": {"title": "ECC Compliance Gap", "text": "Your organization scored 65% on NCA ECC compliance assessment. Management wants to reach 90% within 6 months. How do you prioritize?", "correct_action": "Start with high-priority gaps in Cybersecurity Governance and Defense domains. Focus on quick wins (policies, awareness) while planning longer-term technical controls. Create a remediation roadmap with clear ownership and deadlines."},
                "quiz": [
                    {"q": "How many main domains are in NCA ECC?", "options": ["3", "5", "7", "10"], "answer": 1},
                    {"q": "Which NCA framework applies to national critical infrastructure?", "options": ["ECC", "CSCC", "TCC", "DCC"], "answer": 1},
                    {"q": "What is the relationship between NCA DCC and PDPL?", "options": ["They are unrelated", "DCC provides technical controls to implement PDPL requirements", "PDPL replaces DCC", "DCC is optional"], "answer": 1},
                    {"q": "Who must comply with NCA ECC?", "options": ["Only banks", "Only government entities", "All government entities and critical infrastructure", "Only tech companies"], "answer": 2},
                    {"q": "What is NCA CCC used for?", "options": ["Cloud adoption security requirements", "Cryptography standards", "Crisis management", "Customer communications"], "answer": 0}
                ]
            },
            {
                "id": "grc_pdpl", "title": "Saudi PDPL (Personal Data Protection Law)", "icon": "fas fa-user-shield", "level": "Intermediate", "duration": "20 min",
                "description": "Understanding Saudi Personal Data Protection Law requirements, data subject rights, and compliance obligations.",
                "learning_points": [
                    "PDPL applies to all processing of personal data in Saudi Arabia, regardless of processor location",
                    "Lawful bases: consent, contract, legal obligation, vital interests, public interest",
                    "Data subject rights: access, rectification, deletion, portability, objection",
                    "Data Protection Officer (DPO) requirement for certain organizations",
                    "Cross-border transfer restrictions and adequacy requirements",
                    "Penalties: up to 5 million SAR for violations"
                ],
                "scenario": {"title": "Data Breach Response", "text": "Your organization discovers a data breach affecting 10,000 customer records including national IDs. What are your PDPL obligations?", "correct_action": "Notify SDAIA within 72 hours if the breach poses risk to individuals. Assess the breach impact and document all actions. Notify affected individuals if there's high risk to their rights. Implement measures to prevent recurrence."},
                "quiz": [
                    {"q": "What is the maximum penalty for PDPL violations?", "options": ["1 million SAR", "3 million SAR", "5 million SAR", "10 million SAR"], "answer": 2},
                    {"q": "Within how many hours must you notify SDAIA of a serious data breach?", "options": ["24 hours", "48 hours", "72 hours", "1 week"], "answer": 2},
                    {"q": "Which organization regulates PDPL in Saudi Arabia?", "options": ["NCA", "SDAIA", "SAMA", "CITC"], "answer": 1},
                    {"q": "What is NOT a data subject right under PDPL?", "options": ["Right to access", "Right to deletion", "Right to free products", "Right to rectification"], "answer": 2},
                    {"q": "Can personal data be transferred outside Saudi Arabia?", "options": ["Never", "Always without restriction", "Only with adequate protection measures", "Only to GCC countries"], "answer": 2}
                ]
            },
            {
                "id": "grc_iso27001", "title": "ISO 27001:2022 Deep Dive", "icon": "fas fa-certificate", "level": "Advanced", "duration": "30 min",
                "description": "Detailed understanding of ISO 27001:2022 Information Security Management System requirements and implementation.",
                "learning_points": [
                    "ISO 27001:2022 structure: 10 clauses covering context, leadership, planning, support, operation, evaluation, improvement",
                    "Annex A: 93 controls organized in 4 themes (Organizational, People, Physical, Technological)",
                    "Risk-based approach: systematic identification, assessment, and treatment of information security risks",
                    "Statement of Applicability (SoA): documenting which controls apply and their implementation status",
                    "Certification process: Stage 1 (documentation review), Stage 2 (implementation audit), Surveillance audits"
                ],
                "scenario": {"title": "Certification Preparation", "text": "Your organization wants ISO 27001 certification within 12 months. You have basic security controls but no formal ISMS. Where do you start?", "correct_action": "1) Conduct gap assessment against ISO 27001 requirements. 2) Define ISMS scope and get management commitment. 3) Perform risk assessment and create risk treatment plan. 4) Develop required documentation (policies, procedures, SoA). 5) Implement controls and conduct internal audit. 6) Management review before certification audit."},
                "quiz": [
                    {"q": "How many controls are in ISO 27001:2022 Annex A?", "options": ["114", "93", "50", "27"], "answer": 1},
                    {"q": "What document lists which ISO 27001 controls apply to your organization?", "options": ["Risk Register", "Statement of Applicability", "Security Policy", "Audit Report"], "answer": 1},
                    {"q": "What is Stage 1 of ISO 27001 certification?", "options": ["Implementation audit", "Documentation review", "Surveillance audit", "Gap assessment"], "answer": 1},
                    {"q": "How often are surveillance audits conducted after certification?", "options": ["Monthly", "Quarterly", "Annually", "Every 3 years"], "answer": 2},
                    {"q": "What approach does ISO 27001 require for selecting controls?", "options": ["Implement all controls", "Risk-based selection", "Cost-based selection", "Random selection"], "answer": 1}
                ]
            },
            {
                "id": "grc_gdpr", "title": "GDPR Fundamentals", "icon": "fas fa-globe-europe", "level": "Intermediate", "duration": "20 min",
                "description": "Understanding EU General Data Protection Regulation for organizations serving European customers.",
                "learning_points": [
                    "GDPR applies to any organization processing EU residents' data, regardless of location",
                    "Key principles: lawfulness, fairness, transparency, purpose limitation, data minimization, accuracy, storage limitation, integrity, accountability",
                    "Data subject rights: access, rectification, erasure (right to be forgotten), restriction, portability, objection",
                    "Data Protection Impact Assessment (DPIA) required for high-risk processing",
                    "Penalties: up to €20 million or 4% of global annual turnover"
                ],
                "scenario": {"title": "International Data Transfer", "text": "Your Saudi company wants to process EU customer data in your Riyadh data center. What GDPR requirements apply?", "correct_action": "EU-Saudi transfers require adequate safeguards. Options: Standard Contractual Clauses (SCCs), Binding Corporate Rules, or explicit consent. Conduct a Transfer Impact Assessment. Ensure technical measures (encryption) protect data. Document the legal basis for transfer."},
                "quiz": [
                    {"q": "What is the maximum GDPR fine?", "options": ["€10 million", "€20 million or 4% of global turnover", "€50 million", "€1 million"], "answer": 1},
                    {"q": "Does GDPR apply to a Saudi company with EU customers?", "options": ["No, only EU companies", "Yes, if processing EU residents' data", "Only if they have EU offices", "Only for financial data"], "answer": 1},
                    {"q": "What is the 'right to be forgotten'?", "options": ["Right to delete your data", "Right to anonymity", "Right to encryption", "Right to data backup"], "answer": 0},
                    {"q": "When is a DPIA required?", "options": ["For all data processing", "Only for high-risk processing", "Only for customer data", "Never required"], "answer": 1},
                    {"q": "What mechanism allows EU-Saudi data transfers?", "options": ["Nothing, transfers are banned", "Standard Contractual Clauses", "Verbal agreement", "Email consent"], "answer": 1}
                ]
            },
            {
                "id": "grc_nist_csf", "title": "NIST Cybersecurity Framework 2.0", "icon": "fas fa-shield-virus", "level": "Intermediate", "duration": "20 min",
                "description": "Understanding NIST CSF 2.0 functions, categories, and implementation for cybersecurity program management.",
                "learning_points": [
                    "NIST CSF 2.0 adds GOVERN as the 6th function alongside Identify, Protect, Detect, Respond, Recover",
                    "Framework provides a common language for cybersecurity risk communication",
                    "Implementation Tiers: Partial (1), Risk Informed (2), Repeatable (3), Adaptive (4)",
                    "Framework Profiles: document current state and target state for gap analysis",
                    "Widely referenced by regulators including NCA as a supplementary framework"
                ],
                "scenario": {"title": "Framework Adoption", "text": "Your CISO asks you to map your current security program to NIST CSF 2.0. Where do you start?", "correct_action": "1) Create Current Profile by assessing existing capabilities against CSF categories. 2) Define Target Profile based on business objectives and risk appetite. 3) Identify gaps between current and target. 4) Prioritize actions based on risk and resources. 5) Track progress using CSF tiers."},
                "quiz": [
                    {"q": "What new function was added in NIST CSF 2.0?", "options": ["PROTECT", "GOVERN", "RESPOND", "ANALYZE"], "answer": 1},
                    {"q": "How many core functions does NIST CSF 2.0 have?", "options": ["4", "5", "6", "7"], "answer": 2},
                    {"q": "What does Implementation Tier 4 (Adaptive) mean?", "options": ["No security program", "Basic security", "Organization adapts to changing cyber risks in real-time", "External security only"], "answer": 2},
                    {"q": "What is a Framework Profile?", "options": ["A user account", "Documentation of current/target security state", "A social media page", "An audit report"], "answer": 1},
                    {"q": "Is NIST CSF mandatory?", "options": ["Yes, for all organizations", "Only for US government", "No, it's voluntary but widely adopted", "Only for banks"], "answer": 2}
                ]
            },
            {
                "id": "grc_coso_erm", "title": "COSO ERM Framework", "icon": "fas fa-project-diagram", "level": "Advanced", "duration": "25 min",
                "description": "Comprehensive guide to COSO Enterprise Risk Management framework for organization-wide risk management.",
                "learning_points": [
                    "COSO ERM 2017 integrates strategy and performance with risk management",
                    "5 Components: Governance & Culture, Strategy & Objective-Setting, Performance, Review & Revision, Information & Communication",
                    "20 Principles across the 5 components provide detailed implementation guidance",
                    "Risk appetite defines the amount of risk an organization is willing to accept in pursuit of value",
                    "ERM should be embedded in strategic planning and decision-making processes"
                ],
                "scenario": {"title": "ERM Implementation", "text": "The board wants to implement COSO ERM but is concerned about cost and complexity. How do you make the business case?", "correct_action": "Demonstrate value: better strategic decisions, reduced surprises, improved resource allocation, enhanced stakeholder confidence. Start with high-impact areas, not a full implementation. Show how ERM prevents costly risk events and enables informed risk-taking. Present case studies of ERM benefits."},
                "quiz": [
                    {"q": "How many components are in COSO ERM 2017?", "options": ["3", "4", "5", "8"], "answer": 2},
                    {"q": "What is risk appetite?", "options": ["How hungry the risk team is", "Amount of risk an org will accept to create value", "Maximum number of risks", "Risk register size"], "answer": 1},
                    {"q": "What does COSO ERM integrate risk management with?", "options": ["IT operations only", "Strategy and performance", "Marketing campaigns", "HR processes only"], "answer": 1},
                    {"q": "Who is responsible for ERM according to COSO?", "options": ["Only the risk department", "Only the board", "Everyone with board oversight", "External auditors"], "answer": 2},
                    {"q": "What year was the current COSO ERM framework released?", "options": ["2004", "2013", "2017", "2020"], "answer": 2}
                ]
            },
            {
                "id": "grc_sama_framework", "title": "SAMA Cyber Security Framework", "icon": "fas fa-university", "level": "Intermediate", "duration": "20 min",
                "description": "Understanding Saudi Arabian Monetary Authority cybersecurity requirements for financial institutions.",
                "learning_points": [
                    "SAMA CSF applies to all financial institutions regulated by SAMA (banks, insurance, finance companies)",
                    "Framework aligns with international standards (ISO 27001, NIST) with local requirements",
                    "4 Domains: Cyber Security Leadership & Governance, Cyber Security Risk Management, Cyber Security Operations, Third Party Security",
                    "Annual self-assessment and periodic SAMA examinations required",
                    "BCM (Business Continuity Management) requirements included"
                ],
                "scenario": {"title": "SAMA Examination", "text": "SAMA announces a cybersecurity examination in 3 months. Your bank has gaps in third-party risk management. How do you prepare?", "correct_action": "1) Conduct rapid gap assessment against SAMA CSF third-party requirements. 2) Inventory all critical third parties. 3) Obtain/review vendor security assessments. 4) Update contracts with security requirements. 5) Implement monitoring for critical vendors. 6) Document everything for examination evidence."},
                "quiz": [
                    {"q": "Which institutions must comply with SAMA CSF?", "options": ["All companies in Saudi", "Only banks", "All SAMA-regulated financial institutions", "Only international banks"], "answer": 2},
                    {"q": "How many main domains does SAMA CSF have?", "options": ["3", "4", "5", "7"], "answer": 1},
                    {"q": "What does SAMA require annually?", "options": ["Full external audit only", "Self-assessment submission", "Framework recertification", "Nothing"], "answer": 1},
                    {"q": "Which international frameworks does SAMA CSF align with?", "options": ["PCI DSS only", "ISO 27001 and NIST", "COBIT only", "None"], "answer": 1},
                    {"q": "Does SAMA CSF include business continuity requirements?", "options": ["No, that's separate", "Yes, BCM is included", "Only for large banks", "Optional"], "answer": 1}
                ]
            },
            {
                "id": "grc_eu_regulations", "title": "EU Regulatory Landscape (NIS2, DORA, AI Act)", "icon": "fas fa-landmark", "level": "Advanced", "duration": "25 min",
                "description": "Overview of key EU regulations affecting cybersecurity, digital resilience, and AI governance.",
                "learning_points": [
                    "NIS2 Directive: Enhanced cybersecurity requirements for essential and important entities, 24-hour incident reporting",
                    "DORA (Digital Operational Resilience Act): ICT risk management for financial sector, third-party risk, testing requirements",
                    "EU AI Act: Risk-based regulation of AI systems - prohibited, high-risk, limited-risk, minimal-risk categories",
                    "These regulations affect organizations outside EU if they serve EU customers or markets",
                    "Significant penalties for non-compliance across all three regulations"
                ],
                "scenario": {"title": "EU Compliance Planning", "text": "Your organization plans to offer AI-powered financial services to EU customers. Which regulations apply?", "correct_action": "Multiple regulations apply: 1) GDPR for personal data processing. 2) DORA for digital resilience (if serving financial sector). 3) EU AI Act for AI system compliance - classify the AI risk level. 4) Potentially NIS2 if considered essential/important entity. Conduct comprehensive regulatory mapping and gap assessment."},
                "quiz": [
                    {"q": "What is NIS2?", "options": ["Network Infrastructure Standard", "EU cybersecurity directive for essential entities", "US security framework", "Data privacy law"], "answer": 1},
                    {"q": "What does DORA stand for?", "options": ["Digital Operations Risk Assessment", "Digital Operational Resilience Act", "Data Organization Regulation Act", "Defense Operations Requirements Act"], "answer": 1},
                    {"q": "How does the EU AI Act classify AI systems?", "options": ["By industry", "By company size", "By risk level", "By country of origin"], "answer": 2},
                    {"q": "How quickly must NIS2 incidents be reported?", "options": ["72 hours", "24 hours initial, 72 hours full", "1 week", "Monthly"], "answer": 1},
                    {"q": "Does DORA apply to non-EU financial firms?", "options": ["No, EU only", "Yes, if they serve EU financial entities", "Only to US firms", "Never"], "answer": 1}
                ]
            }
        ],
        "ar": [
            {
                "id": "grc_nca_frameworks", "title": "أطر الأمن السيبراني للهيئة الوطنية (السعودية)", "icon": "fas fa-flag", "level": "متوسط", "duration": "25 دقيقة",
                "description": "دليل شامل لأطر الهيئة الوطنية للأمن السيبراني بما في ذلك ECC و CSCC و DCC وغيرها.",
                "learning_points": [
                    "الضوابط الأساسية للأمن السيبراني (ECC): 114 ضابط في 5 مجالات - إلزامية لجميع الجهات الحكومية",
                    "ضوابط الأنظمة الحساسة (CSCC): ضوابط معززة للبنية التحتية الحرجة الوطنية",
                    "ضوابط الأمن السيبراني للبيانات (DCC): ضوابط حماية البيانات المتوافقة مع نظام حماية البيانات الشخصية",
                    "ضوابط الأمن السيبراني للحوسبة السحابية (CCC): متطلبات أمان تبني الخدمات السحابية",
                    "ضوابط العمل عن بعد (TCC) والتقنيات التشغيلية (OTCC): ضوابط خاصة بكل منها"
                ],
                "scenario": {"title": "فجوة امتثال ECC", "text": "حصلت منظمتك على 65% في تقييم امتثال ECC. الإدارة تريد الوصول إلى 90% خلال 6 أشهر. كيف تحدد الأولويات؟", "correct_action": "ابدأ بالفجوات عالية الأولوية في مجالي حوكمة الأمن السيبراني والدفاع. ركز على المكاسب السريعة (السياسات، التوعية) مع التخطيط للضوابط التقنية طويلة المدى. أنشئ خارطة طريق للمعالجة مع ملكية ومواعيد واضحة."},
                "quiz": [
                    {"q": "كم عدد المجالات الرئيسية في ECC؟", "options": ["3", "5", "7", "10"], "answer": 1},
                    {"q": "أي إطار من الهيئة ينطبق على البنية التحتية الحرجة الوطنية؟", "options": ["ECC", "CSCC", "TCC", "DCC"], "answer": 1},
                    {"q": "ما العلاقة بين DCC و نظام حماية البيانات الشخصية؟", "options": ["غير مرتبطين", "DCC يوفر ضوابط تقنية لتنفيذ متطلبات النظام", "النظام يحل محل DCC", "DCC اختياري"], "answer": 1},
                    {"q": "من يجب أن يمتثل لـ ECC؟", "options": ["البنوك فقط", "الجهات الحكومية فقط", "جميع الجهات الحكومية والبنية التحتية الحرجة", "شركات التقنية فقط"], "answer": 2},
                    {"q": "ما استخدام CCC؟", "options": ["متطلبات أمان تبني الخدمات السحابية", "معايير التشفير", "إدارة الأزمات", "اتصالات العملاء"], "answer": 0}
                ]
            },
            {
                "id": "grc_pdpl", "title": "نظام حماية البيانات الشخصية السعودي (PDPL)", "icon": "fas fa-user-shield", "level": "متوسط", "duration": "20 دقيقة",
                "description": "فهم متطلبات نظام حماية البيانات الشخصية السعودي وحقوق أصحاب البيانات والتزامات الامتثال.",
                "learning_points": [
                    "نظام حماية البيانات الشخصية ينطبق على جميع معالجة البيانات الشخصية في السعودية بغض النظر عن موقع المعالج",
                    "الأسس القانونية: الموافقة، العقد، الالتزام القانوني، المصالح الحيوية، المصلحة العامة",
                    "حقوق صاحب البيانات: الوصول، التصحيح، الحذف، النقل، الاعتراض",
                    "متطلب تعيين مسؤول حماية البيانات (DPO) لبعض المنظمات",
                    "قيود نقل البيانات عبر الحدود ومتطلبات الملاءمة",
                    "العقوبات: حتى 5 ملايين ريال للمخالفات"
                ],
                "scenario": {"title": "الاستجابة لاختراق البيانات", "text": "اكتشفت منظمتك اختراق بيانات يؤثر على 10,000 سجل عميل تشمل أرقام الهوية الوطنية. ما التزاماتك بموجب النظام؟", "correct_action": "إخطار سدايا خلال 72 ساعة إذا كان الاختراق يشكل خطراً على الأفراد. تقييم أثر الاختراق وتوثيق جميع الإجراءات. إخطار الأفراد المتأثرين إذا كان هناك خطر عالٍ على حقوقهم. تنفيذ إجراءات لمنع التكرار."},
                "quiz": [
                    {"q": "ما أقصى عقوبة لمخالفات نظام حماية البيانات الشخصية؟", "options": ["مليون ريال", "3 ملايين ريال", "5 ملايين ريال", "10 ملايين ريال"], "answer": 2},
                    {"q": "خلال كم ساعة يجب إخطار سدايا باختراق بيانات خطير؟", "options": ["24 ساعة", "48 ساعة", "72 ساعة", "أسبوع"], "answer": 2},
                    {"q": "أي جهة تنظم نظام حماية البيانات الشخصية في السعودية؟", "options": ["الهيئة الوطنية للأمن السيبراني", "سدايا", "ساما", "هيئة الاتصالات"], "answer": 1},
                    {"q": "ما الذي ليس من حقوق صاحب البيانات بموجب النظام؟", "options": ["حق الوصول", "حق الحذف", "حق المنتجات المجانية", "حق التصحيح"], "answer": 2},
                    {"q": "هل يمكن نقل البيانات الشخصية خارج السعودية؟", "options": ["أبداً", "دائماً بدون قيود", "فقط مع تدابير حماية كافية", "فقط لدول الخليج"], "answer": 2}
                ]
            },
            {
                "id": "grc_iso27001", "title": "ISO 27001:2022 بالتفصيل", "icon": "fas fa-certificate", "level": "متقدم", "duration": "30 دقيقة",
                "description": "فهم تفصيلي لمتطلبات نظام إدارة أمن المعلومات ISO 27001:2022 وتنفيذه.",
                "learning_points": [
                    "هيكل ISO 27001:2022: 10 بنود تغطي السياق، القيادة، التخطيط، الدعم، التشغيل، التقييم، التحسين",
                    "الملحق A: 93 ضابط منظمة في 4 محاور (تنظيمية، بشرية، مادية، تقنية)",
                    "النهج القائم على المخاطر: تحديد وتقييم ومعالجة منهجية لمخاطر أمن المعلومات",
                    "بيان قابلية التطبيق (SoA): توثيق الضوابط المطبقة وحالة تنفيذها",
                    "عملية الشهادة: المرحلة 1 (مراجعة الوثائق)، المرحلة 2 (تدقيق التنفيذ)، تدقيقات المراقبة"
                ],
                "scenario": {"title": "التحضير للشهادة", "text": "منظمتك تريد شهادة ISO 27001 خلال 12 شهر. لديكم ضوابط أمنية أساسية لكن لا يوجد نظام ISMS رسمي. من أين تبدأ؟", "correct_action": "1) إجراء تقييم الفجوات مقابل متطلبات ISO 27001. 2) تحديد نطاق ISMS والحصول على التزام الإدارة. 3) إجراء تقييم المخاطر وإنشاء خطة معالجة المخاطر. 4) تطوير الوثائق المطلوبة (السياسات، الإجراءات، SoA). 5) تنفيذ الضوابط وإجراء تدقيق داخلي. 6) مراجعة الإدارة قبل تدقيق الشهادة."},
                "quiz": [
                    {"q": "كم عدد الضوابط في الملحق A لـ ISO 27001:2022؟", "options": ["114", "93", "50", "27"], "answer": 1},
                    {"q": "ما الوثيقة التي تسرد ضوابط ISO 27001 المطبقة على منظمتك؟", "options": ["سجل المخاطر", "بيان قابلية التطبيق", "السياسة الأمنية", "تقرير التدقيق"], "answer": 1},
                    {"q": "ما المرحلة 1 من شهادة ISO 27001؟", "options": ["تدقيق التنفيذ", "مراجعة الوثائق", "تدقيق المراقبة", "تقييم الفجوات"], "answer": 1},
                    {"q": "كم مرة تُجرى تدقيقات المراقبة بعد الشهادة؟", "options": ["شهرياً", "ربع سنوياً", "سنوياً", "كل 3 سنوات"], "answer": 2},
                    {"q": "ما النهج الذي يتطلبه ISO 27001 لاختيار الضوابط؟", "options": ["تنفيذ جميع الضوابط", "اختيار قائم على المخاطر", "اختيار قائم على التكلفة", "اختيار عشوائي"], "answer": 1}
                ]
            }
        ]
    }
}

# ============================================================================
# TEMPLATES LIBRARY - Domain Specific
# ============================================================================

DOMAIN_TEMPLATES = {
    "cyber": {
        "policy": {
            "en": {
                "Information Security Policy": {
                    "description": "Comprehensive information security policy covering data protection, access control, and incident management",
                    "framework": "ISO 27001"
                },
                "Access Control Policy": {
                    "description": "Policy governing user access, authentication, and authorization controls",
                    "framework": "NCA ECC"
                },
                "Incident Response Policy": {
                    "description": "Policy for detecting, responding to, and recovering from security incidents",
                    "framework": "NIST CSF"
                },
                "Network Security Policy": {
                    "description": "Policy for securing network infrastructure, firewalls, and communications",
                    "framework": "NCA ECC"
                },
                "Endpoint Security Policy": {
                    "description": "Policy for securing endpoints including workstations, laptops, and mobile devices",
                    "framework": "NCA ECC"
                },
                "Security Awareness Policy": {
                    "description": "Policy for employee security awareness training and phishing prevention",
                    "framework": "ISO 27001"
                },
                "Vulnerability Management Policy": {
                    "description": "Policy for identifying, assessing, and remediating security vulnerabilities",
                    "framework": "NIST CSF"
                },
                "Cloud Security Policy": {
                    "description": "Policy for secure use of cloud services and infrastructure",
                    "framework": "NCA CCC"
                }
            },
            "ar": {
                "سياسة أمن المعلومات": {
                    "description": "سياسة شاملة لأمن المعلومات تغطي حماية البيانات والتحكم بالوصول وإدارة الحوادث",
                    "framework": "ISO 27001"
                },
                "سياسة التحكم بالوصول": {
                    "description": "سياسة تحكم وصول المستخدمين والمصادقة والتفويض",
                    "framework": "NCA ECC"
                },
                "سياسة الاستجابة للحوادث": {
                    "description": "سياسة لاكتشاف الحوادث الأمنية والاستجابة لها والتعافي منها",
                    "framework": "NIST CSF"
                },
                "سياسة أمن الشبكات": {
                    "description": "سياسة لتأمين البنية التحتية للشبكات وجدران الحماية والاتصالات",
                    "framework": "NCA ECC"
                },
                "سياسة أمن النقاط الطرفية": {
                    "description": "سياسة لتأمين الأجهزة الطرفية بما في ذلك محطات العمل والحواسيب المحمولة",
                    "framework": "NCA ECC"
                },
                "سياسة التوعية الأمنية": {
                    "description": "سياسة للتدريب على التوعية الأمنية والوقاية من التصيد الاحتيالي",
                    "framework": "ISO 27001"
                },
                "سياسة إدارة الثغرات": {
                    "description": "سياسة لتحديد وتقييم ومعالجة الثغرات الأمنية",
                    "framework": "NIST CSF"
                },
                "سياسة أمن السحابة": {
                    "description": "سياسة للاستخدام الآمن للخدمات والبنية التحتية السحابية",
                    "framework": "NCA CCC"
                }
            }
        },
        "audit": {
            "en": {
                "ISO 27001 Compliance Audit": {
                    "description": "Audit checklist for ISO 27001 information security management",
                    "framework": "ISO 27001",
                    "controls": 114
                },
                "NCA ECC Compliance Audit": {
                    "description": "Audit against NCA Essential Cybersecurity Controls",
                    "framework": "NCA ECC",
                    "controls": 114
                },
                "NIST CSF Assessment": {
                    "description": "Assessment against NIST Cybersecurity Framework",
                    "framework": "NIST CSF",
                    "controls": 108
                },
                "Penetration Testing Audit": {
                    "description": "Audit for penetration testing and vulnerability assessment results",
                    "framework": "OWASP",
                    "controls": 50
                }
            },
            "ar": {
                "تدقيق الامتثال ISO 27001": {
                    "description": "قائمة تدقيق لنظام إدارة أمن المعلومات",
                    "framework": "ISO 27001",
                    "controls": 114
                },
                "تدقيق الامتثال NCA ECC": {
                    "description": "تدقيق وفق ضوابط الأمن السيبراني الأساسية",
                    "framework": "NCA ECC",
                    "controls": 114
                },
                "تقييم NIST CSF": {
                    "description": "تقييم وفق إطار الأمن السيبراني NIST",
                    "framework": "NIST CSF",
                    "controls": 108
                },
                "تدقيق اختبار الاختراق": {
                    "description": "تدقيق لنتائج اختبار الاختراق وتقييم الثغرات",
                    "framework": "OWASP",
                    "controls": 50
                }
            }
        },
        "risk": {
            "en": {
                "Cybersecurity Risk Assessment": {
                    "description": "Comprehensive cyber risk assessment covering threats, vulnerabilities, and impacts",
                    "categories": ["External Threats", "Internal Threats", "Technical Vulnerabilities"]
                },
                "Ransomware Risk Assessment": {
                    "description": "Focused assessment on ransomware threats and prevention",
                    "categories": ["Prevention", "Detection", "Response", "Recovery"]
                },
                "Third Party Cyber Risk": {
                    "description": "Assessment of cybersecurity risks from vendors and partners",
                    "categories": ["Vendor Security", "Data Sharing", "Access Control"]
                }
            },
            "ar": {
                "تقييم مخاطر الأمن السيبراني": {
                    "description": "تقييم شامل للمخاطر السيبرانية يغطي التهديدات والثغرات والآثار",
                    "categories": ["تهديدات خارجية", "تهديدات داخلية", "ثغرات تقنية"]
                },
                "تقييم مخاطر برامج الفدية": {
                    "description": "تقييم مركز على تهديدات برامج الفدية والوقاية منها",
                    "categories": ["الوقاية", "الكشف", "الاستجابة", "التعافي"]
                },
                "مخاطر الأطراف الثالثة السيبرانية": {
                    "description": "تقييم مخاطر الأمن السيبراني من الموردين والشركاء",
                    "categories": ["أمن المورد", "مشاركة البيانات", "التحكم بالوصول"]
                }
            }
        },
        "procedure": {
            "en": {
                "Password Management Procedure": {
                    "description": "Step-by-step procedure for password creation, rotation, storage, and recovery across all systems",
                    "framework": "ISO 27001"
                },
                "Access Control Procedure": {
                    "description": "Procedure for user provisioning, access reviews, privilege management, and deprovisioning",
                    "framework": "NCA ECC"
                },
                "Incident Response Procedure": {
                    "description": "Detailed procedure for incident detection, triage, containment, eradication, and recovery steps",
                    "framework": "NIST CSF"
                },
                "Vulnerability Management Procedure": {
                    "description": "Procedure for vulnerability scanning, assessment, prioritization, patching, and verification",
                    "framework": "NCA ECC"
                },
                "Change Management Procedure": {
                    "description": "Procedure for requesting, approving, testing, and implementing changes to IT systems",
                    "framework": "ISO 27001"
                },
                "Backup and Recovery Procedure": {
                    "description": "Procedure for data backup scheduling, verification, retention, and disaster recovery testing",
                    "framework": "NCA ECC"
                },
                "Network Security Procedure": {
                    "description": "Procedure for firewall configuration, network segmentation, monitoring, and intrusion detection",
                    "framework": "NCA ECC"
                },
                "Security Awareness Training Procedure": {
                    "description": "Procedure for conducting security awareness programs, phishing simulations, and measuring effectiveness",
                    "framework": "ISO 27001"
                },
                "Third-Party Risk Assessment Procedure": {
                    "description": "Procedure for vendor security assessment, due diligence, contract requirements, and ongoing monitoring",
                    "framework": "NCA ECC"
                },
                "Encryption Key Management Procedure": {
                    "description": "Procedure for key generation, distribution, storage, rotation, and destruction",
                    "framework": "NCA NCS"
                },
                "Physical Security Procedure": {
                    "description": "Procedure for facility access control, visitor management, equipment protection, and environmental controls",
                    "framework": "ISO 27001"
                },
                "Log Management and Monitoring Procedure": {
                    "description": "Procedure for log collection, retention, analysis, alerting, and periodic review",
                    "framework": "NCA ECC"
                }
            },
            "ar": {
                "إجراء إدارة كلمات المرور": {
                    "description": "إجراء تفصيلي لإنشاء وتدوير وتخزين واسترجاع كلمات المرور عبر جميع الأنظمة",
                    "framework": "ISO 27001"
                },
                "إجراء التحكم بالوصول": {
                    "description": "إجراء لتزويد المستخدمين بالصلاحيات ومراجعتها وإدارة الامتيازات وإلغاء الوصول",
                    "framework": "NCA ECC"
                },
                "إجراء الاستجابة للحوادث": {
                    "description": "إجراء تفصيلي لاكتشاف الحوادث وفرزها واحتوائها واستئصالها والتعافي منها",
                    "framework": "NIST CSF"
                },
                "إجراء إدارة الثغرات الأمنية": {
                    "description": "إجراء لفحص الثغرات وتقييمها وترتيب أولوياتها ومعالجتها والتحقق منها",
                    "framework": "NCA ECC"
                },
                "إجراء إدارة التغيير": {
                    "description": "إجراء لطلب واعتماد واختبار وتنفيذ التغييرات على أنظمة تقنية المعلومات",
                    "framework": "ISO 27001"
                },
                "إجراء النسخ الاحتياطي والاسترجاع": {
                    "description": "إجراء لجدولة النسخ الاحتياطي والتحقق منه والاحتفاظ به واختبار التعافي من الكوارث",
                    "framework": "NCA ECC"
                },
                "إجراء أمن الشبكات": {
                    "description": "إجراء لتكوين جدران الحماية وتقسيم الشبكة والمراقبة وكشف التسلل",
                    "framework": "NCA ECC"
                },
                "إجراء التدريب على التوعية الأمنية": {
                    "description": "إجراء لتنفيذ برامج التوعية الأمنية ومحاكاة التصيد وقياس الفعالية",
                    "framework": "ISO 27001"
                },
                "إجراء تقييم مخاطر الأطراف الثالثة": {
                    "description": "إجراء لتقييم أمن الموردين والعناية الواجبة ومتطلبات العقود والمراقبة المستمرة",
                    "framework": "NCA ECC"
                },
                "إجراء إدارة مفاتيح التشفير": {
                    "description": "إجراء لإنشاء وتوزيع وتخزين وتدوير وإتلاف مفاتيح التشفير",
                    "framework": "NCA NCS"
                },
                "إجراء الأمن المادي": {
                    "description": "إجراء للتحكم بالوصول للمرافق وإدارة الزوار وحماية المعدات والضوابط البيئية",
                    "framework": "ISO 27001"
                },
                "إجراء إدارة السجلات والمراقبة": {
                    "description": "إجراء لجمع السجلات والاحتفاظ بها وتحليلها والتنبيه والمراجعة الدورية",
                    "framework": "NCA ECC"
                }
            }
        },
    },
    "data": {
        "policy": {
            "en": {
                "Data Governance Policy": {
                    "description": "Policy establishing data governance framework, roles, and responsibilities",
                    "framework": "NDMO"
                },
                "Data Classification Policy": {
                    "description": "Policy for classifying data based on sensitivity and criticality",
                    "framework": "NDMO"
                },
                "Data Quality Policy": {
                    "description": "Policy ensuring data accuracy, completeness, and consistency",
                    "framework": "NDMO"
                },
                "Data Retention Policy": {
                    "description": "Policy defining data retention periods and disposal procedures",
                    "framework": "PDPL"
                },
                "Personal Data Protection Policy": {
                    "description": "Policy for handling personal data in compliance with PDPL",
                    "framework": "PDPL"
                },
                "Data Sharing Policy": {
                    "description": "Policy governing internal and external data sharing",
                    "framework": "NDMO"
                },
                "Open Data Policy": {
                    "description": "Policy for publishing and managing open government data",
                    "framework": "NDMO"
                },
                "Data Catalog Policy": {
                    "description": "Policy for maintaining enterprise data catalog and metadata",
                    "framework": "NDMO"
                }
            },
            "ar": {
                "سياسة حوكمة البيانات": {
                    "description": "سياسة تأسيس إطار حوكمة البيانات والأدوار والمسؤوليات",
                    "framework": "NDMO"
                },
                "سياسة تصنيف البيانات": {
                    "description": "سياسة لتصنيف البيانات بناءً على الحساسية والأهمية",
                    "framework": "NDMO"
                },
                "سياسة جودة البيانات": {
                    "description": "سياسة لضمان دقة البيانات واكتمالها واتساقها",
                    "framework": "NDMO"
                },
                "سياسة الاحتفاظ بالبيانات": {
                    "description": "سياسة تحدد فترات الاحتفاظ بالبيانات وإجراءات التخلص",
                    "framework": "PDPL"
                },
                "سياسة حماية البيانات الشخصية": {
                    "description": "سياسة للتعامل مع البيانات الشخصية وفقاً لنظام حماية البيانات",
                    "framework": "PDPL"
                },
                "سياسة مشاركة البيانات": {
                    "description": "سياسة تحكم مشاركة البيانات الداخلية والخارجية",
                    "framework": "NDMO"
                },
                "سياسة البيانات المفتوحة": {
                    "description": "سياسة لنشر وإدارة البيانات الحكومية المفتوحة",
                    "framework": "NDMO"
                },
                "سياسة فهرس البيانات": {
                    "description": "سياسة للحفاظ على فهرس بيانات المؤسسة والبيانات الوصفية",
                    "framework": "NDMO"
                }
            }
        },
        "audit": {
            "en": {
                "NDMO Compliance Audit": {
                    "description": "Audit against National Data Management Office standards",
                    "framework": "NDMO",
                    "controls": 85
                },
                "PDPL Compliance Audit": {
                    "description": "Audit for Personal Data Protection Law compliance",
                    "framework": "PDPL",
                    "controls": 45
                },
                "Data Quality Audit": {
                    "description": "Audit assessing data quality dimensions and metrics",
                    "framework": "NDMO",
                    "controls": 30
                },
                "Data Governance Maturity Audit": {
                    "description": "Audit evaluating data governance maturity level",
                    "framework": "NDMO",
                    "controls": 50
                }
            },
            "ar": {
                "تدقيق الامتثال NDMO": {
                    "description": "تدقيق وفق معايير مكتب إدارة البيانات الوطني",
                    "framework": "NDMO",
                    "controls": 85
                },
                "تدقيق الامتثال لنظام حماية البيانات": {
                    "description": "تدقيق للامتثال لنظام حماية البيانات الشخصية",
                    "framework": "PDPL",
                    "controls": 45
                },
                "تدقيق جودة البيانات": {
                    "description": "تدقيق يقيم أبعاد ومقاييس جودة البيانات",
                    "framework": "NDMO",
                    "controls": 30
                },
                "تدقيق نضج حوكمة البيانات": {
                    "description": "تدقيق يقيم مستوى نضج حوكمة البيانات",
                    "framework": "NDMO",
                    "controls": 50
                }
            }
        },
        "risk": {
            "en": {
                "Data Privacy Risk Assessment": {
                    "description": "Assessment of risks related to personal data processing",
                    "categories": ["Collection", "Processing", "Storage", "Sharing", "Disposal"]
                },
                "Data Quality Risk Assessment": {
                    "description": "Assessment of risks affecting data quality",
                    "categories": ["Accuracy", "Completeness", "Timeliness", "Consistency"]
                },
                "Data Breach Risk Assessment": {
                    "description": "Assessment of data breach risks and impacts",
                    "categories": ["Internal Threats", "External Threats", "Technical Failures"]
                }
            },
            "ar": {
                "تقييم مخاطر خصوصية البيانات": {
                    "description": "تقييم المخاطر المتعلقة بمعالجة البيانات الشخصية",
                    "categories": ["الجمع", "المعالجة", "التخزين", "المشاركة", "التخلص"]
                },
                "تقييم مخاطر جودة البيانات": {
                    "description": "تقييم المخاطر المؤثرة على جودة البيانات",
                    "categories": ["الدقة", "الاكتمال", "التوقيت", "الاتساق"]
                },
                "تقييم مخاطر خرق البيانات": {
                    "description": "تقييم مخاطر خرق البيانات وآثارها",
                    "categories": ["تهديدات داخلية", "تهديدات خارجية", "أعطال تقنية"]
                }
            }
        },
        "procedure": {
            "en": {
                "Data Classification Procedure": {
                    "description": "Step-by-step procedure for classifying data assets by sensitivity level, labeling, and handling requirements",
                    "framework": "NDMO"
                },
                "Data Access Request Procedure": {
                    "description": "Procedure for requesting, approving, and revoking access to data repositories and datasets",
                    "framework": "NDMO"
                },
                "Data Quality Assessment Procedure": {
                    "description": "Procedure for measuring data accuracy, completeness, consistency, timeliness, and validity",
                    "framework": "NDMO"
                },
                "Data Retention and Disposal Procedure": {
                    "description": "Procedure for defining retention periods, archival, secure deletion, and disposal verification",
                    "framework": "PDPL"
                },
                "Personal Data Processing Procedure": {
                    "description": "Procedure for lawful processing, consent management, data subject requests, and breach notification",
                    "framework": "PDPL"
                },
                "Data Sharing Agreement Procedure": {
                    "description": "Procedure for establishing data sharing agreements, transfer assessments, and third-party obligations",
                    "framework": "NDMO"
                },
                "Data Breach Response Procedure": {
                    "description": "Procedure for breach detection, assessment, notification to authorities and affected individuals, and remediation",
                    "framework": "PDPL"
                },
                "Master Data Management Procedure": {
                    "description": "Procedure for maintaining golden records, deduplication, data stewardship, and cross-system synchronization",
                    "framework": "NDMO"
                }
            },
            "ar": {
                "إجراء تصنيف البيانات": {
                    "description": "إجراء تفصيلي لتصنيف أصول البيانات حسب مستوى الحساسية ووسمها ومتطلبات التعامل معها",
                    "framework": "NDMO"
                },
                "إجراء طلب الوصول للبيانات": {
                    "description": "إجراء لطلب واعتماد وإلغاء الوصول إلى مستودعات ومجموعات البيانات",
                    "framework": "NDMO"
                },
                "إجراء تقييم جودة البيانات": {
                    "description": "إجراء لقياس دقة البيانات واكتمالها واتساقها وتوقيتها وصلاحيتها",
                    "framework": "NDMO"
                },
                "إجراء الاحتفاظ بالبيانات والتخلص منها": {
                    "description": "إجراء لتحديد فترات الاحتفاظ والأرشفة والحذف الآمن والتحقق من التخلص",
                    "framework": "PDPL"
                },
                "إجراء معالجة البيانات الشخصية": {
                    "description": "إجراء للمعالجة المشروعة وإدارة الموافقات وطلبات أصحاب البيانات وإخطار الاختراقات",
                    "framework": "PDPL"
                },
                "إجراء اتفاقيات مشاركة البيانات": {
                    "description": "إجراء لإبرام اتفاقيات مشاركة البيانات وتقييم النقل والتزامات الأطراف الثالثة",
                    "framework": "NDMO"
                },
                "إجراء الاستجابة لاختراق البيانات": {
                    "description": "إجراء لاكتشاف الاختراق وتقييمه وإخطار الجهات والأفراد المتأثرين والمعالجة",
                    "framework": "PDPL"
                },
                "إجراء إدارة البيانات الرئيسية": {
                    "description": "إجراء للحفاظ على السجلات المرجعية وإزالة التكرار وإشراف البيانات والمزامنة بين الأنظمة",
                    "framework": "NDMO"
                }
            }
        }
    },
    "ai": {
        "policy": {
            "en": {
                "AI Governance Policy": {
                    "description": "Policy establishing AI governance framework and oversight",
                    "framework": "SDAIA"
                },
                "AI Ethics Policy": {
                    "description": "Policy defining ethical principles for AI development and use",
                    "framework": "SDAIA"
                },
                "AI Risk Management Policy": {
                    "description": "Policy for identifying and managing AI-related risks",
                    "framework": "SDAIA"
                },
                "AI Model Development Policy": {
                    "description": "Policy governing AI model development lifecycle",
                    "framework": "SDAIA"
                },
                "AI Data Usage Policy": {
                    "description": "Policy for data usage in AI training and inference",
                    "framework": "SDAIA"
                },
                "AI Transparency Policy": {
                    "description": "Policy ensuring explainability and transparency in AI decisions",
                    "framework": "SDAIA"
                },
                "Generative AI Policy": {
                    "description": "Policy for responsible use of generative AI tools",
                    "framework": "SDAIA"
                },
                "AI Vendor Management Policy": {
                    "description": "Policy for managing AI vendors and third-party AI services",
                    "framework": "SDAIA"
                }
            },
            "ar": {
                "سياسة حوكمة الذكاء الاصطناعي": {
                    "description": "سياسة تأسيس إطار حوكمة الذكاء الاصطناعي والإشراف",
                    "framework": "SDAIA"
                },
                "سياسة أخلاقيات الذكاء الاصطناعي": {
                    "description": "سياسة تحدد المبادئ الأخلاقية لتطوير واستخدام الذكاء الاصطناعي",
                    "framework": "SDAIA"
                },
                "سياسة إدارة مخاطر الذكاء الاصطناعي": {
                    "description": "سياسة لتحديد وإدارة المخاطر المتعلقة بالذكاء الاصطناعي",
                    "framework": "SDAIA"
                },
                "سياسة تطوير نماذج الذكاء الاصطناعي": {
                    "description": "سياسة تحكم دورة حياة تطوير نماذج الذكاء الاصطناعي",
                    "framework": "SDAIA"
                },
                "سياسة استخدام البيانات للذكاء الاصطناعي": {
                    "description": "سياسة لاستخدام البيانات في تدريب واستدلال الذكاء الاصطناعي",
                    "framework": "SDAIA"
                },
                "سياسة شفافية الذكاء الاصطناعي": {
                    "description": "سياسة لضمان قابلية التفسير والشفافية في قرارات الذكاء الاصطناعي",
                    "framework": "SDAIA"
                },
                "سياسة الذكاء الاصطناعي التوليدي": {
                    "description": "سياسة للاستخدام المسؤول لأدوات الذكاء الاصطناعي التوليدي",
                    "framework": "SDAIA"
                },
                "سياسة إدارة موردي الذكاء الاصطناعي": {
                    "description": "سياسة لإدارة موردي الذكاء الاصطناعي وخدمات الطرف الثالث",
                    "framework": "SDAIA"
                }
            }
        },
        "audit": {
            "en": {
                "SDAIA AI Governance Audit": {
                    "description": "Audit against SDAIA AI governance principles",
                    "framework": "SDAIA",
                    "controls": 60
                },
                "AI Ethics Audit": {
                    "description": "Audit assessing AI ethics compliance",
                    "framework": "SDAIA",
                    "controls": 40
                },
                "AI Model Audit": {
                    "description": "Audit of AI model performance, bias, and fairness",
                    "framework": "SDAIA",
                    "controls": 35
                },
                "Generative AI Audit": {
                    "description": "Audit for generative AI usage and compliance",
                    "framework": "SDAIA",
                    "controls": 25
                }
            },
            "ar": {
                "تدقيق حوكمة الذكاء الاصطناعي SDAIA": {
                    "description": "تدقيق وفق مبادئ حوكمة الذكاء الاصطناعي لسدايا",
                    "framework": "SDAIA",
                    "controls": 60
                },
                "تدقيق أخلاقيات الذكاء الاصطناعي": {
                    "description": "تدقيق يقيم الامتثال لأخلاقيات الذكاء الاصطناعي",
                    "framework": "SDAIA",
                    "controls": 40
                },
                "تدقيق نماذج الذكاء الاصطناعي": {
                    "description": "تدقيق أداء نماذج الذكاء الاصطناعي والتحيز والعدالة",
                    "framework": "SDAIA",
                    "controls": 35
                },
                "تدقيق الذكاء الاصطناعي التوليدي": {
                    "description": "تدقيق لاستخدام الذكاء الاصطناعي التوليدي والامتثال",
                    "framework": "SDAIA",
                    "controls": 25
                }
            }
        },
        "risk": {
            "en": {
                "AI Bias Risk Assessment": {
                    "description": "Assessment of bias risks in AI models and decisions",
                    "categories": ["Data Bias", "Algorithm Bias", "Output Bias"]
                },
                "AI Security Risk Assessment": {
                    "description": "Assessment of security risks specific to AI systems",
                    "categories": ["Model Attacks", "Data Poisoning", "Prompt Injection"]
                },
                "AI Operational Risk Assessment": {
                    "description": "Assessment of operational risks in AI deployment",
                    "categories": ["Model Drift", "Performance Degradation", "Integration Failures"]
                }
            },
            "ar": {
                "تقييم مخاطر تحيز الذكاء الاصطناعي": {
                    "description": "تقييم مخاطر التحيز في نماذج وقرارات الذكاء الاصطناعي",
                    "categories": ["تحيز البيانات", "تحيز الخوارزمية", "تحيز المخرجات"]
                },
                "تقييم مخاطر أمن الذكاء الاصطناعي": {
                    "description": "تقييم المخاطر الأمنية الخاصة بأنظمة الذكاء الاصطناعي",
                    "categories": ["هجمات النموذج", "تسميم البيانات", "حقن الموجه"]
                },
                "تقييم المخاطر التشغيلية للذكاء الاصطناعي": {
                    "description": "تقييم المخاطر التشغيلية في نشر الذكاء الاصطناعي",
                    "categories": ["انحراف النموذج", "تدهور الأداء", "فشل التكامل"]
                }
            }
        },
        "procedure": {
            "en": {
                "AI Model Risk Assessment Procedure": {
                    "description": "Procedure for assessing risks of AI/ML models including bias testing, fairness audits, and impact analysis",
                    "framework": "SDAIA"
                },
                "AI Model Development Lifecycle Procedure": {
                    "description": "Procedure for data preparation, model training, validation, deployment, and retirement",
                    "framework": "SDAIA"
                },
                "AI Ethics Review Procedure": {
                    "description": "Procedure for conducting ethics reviews before deploying AI systems including transparency and accountability checks",
                    "framework": "SDAIA"
                },
                "AI Data Handling and Privacy Procedure": {
                    "description": "Procedure for handling training data, ensuring consent, anonymization, and data minimization",
                    "framework": "PDPL"
                },
                "AI Incident Response Procedure": {
                    "description": "Procedure for responding to AI system failures, adversarial attacks, and unintended outputs",
                    "framework": "SDAIA"
                },
                "AI Model Monitoring Procedure": {
                    "description": "Procedure for continuous monitoring of model performance, drift detection, and retraining triggers",
                    "framework": "SDAIA"
                }
            },
            "ar": {
                "إجراء تقييم مخاطر نماذج الذكاء الاصطناعي": {
                    "description": "إجراء لتقييم مخاطر نماذج AI/ML بما في ذلك اختبار التحيز ومراجعات العدالة وتحليل الأثر",
                    "framework": "SDAIA"
                },
                "إجراء دورة حياة تطوير نماذج AI": {
                    "description": "إجراء لإعداد البيانات وتدريب النموذج والتحقق والنشر والإيقاف",
                    "framework": "SDAIA"
                },
                "إجراء مراجعة أخلاقيات الذكاء الاصطناعي": {
                    "description": "إجراء لإجراء مراجعات أخلاقية قبل نشر أنظمة AI بما في ذلك فحوصات الشفافية والمساءلة",
                    "framework": "SDAIA"
                },
                "إجراء التعامل مع بيانات AI والخصوصية": {
                    "description": "إجراء للتعامل مع بيانات التدريب وضمان الموافقة وإخفاء الهوية وتقليل البيانات",
                    "framework": "PDPL"
                },
                "إجراء الاستجابة لحوادث الذكاء الاصطناعي": {
                    "description": "إجراء للاستجابة لأعطال أنظمة AI والهجمات المعادية والمخرجات غير المقصودة",
                    "framework": "SDAIA"
                },
                "إجراء مراقبة نماذج AI": {
                    "description": "إجراء للمراقبة المستمرة لأداء النموذج واكتشاف الانحراف ومحفزات إعادة التدريب",
                    "framework": "SDAIA"
                }
            }
        }
    },
    "dt": {
        "policy": {
            "en": {
                "Digital Transformation Policy": {
                    "description": "Policy guiding organizational digital transformation initiatives",
                    "framework": "DGA"
                },
                "Digital Services Policy": {
                    "description": "Policy for digital service design, delivery, and management",
                    "framework": "DGA"
                },
                "Technology Adoption Policy": {
                    "description": "Policy for evaluating and adopting new technologies",
                    "framework": "DGA"
                },
                "Process Automation Policy": {
                    "description": "Policy governing business process automation initiatives",
                    "framework": "DGA"
                },
                "Digital Identity Policy": {
                    "description": "Policy for digital identity management and authentication",
                    "framework": "DGA"
                },
                "E-Services Policy": {
                    "description": "Policy for electronic government services",
                    "framework": "DGA"
                },
                "Mobile Services Policy": {
                    "description": "Policy for mobile application development and deployment",
                    "framework": "DGA"
                },
                "IT Infrastructure Modernization Policy": {
                    "description": "Policy for modernizing IT infrastructure and systems",
                    "framework": "DGA"
                }
            },
            "ar": {
                "سياسة التحول الرقمي": {
                    "description": "سياسة توجيه مبادرات التحول الرقمي المؤسسي",
                    "framework": "DGA"
                },
                "سياسة الخدمات الرقمية": {
                    "description": "سياسة لتصميم وتقديم وإدارة الخدمات الرقمية",
                    "framework": "DGA"
                },
                "سياسة تبني التقنية": {
                    "description": "سياسة لتقييم وتبني التقنيات الجديدة",
                    "framework": "DGA"
                },
                "سياسة أتمتة العمليات": {
                    "description": "سياسة تحكم مبادرات أتمتة العمليات التجارية",
                    "framework": "DGA"
                },
                "سياسة الهوية الرقمية": {
                    "description": "سياسة لإدارة الهوية الرقمية والمصادقة",
                    "framework": "DGA"
                },
                "سياسة الخدمات الإلكترونية": {
                    "description": "سياسة للخدمات الحكومية الإلكترونية",
                    "framework": "DGA"
                },
                "سياسة خدمات الجوال": {
                    "description": "سياسة لتطوير ونشر تطبيقات الجوال",
                    "framework": "DGA"
                },
                "سياسة تحديث البنية التحتية": {
                    "description": "سياسة لتحديث البنية التحتية والأنظمة التقنية",
                    "framework": "DGA"
                }
            }
        },
        "audit": {
            "en": {
                "Digital Maturity Audit": {
                    "description": "Audit assessing digital transformation maturity level",
                    "framework": "DGA",
                    "controls": 50
                },
                "E-Services Audit": {
                    "description": "Audit of electronic services quality and compliance",
                    "framework": "DGA",
                    "controls": 40
                },
                "Digital Experience Audit": {
                    "description": "Audit of digital user experience and accessibility",
                    "framework": "DGA",
                    "controls": 35
                },
                "IT Modernization Audit": {
                    "description": "Audit of IT infrastructure modernization progress",
                    "framework": "DGA",
                    "controls": 45
                }
            },
            "ar": {
                "تدقيق النضج الرقمي": {
                    "description": "تدقيق يقيم مستوى نضج التحول الرقمي",
                    "framework": "DGA",
                    "controls": 50
                },
                "تدقيق الخدمات الإلكترونية": {
                    "description": "تدقيق جودة الخدمات الإلكترونية والامتثال",
                    "framework": "DGA",
                    "controls": 40
                },
                "تدقيق التجربة الرقمية": {
                    "description": "تدقيق تجربة المستخدم الرقمية وسهولة الوصول",
                    "framework": "DGA",
                    "controls": 35
                },
                "تدقيق تحديث تقنية المعلومات": {
                    "description": "تدقيق تقدم تحديث البنية التحتية التقنية",
                    "framework": "DGA",
                    "controls": 45
                }
            }
        },
        "risk": {
            "en": {
                "Digital Transformation Risk Assessment": {
                    "description": "Assessment of risks in digital transformation programs",
                    "categories": ["Technology", "Change Management", "Budget", "Timeline"]
                },
                "Technology Adoption Risk Assessment": {
                    "description": "Assessment of risks in adopting new technologies",
                    "categories": ["Integration", "Vendor Lock-in", "Skills Gap", "Security"]
                },
                "Legacy System Risk Assessment": {
                    "description": "Assessment of risks from legacy systems",
                    "categories": ["Technical Debt", "Security", "Integration", "Support"]
                }
            },
            "ar": {
                "تقييم مخاطر التحول الرقمي": {
                    "description": "تقييم المخاطر في برامج التحول الرقمي",
                    "categories": ["التقنية", "إدارة التغيير", "الميزانية", "الجدول الزمني"]
                },
                "تقييم مخاطر تبني التقنية": {
                    "description": "تقييم المخاطر في تبني التقنيات الجديدة",
                    "categories": ["التكامل", "الارتباط بالمورد", "فجوة المهارات", "الأمان"]
                },
                "تقييم مخاطر الأنظمة القديمة": {
                    "description": "تقييم المخاطر من الأنظمة القديمة",
                    "categories": ["الديون التقنية", "الأمان", "التكامل", "الدعم"]
                }
            }
        },
        "procedure": {
            "en": {
                "Digital Service Onboarding Procedure": {
                    "description": "Procedure for evaluating, approving, and deploying new digital services and platforms",
                    "framework": "ISO 27001"
                },
                "Cloud Migration Procedure": {
                    "description": "Procedure for assessing, planning, executing, and validating cloud migration projects",
                    "framework": "NCA CCC"
                },
                "API Management Procedure": {
                    "description": "Procedure for API lifecycle management including design, security, versioning, and deprecation",
                    "framework": "ISO 27001"
                },
                "Digital Identity Management Procedure": {
                    "description": "Procedure for managing digital identities, SSO configuration, and federated authentication",
                    "framework": "NIST CSF"
                },
                "Technology Change Request Procedure": {
                    "description": "Procedure for submitting, reviewing, and implementing technology change requests across digital platforms",
                    "framework": "COBIT"
                },
                "Legacy System Decommissioning Procedure": {
                    "description": "Procedure for safe decommissioning of legacy systems including data migration and validation",
                    "framework": "ISO 27001"
                }
            },
            "ar": {
                "إجراء إدراج الخدمات الرقمية": {
                    "description": "إجراء لتقييم واعتماد ونشر الخدمات والمنصات الرقمية الجديدة",
                    "framework": "ISO 27001"
                },
                "إجراء الترحيل السحابي": {
                    "description": "إجراء لتقييم وتخطيط وتنفيذ والتحقق من مشاريع الترحيل السحابي",
                    "framework": "NCA CCC"
                },
                "إجراء إدارة واجهات API": {
                    "description": "إجراء لإدارة دورة حياة API بما في ذلك التصميم والأمان والإصدارات والإيقاف",
                    "framework": "ISO 27001"
                },
                "إجراء إدارة الهوية الرقمية": {
                    "description": "إجراء لإدارة الهويات الرقمية وتكوين SSO والمصادقة الموحدة",
                    "framework": "NIST CSF"
                },
                "إجراء طلب التغيير التقني": {
                    "description": "إجراء لتقديم ومراجعة وتنفيذ طلبات التغيير التقني عبر المنصات الرقمية",
                    "framework": "COBIT"
                },
                "إجراء إيقاف الأنظمة القديمة": {
                    "description": "إجراء للإيقاف الآمن للأنظمة القديمة بما في ذلك ترحيل البيانات والتحقق",
                    "framework": "ISO 27001"
                }
            }
        }
    },
    "global": {
        "policy": {
            "en": {
                "ISO 27001 Policy": {
                    "description": "Policy aligned with ISO 27001 information security standard",
                    "framework": "ISO 27001"
                },
                "ISO 22301 Business Continuity Policy": {
                    "description": "Policy for business continuity management per ISO 22301",
                    "framework": "ISO 22301"
                },
                "ISO 31000 Risk Management Policy": {
                    "description": "Policy for enterprise risk management per ISO 31000",
                    "framework": "ISO 31000"
                },
                "COBIT IT Governance Policy": {
                    "description": "Policy for IT governance based on COBIT framework",
                    "framework": "COBIT"
                },
                "ITIL Service Management Policy": {
                    "description": "Policy for IT service management based on ITIL",
                    "framework": "ITIL"
                },
                "PCI-DSS Compliance Policy": {
                    "description": "Policy for payment card data security compliance",
                    "framework": "PCI-DSS"
                },
                "SOC 2 Compliance Policy": {
                    "description": "Policy for SOC 2 trust service criteria compliance",
                    "framework": "SOC 2"
                },
                "GDPR Compliance Policy": {
                    "description": "Policy for General Data Protection Regulation compliance",
                    "framework": "GDPR"
                }
            },
            "ar": {
                "سياسة ISO 27001": {
                    "description": "سياسة متوافقة مع معيار أمن المعلومات ISO 27001",
                    "framework": "ISO 27001"
                },
                "سياسة استمرارية الأعمال ISO 22301": {
                    "description": "سياسة لإدارة استمرارية الأعمال وفق ISO 22301",
                    "framework": "ISO 22301"
                },
                "سياسة إدارة المخاطر ISO 31000": {
                    "description": "سياسة لإدارة المخاطر المؤسسية وفق ISO 31000",
                    "framework": "ISO 31000"
                },
                "سياسة حوكمة تقنية المعلومات COBIT": {
                    "description": "سياسة لحوكمة تقنية المعلومات بناءً على إطار COBIT",
                    "framework": "COBIT"
                },
                "سياسة إدارة خدمات ITIL": {
                    "description": "سياسة لإدارة خدمات تقنية المعلومات بناءً على ITIL",
                    "framework": "ITIL"
                },
                "سياسة الامتثال PCI-DSS": {
                    "description": "سياسة للامتثال لمعايير أمان بيانات بطاقات الدفع",
                    "framework": "PCI-DSS"
                },
                "سياسة الامتثال SOC 2": {
                    "description": "سياسة للامتثال لمعايير خدمة الثقة SOC 2",
                    "framework": "SOC 2"
                },
                "سياسة الامتثال GDPR": {
                    "description": "سياسة للامتثال للائحة العامة لحماية البيانات",
                    "framework": "GDPR"
                }
            }
        },
        "audit": {
            "en": {
                "ISO 27001 Certification Audit": {
                    "description": "Audit for ISO 27001 certification readiness",
                    "framework": "ISO 27001",
                    "controls": 114
                },
                "ISO 22301 Compliance Audit": {
                    "description": "Audit for business continuity compliance",
                    "framework": "ISO 22301",
                    "controls": 60
                },
                "COBIT Governance Audit": {
                    "description": "Audit for IT governance maturity",
                    "framework": "COBIT",
                    "controls": 40
                },
                "PCI-DSS Compliance Audit": {
                    "description": "Audit for payment card security compliance",
                    "framework": "PCI-DSS",
                    "controls": 250
                }
            },
            "ar": {
                "تدقيق شهادة ISO 27001": {
                    "description": "تدقيق لجاهزية شهادة ISO 27001",
                    "framework": "ISO 27001",
                    "controls": 114
                },
                "تدقيق الامتثال ISO 22301": {
                    "description": "تدقيق للامتثال لاستمرارية الأعمال",
                    "framework": "ISO 22301",
                    "controls": 60
                },
                "تدقيق حوكمة COBIT": {
                    "description": "تدقيق لنضج حوكمة تقنية المعلومات",
                    "framework": "COBIT",
                    "controls": 40
                },
                "تدقيق الامتثال PCI-DSS": {
                    "description": "تدقيق للامتثال لأمان بطاقات الدفع",
                    "framework": "PCI-DSS",
                    "controls": 250
                }
            }
        },
        "risk": {
            "en": {
                "Enterprise Risk Assessment": {
                    "description": "Comprehensive enterprise-wide risk assessment",
                    "categories": ["Strategic", "Operational", "Financial", "Compliance"]
                },
                "Compliance Risk Assessment": {
                    "description": "Assessment of regulatory compliance risks",
                    "categories": ["Legal", "Regulatory", "Contractual", "Industry Standards"]
                },
                "Business Continuity Risk Assessment": {
                    "description": "Assessment of business continuity and disaster recovery risks",
                    "categories": ["Natural Disasters", "Technical Failures", "Human Factors"]
                }
            },
            "ar": {
                "تقييم المخاطر المؤسسية": {
                    "description": "تقييم شامل للمخاطر على مستوى المؤسسة",
                    "categories": ["استراتيجية", "تشغيلية", "مالية", "امتثال"]
                },
                "تقييم مخاطر الامتثال": {
                    "description": "تقييم مخاطر الامتثال التنظيمي",
                    "categories": ["قانونية", "تنظيمية", "تعاقدية", "معايير الصناعة"]
                },
                "تقييم مخاطر استمرارية الأعمال": {
                    "description": "تقييم مخاطر استمرارية الأعمال والتعافي من الكوارث",
                    "categories": ["كوارث طبيعية", "أعطال تقنية", "عوامل بشرية"]
                }
            }
        },
        "procedure": {
            "en": {
                "Internal Audit Procedure": {
                    "description": "Procedure for planning, executing, reporting, and following up on internal compliance audits",
                    "framework": "ISO 19011"
                },
                "Corrective Action Procedure": {
                    "description": "Procedure for identifying root causes, implementing corrections, and verifying effectiveness",
                    "framework": "ISO 27001"
                },
                "Document Control Procedure": {
                    "description": "Procedure for creating, reviewing, approving, distributing, and retiring controlled documents",
                    "framework": "ISO 27001"
                },
                "Management Review Procedure": {
                    "description": "Procedure for conducting management reviews including agenda, inputs, decisions, and action tracking",
                    "framework": "ISO 27001"
                },
                "Compliance Monitoring Procedure": {
                    "description": "Procedure for monitoring regulatory changes, assessing impact, and updating compliance posture",
                    "framework": "ISO 37301"
                },
                "Continuous Improvement Procedure": {
                    "description": "Procedure for Plan-Do-Check-Act (PDCA) cycle implementation across all management system domains",
                    "framework": "ISO 27001"
                }
            },
            "ar": {
                "إجراء التدقيق الداخلي": {
                    "description": "إجراء لتخطيط وتنفيذ والإبلاغ عن ومتابعة عمليات التدقيق الداخلي للامتثال",
                    "framework": "ISO 19011"
                },
                "إجراء الإجراءات التصحيحية": {
                    "description": "إجراء لتحديد الأسباب الجذرية وتنفيذ التصحيحات والتحقق من الفعالية",
                    "framework": "ISO 27001"
                },
                "إجراء ضبط الوثائق": {
                    "description": "إجراء لإنشاء ومراجعة واعتماد وتوزيع وإيقاف الوثائق المضبوطة",
                    "framework": "ISO 27001"
                },
                "إجراء مراجعة الإدارة": {
                    "description": "إجراء لإجراء مراجعات الإدارة بما في ذلك جدول الأعمال والمدخلات والقرارات وتتبع الإجراءات",
                    "framework": "ISO 27001"
                },
                "إجراء مراقبة الامتثال": {
                    "description": "إجراء لمراقبة التغييرات التنظيمية وتقييم الأثر وتحديث وضع الامتثال",
                    "framework": "ISO 37301"
                },
                "إجراء التحسين المستمر": {
                    "description": "إجراء لتنفيذ دورة PDCA (خطط-نفذ-تحقق-صحح) عبر جميع مجالات نظام الإدارة",
                    "framework": "ISO 27001"
                }
            }
        }
    },
    "erm": {
        "policy": {
            "en": {
                "Enterprise Risk Management Policy": {
                    "description": "Comprehensive ERM policy defining risk management framework, governance, and accountability across the organization",
                    "framework": "ISO 31000"
                },
                "Risk Appetite & Tolerance Policy": {
                    "description": "Policy establishing the organization's risk appetite statement, tolerance levels, and escalation thresholds",
                    "framework": "COSO ERM"
                },
                "Risk Assessment & Treatment Policy": {
                    "description": "Policy for systematic risk identification, analysis, evaluation, and treatment selection",
                    "framework": "ISO 31000"
                },
                "Business Continuity & Crisis Management Policy": {
                    "description": "Policy for business continuity planning, disaster recovery, and organizational crisis management",
                    "framework": "ISO 22301"
                },
                "Third-Party & Supply Chain Risk Policy": {
                    "description": "Policy governing risk management for vendors, suppliers, partners, and outsourced services",
                    "framework": "COSO ERM"
                },
                "Operational Risk Management Policy": {
                    "description": "Policy for managing risks arising from internal processes, people, systems, and external events",
                    "framework": "Basel III"
                },
                "Insurance & Risk Transfer Policy": {
                    "description": "Policy for risk financing, insurance procurement, and risk transfer mechanisms",
                    "framework": "IRM Standard"
                },
                "Risk Culture & Communication Policy": {
                    "description": "Policy promoting risk-aware culture, risk reporting, and stakeholder communication",
                    "framework": "COSO ERM"
                }
            },
            "ar": {
                "سياسة إدارة المخاطر المؤسسية": {
                    "description": "سياسة شاملة لإدارة المخاطر تحدد إطار العمل والحوكمة والمساءلة عبر المنظمة",
                    "framework": "ISO 31000"
                },
                "سياسة شهية المخاطر والتحمل": {
                    "description": "سياسة تحدد بيان شهية المخاطر ومستويات التحمل وعتبات التصعيد",
                    "framework": "COSO ERM"
                },
                "سياسة تقييم ومعالجة المخاطر": {
                    "description": "سياسة للتحديد والتحليل والتقييم واختيار معالجة المخاطر بشكل منهجي",
                    "framework": "ISO 31000"
                },
                "سياسة استمرارية الأعمال وإدارة الأزمات": {
                    "description": "سياسة لتخطيط استمرارية الأعمال والتعافي من الكوارث وإدارة الأزمات",
                    "framework": "ISO 22301"
                },
                "سياسة مخاطر الأطراف الثالثة وسلسلة التوريد": {
                    "description": "سياسة لإدارة مخاطر الموردين والشركاء ومقدمي الخدمات الخارجيين",
                    "framework": "COSO ERM"
                },
                "سياسة إدارة المخاطر التشغيلية": {
                    "description": "سياسة لإدارة المخاطر الناشئة من العمليات الداخلية والأفراد والأنظمة والأحداث الخارجية",
                    "framework": "Basel III"
                },
                "سياسة التأمين ونقل المخاطر": {
                    "description": "سياسة لتمويل المخاطر وشراء التأمين وآليات نقل المخاطر",
                    "framework": "IRM Standard"
                },
                "سياسة ثقافة المخاطر والتواصل": {
                    "description": "سياسة لتعزيز ثقافة الوعي بالمخاطر والإبلاغ عنها والتواصل مع أصحاب المصلحة",
                    "framework": "COSO ERM"
                }
            }
        },
        "audit": {
            "en": {
                "ERM Framework Maturity Audit": {
                    "description": "Comprehensive audit of ERM framework maturity, governance structures, and integration across the organization",
                    "framework": "COSO ERM"
                },
                "Risk Assessment Process Audit": {
                    "description": "Audit of risk identification, analysis, evaluation, and treatment processes and their effectiveness",
                    "framework": "ISO 31000"
                },
                "Business Continuity Readiness Audit": {
                    "description": "Audit of BCP/DRP readiness, testing frequency, recovery capabilities, and crisis management procedures",
                    "framework": "ISO 22301"
                },
                "Three Lines Model Audit": {
                    "description": "Audit of the three lines of defense model implementation, roles, and assurance coverage",
                    "framework": "IIA Standards"
                }
            },
            "ar": {
                "تدقيق نضج إطار إدارة المخاطر المؤسسية": {
                    "description": "تدقيق شامل لنضج إطار إدارة المخاطر وهياكل الحوكمة والتكامل عبر المنظمة",
                    "framework": "COSO ERM"
                },
                "تدقيق عمليات تقييم المخاطر": {
                    "description": "تدقيق عمليات تحديد وتحليل وتقييم ومعالجة المخاطر وفعاليتها",
                    "framework": "ISO 31000"
                },
                "تدقيق جاهزية استمرارية الأعمال": {
                    "description": "تدقيق جاهزية خطط استمرارية الأعمال وتكرار الاختبار وقدرات التعافي وإجراءات إدارة الأزمات",
                    "framework": "ISO 22301"
                },
                "تدقيق نموذج الخطوط الثلاثة": {
                    "description": "تدقيق تطبيق نموذج خطوط الدفاع الثلاثة والأدوار وتغطية التأكيد",
                    "framework": "IIA Standards"
                }
            }
        },
        "risk": {
            "en": {
                "Enterprise-Wide Risk Assessment": {
                    "description": "Comprehensive risk assessment covering strategic, operational, financial, compliance, and emerging risks",
                    "categories": ["Strategic", "Operational", "Financial", "Compliance", "Reputational", "Emerging"]
                },
                "Risk Appetite Calibration Assessment": {
                    "description": "Assessment to calibrate and validate the organization's risk appetite and tolerance thresholds",
                    "categories": ["Risk Appetite", "Tolerance", "Capacity", "Board Expectations"]
                },
                "Interconnected & Emerging Risk Assessment": {
                    "description": "Assessment of emerging risks, systemic risks, and interconnected risk cascades",
                    "categories": ["Geopolitical", "Climate", "Technology", "Pandemic", "Systemic"]
                }
            },
            "ar": {
                "تقييم المخاطر على مستوى المؤسسة": {
                    "description": "تقييم شامل للمخاطر يغطي المخاطر الاستراتيجية والتشغيلية والمالية والامتثال والناشئة",
                    "categories": ["استراتيجية", "تشغيلية", "مالية", "امتثال", "سمعة", "ناشئة"]
                },
                "تقييم معايرة شهية المخاطر": {
                    "description": "تقييم لمعايرة والتحقق من شهية المخاطر وعتبات التحمل في المنظمة",
                    "categories": ["شهية المخاطر", "التحمل", "القدرة", "توقعات مجلس الإدارة"]
                },
                "تقييم المخاطر المترابطة والناشئة": {
                    "description": "تقييم المخاطر الناشئة والمخاطر النظامية وتتابع المخاطر المترابطة",
                    "categories": ["جيوسياسية", "مناخية", "تقنية", "جائحة", "نظامية"]
                }
            }
        },
        "procedure": {
            "en": {
                "Risk Identification Procedure": {
                    "description": "Procedure for systematic identification of strategic, operational, financial, and compliance risks",
                    "framework": "ISO 31000"
                },
                "Risk Assessment and Scoring Procedure": {
                    "description": "Procedure for evaluating risk likelihood, impact, scoring methodology, and risk matrix application",
                    "framework": "ISO 31000"
                },
                "Risk Treatment Selection Procedure": {
                    "description": "Procedure for selecting treatment options (accept, mitigate, transfer, avoid) with cost-benefit analysis",
                    "framework": "COSO ERM"
                },
                "Risk Monitoring and Reporting Procedure": {
                    "description": "Procedure for ongoing risk monitoring, KRI tracking, escalation triggers, and periodic reporting",
                    "framework": "COSO ERM"
                },
                "Business Impact Analysis Procedure": {
                    "description": "Procedure for conducting BIA including critical process identification, RTO/RPO definition, and dependency mapping",
                    "framework": "ISO 22301"
                },
                "Risk Appetite Review Procedure": {
                    "description": "Procedure for defining, documenting, communicating, and periodically reviewing organizational risk appetite",
                    "framework": "COSO ERM"
                },
                "Emerging Risk Assessment Procedure": {
                    "description": "Procedure for scanning, identifying, and evaluating emerging and horizon risks",
                    "framework": "ISO 31000"
                },
                "Risk Committee Governance Procedure": {
                    "description": "Procedure for risk committee operations, meeting cadence, decision authority, and escalation protocols",
                    "framework": "COSO ERM"
                }
            },
            "ar": {
                "إجراء تحديد المخاطر": {
                    "description": "إجراء للتحديد المنهجي للمخاطر الاستراتيجية والتشغيلية والمالية ومخاطر الامتثال",
                    "framework": "ISO 31000"
                },
                "إجراء تقييم وتسجيل المخاطر": {
                    "description": "إجراء لتقييم احتمالية وأثر المخاطر ومنهجية التسجيل وتطبيق مصفوفة المخاطر",
                    "framework": "ISO 31000"
                },
                "إجراء اختيار معالجة المخاطر": {
                    "description": "إجراء لاختيار خيارات المعالجة (قبول، تخفيف، نقل، تجنب) مع تحليل التكلفة والعائد",
                    "framework": "COSO ERM"
                },
                "إجراء مراقبة المخاطر والإبلاغ": {
                    "description": "إجراء للمراقبة المستمرة للمخاطر وتتبع KRI ومحفزات التصعيد والتقارير الدورية",
                    "framework": "COSO ERM"
                },
                "إجراء تحليل أثر الأعمال": {
                    "description": "إجراء لإجراء BIA بما في ذلك تحديد العمليات الحرجة وتحديد RTO/RPO وخريطة الاعتماديات",
                    "framework": "ISO 22301"
                },
                "إجراء مراجعة شهية المخاطر": {
                    "description": "إجراء لتعريف وتوثيق وإبلاغ ومراجعة شهية المخاطر المؤسسية دورياً",
                    "framework": "COSO ERM"
                },
                "إجراء تقييم المخاطر الناشئة": {
                    "description": "إجراء لمسح وتحديد وتقييم المخاطر الناشئة ومخاطر الأفق",
                    "framework": "ISO 31000"
                },
                "إجراء حوكمة لجنة المخاطر": {
                    "description": "إجراء لعمليات لجنة المخاطر ووتيرة الاجتماعات وصلاحيات القرار وبروتوكولات التصعيد",
                    "framework": "COSO ERM"
                }
            }
        }
    }
}
POLICY_TEMPLATES = {"en": {}, "ar": {}}
STRATEGY_TEMPLATES = {"en": {}, "ar": {}}
AUDIT_TEMPLATES = {"en": {}, "ar": {}}
RISK_TEMPLATES = {"en": {}, "ar": {}}

# ── Execute framework awareness merge (must run after AWARENESS_MODULES is defined) ──
merge_framework_awareness()

# ============================================================================
# AI SERVICE
# ============================================================================

def check_ai_available():
    """Check if any AI API is available."""
    return bool(config.OPENAI_API_KEY or config.ANTHROPIC_API_KEY or config.GOOGLE_API_KEY or config.GROQ_API_KEY)

def get_available_providers():
    """Get list of available AI providers based on configured API keys."""
    providers = []
    if config.ANTHROPIC_API_KEY:
        providers.append({'id': 'anthropic', 'name': 'Anthropic Claude', 'model': 'claude-sonnet-4-20250514'})
    if config.OPENAI_API_KEY:
        providers.append({'id': 'openai', 'name': 'OpenAI GPT', 'model': 'gpt-4-turbo'})
    if config.GOOGLE_API_KEY:
        providers.append({'id': 'google', 'name': 'Google Gemini', 'model': 'gemini-2.0-flash'})
    if config.GROQ_API_KEY:
        providers.append({'id': 'groq', 'name': 'Groq (Llama 3.1)', 'model': 'llama-3.1-70b-versatile'})
    return providers

def get_user_ai_preference(task_type='generate'):
    """Get user's preferred AI provider for a specific task type.
    
    Args:
        task_type: 'generate' for document generation, 'review' for document review
    
    Returns:
        Provider string ('anthropic', 'openai', 'google', 'groq', 'auto') or None
    """
    # Map frontend values to backend values (backward compatibility)
    _provider_normalize = {
        'claude': 'anthropic', 'gpt': 'openai', 'gemini': 'google',
        'groq': 'groq', 'auto': 'auto',
        'anthropic': 'anthropic', 'openai': 'openai', 'google': 'google',
    }
    try:
        from flask import session
        if 'user_id' not in session:
            return 'auto'
        
        conn = get_db()
        col = 'ai_provider_generate' if task_type == 'generate' else 'ai_provider_review'
        user = conn.execute(f'SELECT {col} FROM users WHERE id = ?', (session['user_id'],)).fetchone()
        
        if user and user[0]:
            return _provider_normalize.get(user[0], user[0])
    except Exception:
        pass
    return 'auto'

def get_ai_provider(task_type='generate'):
    """Determine which AI provider to use based on user preference and available keys.
    
    Args:
        task_type: 'generate' for document generation, 'review' for document review
    """
    # First check user preference
    user_pref = get_user_ai_preference(task_type)
    
    # If user has a specific preference (not auto), try to use it
    if user_pref and user_pref != 'auto':
        key_map = {
            'anthropic': config.ANTHROPIC_API_KEY,
            'openai': config.OPENAI_API_KEY,
            'google': config.GOOGLE_API_KEY,
            'groq': config.GROQ_API_KEY
        }
        if key_map.get(user_pref):
            return user_pref
    
    # Fall back to system config or auto-detect
    provider = config.AI_PROVIDER.lower()
    
    if provider == 'openai' and config.OPENAI_API_KEY:
        return 'openai'
    elif provider == 'anthropic' and config.ANTHROPIC_API_KEY:
        return 'anthropic'
    elif provider == 'google' and config.GOOGLE_API_KEY:
        return 'google'
    elif provider == 'groq' and config.GROQ_API_KEY:
        return 'groq'
    elif provider == 'auto' or user_pref == 'auto':
        # Priority: Anthropic > OpenAI > Groq > Google
        if config.ANTHROPIC_API_KEY:
            return 'anthropic'
        elif config.OPENAI_API_KEY:
            return 'openai'
        elif config.GROQ_API_KEY:
            return 'groq'
        elif config.GOOGLE_API_KEY:
            return 'google'
    return None

def get_ai_status():
    """Get current AI provider status for display."""
    provider = get_ai_provider()
    if not provider:
        return {'provider': 'simulation', 'model': 'Built-in', 'connected': False}
    
    models = {
        'openai': config.AI_MODEL or 'gpt-4-turbo',
        'anthropic': config.AI_MODEL or 'claude-sonnet-4-20250514',
        'google': config.AI_MODEL or 'gemini-2.0-flash',
        'groq': config.AI_MODEL or 'llama-3.1-70b-versatile',
    }
    names = {
        'openai': 'OpenAI', 
        'anthropic': 'Anthropic Claude', 
        'google': 'Google Gemini',
        'groq': 'Groq (Llama 3.1)'
    }
    return {'provider': names.get(provider, provider), 'model': models.get(provider, ''), 'connected': True}

def clean_ai_response(text):
    """Remove instruction artifacts that AI might have included in output.
    
    This function removes:
    - FINAL CRITICAL REMINDER sections
    - Instruction warnings (⚠️ CRITICAL:...)
    - Section markers like [SECTION]
    - Meta-commentary about the document
    - AI closing statements (شكراً لكم, نتطلع للعمل, etc.)
    - Confidence Level Interpretation sections
    """
    import re
    
    if not text:
        return text
    
    # Patterns to remove
    patterns = [
        # English FINAL CRITICAL REMINDER and ALL text after it until next section or end
        r'⚠️⚠️⚠️\s*FINAL CRITICAL REMINDER\s*⚠️⚠️⚠️.*?(?=\n##|\Z)',
        # Arabic FINAL CRITICAL REMINDER
        r'⚠️⚠️⚠️\s*تذكير نهائي مهم جداً\s*⚠️⚠️⚠️.*?(?=\n##|\Z)',
        r'⚠️⚠️⚠️\s*تذكير نهائي\s*⚠️⚠️⚠️.*?(?=\n##|\Z)',
        # "The above sections provide..." meta-commentary (multi-line)
        r'The above sections provide[^#]*?(?=\n##|\n\n##|\Z)',
        r'الأقسام أعلاه توفر[^#]*?(?=\n##|\n\n##|\Z)',
        # Remove any paragraph mentioning "detailed implementation guides for each identified gap"
        r'[^\n]*detailed implementation guides for each identified gap[^\n]*(?:\n[^\n#]*)*?(?=\n\n|\n##|\Z)',
        r'[^\n]*ensuring a structured and comprehensive approach[^\n]*(?:\n[^\n#]*)*?(?=\n\n|\n##|\Z)',
        r'[^\n]*Each gap\'s implementation steps have been outlined[^\n]*(?:\n[^\n#]*)*?(?=\n\n|\n##|\Z)',
        r'[^\n]*adhering to the requirements for specificity[^\n]*(?:\n[^\n#]*)*?(?=\n\n|\n##|\Z)',
        # NOTE: Do NOT remove [SECTION] markers here - they are needed for strategy parsing
        # Remove instruction lines starting with ⚠️ CRITICAL or ⚠️ مهم
        r'\n⚠️\s*CRITICAL[^\n]*',
        r'\n⚠️\s*مهم جداً[^\n]*',
        r'\n⚠️\s*تنبيه[^\n]*',
        r'\n⚠️\s*تعليمات[^\n]*',
        r'\n⚠️\s*FRAMEWORK RESTRICTION[^\n]*',
        r'\nCRITICAL INSTRUCTION - SELECTED FRAMEWORK[^\n]*',
        r'\n⚠️\s*قيد الأطر التنظيمية[^\n]*',
        # Remove leaked prompt instruction phrases from within content
        r'The user has selected these specific frameworks?:\s*',
        r'Selected Frameworks \(USE ONLY THESE\):\s*',
        r'Allowed frameworks:\s*',
        r'The selected framework for this audit is:\s*',
        r'Reference framework:\s*',
        # Remove "per {raw_instruction}" patterns where AI echoed prompt labels
        r'\bper\s+The user has selected[^|.\n]*',
        r'\bper\s+Allowed frameworks[^|.\n]*',
        r'\bper\s+Selected Frameworks[^|.\n]*',
        r'\bوفق\s+الأطر المحددة[^|.\n]*',
        # Remove FRAMEWORK/OUTPUT RULES sections if echoed
        r'⚠\s*FRAMEWORK RULES[^\n]*(?:\n-[^\n]*)*',
        r'⚠\s*OUTPUT RULES[^\n]*(?:\n\d+\.[^\n]*)*',
        # Remove instruction blocks with warning emojis
        r'⚠️⚠️⚠️[^⚠]*?⚠️⚠️⚠️',
        r'⚠️ CRITICAL RULES[^\n]*(?:\n-[^\n]*)*',
        r'⚠️ تعليمات مهمة[^\n]*(?:\n-[^\n]*)*',
        r'⚠️ CRITICAL:[^\n]*',
        # Remove lines that are just instructions in parentheses
        r'\n\(?\d+-\d+\s+(steps|خطوات)[^\)]*\)?',
        r'\n\(?\d+-\d+\s+(KPIs?|مؤشر)[^\)]*\)?',
        r'\n\(?\d+-\d+\s+(gaps?|فجوات?)[^\)]*\)?',
        r'\n\(?\d+-\d+\s+(risks?|مخاطر)[^\)]*\)?',
        r'\n\(?\d+-\d+\s+(objectives?|أهداف)[^\)]*\)?',
        r'\n\(?\d+-\d+\s+(activities?|أنشطة)[^\)]*\)?',
        # Remove standalone parenthetical instructions
        r'\(\d+-\d+\s+[^\)]+\)',
        # Remove meta-commentary sentences
        r'This document provides a comprehensive[^\n\.]*\.',
        r'The strategy ensures alignment with[^\n\.]*selected frameworks[^\n\.]*\.',
        # Remove Arabic meta-commentary — ONLY specific boilerplate phrases
        r'هذه الوثيقة توفر نظرة شاملة[^\n\.]*\.',
        r'تضمن هذه الاستراتيجية التوافق مع الأطر[^\n\.]*\.',
        # Remove AI closing statements (Arabic) — ONLY when clearly conversational chatter
        # These patterns are anchored to specific AI-chatter phrases to avoid
        # stripping legitimate section conclusions
        r'شكراً لكم على ثقتكم[^\n]*',
        r'شكرا لكم على ثقتكم[^\n]*',
        r'شكراً لكم\.\s*$',
        r'شكرا لكم\.\s*$',
        r'نتطلع للعمل معكم[^\n]*',
        r'نأمل أن يكون هذا (?:التقرير|المستند|العمل) مفيداً[^\n]*',
        r'نأمل أن تكون هذه (?:الاستراتيجية|الوثيقة) مفيدة[^\n]*',
        # Only strip "في الختام" when followed by AI sign-off patterns, NOT in strategy sections
        r'في الختام،?\s*(?:نأمل|نتمنى|نؤكد على التزامنا|يسعدنا)[^\n]*',
        r'ختاماً،?\s*نتمنى[^\n]*',
        # Remove AI closing statements (English)
        r'Thank you for.*?(?=\n\n|\n---|\Z)',
        r'We look forward to.*?(?=\n\n|\n---|\Z)',
        r'Please feel free to.*?(?=\n\n|\n---|\Z)',
        r'If you have any questions.*?(?=\n\n|\n---|\Z)',
        # Remove Confidence Level Interpretation sections
        r'\*\*Confidence Level Interpretation:\*\*.*?(?=\n\n---|\n##|\Z)',
        r'\*\*تفسير مستوى الثقة:\*\*.*?(?=\n\n---|\n##|\Z)',
        r'Confidence Level Interpretation:.*?(?=\n\n---|\n##|\Z)',
        # Remove bullet lists explaining confidence levels
        r'-\s*90-100%: High confidence[^\n]*\n-\s*70-89%[^\n]*\n-\s*50-69%[^\n]*\n-\s*Below 50%[^\n]*',
        r'[-•]\s*90-100%.*?[-•]\s*أقل من 50%[^\n]*',
        # Remove entire Confidence Assessment sections — ONLY when they are standalone meta-commentary
        # (NOT when they are numbered strategy section 6: "## 6. تقييم الثقة")
        r'## Confidence Assessment\s*\n\*\*Confidence Level:\*\*[^#]*?(?=\n---|\n##|\Z)',
        # SAFE: Only strip "## مستوى الثقة" when NOT preceded by a section number (protects "## 6. ...")
        r'## (?!\d+\.?\s)مستوى الثقة[^#]*?(?=\n---|\n##|\Z)',
        r'\*\*Key Factors:\*\*\s*\n\|[^#]*?(?=\n---|\n##|\Z)',
        # SAFE: Only strip "### عوامل تقييم الثقة" followed by boilerplate interpretation text
        r'### عوامل تقييم الثقة\s*:?\s*\n(?:\s*[-•].*(?:90|80|70|50)%[^\n]*\n?)+',
        # Remove prompt text that leaked into output
        r'The user has selected these specific frameworks?:[^\n|]*',
        r'selected these specific frameworks?:[^\n|]*',
        r'FRAMEWORK RESTRICTION[^\n|]*',
        r'CRITICAL INSTRUCTION[^\n|]*',
        r'(?:^|\|)\s*OUTPUT RULES[^\n|]*',
        r'Do NOT echo[^\n|]*',
        r'NEVER include phrases like[^\n|]*',
        # Clean leaked instruction fragments inside table cells
        r'(?<=\|)\s*The user has selected[^|]*(?=\|)',
    ]
    
    for pattern in patterns:
        text = re.sub(pattern, '', text, flags=re.DOTALL | re.IGNORECASE | re.MULTILINE)
    
    # Clean up multiple consecutive newlines (more than 3)
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    
    # Clean up trailing whitespace on lines
    text = re.sub(r'[ \t]+\n', '\n', text)
    
    # Shorten repetitive framework full names to abbreviations
    # Replace "NCA ECC (Essential Cybersecurity Controls)" → "NCA ECC"
    text = re.sub(r'NCA ECC\s*\(Essential Cybersecurity Controls\)', 'NCA ECC', text)
    text = re.sub(r'NCA CSCC\s*\(Critical Systems Cybersecurity Controls\)', 'NCA CSCC', text)
    text = re.sub(r'NCA DCC\s*\(Data Cybersecurity Controls\)', 'NCA DCC', text)
    text = re.sub(r'NCA OTCC\s*\(Operational Technology Cybersecurity Controls\)', 'NCA OTCC', text)
    text = re.sub(r'NCA TCC\s*\(Telework Cybersecurity Controls\)', 'NCA TCC', text)
    text = re.sub(r'NCA CCC\s*\(Cloud Cybersecurity Controls\)', 'NCA CCC', text)
    text = re.sub(r'NCA NCS\s*\(National Cryptographic Standards\)', 'NCA NCS', text)
    text = re.sub(r'SAMA CSF\s*\(Cybersecurity Framework\)', 'SAMA CSF', text)
    text = re.sub(r'SAMA BCM\s*\(Business Continuity Management\)', 'SAMA BCM', text)
    text = re.sub(r'PDPL\s*\(Personal Data Protection Law\)', 'PDPL', text)
    # Also clean "Essential Cybersecurity Controls" alone
    text = re.sub(r'Essential Cybersecurity Controls\s*\(NCA ECC\)', 'NCA ECC', text)
    
    return text.strip()


def ensure_markdown_formatting(text):
    """Post-process AI output to ensure proper Markdown structure.
    
    Fixes common AI output issues:
    - Headings merged with content on same line
    - Missing blank lines before/after tables
    - Missing table separator rows (|---|---|)
    - ### sub-headings stuck to ## headings
    - **bold:** labels merged with preceding headings
    - --- separators not on their own lines
    - Table rows merged onto heading lines
    - Broken bold markers (Score:** 78% **Justification:)
    """
    import re
    if not text:
        return text
    
    # 0. CRITICAL: Split heading appearing in the MIDDLE of a line (text ### Heading)
    # e.g., "...across all service offerings. ### Strategic Objectives:" → split them
    text = re.sub(r'([^\n#])\s+(#{2,4}\s+[A-Za-z\u0600-\u06FF])', r'\1\n\n\2', text)
    
    # 0a. CRITICAL: Split heading + table merged on same line
    # e.g., "### Strategic Objectives: | # | Objective |..." → heading then table on next line
    text = re.sub(r'(#{2,4}\s+[^|\n]+?:?)\s*(\|\s*#?\s*\|)', r'\1\n\n\2', text)
    
    # 0b. Split any text ending with ":" followed immediately by a pipe table
    text = re.sub(r'([^\n|]+:)\s*(\|\s*[^|]+\|)', lambda m: m.group(1) + '\n\n' + m.group(2) if '|' not in m.group(1) else m.group(0), text)
    
    # 0c. Fix broken bold markers in confidence section
    # Pattern: "Confidence Score:** 78% **Score Justification:**" 
    # → "**Confidence Score:** 78%\n\n**Score Justification:**"
    # Detect: "text:** value **text:**" where ** are misaligned
    text = re.sub(
        r'([A-Za-z\u0600-\u06FF][\w\s]+:)\*\*\s+(\d+%?)\s+\*\*([A-Za-z\u0600-\u06FF][\w\s]+:)\*\*',
        r'**\1** \2\n\n**\3**',
        text
    )
    # Also handle single misplaced bold: "word:**" at start of a line → "**word:**"
    text = re.sub(r'^([A-Za-z\u0600-\u06FF][\w\s]+:)\*\*\s*$', r'**\1**', text, flags=re.MULTILINE)
    
    # 1. Split ## heading from ### sub-heading merged on same line
    text = re.sub(r'(##\s+\d*\.?\s*[^\n#]+?)\s+(###\s+)', r'\1\n\n\2', text)
    
    # 2. Split ## heading from **bold:** merged on same line
    text = re.sub(r'(##\s+\d*\.?\s*[^\n*]+?)\s+(\*\*[^*]+\*\*)', r'\1\n\n\2', text)
    
    # 3. Ensure blank line before table start (non-table line followed by | line)
    text = re.sub(r'([^\n|])\n(\|[^\n]+\|)', r'\1\n\n\2', text)
    
    # 4. Ensure blank line after table block (last | row followed by non-| non-empty line)
    text = re.sub(r'(\|[^\n]+\|)\n([^|\n\s])', r'\1\n\n\2', text)
    
    # 5. Ensure --- separators are on their own line with blank lines around them
    text = re.sub(r'([^\n-])\s*---\s*\n', r'\1\n\n---\n\n', text)
    text = re.sub(r'\n---\s*([^\n\-\s])', r'\n---\n\n\1', text)
    
    # 6. Ensure blank line before ## and ### and #### headings
    text = re.sub(r'([^\n])\n(#{2,4}\s+)', r'\1\n\n\2', text)
    
    # 7. Ensure blank line after heading before non-heading non-table content
    text = re.sub(r'(#{2,4}\s+[^\n]+)\n([^#\n|\s])', r'\1\n\n\2', text)
    
    # 8. Split ## / ### / #### heading from preceding content merged on same line
    text = re.sub(r'([^\n])\s+(#{2,4}\s+)', r'\1\n\n\2', text)
    
    # 9. Ensure **Evidence Required:** is on its own line after tables
    text = re.sub(r'(\|[^\n]+\|)\n(\*\*(?:Evidence|الأدلة)[^\n]+)', r'\1\n\n\2', text)
    
    # 10. Fix bullet points: ensure • is preceded by blank line if after table/heading
    text = re.sub(r'(\|[^\n]+\|)\n(•\s)', r'\1\n\n\2', text)
    
    # 11. Fix table rows that got joined onto one line (pipe-separated content without newlines)
    lines = text.split('\n')
    fixed_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|') and stripped.count('|') > 10:
            parts = re.split(r'\|\s*\|\s*(?=\d+\s*\||\s*#\s*\|)', stripped)
            if len(parts) > 1:
                for p in parts:
                    p = p.strip()
                    if p and not p.startswith('|'):
                        p = '| ' + p
                    if p and not p.endswith('|'):
                        p = p + ' |'
                    if p.strip():
                        fixed_lines.append(p)
                continue
        fixed_lines.append(line)
    text = '\n'.join(fixed_lines)
    
    # 12. CRITICAL: Insert missing table separator rows (|---|---|)
    # marked.js requires a separator row after the header to render a table
    lines = text.split('\n')
    repaired_lines = []
    for i, line in enumerate(lines):
        repaired_lines.append(line)
        stripped = line.strip()
        # Check if this is a potential table header row (starts with |, has multiple |)
        if stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 3:
            # Only process if this is the FIRST row of a table block
            # (previous non-empty line does NOT start with |)
            prev_idx = i - 1
            while prev_idx >= 0 and lines[prev_idx].strip() == '':
                prev_idx -= 1
            prev_is_table = prev_idx >= 0 and lines[prev_idx].strip().startswith('|')
            
            if not prev_is_table:
                # This is a header row - check if next non-empty line is a separator
                next_idx = i + 1
                while next_idx < len(lines) and lines[next_idx].strip() == '':
                    next_idx += 1
                if next_idx < len(lines):
                    next_stripped = lines[next_idx].strip()
                    is_separator = bool(re.match(r'^\|[\s\-:|]+\|$', next_stripped))
                    if not is_separator and next_stripped.startswith('|') and next_stripped.endswith('|'):
                        # Header → data with no separator - insert one
                        cols = stripped.count('|') - 1
                        separator = '|' + '|'.join(['---'] * cols) + '|'
                        repaired_lines.append(separator)
    text = '\n'.join(repaired_lines)
    
    # 13. Fix consecutive blank lines (max 2)
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    
    return text.strip()


def generate_ai_content(prompt, language='en', task_type='generate', content_type=None):
    """Generate content using the user's preferred AI provider.
    
    Args:
        prompt: The prompt to send to the AI
        language: 'en' or 'ar'
        task_type: 'generate' for document generation, 'review' for document review
        content_type: Explicit type for simulation fallback routing
    """
    # ── Pillar 3: Token quota enforcement ──
    user_id = None
    try:
        user_id = session.get('user_id')
    except RuntimeError:
        pass  # Outside request context (background thread)
    
    if user_id:
        allowed, used, limit = check_token_quota(user_id)
        if not allowed:
            print(f"TOKEN LIMIT: user {user_id} at {used}/{limit}", flush=True)
            if language == 'ar':
                return f"## ⚠️ تم الوصول إلى حد الاستخدام\n\nلقد استخدمت {used:,} من أصل {limit:,} رمز. يرجى التواصل مع المسؤول لزيادة الحد.\n\n**معرف المستخدم:** {user_id}"
            return f"## ⚠️ Token Quota Exceeded\n\nYou have used {used:,} of {limit:,} tokens. Please contact your administrator to increase the limit.\n\n**User ID:** {user_id}"
    
    provider = get_ai_provider(task_type)
    
    if not provider:
        print("DEBUG: No AI provider available, using simulation", flush=True)
        return generate_simulation_content(prompt, language, content_type=content_type)
    
    # ── Domain-specific system prompts to prevent cross-domain leakage ──
    domain_system_prompts = {
        'en': {
            'Cyber Security': (
                "You are a Cybersecurity Strategy Expert. Focus STRICTLY on cybersecurity governance, "
                "threat management, SOC operations, CSIRT, and CISO-level organizational structures. "
                "Do NOT discuss data governance, AI ethics, or digital transformation unless explicitly asked."
            ),
            'Data Management': (
                "You are a Data Privacy Officer and Data Governance Expert. Focus STRICTLY on data governance, "
                "data quality, data lifecycle management, NDMO compliance, data stewardship, CDO structures, "
                "and data protection (PDPL). "
                "SYSTEM RULE — FORBIDDEN TERMS: Do NOT mention 'CISO', 'Cybersecurity Department', 'SOC', "
                "'CSIRT', or 'Security Operations Center' anywhere in the output. These are cybersecurity-specific "
                "roles that do not belong in a Data Management strategy. "
                "REQUIRED REPLACEMENTS: Use 'Data Protection Officer (DPO)' instead of CISO, "
                "'Data Governance Committee' instead of Cybersecurity Department, "
                "'Data Management Office (DMO)' with CDO, Data Stewards, and Data Quality Analysts "
                "for any organizational structure recommendations."
            ),
            'Artificial Intelligence': (
                "You are an AI Governance and Ethics Expert. Focus STRICTLY on AI ethics, model transparency, "
                "bias assessment, SDAIA regulations, responsible AI practices, and AI risk management. "
                "SYSTEM RULE — FORBIDDEN TERMS: Do NOT mention 'CISO', 'Cybersecurity Department', 'SOC', "
                "'CSIRT', or 'Security Operations Center' anywhere in the output. These are cybersecurity-specific "
                "roles that do not belong in an AI Governance strategy. "
                "REQUIRED REPLACEMENTS: Use 'AI Ethics Board' instead of cybersecurity steering committee, "
                "'AI Governance Committee' or 'AI Ethics Office' with AI Ethics Officer, Model Risk team, "
                "and AI Compliance roles for organizational structure recommendations."
            ),
            'Digital Transformation': (
                "You are a Digital Transformation Strategy Expert. Focus STRICTLY on digitization strategy, "
                "emerging technologies, change management, digital skills, and innovation governance. "
                "SYSTEM RULE — FORBIDDEN TERMS: Do NOT mention 'CISO', 'Cybersecurity Department', 'SOC', "
                "'CSIRT', or 'Security Operations Center' anywhere in the output. "
                "REQUIRED REPLACEMENTS: Use 'Chief Digital Officer (CDO)', 'Digital Transformation Office', "
                "and 'Innovation Team' for organizational structure recommendations."
            ),
            'Enterprise Risk Management': (
                "You are an Enterprise Risk Management Expert specializing in COSO ERM and ISO 31000. "
                "Focus STRICTLY on risk identification, assessment, treatment, monitoring, risk appetite, "
                "and organizational resilience. "
                "SYSTEM RULE — FORBIDDEN TERMS: Do NOT mention 'CISO', 'Cybersecurity Department', 'SOC', "
                "'CSIRT', or 'Security Operations Center' anywhere in the output. "
                "REQUIRED REPLACEMENTS: Use 'Chief Risk Officer (CRO)', 'Risk Owners', "
                "'Risk Management Committee', and 'Risk Analysts' for organizational structure recommendations."
            ),
            'Global Standards': (
                "You are a Global Standards and Compliance Expert. Focus STRICTLY on international standards "
                "compliance (ISO, NIST, COBIT), certification processes, and standards harmonization. "
                "SYSTEM RULE — FORBIDDEN TERMS: Do NOT mention 'CISO', 'Cybersecurity Department', or 'SOC' "
                "unless the specific standard being discussed is a cybersecurity standard. "
                "REQUIRED REPLACEMENTS: Use 'Chief Compliance Officer (CCO)', 'Standards Liaisons', "
                "and 'Audit Coordinators' for organizational structure recommendations."
            ),
        },
        'ar': {
            'Cyber Security': (
                "أنت خبير في استراتيجيات الأمن السيبراني. ركّز حصرياً على حوكمة الأمن السيبراني، "
                "وإدارة التهديدات، وعمليات SOC، وفريق الاستجابة CSIRT، وهيكل CISO التنظيمي. "
                "لا تناقش حوكمة البيانات أو أخلاقيات الذكاء الاصطناعي أو التحول الرقمي إلا إذا طُلب صراحةً."
            ),
            'Data Management': (
                "أنت خبير في حوكمة البيانات وحماية الخصوصية. ركّز حصرياً على حوكمة البيانات، "
                "وجودة البيانات، وإدارة دورة حياة البيانات، والامتثال لـ NDMO، وأمناء البيانات، "
                "وحماية البيانات الشخصية (PDPL). "
                "قاعدة نظام — مصطلحات محظورة: يُمنع منعاً باتاً ذكر 'CISO' أو 'إدارة الأمن السيبراني' أو 'SOC' "
                "أو 'مركز عمليات الأمن' في أي مكان من المخرجات. هذه أدوار خاصة بالأمن السيبراني. "
                "البدائل المطلوبة: استخدم 'مسؤول حماية البيانات (DPO)' بدلاً من CISO، "
                "'لجنة حوكمة البيانات' بدلاً من إدارة الأمن السيبراني، "
                "'مكتب إدارة البيانات (DMO)' مع مدير البيانات CDO وأمناء البيانات وفريق جودة البيانات."
            ),
            'Artificial Intelligence': (
                "أنت خبير في حوكمة الذكاء الاصطناعي وأخلاقياته. ركّز حصرياً على أخلاقيات AI، "
                "وشفافية النماذج، وتقييم التحيز، ولوائح سدايا SDAIA، وممارسات AI المسؤولة. "
                "قاعدة نظام — مصطلحات محظورة: يُمنع منعاً باتاً ذكر 'CISO' أو 'إدارة الأمن السيبراني' أو 'SOC'. "
                "البدائل المطلوبة: استخدم 'مجلس أخلاقيات AI' أو 'لجنة حوكمة AI' أو 'مكتب أخلاقيات AI' "
                "مع مسؤول أخلاقيات AI وفريق مخاطر النماذج وفريق امتثال AI."
            ),
            'Digital Transformation': (
                "أنت خبير في استراتيجيات التحول الرقمي. ركّز حصرياً على استراتيجية الرقمنة "
                "والتقنيات الناشئة وإدارة التغيير والمهارات الرقمية وحوكمة الابتكار. "
                "قاعدة نظام — مصطلحات محظورة: يُمنع ذكر 'CISO' أو 'إدارة الأمن السيبراني' أو 'SOC'. "
                "البدائل المطلوبة: استخدم 'مدير التحول الرقمي' و'مكتب التحول الرقمي' و'فريق الابتكار'."
            ),
            'Enterprise Risk Management': (
                "أنت خبير في إدارة المخاطر المؤسسية وفق COSO ERM وISO 31000. ركّز حصرياً على "
                "تحديد وتقييم ومعالجة ومراقبة المخاطر وشهية المخاطر والمرونة المؤسسية. "
                "قاعدة نظام — مصطلحات محظورة: يُمنع ذكر 'CISO' أو 'إدارة الأمن السيبراني' أو 'SOC'. "
                "البدائل المطلوبة: استخدم 'مدير المخاطر (CRO)' و'ملاك المخاطر' و'لجنة المخاطر'."
            ),
            'Global Standards': (
                "أنت خبير في المعايير العالمية والامتثال. ركّز حصرياً على الامتثال للمعايير الدولية "
                "(ISO, NIST, COBIT) وعمليات الاعتماد والشهادات. "
                "قاعدة نظام — مصطلحات محظورة: لا تذكر 'CISO' أو 'SOC' إلا إذا كان المعيار المناقش معياراً للأمن السيبراني."
            ),
        }
    }
    
    # ── Framework adherence rule ──
    framework_rule_en = (
        "\n\nSTRICT ADHERENCE: Use ONLY the selected framework provided in the user input. "
        "Do not suggest or mention other regulations (e.g., ISO, NIST, COBIT, NCA ECC, SAMA CSF) "
        "unless they are specifically listed in the user's selected frameworks."
    )
    framework_rule_ar = (
        "\n\nالتزام صارم: استخدم فقط الإطار التنظيمي المحدد في مدخلات المستخدم. "
        "لا تقترح أو تذكر أي أنظمة أخرى (مثل ISO، NIST، COBIT، NCA ECC، SAMA CSF) "
        "إلا إذا كانت مدرجة في الأطر المختارة من المستخدم."
    )
    
    # ── Technical Translation Glossary ──
    translation_glossary = (
        "\n\nTechnical Translation Glossary (USE THESE EXACT TRANSLATIONS):\n"
        "- Zero-day → الثغرات الصفرية (NOT يوم الصفر or استغلال يوم الصفر)\n"
        "- Phishing → التصيد الاحتيالي\n"
        "- Ransomware → برمجيات الفدية\n"
        "- Malware → البرمجيات الخبيثة\n"
        "- Social Engineering → الهندسة الاجتماعية\n"
        "- Penetration Testing → اختبار الاختراق\n"
        "- Vulnerability → الثغرة الأمنية\n"
        "- Firewall → جدار الحماية\n"
        "- Encryption → التشفير\n"
        "- Data Breach → اختراق البيانات\n"
        "- Threat Intelligence → استخبارات التهديدات\n"
        "- Incident Response → الاستجابة للحوادث\n"
        "- Access Control → التحكم في الوصول\n"
        "- Authentication → المصادقة\n"
        "- Authorization → التفويض\n"
        "- Data Loss Prevention (DLP) → منع فقدان البيانات\n"
        "- Business Continuity → استمرارية الأعمال\n"
        "- Disaster Recovery → التعافي من الكوارث\n"
        "- Governance → الحوكمة (NOT الإدارة or الحكم)\n"
        "- Compliance → الامتثال (NOT المطابقة or التوافق)\n"
        "- Risk Appetite → شهية المخاطر\n"
        "- Stakeholder → أصحاب المصلحة\n"
        "- Framework → الإطار التنظيمي"
    )
    
    # Detect domain from prompt content
    detected_domain = None
    domain_keywords = {
        'Cyber Security': ['cybersecurity', 'cyber security', 'الأمن السيبراني', 'siem', 'soc ', 'csirt'],
        'Data Management': ['data management', 'data governance', 'إدارة البيانات', 'حوكمة البيانات', 'ndmo', 'pdpl', 'data quality'],
        'Artificial Intelligence': ['artificial intelligence', 'ai governance', 'الذكاء الاصطناعي', 'sdaia', 'ai ethics'],
        'Digital Transformation': ['digital transformation', 'التحول الرقمي', 'digitization'],
        'Enterprise Risk Management': ['enterprise risk', 'إدارة المخاطر المؤسسية', 'coso', 'risk management framework'],
        'Global Standards': ['global standards', 'المعايير العالمية', 'iso 27001', 'cobit'],
    }
    prompt_lower = prompt.lower()[:500]
    for dom, kws in domain_keywords.items():
        if any(kw in prompt_lower for kw in kws):
            detected_domain = dom
            break
    
    lang_key = 'ar' if language == 'ar' else 'en'
    
    # Build system prompt: domain-specific + framework rule + glossary
    if detected_domain and detected_domain in domain_system_prompts[lang_key]:
        system_prompt = domain_system_prompts[lang_key][detected_domain]
    else:
        system_prompt = ("You are an expert GRC consultant. Provide professional, detailed responses."
                        if lang_key == 'en' else
                        "أنت مستشار خبير في الحوكمة والمخاطر والامتثال. قدم ردوداً مهنية ومفصلة باللغة العربية.")
    
    system_prompt += framework_rule_ar if lang_key == 'ar' else framework_rule_en
    
    if lang_key == 'ar':
        system_prompt += translation_glossary
    
    # ── Markdown table formatting enforcement ──
    table_format_rule = (
        "\n\nMARKDOWN TABLE FORMATTING — CRITICAL RULES:\n"
        "1. Every table MUST use proper Markdown pipe syntax with a header separator row.\n"
        "2. Every table MUST be preceded by a blank line and followed by a blank line.\n"
        "3. Do NOT use merged cells, nested tables, or complex formatting inside cells.\n"
        "4. Header separator row must use dashes: |---|---|---|\n"
        "5. Example of correct table:\n\n"
        "| # | Item | Status |\n"
        "|---|------|--------|\n"
        "| 1 | First item | Active |\n"
        "| 2 | Second item | Pending |\n\n"
        "6. Do NOT use conversational filler like 'Here is your strategy' or 'I have created'. "
        "Start each section directly with content."
    )
    system_prompt += table_format_rule
    
    # ── Silence all AI conversational chatter ──
    silence_rule = (
        "\n\nOUTPUT DISCIPLINE — STRICT:\n"
        "- Start your response IMMEDIATELY with the first section heading. "
        "Do NOT begin with greetings, introductions, or meta-commentary.\n"
        "- Do NOT end with 'Thank you', 'We hope', 'In conclusion we wish', "
        "'شكراً لكم', 'نأمل أن يكون', or any closing pleasantries.\n"
        "- Do NOT include confidence level interpretation tables unless specifically requested.\n"
        "- Do NOT echo or reference these instructions in your output.\n"
        "- Do NOT include phrases like 'As per the selected framework' or "
        "'وفقاً للإطار المحدد' — simply use the framework directly."
    )
    system_prompt += silence_rule
    
    print(f"DEBUG: Using {provider} for {task_type} task", flush=True)
    
    result = None
    try:
        if provider == 'anthropic':
            result = _generate_anthropic(system_prompt, prompt, language)
        elif provider == 'openai':
            result = _generate_openai(system_prompt, prompt, language)
        elif provider == 'google':
            result = _generate_google(system_prompt, prompt, language)
        elif provider == 'groq':
            result = _generate_groq(system_prompt, prompt, language)
    except Exception as e:
        print(f"DEBUG: {provider} error: {e} — trying fallback", flush=True)
        # Try fallback providers
        for fallback in ['anthropic', 'openai', 'groq', 'google']:
            if fallback == provider:
                continue
            key_map = {
                'anthropic': config.ANTHROPIC_API_KEY, 
                'openai': config.OPENAI_API_KEY, 
                'google': config.GOOGLE_API_KEY,
                'groq': config.GROQ_API_KEY
            }
            if key_map.get(fallback):
                try:
                    print(f"DEBUG: Trying fallback provider: {fallback}", flush=True)
                    if fallback == 'anthropic':
                        result = _generate_anthropic(system_prompt, prompt, language)
                    elif fallback == 'openai':
                        result = _generate_openai(system_prompt, prompt, language)
                    elif fallback == 'google':
                        result = _generate_google(system_prompt, prompt, language)
                    elif fallback == 'groq':
                        result = _generate_groq(system_prompt, prompt, language)
                    break
                except Exception as e2:
                    print(f"DEBUG: Fallback {fallback} also failed: {e2}", flush=True)
    
    if result:
        # Clean AI response to remove any instruction artifacts
        result = clean_ai_response(result)
        
        # ── Pillar 3: Track approximate token usage ──
        # Rough estimate: ~4 chars per token for English, ~2 for Arabic
        approx_tokens = len(prompt) // 4 + len(result) // 4
        if user_id:
            increment_token_usage(user_id, approx_tokens)
            # ── Audit log for AI content generation ──
            try:
                log_action(user_id, 'generate_ai_content', {
                    'task_type': task_type,
                    'content_type': content_type or 'general',
                    'provider': provider,
                    'language': language,
                    'domain': detected_domain or 'unknown',
                    'tokens_approx': approx_tokens
                })
            except Exception:
                pass  # Non-fatal: don't break generation if audit fails
        
        return result
    
    print("DEBUG: All AI providers failed, using simulation", flush=True)
    return generate_simulation_content(prompt, language, content_type=content_type)

def _generate_openai(system_prompt, prompt, language='en'):
    """Generate using OpenAI API."""
    import openai
    client = openai.OpenAI(api_key=config.OPENAI_API_KEY)
    model = config.AI_MODEL if config.AI_MODEL and 'gpt' in config.AI_MODEL.lower() else 'gpt-4-turbo'
    
    print(f"DEBUG: Calling OpenAI ({model})...", flush=True)
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ],
        max_tokens=8000,
        temperature=0.7
    )
    result = response.choices[0].message.content
    print(f"DEBUG: OpenAI success, length: {len(result)}", flush=True)
    return result

def _generate_anthropic(system_prompt, prompt, language='en'):
    """Generate using Anthropic Claude API."""
    import anthropic
    client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
    model = config.AI_MODEL if config.AI_MODEL and 'claude' in config.AI_MODEL.lower() else 'claude-sonnet-4-20250514'
    
    print(f"DEBUG: Calling Anthropic ({model})...", flush=True)
    response = client.messages.create(
        model=model,
        max_tokens=8000,
        system=system_prompt,
        messages=[{"role": "user", "content": prompt}]
    )
    result = response.content[0].text
    print(f"DEBUG: Anthropic success, length: {len(result)}", flush=True)
    return result

def _generate_google(system_prompt, prompt, language='en'):
    """Generate using Google Gemini API."""
    import google.generativeai as genai
    genai.configure(api_key=config.GOOGLE_API_KEY)
    model_name = config.AI_MODEL if config.AI_MODEL and 'gemini' in config.AI_MODEL.lower() else 'gemini-2.0-flash'
    
    print(f"DEBUG: Calling Google Gemini ({model_name})...", flush=True)
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(f"{system_prompt}\n\n{prompt}")
    result = response.text
    print(f"DEBUG: Google Gemini success, length: {len(result)}", flush=True)
    return result

def _generate_groq(system_prompt, prompt, language='en'):
    """Generate using Groq API (Llama 3.1, Mistral, etc.)."""
    from openai import OpenAI
    
    # Groq uses OpenAI-compatible API
    client = OpenAI(
        api_key=config.GROQ_API_KEY,
        base_url="https://api.groq.com/openai/v1"
    )
    
    # Available models: llama-3.1-70b-versatile, llama-3.1-8b-instant, mixtral-8x7b-32768
    model = config.AI_MODEL if config.AI_MODEL and ('llama' in config.AI_MODEL.lower() or 'mixtral' in config.AI_MODEL.lower()) else 'llama-3.1-70b-versatile'
    
    print(f"DEBUG: Calling Groq ({model})...", flush=True)
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ],
        max_tokens=8000,
        temperature=0.7
    )
    result = response.choices[0].message.content
    print(f"DEBUG: Groq success, length: {len(result)}", flush=True)
    return result


def generate_simulation_content(prompt, language='en', content_type=None):
    """Minimal fallback when AI service is unavailable.
    Returns a clear message directing users to try again when AI is restored."""
    
    type_labels = {
        'strategy': ('Cybersecurity Strategy', 'الاستراتيجية السيبرانية'),
        'policy': ('Policy Document', 'وثيقة السياسة'),
        'audit': ('Audit Report', 'تقرير التدقيق'),
        'risk': ('Risk Analysis', 'تحليل المخاطر'),
        'review': ('Policy Review', 'مراجعة السياسة'),
        'modify_policy': ('Policy Modification', 'تعديل السياسة'),
        'gap_remediation': ('Gap Remediation Plan', 'خطة معالجة الفجوات'),
        'risk_appetite': ('Risk Appetite Statement', 'بيان شهية المخاطر'),
        'chat': ('Document Chat', 'محادثة الوثيقة'),
    }
    
    label = type_labels.get(content_type, ('Document', 'الوثيقة'))
    
    if language == 'ar':
        return f"""## ⚠️ خدمة الذكاء الاصطناعي غير متاحة حالياً

عذراً، لا يمكن إنشاء **{label[1]}** في الوقت الحالي لأن خدمة الذكاء الاصطناعي غير متصلة.

### ما يمكنك فعله:
- **إعادة المحاولة** بعد بضع دقائق
- **التحقق من إعدادات API** في لوحة الإدارة
- **التواصل مع مسؤول النظام** إذا استمرت المشكلة

> 💡 هذا التطبيق يعتمد على الذكاء الاصطناعي لإنشاء محتوى احترافي ومخصص لمتطلباتك. المحتوى الذي يُنشئه الذكاء الاصطناعي يكون مبنياً على تحليل فعلي لمدخلاتك ومتوافقاً مع الإطار التنظيمي المحدد."""
    else:
        return f"""## ⚠️ AI Service Currently Unavailable

Sorry, the **{label[0]}** cannot be generated right now because the AI service is not connected.

### What you can do:
- **Try again** in a few minutes
- **Check API settings** in the admin panel
- **Contact your system administrator** if the issue persists

> 💡 This application uses AI to generate professional, tailored content based on your specific inputs and regulatory framework requirements. AI-generated content is built from actual analysis of your data, not from templates."""

# ROUTES - LANDING PAGE & AUTHENTICATION
# ============================================================================

@app.route('/')
def index():
    """Landing page for non-logged-in users, dashboard for logged-in users."""
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    lang = request.args.get('lang', 'en')
    txt = get_text(lang)
    return render_template('landing.html', txt=txt, lang=lang, config=config, is_rtl=(lang == 'ar'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page with rate limiting."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    session['lang'] = lang
    txt = get_text(lang)
    
    if request.method == 'POST':
        # Rate limiting - 5 attempts per minute per IP
        client_ip = request.environ.get('HTTP_X_FORWARDED_FOR', request.remote_addr)
        if not check_rate_limit(f'login_{client_ip}', max_requests=5, window_seconds=60):
            flash('Too many login attempts. Please wait a minute.', 'error')
            return render_template('login.html', txt=txt, lang=lang, config=config, is_rtl=(lang == 'ar'))
        
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        
        # Input validation
        if not validate_username(username):
            flash('Invalid username format', 'error')
            return render_template('login.html', txt=txt, lang=lang, config=config, is_rtl=(lang == 'ar'))
        
        conn = get_db()
        user = conn.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
        
        if user and verify_password(password, user['password_hash']):
            # Check if user is active
            if user['is_active'] == 0:
                flash('Account is deactivated. Contact admin.', 'error')
                return render_template('login.html', txt=txt, lang=lang, config=config, is_rtl=(lang == 'ar'))
            
            # Update last login
            conn.execute('UPDATE users SET last_login = ? WHERE id = ?', (datetime.now(), user['id']))
            conn.commit()
            
            # Regenerate session to prevent session fixation
            session.clear()
            session['lang'] = lang
            session.permanent = True
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['role'] = user['role'] if user['role'] else 'user'
            generate_csrf_token()  # Generate new CSRF token
            
            # Redirect admin to admin panel
            if session['role'] == 'admin':
                return redirect(url_for('admin_dashboard'))
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password', 'error')
    
    return render_template('login.html', 
                          txt=txt, 
                          lang=lang, 
                          config=config,
                          is_rtl=(lang == 'ar'))

@app.route('/register', methods=['POST'])
def register():
    """Register new user with limit check and rate limiting."""
    lang = session.get('lang', 'en')
    reg_url = url_for('login', lang=lang) + '#register'
    
    # Rate limiting - 3 registration attempts per minute per IP
    client_ip = request.environ.get('HTTP_X_FORWARDED_FOR', request.remote_addr)
    if not check_rate_limit(f'register_{client_ip}', max_requests=3, window_seconds=60):
        flash('Too many registration attempts. Please wait.', 'error')
        return redirect(reg_url)
    
    username = request.form.get('username', '').strip()
    email = request.form.get('email', '').strip().lower()
    password = request.form.get('password', '')
    
    # Input validation using helper functions
    if not validate_username(username):
        flash('Username: 3-50 characters, letters/numbers/underscore only', 'error')
        return redirect(reg_url)
    
    if not validate_email(email):
        flash('Invalid email format', 'error')
        return redirect(reg_url)
    
    valid, msg = validate_password(password)
    if not valid:
        flash(msg, 'error')
        return redirect(reg_url)
    
    conn = get_db()
    
    # Check user limit
    user_count = conn.execute('SELECT COUNT(*) FROM users').fetchone()[0]
    if user_count >= MAX_USERS:
        flash('Registration limit reached (2000 users). Contact admin.', 'error')
        return redirect(reg_url)
    
    try:
        conn.execute('INSERT INTO users (username, email, password_hash, role) VALUES (?, ?, ?, ?)',
                    (username, email if email else None, hash_password(password), 'user'))
        conn.commit()
        flash('Account created successfully! Please login.', 'success')
    except sqlite3.IntegrityError:
        flash('Username or email already exists', 'error')
        return redirect(reg_url)
    
    return redirect(url_for('login', lang=lang))

@app.route('/logout')
def logout():
    """Logout user."""
    lang = session.get('lang', 'en')
    session.clear()
    session['lang'] = lang
    return redirect(url_for('login', lang=lang))

# ============================================================================
# ROUTES - PROFILE
# ============================================================================

@app.route('/profile')
@login_required
def profile_page():
    """User profile page."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    session['lang'] = lang
    txt = get_text(lang)
    
    conn = get_db()
    user = conn.execute('SELECT * FROM users WHERE id = ?', (session['user_id'],)).fetchone()
    
    # Get per-domain usage stats
    domain_stats = {}
    en_domains = TRANSLATIONS['en']['domains']
    for domain in en_domains:
        s_count = conn.execute('SELECT COUNT(*) FROM strategies WHERE user_id = ? AND domain = ?', (session['user_id'], domain)).fetchone()[0]
        p_count = conn.execute('SELECT COUNT(*) FROM policies WHERE user_id = ? AND domain = ?', (session['user_id'], domain)).fetchone()[0]
        a_count = conn.execute('SELECT COUNT(*) FROM audits WHERE user_id = ? AND domain = ?', (session['user_id'], domain)).fetchone()[0]
        r_count = conn.execute('SELECT COUNT(*) FROM risks WHERE user_id = ? AND domain = ?', (session['user_id'], domain)).fetchone()[0]
        total = s_count + p_count + a_count + r_count
        if total > 0:
            domain_stats[domain] = {
                'strategies': s_count, 'policies': p_count,
                'audits': a_count, 'risks': r_count, 'total': total,
                'strategy_limit': 1, 'policy_limit': 2, 'audit_limit': 2, 'risk_limit': 2
            }
    
    total_docs = conn.execute('''
        SELECT 
            (SELECT COUNT(*) FROM strategies WHERE user_id = ?) +
            (SELECT COUNT(*) FROM policies WHERE user_id = ?) +
            (SELECT COUNT(*) FROM audits WHERE user_id = ?) +
            (SELECT COUNT(*) FROM risks WHERE user_id = ?)
    ''', (session['user_id'], session['user_id'], session['user_id'], session['user_id'])).fetchone()[0]
    
    # Get user's AI preferences
    ai_pref_generate = 'auto'
    ai_pref_review = 'auto'
    try:
        ai_pref_generate = user['ai_provider_generate'] or 'auto'
        ai_pref_review = user['ai_provider_review'] or 'auto'
    except:
        pass
    
    # Map backend → frontend provider values for dropdown selection
    backend_to_frontend = {
        'anthropic': 'claude', 'openai': 'gpt', 'google': 'gemini',
        'groq': 'groq', 'auto': 'auto',
        # Already frontend values (backward compat)
        'claude': 'claude', 'gpt': 'gpt', 'gemini': 'gemini',
    }
    ai_pref_generate_ui = backend_to_frontend.get(ai_pref_generate, 'claude')
    ai_pref_review_ui = backend_to_frontend.get(ai_pref_review, 'gpt')
    
    # ── Token quota data for progress bar ──
    token_usage = 0
    token_limit = 500000
    try:
        token_usage = user['token_usage'] or 0
        token_limit = user['token_limit'] or 500000
    except (KeyError, TypeError):
        pass
    
    # ── User Command Center data ──
    # Tasks
    user_tasks = conn.execute('''
        SELECT * FROM project_tasks WHERE user_id = ? ORDER BY created_at DESC LIMIT 50
    ''', (session['user_id'],)).fetchall()
    
    task_stats = {
        'total': len(user_tasks),
        'todo': sum(1 for t in user_tasks if t['status'] == 'todo'),
        'in_progress': sum(1 for t in user_tasks if t['status'] == 'in_progress'),
        'done': sum(1 for t in user_tasks if t['status'] == 'done'),
    }
    
    # Risk Register — include computed score for template
    user_risks = conn.execute('''
        SELECT *, name AS risk_name, (likelihood * impact) AS risk_score
        FROM risk_register WHERE user_id = ? ORDER BY (likelihood * impact) DESC, created_at DESC LIMIT 50
    ''', (session['user_id'],)).fetchall()
    
    # KPIs
    user_kpis = []
    try:
        user_kpis = conn.execute('''
            SELECT * FROM kpis WHERE user_id = ? ORDER BY created_at DESC LIMIT 50
        ''', (session['user_id'],)).fetchall()
    except Exception:
        pass  # Table may not exist yet
    
    # Audit logs (recent 50)
    user_audit = conn.execute('''
        SELECT * FROM audit_logs WHERE user_id = ? ORDER BY created_at DESC LIMIT 50
    ''', (session['user_id'],)).fetchall()
    
    
    return render_template('profile.html',
                          txt=txt, lang=lang, config=config,
                          is_rtl=(lang == 'ar'),
                          username=session.get('username'),
                          user=user,
                          domain_stats=domain_stats,
                          total_docs=total_docs,
                          ai_available=check_ai_available(),
                          ai_providers=get_available_providers(),
                          ai_pref_generate=ai_pref_generate_ui,
                          ai_pref_review=ai_pref_review_ui,
                          ai_prefs={'generate': ai_pref_generate_ui, 'review': ai_pref_review_ui},
                          domains=txt['domains'],
                          token_usage=token_usage,
                          token_limit=token_limit,
                          user_tasks=user_tasks,
                          task_stats=task_stats,
                          user_risks=user_risks,
                          user_kpis=user_kpis,
                          user_audit=user_audit)

@app.route('/profile/change-password', methods=['POST'])
@login_required
def change_password():
    """Change user password."""
    lang = session.get('lang', 'en')
    txt = get_text(lang)
    
    current_password = request.form.get('current_password', '')
    new_password = request.form.get('new_password', '')
    confirm_password = request.form.get('confirm_password', '')
    
    if new_password != confirm_password:
        flash(txt['password_mismatch'], 'error')
        return redirect(url_for('profile_page', lang=lang))
    
    valid, msg = validate_password(new_password)
    if not valid:
        flash(msg, 'error')
        return redirect(url_for('profile_page', lang=lang))
    
    conn = get_db()
    user = conn.execute('SELECT * FROM users WHERE id = ?', (session['user_id'],)).fetchone()
    
    if not verify_password(current_password, user['password_hash']):
        flash(txt['wrong_password'], 'error')
        return redirect(url_for('profile_page', lang=lang))
    
    conn.execute('UPDATE users SET password_hash = ? WHERE id = ?',
                (hash_password(new_password), session['user_id']))
    conn.commit()
    
    flash(txt['password_updated'], 'success')
    return redirect(url_for('profile_page', lang=lang))

@app.route('/profile/ai-preferences', methods=['POST'])
@login_required
def save_ai_preferences():
    """Save user's AI provider preferences."""
    lang = session.get('lang', 'en')
    
    # Accept both JSON and form data
    if request.is_json:
        data = request.get_json()
        generate_provider = data.get('ai_provider_generate', 'auto')
        review_provider = data.get('ai_provider_review', 'auto')
    else:
        generate_provider = request.form.get('ai_provider_generate', 'auto')
        review_provider = request.form.get('ai_provider_review', 'auto')
    
    # Map frontend values to backend values
    provider_map = {
        'claude': 'anthropic',
        'gpt': 'openai',
        'gemini': 'google',
        'groq': 'groq',
        'auto': 'auto',
        # Also accept backend values directly
        'anthropic': 'anthropic',
        'openai': 'openai',
        'google': 'google',
    }
    generate_provider = provider_map.get(generate_provider, 'auto')
    review_provider = provider_map.get(review_provider, 'auto')
    
    try:
        conn = get_db()
        conn.execute('''
            UPDATE users 
            SET ai_provider_generate = ?, ai_provider_review = ?
            WHERE id = ?
        ''', (generate_provider, review_provider, session['user_id']))
        conn.commit()
        
        if request.is_json:
            return jsonify({'success': True, 'message': 'AI preferences updated successfully' if lang == 'en' else 'تم تحديث تفضيلات الذكاء الاصطناعي بنجاح'})
        
        flash('AI preferences updated successfully' if lang == 'en' else 'تم تحديث تفضيلات الذكاء الاصطناعي بنجاح', 'success')
        return redirect(url_for('profile_page', lang=lang))
    except Exception as e:
        if request.is_json:
            return jsonify({'success': False, 'error': str(e)}), 500
        flash('Error saving preferences' if lang == 'en' else 'خطأ في حفظ التفضيلات', 'error')
        return redirect(url_for('profile_page', lang=lang))

# ============================================================================
# ROUTES - EXPORT ALL DOCUMENTS
# ============================================================================

@app.route('/api/export-all', methods=['GET'])
@login_required
def export_all_documents():
    """Export all user documents as a ZIP file."""
    import zipfile
    from io import BytesIO
    
    lang = session.get('lang', 'en')
    user_id = session['user_id']
    username = session.get('username', 'user')
    
    conn = get_db()
    
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        doc_count = 0
        
        # Export strategies
        strategies = conn.execute('SELECT * FROM strategies WHERE user_id = ? ORDER BY created_at DESC', (user_id,)).fetchall()
        for s in strategies:
            fname = f"strategies/{s['domain']}_{s['language']}_{s['id']}.md"
            zf.writestr(fname, s['content'] or '')
            doc_count += 1
        
        # Export policies
        policies = conn.execute('SELECT * FROM policies WHERE user_id = ? ORDER BY created_at DESC', (user_id,)).fetchall()
        for p in policies:
            fname = f"policies/{p['domain']}_{p['policy_name']}_{p['language']}_{p['id']}.md"
            zf.writestr(fname, p['content'] or '')
            doc_count += 1
        
        # Export audits
        audits = conn.execute('SELECT * FROM audits WHERE user_id = ? ORDER BY created_at DESC', (user_id,)).fetchall()
        for a in audits:
            fname = f"audits/{a['domain']}_{a['language']}_{a['id']}.md"
            zf.writestr(fname, a['content'] or '')
            doc_count += 1
        
        # Export risks
        risks = conn.execute('SELECT * FROM risks WHERE user_id = ? ORDER BY created_at DESC', (user_id,)).fetchall()
        for r in risks:
            fname = f"risk_analyses/{r['domain']}_{r['language']}_{r['id']}.md"
            zf.writestr(fname, r['analysis'] or '')
            doc_count += 1
        
        # Add summary file
        summary = f"# Mizan GRC - Document Export\n"
        summary += f"**User:** {username}\n"
        summary += f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        summary += f"**Total Documents:** {doc_count}\n\n"
        summary += f"- Strategies: {len(strategies)}\n"
        summary += f"- Policies: {len(policies)}\n"
        summary += f"- Audits: {len(audits)}\n"
        summary += f"- Risk Analyses: {len(risks)}\n"
        zf.writestr('README.md', summary)
    
    
    if doc_count == 0:
        return jsonify({'error': 'No documents to export'}), 404
    
    buffer.seek(0)
    from flask import send_file
    return send_file(buffer, mimetype='application/zip',
                    as_attachment=True,
                    download_name=f'mizan_documents_{username}_{datetime.now().strftime("%Y%m%d")}.zip')

# ============================================================================
# ROUTES - DASHBOARD
# ============================================================================

@app.route('/dashboard')
@login_required
def dashboard():
    """Main dashboard with widgets."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    session['lang'] = lang
    txt = get_text(lang)
    
    conn = get_db()
    user_id = session['user_id']
    
    # Get user stats
    strategies_count = conn.execute('SELECT COUNT(*) FROM strategies WHERE user_id = ?', (user_id,)).fetchone()[0]
    policies_count = conn.execute('SELECT COUNT(*) FROM policies WHERE user_id = ?', (user_id,)).fetchone()[0]
    audits_count = conn.execute('SELECT COUNT(*) FROM audits WHERE user_id = ?', (user_id,)).fetchone()[0]
    risks_count = conn.execute('SELECT COUNT(*) FROM risks WHERE user_id = ?', (user_id,)).fetchone()[0]
    
    # Get recent documents (last 5)
    recent_docs = []
    for doc_type, table, content_field in [
        ('strategy', 'strategies', 'content'), ('policy', 'policies', 'content'),
        ('audit', 'audits', 'content'), ('risk', 'risks', 'analysis')
    ]:
        rows = conn.execute(f'SELECT id, domain, language, created_at FROM {table} WHERE user_id = ? ORDER BY created_at DESC LIMIT 3', (user_id,)).fetchall()
        for row in rows:
            recent_docs.append({
                'type': doc_type, 'id': row['id'], 'domain': row['domain'],
                'language': row['language'], 'created_at': row['created_at']
            })
    
    # Sort by date, take latest 5
    recent_docs.sort(key=lambda x: x['created_at'] or '', reverse=True)
    recent_docs = recent_docs[:5]
    
    # Count active domains
    active_domains = set()
    for table in ['strategies', 'policies', 'audits', 'risks']:
        rows = conn.execute(f'SELECT DISTINCT domain FROM {table} WHERE user_id = ?', (user_id,)).fetchall()
        for row in rows:
            if row['domain']:
                active_domains.add(row['domain'])
    
    # Get risk register data for unified RR on dashboard
    risk_register = []
    try:
        risk_register = [dict(r) for r in conn.execute(
            'SELECT * FROM risk_register WHERE user_id = ? ORDER BY (likelihood * impact) DESC, created_at DESC LIMIT 50',
            (user_id,)
        ).fetchall()]
    except Exception:
        pass
    
    # Get tasks data for unified Tasks Register on dashboard
    dashboard_tasks = []
    task_counts = {}
    try:
        dashboard_tasks = [dict(t) for t in conn.execute(
            'SELECT * FROM project_tasks WHERE user_id = ? ORDER BY CASE priority WHEN "critical" THEN 0 WHEN "high" THEN 1 WHEN "medium" THEN 2 WHEN "low" THEN 3 END, created_at DESC LIMIT 100',
            (user_id,)
        ).fetchall()]
        for t in dashboard_tasks:
            cat = t.get('category', 'implementation')
            task_counts[cat] = task_counts.get(cat, 0) + 1
    except Exception:
        pass
    
    return render_template('dashboard.html',
                          txt=txt, lang=lang, config=config,
                          is_rtl=(lang == 'ar'),
                          username=session.get('username'),
                          ai_available=check_ai_available(),
                          ai_status=get_ai_status(),
                          stats={
                              'strategies': strategies_count,
                              'policies': policies_count,
                              'audits': audits_count,
                              'risks': risks_count,
                              'total': strategies_count + policies_count + audits_count + risks_count
                          },
                          recent_docs=recent_docs,
                          active_domains=len(active_domains),
                          domains=txt['domains'],
                          domain_codes=DOMAIN_CODES,
                          frameworks=DOMAIN_FRAMEWORKS,
                          risk_register=risk_register,
                          dashboard_tasks=dashboard_tasks,
                          task_counts=task_counts)

@app.route('/domain/<domain_name>')
@login_required
def domain_page(domain_name):
    """Domain-specific page."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    session['lang'] = lang
    txt = get_text(lang)
    
    domain_code = DOMAIN_CODES.get(domain_name, 'global')
    
    # Get frameworks — prefer DB (grc_frameworks) with fallback to hardcoded dict
    domain_frameworks = {}
    frameworks_flat = []
    try:
        conn_fw = get_db()
        db_fw_rows = conn_fw.execute(
            'SELECT region, name FROM grc_frameworks WHERE domain = ? AND is_active = 1 ORDER BY sort_order',
            (domain_code,)
        ).fetchall()
        if db_fw_rows:
            for row in db_fw_rows:
                region = row['region']
                if region not in domain_frameworks:
                    domain_frameworks[region] = []
                domain_frameworks[region].append(row['name'])
                frameworks_flat.append(row['name'])
        else:
            # Fallback to hardcoded dict if DB is empty for this domain
            domain_frameworks = DOMAIN_FRAMEWORKS.get(domain_code, {})
            frameworks_flat = DOMAIN_FRAMEWORKS_FLAT.get(domain_code, [])
    except Exception:
        domain_frameworks = DOMAIN_FRAMEWORKS.get(domain_code, {})
        frameworks_flat = DOMAIN_FRAMEWORKS_FLAT.get(domain_code, [])
    
    # Get domain-specific technologies
    lang_key = 'ar' if lang == 'ar' else 'en'
    technologies = DOMAIN_TECHNOLOGIES.get(domain_code, {}).get(lang_key, {})
    
    # Get risk categories with scenarios
    risk_data = RISK_CATEGORIES.get(domain_code, {}).get(lang_key, {})
    
    # Get user's remaining usage FOR THIS DOMAIN
    usage_info = get_remaining_usage(session['user_id'], domain_name)
    
    # Get domain-specific awareness modules
    domain_awareness = AWARENESS_MODULES.get(domain_code, {}).get('ar' if lang == 'ar' else 'en', [])
    
    # Get GRC Professional modules - filter by domain relevance
    all_grc_modules = AWARENESS_MODULES.get('grc_professional', {}).get('ar' if lang == 'ar' else 'en', [])
    
    # Domain-specific GRC module mapping
    grc_domain_relevance = {
        'cyber': ['grc_nca_frameworks', 'grc_iso27001', 'grc_nist_csf', 'grc_sama_framework', 'grc_eu_regulations'],
        'data': ['grc_pdpl', 'grc_gdpr', 'grc_nca_frameworks', 'grc_eu_regulations'],
        'ai': ['grc_eu_regulations', 'grc_nist_csf', 'grc_iso27001'],
        'dt': ['grc_nist_csf', 'grc_iso27001', 'grc_coso_erm', 'grc_eu_regulations'],
        'global': ['grc_iso27001', 'grc_nist_csf', 'grc_coso_erm', 'grc_gdpr', 'grc_eu_regulations'],
        'erm': ['grc_coso_erm', 'grc_nist_csf', 'grc_iso27001', 'grc_sama_framework']
    }
    
    # Filter GRC modules by domain
    relevant_module_ids = grc_domain_relevance.get(domain_code, [])
    grc_professional_modules = [m for m in all_grc_modules if m.get('id') in relevant_module_ids]
    
    return render_template('domain.html',
                          txt=txt,
                          lang=lang,
                          config=config,
                          is_rtl=(lang == 'ar'),
                          username=session.get('username'),
                          ai_available=check_ai_available(),
                          domain_name=domain_name,
                          domain_code=domain_code,
                          frameworks=domain_frameworks,
                          frameworks_flat=frameworks_flat,
                          technologies=technologies,
                          risk_categories=risk_data,
                          usage_info=usage_info,
                          usage_limits=USAGE_LIMITS,
                          awareness_modules=domain_awareness,
                          grc_professional_modules=grc_professional_modules)

# ============================================================================
# ANALYTICS & INSIGHTS
# ============================================================================

def calculate_compliance_score(user_id):
    """Calculate compliance score based on user's documents."""
    conn = get_db()
    
    # Get counts
    strategies = conn.execute('SELECT COUNT(*) FROM strategies WHERE user_id = ?', (user_id,)).fetchone()[0]
    policies = conn.execute('SELECT COUNT(*) FROM policies WHERE user_id = ?', (user_id,)).fetchone()[0]
    audits = conn.execute('SELECT COUNT(*) FROM audits WHERE user_id = ?', (user_id,)).fetchone()[0]
    risks = conn.execute('SELECT COUNT(*) FROM risks WHERE user_id = ?', (user_id,)).fetchone()[0]
    
    # Get unique domains covered
    domains_covered = set()
    for table in ['strategies', 'policies', 'audits', 'risks']:
        rows = conn.execute(f'SELECT DISTINCT domain FROM {table} WHERE user_id = ?', (user_id,)).fetchall()
        for row in rows:
            if row['domain']:
                domains_covered.add(row['domain'])
    
    
    # Calculate score (weighted)
    # Max expected: 6 strategies, 12 policies, 6 audits, 12 risks, 6 domains
    strategy_score = min(strategies / 6, 1) * 20  # 20%
    policy_score = min(policies / 10, 1) * 25     # 25%
    audit_score = min(audits / 6, 1) * 25         # 25%
    risk_score = min(risks / 10, 1) * 20          # 20%
    domain_score = min(len(domains_covered) / 5, 1) * 10  # 10%
    
    total_score = strategy_score + policy_score + audit_score + risk_score + domain_score
    
    return {
        'score': round(total_score, 1),
        'strategies': strategies,
        'policies': policies,
        'audits': audits,
        'risks': risks,
        'domains_covered': len(domains_covered)
    }

def calculate_domain_compliance(user_id):
    """Calculate compliance scores per domain for radar chart."""
    conn = get_db()
    
    domain_codes = {
        'Cyber Security': 'cyber', 'الأمن السيبراني': 'cyber',
        'Data Management': 'data', 'إدارة البيانات': 'data',
        'Artificial Intelligence': 'ai', 'الذكاء الاصطناعي': 'ai',
        'Digital Transformation': 'dt', 'التحول الرقمي': 'dt',
        'Global Standards': 'global', 'المعايير العالمية': 'global',
        'Enterprise Risk Management': 'erm', 'إدارة المخاطر المؤسسية': 'erm'
    }
    
    domain_scores = {}
    for domain_name, code in domain_codes.items():
        if code in domain_scores:
            continue
        
        # Count documents per domain
        strategies = conn.execute(
            'SELECT COUNT(*) FROM strategies WHERE user_id=? AND domain IN (?,?)',
            (user_id, domain_name, [k for k,v in domain_codes.items() if v==code][0])
        ).fetchone()[0]
        policies = conn.execute(
            'SELECT COUNT(*) FROM policies WHERE user_id=? AND domain LIKE ?',
            (user_id, f'%{code[:3]}%' if len(code) < 4 else f'%{domain_name[:4]}%')
        ).fetchone()[0]
        audits = conn.execute(
            'SELECT COUNT(*) FROM audits WHERE user_id=? AND domain LIKE ?',
            (user_id, f'%{domain_name[:6]}%')
        ).fetchone()[0]
        risks = conn.execute(
            'SELECT COUNT(*) FROM risks WHERE user_id=? AND domain LIKE ?',
            (user_id, f'%{domain_name[:6]}%')
        ).fetchone()[0]
        
        # Calculate domain score (0-100)
        score = min(100, (
            min(strategies, 2) * 15 +  # 2 strategies = 30 points
            min(policies, 3) * 10 +    # 3 policies = 30 points
            min(audits, 2) * 10 +      # 2 audits = 20 points
            min(risks, 2) * 10         # 2 risks = 20 points
        ))
        
        domain_scores[code] = {
            'score': score,
            'strategies': strategies,
            'policies': policies,
            'audits': audits,
            'risks': risks
        }
    
    return domain_scores

def get_framework_coverage(user_id):
    """Calculate which frameworks are covered by user's policies."""
    conn = get_db()
    
    # Get all policies with their frameworks
    policies = conn.execute(
        'SELECT framework, domain FROM policies WHERE user_id = ?',
        (user_id,)
    ).fetchall()
    
    # Get all audits with frameworks
    audits = conn.execute(
        'SELECT framework, domain FROM audits WHERE user_id = ?',
        (user_id,)
    ).fetchall()
    
    
    # Framework categories
    all_frameworks = {
        'cyber': ['NCA ECC', 'NCA CSCC', 'NCA DCC', 'SAMA CSF', 'ISO 27001', 'NIST CSF', 'CIS Controls', 'PCI DSS'],
        'data': ['PDPL', 'NDMO', 'ISO 27701', 'GDPR', 'Data Governance Framework'],
        'ai': ['SDAIA AI Ethics', 'UNESCO AI', 'EU AI Act', 'NIST AI RMF'],
        'global': ['ISO 27001', 'ISO 22301', 'ISO 31000', 'COBIT', 'ITIL'],
        'erm': ['ISO 31000', 'COSO ERM', 'SAMA Regulations', 'Basel III']
    }
    
    # Count covered frameworks
    covered = set()
    for p in policies:
        if p['framework']:
            for fw in p['framework'].split(','):
                covered.add(fw.strip())
    for a in audits:
        if a['framework']:
            for fw in a['framework'].split(','):
                covered.add(fw.strip())
    
    # Calculate coverage per category
    coverage = {}
    for cat, frameworks in all_frameworks.items():
        cat_covered = sum(1 for fw in frameworks if any(fw.lower() in c.lower() for c in covered))
        coverage[cat] = {
            'covered': cat_covered,
            'total': len(frameworks),
            'percentage': round(cat_covered / len(frameworks) * 100) if frameworks else 0,
            'frameworks': frameworks,
            'covered_list': [fw for fw in frameworks if any(fw.lower() in c.lower() for c in covered)]
        }
    
    return coverage

def get_compliance_trend(user_id, days=30):
    """Get compliance score trend over time."""
    conn = get_db()
    
    # Record current score if not recorded today
    today = conn.execute("SELECT DATE('now')").fetchone()[0]
    last_record = conn.execute(
        "SELECT DATE(recorded_at) as d FROM compliance_history WHERE user_id=? ORDER BY recorded_at DESC LIMIT 1",
        (user_id,)
    ).fetchone()
    
    if not last_record or last_record['d'] != today:
        # Record today's score
        current = calculate_compliance_score(user_id)
        maturity = calculate_maturity_levels(user_id)
        conn.execute('''
            INSERT INTO compliance_history 
            (user_id, score, maturity_avg, strategies, policies, audits, risks, domains_covered)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (user_id, current['score'], maturity['average'], current['strategies'],
              current['policies'], current['audits'], current['risks'], current['domains_covered']))
        conn.commit()
    
    # Get history
    history = conn.execute('''
        SELECT DATE(recorded_at) as date, score, maturity_avg, strategies, policies, audits, risks
        FROM compliance_history 
        WHERE user_id = ? AND recorded_at >= DATE('now', ?)
        ORDER BY recorded_at ASC
    ''', (user_id, f'-{days} days')).fetchall()
    
    
    return [dict(h) for h in history]

def generate_gap_recommendations(user_id, lang='en'):
    """Generate AI-powered gap analysis recommendations."""
    compliance = calculate_compliance_score(user_id)
    maturity = calculate_maturity_levels(user_id)
    domain_scores = calculate_domain_compliance(user_id)
    frameworks = get_framework_coverage(user_id)
    
    recommendations = []
    
    # Check overall compliance
    if compliance['score'] < 50:
        if lang == 'ar':
            recommendations.append({
                'priority': 'critical',
                'area': 'الامتثال العام',
                'issue': f"درجة الامتثال منخفضة ({compliance['score']}%)",
                'action': 'ابدأ بإنشاء استراتيجية شاملة لكل مجال، ثم طور السياسات الأساسية',
                'impact': 'تحسين الامتثال بنسبة 20-30%'
            })
        else:
            recommendations.append({
                'priority': 'critical',
                'area': 'Overall Compliance',
                'issue': f"Low compliance score ({compliance['score']}%)",
                'action': 'Start by creating a comprehensive strategy for each domain, then develop core policies',
                'impact': 'Improve compliance by 20-30%'
            })
    
    # Check strategies
    if compliance['strategies'] < 3:
        if lang == 'ar':
            recommendations.append({
                'priority': 'high',
                'area': 'التخطيط الاستراتيجي',
                'issue': f"عدد الاستراتيجيات غير كافٍ ({compliance['strategies']} من 6)",
                'action': 'أنشئ استراتيجيات للمجالات الحرجة: الأمن السيبراني، إدارة البيانات، إدارة المخاطر',
                'impact': 'تغطية استراتيجية شاملة'
            })
        else:
            recommendations.append({
                'priority': 'high',
                'area': 'Strategic Planning',
                'issue': f"Insufficient strategies ({compliance['strategies']} of 6)",
                'action': 'Create strategies for critical domains: Cyber Security, Data Management, ERM',
                'impact': 'Comprehensive strategic coverage'
            })
    
    # Check policies
    if compliance['policies'] < 6:
        if lang == 'ar':
            recommendations.append({
                'priority': 'high',
                'area': 'السياسات',
                'issue': f"نقص في السياسات ({compliance['policies']} من 10 موصى به)",
                'action': 'طور سياسات أمن المعلومات، حماية البيانات، الاستجابة للحوادث',
                'impact': 'تحسين الامتثال التنظيمي'
            })
        else:
            recommendations.append({
                'priority': 'high',
                'area': 'Policies',
                'issue': f"Policy gap ({compliance['policies']} of 10 recommended)",
                'action': 'Develop Information Security, Data Protection, and Incident Response policies',
                'impact': 'Improved regulatory compliance'
            })
    
    # Check audits
    if compliance['audits'] < 2:
        if lang == 'ar':
            recommendations.append({
                'priority': 'medium',
                'area': 'التدقيق',
                'issue': 'لا توجد تقارير تدقيق كافية',
                'action': 'أجرِ تدقيق على إطار NCA ECC أو ISO 27001',
                'impact': 'تحديد فجوات الامتثال الفعلية'
            })
        else:
            recommendations.append({
                'priority': 'medium',
                'area': 'Auditing',
                'issue': 'Insufficient audit coverage',
                'action': 'Conduct an audit against NCA ECC or ISO 27001 framework',
                'impact': 'Identify actual compliance gaps'
            })
    
    # Check domain coverage
    weak_domains = [code for code, data in domain_scores.items() if data['score'] < 30]
    if weak_domains:
        domain_names = {'cyber': 'Cyber Security', 'data': 'Data Management', 'ai': 'AI Governance',
                       'dt': 'Digital Transformation', 'global': 'Global Standards', 'erm': 'Enterprise Risk'}
        if lang == 'ar':
            domain_names_ar = {'cyber': 'الأمن السيبراني', 'data': 'إدارة البيانات', 'ai': 'حوكمة الذكاء الاصطناعي',
                              'dt': 'التحول الرقمي', 'global': 'المعايير العالمية', 'erm': 'إدارة المخاطر المؤسسية'}
            recommendations.append({
                'priority': 'medium',
                'area': 'تغطية المجالات',
                'issue': f"مجالات ضعيفة: {', '.join(domain_names_ar.get(d, d) for d in weak_domains[:3])}",
                'action': 'ركز على تطوير محتوى لهذه المجالات',
                'impact': 'تغطية شاملة لجميع مجالات الحوكمة'
            })
        else:
            recommendations.append({
                'priority': 'medium',
                'area': 'Domain Coverage',
                'issue': f"Weak domains: {', '.join(domain_names.get(d, d) for d in weak_domains[:3])}",
                'action': 'Focus on developing content for these domains',
                'impact': 'Comprehensive GRC coverage'
            })
    
    # Check framework coverage
    for cat, data in frameworks.items():
        if data['percentage'] < 25 and cat in ['cyber', 'global']:
            if lang == 'ar':
                recommendations.append({
                    'priority': 'high',
                    'area': 'تغطية الأطر التنظيمية',
                    'issue': f"تغطية منخفضة لأطر {cat} ({data['percentage']}%)",
                    'action': f"طور سياسات وتدقيق لـ: {', '.join(data['frameworks'][:3])}",
                    'impact': 'تحسين الامتثال للأطر التنظيمية'
                })
            else:
                recommendations.append({
                    'priority': 'high',
                    'area': 'Framework Coverage',
                    'issue': f"Low {cat} framework coverage ({data['percentage']}%)",
                    'action': f"Develop policies and audits for: {', '.join(data['frameworks'][:3])}",
                    'impact': 'Improved regulatory framework alignment'
                })
    
    # Sort by priority
    priority_order = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}
    recommendations.sort(key=lambda x: priority_order.get(x['priority'], 4))
    
    return recommendations[:6]  # Return top 6 recommendations

def calculate_maturity_levels(user_id):
    """Calculate maturity levels for radar chart."""
    conn = get_db()
    
    # Get counts by type
    strategies = conn.execute('SELECT COUNT(*) FROM strategies WHERE user_id = ?', (user_id,)).fetchone()[0]
    policies = conn.execute('SELECT COUNT(*) FROM policies WHERE user_id = ?', (user_id,)).fetchone()[0]
    audits = conn.execute('SELECT COUNT(*) FROM audits WHERE user_id = ?', (user_id,)).fetchone()[0]
    risks = conn.execute('SELECT COUNT(*) FROM risks WHERE user_id = ?', (user_id,)).fetchone()[0]
    
    # Get unique domains
    domains = set()
    for table in ['strategies', 'policies', 'audits', 'risks']:
        rows = conn.execute(f'SELECT DISTINCT domain FROM {table} WHERE user_id = ?', (user_id,)).fetchall()
        for row in rows:
            if row['domain']:
                domains.add(row['domain'])
    
    
    # Calculate maturity (1-5 scale)
    # Governance: based on strategies
    governance = min(1 + (strategies * 0.8), 5)
    
    # Risk Management: based on risk assessments
    risk_mgmt = min(1 + (risks * 0.4), 5)
    
    # Compliance: based on policies and audits
    compliance = min(1 + ((policies + audits) * 0.3), 5)
    
    # Technology: based on domain coverage (Cyber, AI, DT)
    tech_domains = len([d for d in domains if any(x in d for x in ['Cyber', 'AI', 'Digital', 'سيبراني', 'ذكاء', 'رقمي'])])
    technology = min(1 + (tech_domains * 1.3), 5)
    
    # Process: overall completeness
    total_docs = strategies + policies + audits + risks
    process = min(1 + (total_docs * 0.15), 5)
    
    return {
        'governance': round(governance, 1),
        'risk_mgmt': round(risk_mgmt, 1),
        'compliance': round(compliance, 1),
        'technology': round(technology, 1),
        'process': round(process, 1),
        'average': round((governance + risk_mgmt + compliance + technology + process) / 5, 1)
    }

def get_risk_heatmap_data(user_id):
    """Get risk distribution for heatmap."""
    conn = get_db()
    risks = conn.execute('SELECT risk_level, threat FROM risks WHERE user_id = ?', (user_id,)).fetchall()
    
    # Initialize heatmap matrix (4x4: Low, Medium, High, Critical)
    # For simplicity, we'll estimate likelihood from threat description
    heatmap = {
        'low': {'low': 0, 'medium': 0, 'high': 0, 'critical': 0},
        'medium': {'low': 0, 'medium': 0, 'high': 0, 'critical': 0},
        'high': {'low': 0, 'medium': 0, 'high': 0, 'critical': 0},
        'critical': {'low': 0, 'medium': 0, 'high': 0, 'critical': 0}
    }
    
    for risk in risks:
        impact = (risk['risk_level'] or 'medium').lower()
        if impact not in heatmap:
            impact = 'medium'
        
        # Estimate likelihood based on common keywords
        threat = (risk['threat'] or '').lower()
        if any(x in threat for x in ['common', 'frequent', 'daily', 'شائع', 'متكرر']):
            likelihood = 'high'
        elif any(x in threat for x in ['rare', 'unlikely', 'نادر']):
            likelihood = 'low'
        elif any(x in threat for x in ['critical', 'severe', 'حرج']):
            likelihood = 'critical'
        else:
            likelihood = 'medium'
        
        heatmap[likelihood][impact] += 1
    
    return heatmap

def get_benchmark_comparison(user_id, sector):
    """Compare user's metrics with industry benchmark."""
    conn = get_db()
    
    # Get user's metrics
    user_metrics = calculate_compliance_score(user_id)
    user_maturity = calculate_maturity_levels(user_id)
    
    # Get benchmark for sector
    benchmark = conn.execute('SELECT * FROM benchmarks WHERE sector = ?', (sector,)).fetchone()
    
    if not benchmark:
        return None
    
    comparison = {
        'sector': sector,
        'source': benchmark['source'],
        'source_year': benchmark['source_year'],
        'metrics': {
            'compliance_score': {
                'user': user_metrics['score'],
                'benchmark': benchmark['compliance_score_avg'],
                'gap': round(user_metrics['score'] - benchmark['compliance_score_avg'], 1)
            },
            'maturity_level': {
                'user': user_maturity['average'],
                'benchmark': benchmark['maturity_level_avg'],
                'gap': round(user_maturity['average'] - benchmark['maturity_level_avg'], 1)
            },
            'policy_count': {
                'user': user_metrics['policies'],
                'benchmark': benchmark['policy_count_avg'],
                'gap': user_metrics['policies'] - benchmark['policy_count_avg']
            },
            'audit_count': {
                'user': user_metrics['audits'],
                'benchmark': benchmark['audit_count_avg'],
                'gap': user_metrics['audits'] - benchmark['audit_count_avg']
            },
            'risk_assessments': {
                'user': user_metrics['risks'],
                'benchmark': benchmark['risk_assessment_avg'],
                'gap': user_metrics['risks'] - benchmark['risk_assessment_avg']
            }
        }
    }
    
    return comparison

@app.route('/analytics')
@login_required
def analytics_page():
    """Analytics and Insights dashboard."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    session['lang'] = lang
    txt = get_text(lang)
    
    user_id = session['user_id']
    sector = request.args.get('sector', 'Government')
    
    # Calculate all analytics
    compliance = calculate_compliance_score(user_id)
    maturity = calculate_maturity_levels(user_id)
    heatmap = get_risk_heatmap_data(user_id)
    benchmark = get_benchmark_comparison(user_id, sector)
    
    # Enhanced analytics
    domain_compliance = calculate_domain_compliance(user_id)
    framework_coverage = get_framework_coverage(user_id)
    compliance_trend = get_compliance_trend(user_id, days=30)
    recommendations = generate_gap_recommendations(user_id, lang)
    
    # Get available sectors for dropdown
    conn = get_db()
    sectors = conn.execute('SELECT sector, sector_ar FROM benchmarks ORDER BY sector').fetchall()
    
    has_data = compliance['score'] > 0
    
    return render_template('analytics.html',
                          txt=txt,
                          lang=lang,
                          is_rtl=(lang == 'ar'),
                          username=session.get('username'),
                          compliance=compliance,
                          maturity=maturity,
                          heatmap=heatmap,
                          benchmark=benchmark,
                          sectors=sectors,
                          selected_sector=sector,
                          has_data=has_data,
                          domain_compliance=domain_compliance,
                          framework_coverage=framework_coverage,
                          compliance_trend=compliance_trend,
                          recommendations=recommendations)

@app.route('/api/analytics/benchmark/<sector>')
@login_required
def api_get_benchmark(sector):
    """Get benchmark comparison for a sector."""
    try:
        user_id = session['user_id']
        comparison = get_benchmark_comparison(user_id, sector)
        if comparison:
            return jsonify({'success': True, 'data': comparison})
        return jsonify({'success': False, 'error': 'Sector not found'}), 404
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/analytics/domain-scores')
@login_required
def api_domain_scores():
    """Get compliance scores per domain."""
    try:
        user_id = session['user_id']
        scores = calculate_domain_compliance(user_id)
        return jsonify({'success': True, 'data': scores})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/analytics/framework-coverage')
@login_required
def api_framework_coverage():
    """Get framework coverage analysis."""
    try:
        user_id = session['user_id']
        coverage = get_framework_coverage(user_id)
        return jsonify({'success': True, 'data': coverage})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/analytics/trend')
@login_required
def api_compliance_trend():
    """Get compliance score trend over time."""
    try:
        user_id = session['user_id']
        days = request.args.get('days', 30, type=int)
        trend = get_compliance_trend(user_id, days)
        return jsonify({'success': True, 'data': trend})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/analytics/recommendations')
@login_required
def api_recommendations():
    """Get gap analysis recommendations."""
    try:
        user_id = session['user_id']
        lang = request.args.get('lang', session.get('lang', 'en'))
        recommendations = generate_gap_recommendations(user_id, lang)
        return jsonify({'success': True, 'data': recommendations})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/history')
@login_required
def document_history():
    """Document history page - view all generated documents."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    session['lang'] = lang
    txt = get_text(lang)
    
    user_id = session['user_id']
    conn = get_db()
    
    # Get all documents for this user
    strategies = conn.execute('''
        SELECT id, domain, org_name, sector, language, created_at, 'strategy' as doc_type
        FROM strategies WHERE user_id = ? ORDER BY created_at DESC
    ''', (user_id,)).fetchall()
    
    policies = conn.execute('''
        SELECT id, domain, policy_name, framework, language, created_at, 'policy' as doc_type
        FROM policies WHERE user_id = ? ORDER BY created_at DESC
    ''', (user_id,)).fetchall()
    
    audits = conn.execute('''
        SELECT id, domain, framework, scope, language, created_at, 'audit' as doc_type
        FROM audits WHERE user_id = ? ORDER BY created_at DESC
    ''', (user_id,)).fetchall()
    
    risks = conn.execute('''
        SELECT id, domain, asset_name, threat, risk_level, created_at, 'risk' as doc_type
        FROM risks WHERE user_id = ? ORDER BY created_at DESC
    ''', (user_id,)).fetchall()
    
    
    # Combine and sort all documents
    all_documents = []
    
    for s in strategies:
        all_documents.append({
            'id': s['id'],
            'type': 'strategy',
            'type_label': txt.get('strategies', 'Strategy'),
            'domain': s['domain'],
            'title': s['org_name'] or 'Strategy',
            'subtitle': s['sector'] or '',
            'language': s['language'],
            'created_at': s['created_at']
        })
    
    for p in policies:
        all_documents.append({
            'id': p['id'],
            'type': 'policy',
            'type_label': txt.get('policies', 'Policy'),
            'domain': p['domain'],
            'title': p['policy_name'] or 'Policy',
            'subtitle': p['framework'] or '',
            'language': p['language'],
            'created_at': p['created_at']
        })
    
    for a in audits:
        all_documents.append({
            'id': a['id'],
            'type': 'audit',
            'type_label': txt.get('audits', 'Audit'),
            'domain': a['domain'],
            'title': a['framework'] or 'Audit Report',
            'subtitle': a['scope'] or '',
            'language': a['language'],
            'created_at': a['created_at']
        })
    
    for r in risks:
        all_documents.append({
            'id': r['id'],
            'type': 'risk',
            'type_label': txt.get('risks', 'Risk Analysis'),
            'domain': r['domain'],
            'title': r['asset_name'] or 'Risk Analysis',
            'subtitle': r['risk_level'] or '',
            'language': r['language'] if 'language' in r.keys() else 'en',
            'created_at': r['created_at']
        })
    
    # Sort by date descending
    all_documents.sort(key=lambda x: x['created_at'] or '', reverse=True)
    
    # Get unique domains for filter
    domains = list(set([d['domain'] for d in all_documents if d['domain']]))
    
    # Get audit logs for timeline
    audit_logs = conn.execute('''
        SELECT * FROM audit_logs WHERE user_id = ? ORDER BY created_at DESC LIMIT 100
    ''', (session['user_id'],)).fetchall()
    
    # Count actions by type for summary stats
    action_counts = {}
    for log in audit_logs:
        act = log['action'] or ''
        base_act = act.split('_')[0] if '_' in act else act
        action_counts[base_act] = action_counts.get(base_act, 0) + 1

    return render_template('history.html',
                          txt=txt,
                          lang=lang,
                          is_rtl=(lang == 'ar'),
                          username=session.get('username'),
                          documents=all_documents,
                          domains=domains,
                          total_count=len(all_documents),
                          audit_logs=audit_logs,
                          action_counts=action_counts)

# ============================================================================
# PROCEDURES LIST API - AI-generated mandatory procedures per framework
# ============================================================================

@app.route('/api/procedures/list', methods=['POST'])
@login_required
def api_procedures_list():
    """Use AI to list all mandatory procedures for a selected framework/regulation."""
    try:
        data = request.json
        domain = data.get('domain', '')
        language = data.get('language', 'en')
        framework = data.get('framework', '')
        
        if not framework:
            return jsonify({'success': False, 'error': 'Framework not specified'}), 400
        
        lang = 'ar' if language == 'ar' else 'en'
        
        if lang == 'ar':
            prompt = f"""أنت خبير في الحوكمة والمخاطر والامتثال (GRC). قم بتحليل "{framework}" وأدرج جميع الإجراءات التشغيلية الإلزامية المطلوبة للامتثال الكامل.

لكل إجراء، قدم:
- العنوان: اسم الإجراء واضح ومحدد
- الوصف: جملة واحدة تصف الغرض والنطاق

يجب أن تغطي الإجراءات جميع الضوابط والمتطلبات الإلزامية في "{framework}".

أجب بتنسيق JSON فقط بدون أي نص إضافي:
[
  {{"title": "عنوان الإجراء", "description": "وصف موجز"}},
  ...
]"""
        else:
            prompt = f"""You are a GRC (Governance, Risk, and Compliance) expert. Analyze "{framework}" and list ALL mandatory operational procedures required for full compliance.

For each procedure, provide:
- title: Clear, specific procedure name
- description: One sentence describing purpose and scope

The procedures must cover EVERY mandatory control and requirement in "{framework}". Include procedures for governance, technical implementation, monitoring, incident response, audit, training, and any other areas mandated by the framework.

Respond ONLY with a JSON array, no additional text:
[
  {{"title": "Procedure Name", "description": "Brief description"}},
  ...
]"""
        
        # Generate using AI
        content = generate_ai_content(prompt, lang, content_type='policy')
        
        if not content:
            return jsonify({'success': False, 'error': 'AI service unavailable'}), 503
        
        # Parse JSON from AI response
        import json as json_lib
        
        # Clean up AI response - extract JSON array
        content = content.strip()
        # Remove markdown code blocks if present
        if content.startswith('```'):
            content = content.split('\n', 1)[1] if '\n' in content else content[3:]
        if content.endswith('```'):
            content = content.rsplit('```', 1)[0]
        content = content.strip()
        
        # Find the JSON array
        start_idx = content.find('[')
        end_idx = content.rfind(']')
        if start_idx >= 0 and end_idx > start_idx:
            json_str = content[start_idx:end_idx + 1]
            try:
                procedures = json_lib.loads(json_str)
                if isinstance(procedures, list) and len(procedures) > 0:
                    # Validate each item has title and description
                    valid_procedures = []
                    for p in procedures:
                        if isinstance(p, dict) and 'title' in p and 'description' in p:
                            valid_procedures.append({
                                'title': str(p['title']).strip(),
                                'description': str(p['description']).strip()
                            })
                    
                    if valid_procedures:
                        return jsonify({'success': True, 'procedures': valid_procedures, 'framework': framework})
            except json_lib.JSONDecodeError:
                pass
        
        # Fallback: If JSON parsing fails, try to extract procedure names from text
        fallback_procedures = []
        for line in content.split('\n'):
            line = line.strip()
            if line and (line[0].isdigit() or line.startswith('-') or line.startswith('•')):
                # Clean up the line
                clean = line.lstrip('0123456789.-•) ').strip()
                if clean and len(clean) > 5:
                    # Split on first colon or dash for title/description
                    if ':' in clean:
                        parts = clean.split(':', 1)
                        fallback_procedures.append({'title': parts[0].strip(), 'description': parts[1].strip() if len(parts) > 1 else ''})
                    elif ' - ' in clean:
                        parts = clean.split(' - ', 1)
                        fallback_procedures.append({'title': parts[0].strip(), 'description': parts[1].strip() if len(parts) > 1 else ''})
                    else:
                        fallback_procedures.append({'title': clean, 'description': ''})
        
        if fallback_procedures:
            return jsonify({'success': True, 'procedures': fallback_procedures, 'framework': framework})
        
        return jsonify({'success': False, 'error': 'Could not parse procedures from AI response'}), 500
        
    except Exception as e:
        print(f"Procedures list error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/domain-stats')
@login_required
def api_domain_stats():
    """Get domain-specific statistics for analytics charts."""
    try:
        domain = request.args.get('domain', '')
        user_id = session['user_id']
        conn = get_db()
        
        # Risk register stats by severity for this domain
        risk_critical = 0
        risk_high = 0
        risk_medium = 0
        risk_low = 0
        
        try:
            risks = conn.execute(
                'SELECT likelihood, impact FROM risk_register WHERE user_id = ? AND domain = ?',
                (user_id, domain)
            ).fetchall()
            
            for r in risks:
                score = (r['likelihood'] or 1) * (r['impact'] or 1)
                if score >= 20:
                    risk_critical += 1
                elif score >= 12:
                    risk_high += 1
                elif score >= 6:
                    risk_medium += 1
                else:
                    risk_low += 1
        except Exception:
            pass
        
        # Task stats
        task_total = 0
        task_done = 0
        try:
            task_rows = conn.execute(
                'SELECT status FROM project_tasks WHERE user_id = ? AND domain = ?',
                (user_id, domain)
            ).fetchall()
            task_total = len(task_rows)
            task_done = sum(1 for t in task_rows if t['status'] == 'done')
        except Exception:
            pass
        
        return jsonify({
            'success': True,
            'risk_critical': risk_critical,
            'risk_high': risk_high,
            'risk_medium': risk_medium,
            'risk_low': risk_low,
            'task_total': task_total,
            'task_done': task_done
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/templates/<template_type>')
@login_required
def api_get_templates(template_type):
    """Get templates for a specific type and domain."""
    lang = request.args.get('lang', 'en')
    domain = request.args.get('domain', '')
    
    # Map domain name to domain code
    domain_code = DOMAIN_CODES.get(domain, 'cyber')
    
    # Get domain-specific templates
    domain_templates = DOMAIN_TEMPLATES.get(domain_code, {})
    templates = domain_templates.get(template_type, {}).get(lang, {})
    
    return jsonify({'success': True, 'templates': templates, 'domain': domain_code})

@app.route('/api/document/<doc_type>/<int:doc_id>')
@login_required
def api_get_document(doc_type, doc_id):
    """Get a specific document content."""
    user_id = session['user_id']
    conn = get_db()
    
    if doc_type == 'strategy':
        doc = conn.execute('SELECT * FROM strategies WHERE id = ? AND user_id = ?', (doc_id, user_id)).fetchone()
        if doc:
            return jsonify({'success': True, 'content': doc['content'], 'domain': doc['domain'], 'language': doc['language']})
    elif doc_type == 'policy':
        doc = conn.execute('SELECT * FROM policies WHERE id = ? AND user_id = ?', (doc_id, user_id)).fetchone()
        if doc:
            return jsonify({'success': True, 'content': doc['content'], 'domain': doc['domain'], 'language': doc['language']})
    elif doc_type == 'audit':
        doc = conn.execute('SELECT * FROM audits WHERE id = ? AND user_id = ?', (doc_id, user_id)).fetchone()
        if doc:
            return jsonify({'success': True, 'content': doc['content'], 'domain': doc['domain'], 'language': doc['language']})
    elif doc_type == 'risk':
        doc = conn.execute('SELECT * FROM risks WHERE id = ? AND user_id = ?', (doc_id, user_id)).fetchone()
        if doc:
            return jsonify({'success': True, 'content': doc['analysis'], 'domain': doc['domain'], 'language': doc['language'] if 'language' in doc.keys() else 'en'})
    
    return jsonify({'success': False, 'error': 'Document not found'}), 404

@app.route('/api/document/<doc_type>/<int:doc_id>', methods=['DELETE'])
@login_required
def api_delete_document(doc_type, doc_id):
    """Delete a specific document."""
    user_id = session['user_id']
    conn = get_db()
    
    table_map = {
        'strategy': 'strategies',
        'policy': 'policies',
        'audit': 'audits',
        'risk': 'risks'
    }
    
    table = table_map.get(doc_type)
    if not table:
        return jsonify({'success': False, 'error': 'Invalid document type'}), 400
    
    try:
        conn.execute(f'DELETE FROM {table} WHERE id = ? AND user_id = ?', (doc_id, user_id))
        conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/share/<doc_type>/<int:doc_id>', methods=['POST'])
@login_required
def api_share_document(doc_type, doc_id):
    """Create a shareable link for a document with optional OTP protection."""
    import uuid
    import random
    user_id = session['user_id']
    data = request.json or {}
    
    # Check if secure share (OTP) is requested
    secure_share = data.get('secure', False)
    recipient_email = data.get('recipient_email', '').strip()
    
    if secure_share and not recipient_email:
        return jsonify({'success': False, 'error': 'Recipient email required for secure share'}), 400
    
    if secure_share and not validate_email(recipient_email):
        return jsonify({'success': False, 'error': 'Invalid email format'}), 400
    
    conn = get_db()
    
    # Get the document content
    content = None
    title = None
    domain = None
    language = 'en'
    
    if doc_type == 'strategy':
        doc = conn.execute('SELECT * FROM strategies WHERE id = ? AND user_id = ?', (doc_id, user_id)).fetchone()
        if doc:
            content = doc['content']
            title = doc['org_name'] or 'Strategy'
            domain = doc['domain']
            language = doc['language']
    elif doc_type == 'policy':
        doc = conn.execute('SELECT * FROM policies WHERE id = ? AND user_id = ?', (doc_id, user_id)).fetchone()
        if doc:
            content = doc['content']
            title = doc['policy_name'] or 'Policy'
            domain = doc['domain']
            language = doc['language']
    elif doc_type == 'audit':
        doc = conn.execute('SELECT * FROM audits WHERE id = ? AND user_id = ?', (doc_id, user_id)).fetchone()
        if doc:
            content = doc['content']
            title = doc['framework'] or 'Audit Report'
            domain = doc['domain']
            language = doc['language']
    elif doc_type == 'risk':
        doc = conn.execute('SELECT * FROM risks WHERE id = ? AND user_id = ?', (doc_id, user_id)).fetchone()
        if doc:
            content = doc['analysis']
            title = doc['asset_name'] or 'Risk Analysis'
            domain = doc['domain']
            language = 'en'
    
    if not content:
        return jsonify({'success': False, 'error': 'Document not found'}), 404
    
    # Generate unique share ID
    share_id = str(uuid.uuid4())[:8]
    
    # For non-secure shares, check if already shared
    if not secure_share:
        existing = conn.execute('SELECT share_id FROM shared_documents WHERE doc_type = ? AND doc_id = ? AND user_id = ? AND is_active = 1 AND requires_otp = 0', 
                               (doc_type, doc_id, user_id)).fetchone()
        if existing:
            return jsonify({
                'success': True, 
                'share_id': existing['share_id'],
                'share_url': f"/shared/{existing['share_id']}",
                'already_shared': True,
                'secure': False
            })
    
    # Generate OTP if secure share
    otp_code = None
    otp_expires = None
    if secure_share:
        otp_code = ''.join([str(random.randint(0, 9)) for _ in range(6)])
        otp_expires = (datetime.now() + timedelta(minutes=10)).isoformat()
    
    # Create share record
    try:
        log_action(session['user_id'], 'share_document', {'doc_type': doc_type, 'doc_id': doc_id, 'recipient': recipient_email})
        conn.execute('''INSERT INTO shared_documents 
                        (share_id, user_id, doc_type, doc_id, title, domain, content, language, requires_otp, recipient_email, otp_code, otp_expires_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                    (share_id, user_id, doc_type, doc_id, title, domain, content, language, 
                     1 if secure_share else 0, recipient_email if secure_share else None, 
                     otp_code, otp_expires))
        conn.commit()
        
        # Send OTP email if secure share
        if secure_share:
            username = session.get('username', 'A Mizan user')
            email_sent, email_msg = send_otp_email(recipient_email, otp_code, title, username)
            
            return jsonify({
                'success': True,
                'share_id': share_id,
                'share_url': f"/shared/{share_id}",
                'secure': True,
                'email_sent': email_sent,
                'recipient': recipient_email
            })
        
        return jsonify({
            'success': True,
            'share_id': share_id,
            'share_url': f"/shared/{share_id}",
            'secure': False
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/share/<share_id>/resend-otp', methods=['POST'])
def api_resend_otp(share_id):
    """Resend OTP code for a secure share."""
    import random
    conn = get_db()
    
    doc = conn.execute('SELECT * FROM shared_documents WHERE share_id = ? AND is_active = 1 AND requires_otp = 1', (share_id,)).fetchone()
    if not doc:
        return jsonify({'success': False, 'error': 'Share not found'}), 404
    
    # Generate new OTP
    otp_code = ''.join([str(random.randint(0, 9)) for _ in range(6)])
    otp_expires = (datetime.now() + timedelta(minutes=10)).isoformat()
    
    conn.execute('UPDATE shared_documents SET otp_code = ?, otp_expires_at = ?, otp_verified = 0 WHERE share_id = ?',
                (otp_code, otp_expires, share_id))
    conn.commit()
    
    # Get sharer username
    user = conn.execute('SELECT username FROM users WHERE id = ?', (doc['user_id'],)).fetchone()
    username = user['username'] if user else 'A Mizan user'
    
    # Send email
    email_sent, email_msg = send_otp_email(doc['recipient_email'], otp_code, doc['title'], username)
    
    return jsonify({
        'success': email_sent,
        'message': 'OTP resent' if email_sent else email_msg
    })

@app.route('/api/share/<share_id>/verify-otp', methods=['POST'])
def api_verify_otp(share_id):
    """Verify OTP code for accessing a secure shared document."""
    data = request.json or {}
    otp_input = data.get('otp', '').strip()
    
    if not otp_input or len(otp_input) != 6:
        return jsonify({'success': False, 'error': 'Invalid OTP format'}), 400
    
    conn = get_db()
    doc = conn.execute('SELECT * FROM shared_documents WHERE share_id = ? AND is_active = 1 AND requires_otp = 1', (share_id,)).fetchone()
    
    if not doc:
        return jsonify({'success': False, 'error': 'Share not found'}), 404
    
    # Check if OTP expired
    if doc['otp_expires_at']:
        expires = datetime.fromisoformat(doc['otp_expires_at'])
        if datetime.now() > expires:
            return jsonify({'success': False, 'error': 'OTP expired', 'expired': True}), 400
    
    # Verify OTP
    if doc['otp_code'] != otp_input:
        return jsonify({'success': False, 'error': 'Invalid OTP'}), 400
    
    # Mark as verified and increment view count
    conn.execute('UPDATE shared_documents SET otp_verified = 1, view_count = view_count + 1 WHERE share_id = ?', (share_id,))
    conn.commit()
    
    # Get sharer username
    user = conn.execute('SELECT username FROM users WHERE id = ?', (doc['user_id'],)).fetchone()
    username = user['username'] if user else 'Unknown'
    
    return jsonify({
        'success': True,
        'document': {
            'title': doc['title'],
            'domain': doc['domain'],
            'doc_type': doc['doc_type'],
            'content': doc['content'],
            'language': doc['language'],
            'shared_by': username,
            'view_count': doc['view_count'] + 1
        }
    })

@app.route('/api/share/<share_id>', methods=['DELETE'])
@login_required
def api_unshare_document(share_id):
    """Stop sharing a document."""
    user_id = session['user_id']
    conn = get_db()
    
    try:
        conn.execute('UPDATE shared_documents SET is_active = 0 WHERE share_id = ? AND user_id = ?', (share_id, user_id))
        conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/my-shares')
@login_required
def api_my_shares():
    """Get all shared documents for current user."""
    user_id = session['user_id']
    conn = get_db()
    
    shares = conn.execute('''SELECT share_id, doc_type, title, domain, view_count, created_at 
                            FROM shared_documents WHERE user_id = ? AND is_active = 1
                            ORDER BY created_at DESC''', (user_id,)).fetchall()
    
    result = []
    for s in shares:
        result.append({
            'share_id': s['share_id'],
            'doc_type': s['doc_type'],
            'title': s['title'],
            'domain': s['domain'],
            'view_count': s['view_count'],
            'created_at': s['created_at']
        })
    
    return jsonify({'success': True, 'shares': result})

@app.route('/shared/<share_id>')
def view_shared_document(share_id):
    """Public page to view a shared document."""
    lang = request.args.get('lang', 'en')
    txt = get_text(lang)
    
    conn = get_db()
    doc = conn.execute('SELECT * FROM shared_documents WHERE share_id = ? AND is_active = 1', (share_id,)).fetchone()
    
    if not doc:
        return render_template('shared.html', 
                              txt=txt, 
                              lang=lang,
                              is_rtl=(lang == 'ar'),
                              error=True,
                              error_message=txt.get('document_not_found', 'Document not found'))
    
    # Check if OTP is required
    if doc['requires_otp'] == 1:
        # Get username for display
        user = conn.execute('SELECT username FROM users WHERE id = ?', (doc['user_id'],)).fetchone()
        username = user['username'] if user else 'Unknown'
        
        return render_template('shared.html',
                              txt=txt,
                              lang=doc['language'] or lang,
                              is_rtl=(doc['language'] == 'ar'),
                              error=False,
                              requires_otp=True,
                              share_id=share_id,
                              doc_title=doc['title'],
                              shared_by=username,
                              recipient_email=doc['recipient_email'])
    
    # Public share - increment view count and show document
    conn.execute('UPDATE shared_documents SET view_count = view_count + 1 WHERE share_id = ?', (share_id,))
    conn.commit()
    
    # Get username
    user = conn.execute('SELECT username FROM users WHERE id = ?', (doc['user_id'],)).fetchone()
    username = user['username'] if user else 'Unknown'
    
    return render_template('shared.html',
                          txt=txt,
                          lang=doc['language'] or lang,
                          is_rtl=(doc['language'] == 'ar'),
                          error=False,
                          requires_otp=False,
                          document={
                              'title': doc['title'],
                              'domain': doc['domain'],
                              'doc_type': doc['doc_type'],
                              'content': doc['content'],
                              'view_count': doc['view_count'] + 1,
                              'created_at': doc['created_at'],
                              'shared_by': username
                          },
                          share_id=share_id)

@app.route('/api/usage/<domain>')
@login_required
def api_get_usage(domain):
    """Get user's remaining usage for a specific domain."""
    usage = get_remaining_usage(session['user_id'], domain)
    return jsonify({'success': True, 'usage': usage, 'domain': domain})

# ============================================================================
# ROUTES - API ENDPOINTS
# ============================================================================

@app.route('/api/generate-strategy', methods=['POST'])
@login_required
def api_generate_strategy():
    """Generate strategy via AI."""
    import sys
    print("=" * 60, flush=True)
    print("STRATEGY GENERATION STARTED", flush=True)
    print("=" * 60, flush=True)
    
    try:
        data = request.json
        domain = data.get('domain', 'Cyber Security')
        print(f"DEBUG: Domain = {domain}", flush=True)
        
        # Check usage limit for this domain
        can_generate, used, limit = check_usage_limit(session['user_id'], 'strategies', domain)
        if not can_generate:
            return jsonify({
                'success': False,
                'error': f'Usage limit reached for {domain}. You have used {used}/{limit} strategies in this domain.',
                'limit_reached': True
            }), 429
        
        lang = data.get('language', 'en')
        print(f"DEBUG: Language = {lang}", flush=True)
        
        # Get current state info
        org_structure = data.get('org_structure', 'Not specified')
        technologies = data.get('technologies', [])
        maturity = data.get('maturity_level', 'initial')
        tech_list = ', '.join(technologies) if technologies else 'None specified'
        frameworks_list = ', '.join(data.get('frameworks', [])) if data.get('frameworks') else 'Not specified'
        
        if lang == 'ar':
            # Map domain to Arabic specialized context
            domain_context = {
                'Cyber Security': 'الأمن السيبراني - يشمل حماية الأنظمة والشبكات والبيانات من التهديدات السيبرانية',
                'الأمن السيبراني': 'الأمن السيبراني - يشمل حماية الأنظمة والشبكات والبيانات من التهديدات السيبرانية',
                'Data Management': 'إدارة البيانات - يشمل حوكمة البيانات وجودتها وحمايتها وإدارة دورة حياتها',
                'إدارة البيانات': 'إدارة البيانات - يشمل حوكمة البيانات وجودتها وحمايتها وإدارة دورة حياتها',
                'Artificial Intelligence': 'الذكاء الاصطناعي - يشمل حوكمة الذكاء الاصطناعي وأخلاقياته ومخاطره',
                'الذكاء الاصطناعي': 'الذكاء الاصطناعي - يشمل حوكمة الذكاء الاصطناعي وأخلاقياته ومخاطره',
                'Digital Transformation': 'التحول الرقمي - يشمل استراتيجية الرقمنة والتقنيات الناشئة وإدارة التغيير',
                'التحول الرقمي': 'التحول الرقمي - يشمل استراتيجية الرقمنة والتقنيات الناشئة وإدارة التغيير',
                'Global Standards': 'المعايير العالمية - يشمل الامتثال للمعايير الدولية مثل ISO وNIST وCOBIT',
                'المعايير العالمية': 'المعايير العالمية - يشمل الامتثال للمعايير الدولية مثل ISO وNIST وCOBIT',
                'Enterprise Risk Management': 'إدارة المخاطر المؤسسية - يشمل تحديد وتقييم ومعالجة ومراقبة المخاطر على مستوى المؤسسة وفق COSO ERM وISO 31000',
                'إدارة المخاطر المؤسسية': 'إدارة المخاطر المؤسسية - يشمل تحديد وتقييم ومعالجة ومراقبة المخاطر على مستوى المؤسسة وفق COSO ERM وISO 31000',
            }
            domain_desc = domain_context.get(domain, domain)
            
            # Domain-specific gaps for Arabic
            domain_gaps_ar = {
                'Cyber Security': [
                    ('غياب سياسة الأمن السيبراني', 'عدم وجود سياسة أمن سيبراني معتمدة وشاملة'),
                    ('ضعف إدارة الثغرات الأمنية', 'عدم وجود برنامج منظم لاكتشاف ومعالجة الثغرات'),
                    ('نقص في المراقبة الأمنية', 'محدودية قدرات الكشف والاستجابة للتهديدات'),
                    ('ضعف حماية نقاط النهاية', 'عدم كفاية حلول حماية الأجهزة والخوادم'),
                    ('عدم اكتمال خطة الاستجابة للحوادث', 'نقص في إجراءات الاستجابة للحوادث السيبرانية'),
                ],
                'الأمن السيبراني': [
                    ('غياب سياسة الأمن السيبراني', 'عدم وجود سياسة أمن سيبراني معتمدة وشاملة'),
                    ('ضعف إدارة الثغرات الأمنية', 'عدم وجود برنامج منظم لاكتشاف ومعالجة الثغرات'),
                    ('نقص في المراقبة الأمنية', 'محدودية قدرات الكشف والاستجابة للتهديدات'),
                    ('ضعف حماية نقاط النهاية', 'عدم كفاية حلول حماية الأجهزة والخوادم'),
                    ('عدم اكتمال خطة الاستجابة للحوادث', 'نقص في إجراءات الاستجابة للحوادث السيبرانية'),
                ],
                'Data Management': [
                    ('غياب سياسة حوكمة البيانات', 'عدم وجود إطار شامل لحوكمة البيانات المؤسسية'),
                    ('ضعف تصنيف البيانات', 'عدم تصنيف البيانات حسب الحساسية والأهمية'),
                    ('نقص في جودة البيانات', 'غياب معايير ومقاييس جودة البيانات'),
                    ('ضعف إدارة دورة حياة البيانات', 'عدم وجود سياسات للاحتفاظ والإتلاف'),
                    ('عدم الامتثال لنظام حماية البيانات الشخصية', 'فجوات في تطبيق متطلبات PDPL'),
                ],
                'إدارة البيانات': [
                    ('غياب سياسة حوكمة البيانات', 'عدم وجود إطار شامل لحوكمة البيانات المؤسسية'),
                    ('ضعف تصنيف البيانات', 'عدم تصنيف البيانات حسب الحساسية والأهمية'),
                    ('نقص في جودة البيانات', 'غياب معايير ومقاييس جودة البيانات'),
                    ('ضعف إدارة دورة حياة البيانات', 'عدم وجود سياسات للاحتفاظ والإتلاف'),
                    ('عدم الامتثال لنظام حماية البيانات الشخصية', 'فجوات في تطبيق متطلبات PDPL'),
                ],
                'Artificial Intelligence': [
                    ('غياب إطار حوكمة الذكاء الاصطناعي', 'عدم وجود سياسات وإجراءات لحوكمة أنظمة الذكاء الاصطناعي'),
                    ('ضعف إدارة مخاطر الذكاء الاصطناعي', 'عدم تقييم المخاطر المرتبطة بنماذج AI'),
                    ('نقص في الشفافية والتفسير', 'غياب آليات لتفسير قرارات نماذج AI'),
                    ('ضعف ضمان العدالة والحياد', 'عدم اختبار التحيز في نماذج الذكاء الاصطناعي'),
                    ('عدم الامتثال لمبادئ أخلاقيات AI', 'فجوات في تطبيق مبادئ SDAIA الأخلاقية'),
                ],
                'الذكاء الاصطناعي': [
                    ('غياب إطار حوكمة الذكاء الاصطناعي', 'عدم وجود سياسات وإجراءات لحوكمة أنظمة الذكاء الاصطناعي'),
                    ('ضعف إدارة مخاطر الذكاء الاصطناعي', 'عدم تقييم المخاطر المرتبطة بنماذج AI'),
                    ('نقص في الشفافية والتفسير', 'غياب آليات لتفسير قرارات نماذج AI'),
                    ('ضعف ضمان العدالة والحياد', 'عدم اختبار التحيز في نماذج الذكاء الاصطناعي'),
                    ('عدم الامتثال لمبادئ أخلاقيات AI', 'فجوات في تطبيق مبادئ SDAIA الأخلاقية'),
                ],
                'Digital Transformation': [
                    ('غياب استراتيجية التحول الرقمي', 'عدم وجود خارطة طريق واضحة للتحول الرقمي'),
                    ('ضعف إدارة التغيير', 'محدودية برامج إدارة التغيير المؤسسي'),
                    ('نقص في الكفاءات الرقمية', 'فجوة في المهارات الرقمية للموظفين'),
                    ('ضعف التكامل بين الأنظمة', 'عدم ترابط الأنظمة الرقمية المختلفة'),
                    ('عدم قياس العائد من التحول الرقمي', 'غياب مؤشرات قياس نجاح التحول'),
                ],
                'التحول الرقمي': [
                    ('غياب استراتيجية التحول الرقمي', 'عدم وجود خارطة طريق واضحة للتحول الرقمي'),
                    ('ضعف إدارة التغيير', 'محدودية برامج إدارة التغيير المؤسسي'),
                    ('نقص في الكفاءات الرقمية', 'فجوة في المهارات الرقمية للموظفين'),
                    ('ضعف التكامل بين الأنظمة', 'عدم ترابط الأنظمة الرقمية المختلفة'),
                    ('عدم قياس العائد من التحول الرقمي', 'غياب مؤشرات قياس نجاح التحول'),
                ],
                'Enterprise Risk Management': [
                    ('غياب إطار إدارة المخاطر المؤسسية', 'عدم وجود إطار ERM متكامل'),
                    ('ضعف تحديد المخاطر', 'عدم شمولية عملية تحديد المخاطر'),
                    ('نقص في تقييم المخاطر', 'غياب منهجية موحدة لتقييم المخاطر'),
                    ('ضعف مراقبة المخاطر', 'محدودية آليات المتابعة والرقابة'),
                    ('عدم ربط المخاطر بالأهداف الاستراتيجية', 'فجوة بين إدارة المخاطر والتخطيط الاستراتيجي'),
                ],
            }
            
            # Get domain-specific gaps or use default
            gaps = domain_gaps_ar.get(domain, [
                ('غياب السياسات المعتمدة', 'عدم وجود سياسات موثقة ومعتمدة'),
                ('ضعف برنامج التوعية', 'عدم كفاية برامج التدريب والتوعية'),
                ('نقص في إدارة المخاطر', 'غياب منهجية واضحة لإدارة المخاطر'),
                ('ضعف آليات المراقبة', 'محدودية قدرات الرقابة والمتابعة'),
                ('عدم اكتمال خطة الاستمرارية', 'نقص في خطط استمرارية الأعمال'),
            ])
            
            # Build gaps table
            gaps_table = ""
            # If no organizational structure, add domain-specific structure gap as first confirmed gap
            if org_structure and ('no' in org_structure.lower() or 'لا' in org_structure or 'none' in org_structure.lower()):
                domain_structure_gap_ar = {
                    'Cyber Security': ('غياب الهيكل التنظيمي للأمن السيبراني', 'عدم وجود إدارة مختصة بالأمن السيبراني - أساس تطبيق جميع الضوابط'),
                    'Data Management': ('غياب الهيكل التنظيمي لإدارة البيانات', 'عدم وجود مكتب إدارة بيانات أو أدوار حوكمة البيانات - أساس تطبيق جميع ضوابط البيانات'),
                    'Artificial Intelligence': ('غياب الهيكل التنظيمي لحوكمة الذكاء الاصطناعي', 'عدم وجود لجنة أو مكتب حوكمة الذكاء الاصطناعي - أساس الإشراف على أنظمة AI'),
                    'Digital Transformation': ('غياب الهيكل التنظيمي للتحول الرقمي', 'عدم وجود مكتب أو فريق للتحول الرقمي - أساس قيادة مبادرات التحول'),
                    'Enterprise Risk Management': ('غياب الهيكل التنظيمي لإدارة المخاطر', 'عدم وجود إدارة مختصة بالمخاطر المؤسسية - أساس تطبيق إطار ERM'),
                    'Global Standards': ('غياب الهيكل التنظيمي للامتثال', 'عدم وجود إدارة مختصة بالامتثال والمعايير - أساس تطبيق جميع المعايير'),
                }
                gap_name, gap_desc = domain_structure_gap_ar.get(domain, ('غياب الهيكل التنظيمي المختص', 'عدم وجود إدارة مختصة - أساس تطبيق جميع الضوابط'))
                gaps_table += f"| 1 | {gap_name} | {gap_desc} | حرجة | مفتوحة - مؤكدة |\n"
                start_idx = 2
            else:
                start_idx = 1
            for i, (gap_name, gap_desc) in enumerate(gaps, start_idx):
                priority = "عالية" if i <= 2 else ("متوسطة" if i <= 4 else "منخفضة")
                gaps_table += f"| {i} | {gap_name} | {gap_desc} | {priority} | مفتوحة |\n"
            
            # Translate framework names for Arabic
            frameworks_ar = translate_frameworks_list_ar(frameworks_list)
            
            # Pre-compute Arabic domain-specific text
            _domain_isolation_ar = {
                'Cyber Security': 'أمن سيبراني. لا تذكر لجان حوكمة بيانات أو أخلاقيات AI أو مكاتب تحول رقمي',
                'Data Management': 'إدارة بيانات. لا تذكر CISO أو إدارة الأمن السيبراني أو SOC أو CSIRT. استخدم بدلاً منها: مسؤول حماية البيانات (DPO)، لجنة حوكمة البيانات، مدير البيانات (CDO)، مكتب خصوصية البيانات',
                'Artificial Intelligence': 'ذكاء اصطناعي. لا تذكر CISO أو إدارة الأمن السيبراني أو SOC. استخدم بدلاً منها: مجلس أخلاقيات AI، لجنة حوكمة AI، مسؤول مخاطر النماذج',
                'Enterprise Risk Management': 'إدارة مخاطر مؤسسية. لا تذكر CISO أو SOC. استخدم: مدير المخاطر (CRO)، ملاك المخاطر، لجنة المخاطر',
                'Digital Transformation': 'تحول رقمي. لا تذكر CISO أو SOC. استخدم: مدير التحول الرقمي، مكتب التحول الرقمي، فريق الابتكار',
                'Global Standards': 'معايير عالمية. التزم بأدوار الامتثال والمعايير فقط',
            }.get(domain, 'ابقَ ضمن نطاق المجال المحدد')
            domain_isolation_text_ar = f"عزل المجال — حاسم: هذه استراتيجية {_domain_isolation_ar}. التزم حصرياً بنطاق {domain}."
            
            _has_no_org_ar = org_structure and ('no' in org_structure.lower() or 'لا' in org_structure or 'not' in org_structure.lower() or 'none' in org_structure.lower())
            
            _org_struct_detail_ar = {
                'Cyber Security': 'إدارة أو قسم مختص بالأمن السيبراني. يجب أن تتضمن الركيزة الأولى تفاصيل الهيكل المقترح (مدير الأمن السيبراني، فريق العمليات SOC، فريق الحوكمة، فريق الاستجابة CSIRT) مع تحديد الصلاحيات وخطوط التقارير.',
                'Data Management': 'مكتب إدارة البيانات. يجب أن تتضمن الركيزة الأولى تفاصيل الهيكل المقترح (مدير البيانات CDO، أمناء البيانات، فريق جودة البيانات، فريق حوكمة البيانات) مع تحديد الصلاحيات وخطوط التقارير.',
                'Artificial Intelligence': 'لجنة أو مكتب حوكمة الذكاء الاصطناعي. يجب أن تتضمن الركيزة الأولى تفاصيل الهيكل المقترح (مسؤول أخلاقيات AI، فريق مخاطر النماذج، فريق امتثال AI) مع تحديد الصلاحيات وخطوط التقارير.',
                'Digital Transformation': 'مكتب أو فريق للتحول الرقمي. يجب أن تتضمن الركيزة الأولى تفاصيل الهيكل المقترح (مدير التحول الرقمي، فريق الابتكار، فريق إدارة التغيير) مع تحديد الصلاحيات وخطوط التقارير.',
                'Enterprise Risk Management': 'إدارة مختصة بالمخاطر المؤسسية. يجب أن تتضمن الركيزة الأولى تفاصيل الهيكل المقترح (مدير المخاطر CRO، محللو المخاطر، ملاك المخاطر) مع تحديد الصلاحيات وخطوط التقارير.',
                'Global Standards': 'إدارة مختصة بالامتثال والمعايير. يجب أن تتضمن الركيزة الأولى تفاصيل الهيكل المقترح (مدير الامتثال، منسقو المعايير، فريق التدقيق) مع تحديد الصلاحيات وخطوط التقارير.',
            }.get(domain, 'الإدارة المختصة مع تفاصيل الهيكل المقترح وتحديد الصلاحيات وخطوط التقارير.')
            org_struct_critical_ar = f"تنبيه حاسم: الهيكل التنظيمي الحالي لا يتضمن إدارة مختصة. يجب أن يكون الهدف الأول في الاستراتيجية هو إنشاء {_org_struct_detail_ar} هذا هو الأساس الذي تُبنى عليه جميع الأهداف الأخرى." if _has_no_org_ar else ""
            
            prompt = f"""أنت خبير استراتيجي في مجال {domain_desc}.

أنشئ وثيقة استراتيجية احترافية لمجال **{domain_desc}** وفق إطار **{frameworks_ar}**.

المنظمة: {data.get('org_name', 'المنظمة')}
القطاع: {data.get('sector', 'حكومي')}
الحجم: {data.get('size', 'متوسطة')}
مستوى النضج: {maturity}
الهيكل التنظيمي الحالي: {org_structure}

تقييد الإطار التنظيمي:
- الإطار المحدد هو: {frameworks_ar}
- يجب الإشارة إلى {frameworks_ar} بالاسم في الرؤية والأهداف والركائز والمؤشرات
- عند تكرار اسم الإطار، استخدم الاسم المختصر فقط (مثلاً: NCA ECC بدلاً من NCA ECC Essential Cybersecurity Controls)
- لا تكتب الاسم الكامل بالإنجليزية - استخدم الاختصار العربي أو الاختصار الإنجليزي القصير فقط
- لا تذكر أي إطار آخر مثل ISO أو NIST أو COBIT إلا إذا كان مذكوراً في القائمة أعلاه
- اربط كل هدف وكل ركيزة بمتطلبات محددة من {frameworks_ar}

{org_struct_critical_ar}

قواعد صارمة:
- المحتوى خاص بمجال {domain} فقط
- القيم الحالية: يجب أن تكون جميعها "يُحدد بعد التقييم" - لا تكتب أي نسب مئوية أو أرقام في عمود القيمة الحالية
- القيم المستهدفة فقط يمكن أن تحتوي على نسب مئوية مع مبرر
- لا تكتب أي تعليمات أو تنبيهات في المخرجات
- لا تنسخ أي نص من التعليمات إلى المخرجات - المخرجات يجب أن تكون وثيقة احترافية فقط
- استخدم [SECTION] بين الأقسام الستة
- يجب تضمين جميع الأقسام الستة كاملة
- لا تبدأ أي قسم بعبارة "إليك" أو "فيما يلي" أو أي عبارة حوارية

أسلوب الوثيقة (معيار شركات الاستشارات الكبرى Big-4):
- اكتب بأسلوب وثائق الاستشارات الاحترافية (مثل ديلويت، PwC، EY، KPMG)
- استخدم لغة رسمية موثوقة على مستوى تنفيذي — بدون أسلوب غير رسمي
- الجداول يجب أن تكون شاملة بعناوين واضحة وبيانات قابلة للتنفيذ
- كل قسم يجب أن يكون مفصلاً وتحليلياً وليس ملخصات سطحية
- استخدم رموز الضوابط المحددة ومراجع البنود وأهداف قابلة للقياس
- استخدم "المنظمة" أو "الجهة" — لا تستخدم "أنتم" أو "شركتكم"
- كل توصية يجب أن تكون محددة وقابلة للتنفيذ مع مسؤول واضح

قواعد التنسيق (صارمة — يجب أن يكون المخرج Markdown صالح):
- استخدم ## لعناوين الأقسام و ### لعناوين الجداول الفرعية
- جميع الجداول يجب أن تستخدم صيغة جداول Markdown القياسية بالضبط بهذا الترتيب:
  1. صف الرأس:    | العمود 1 | العمود 2 | العمود 3 |
  2. صف الفاصل:   |----------|----------|----------|
  3. صفوف البيانات: | بيانات   | بيانات   | بيانات   |
- صف الفاصل (السطر 2 بالشرطات) إلزامي — بدونه الجدول لن يظهر
- كل صف في الجدول يجب أن يكون في سطر مستقل — لا تدمج صفوف متعددة في سطر واحد
- استخدم • (نقطة) للمبادرات في الركائز الاستراتيجية فقط
- ضع سطر فارغ قبل وبعد كل جدول
- استخدم **عريض** للتأكيد
- لا تستخدم نصوصاً حوارية مثل "إليك" أو "فيما يلي" أو "بناءً على طلبك"
- ابدأ كل قسم مباشرة بالمحتوى بدون مقدمة

قاموس المصطلحات (استخدم هذه الترجمات بالضبط):
- Governance → الحوكمة
- Compliance → الامتثال
- Zero-day → الثغرات الصفرية (وليس: يوم الصفر)
- Phishing → التصيد الاحتيالي
- Ransomware → برمجيات الفدية
- Penetration Testing → اختبار الاختراق
- Vulnerability → الثغرة الأمنية
- Incident Response → الاستجابة للحوادث
- Access Control → التحكم في الوصول
- Data Breach → اختراق البيانات
- Business Continuity → استمرارية الأعمال
- Disaster Recovery → التعافي من الكوارث

{domain_isolation_text_ar}

قاعدة حاسمة - عمود المبررات:
- المبرر يجب أن يشرح لماذا الهدف مهم وكيف يرتبط بمتطلب محدد من {frameworks_ar}
- مثال جيد: "الخطأ البشري مسبب رئيسي للحوادث؛ {frameworks_ar} يلزم بتدريب أمني منتظم"
- مثال جيد: "بدون فريق مختص لا يمكن تطبيق أي ضابط بشكل مستدام"
- مثال سيء: "ضوابط الحماية التقنية" (هذا وصف وليس مبرر)
- مثال سيء: "متطلبات التوعية والتدريب" (هذا عنوان وليس مبرر)
- لا تنسخ عبارات التعليمات مثل "selected frameworks" أو "FRAMEWORK" إلى المبررات

## 1. الرؤية والأهداف الاستراتيجية
(هذا القسم يجيب على: أين نريد أن نكون؟ — رؤية طموحة وأهداف قابلة للقياس. لا تذكر أي مبادرات أو مشاريع أو خطط تشغيلية — تلك في القسم الثالث فقط)

**الرؤية الاستراتيجية:**
(فقرة أو فقرتان تصفان الحالة المستقبلية المرجوة للمنظمة فيما يتعلق بالامتثال لإطار {frameworks_ar}. ركّز على النتائج والطموحات وليس المشاريع أو المبادرات)

### الأهداف الاستراتيجية:
| # | الهدف الاستراتيجي | المؤشر المستهدف | المبرر (لماذا هذا مهم وما علاقته بـ {frameworks_ar}) | الإطار الزمني |
|---|-------------------|-----------------|------------------------------------------------------|---------------|
(5 أهداف - كل مبرر يشرح السبب الحقيقي وراء الهدف وعلاقته بمتطلب محدد من {frameworks_ar}، وليس مجرد اسم ضابط)

[SECTION]

## 2. الركائز الاستراتيجية
(هذا القسم يجيب على: كيف سنحقق الرؤية؟ — محاور عمل تشغيلية بمبادرات ملموسة. لا تكرر جدول الأهداف من القسم الأول. ابدأ مباشرة بالركائز والمبادرات التشغيلية)

تعليمات حاسمة: كل ركيزة يجب أن تحتوي على 3-4 مبادرات تبدأ بـ "• مبادرة:". المبادرات هي جوهر الاستراتيجية — بدونها الاستراتيجية بلا قيمة تشغيلية. كل مبادرة يجب أن تصف إجراءً محدداً وقابلاً للقياس مع مخرج واضح.

تنبيه عزل الركائز — يجب أن تكون كل ركيزة مستقلة بلا تداخل في المحتوى:
- ركيزة الحوكمة: سياسات، لجان، هياكل تنظيمية، إدارة مخاطر فقط
- ركيزة التقنية: أدوات، أنظمة، بنية تحتية فقط — لا تذكر السياسات أو التدريب
- ركيزة الكوادر: تدريب، توعية، توظيف فقط — لا تذكر الأدوات أو السياسات
- ركيزة الامتثال: تدقيق، تقييم، ضمان فقط — لا تذكر التقنيات أو التوظيف

### الركيزة 1: (الاسم — مثال: الحوكمة وإدارة المخاطر)
• مبادرة: [إجراء تشغيلي ملموس مع مخرج واضح — مثال: "إنشاء لجنة توجيهية للحوكمة تجتمع ربع سنوياً"]
• مبادرة: [إجراء تشغيلي ملموس]
• مبادرة: [إجراء تشغيلي ملموس]

### الركيزة 2: (الاسم — مثال: التقنية والبنية التحتية)
(وصف وأهمية)
• مبادرة 1 — [إجراء تشغيلي ملموس]
• مبادرة 2 — [إجراء تشغيلي ملموس]
• مبادرة 3 — [إجراء تشغيلي ملموس]

### الركيزة 3: (الاسم — مثال: الكوادر والثقافة)
(وصف وأهمية)
• مبادرة 1 — [إجراء تشغيلي ملموس]
• مبادرة 2 — [إجراء تشغيلي ملموس]
• مبادرة 3 — [إجراء تشغيلي ملموس]

### الركيزة 4: (الاسم — مثال: الامتثال والضمان)
(وصف وأهمية)
• مبادرة 1 — [إجراء تشغيلي ملموس]
• مبادرة 2 — [إجراء تشغيلي ملموس]
• مبادرة 3 — [إجراء تشغيلي ملموس]

[SECTION]

## 3. تحليل الفجوات

### الفجوات المحددة:
| # | الفجوة | الوصف | الأولوية | الحالة |
|---|--------|-------|---------|--------|
{gaps_table}

### أدلة معالجة الفجوات التفصيلية:

قدّم دليل تنفيذ مستقل وكامل لكل فجوة أعلاه. لا تكتب "نفس الخطوات السابقة" أو "كرر".

---
#### دليل معالجة الفجوة رقم 1: [اسم الفجوة الأولى]
| المرحلة | الخطوة | الوصف | المسؤول | المخرج |
|---------|--------|-------|---------|--------|
| التخطيط | 1.1 | [خطوة محددة] | [الفريق] | [المخرج] |
| التخطيط | 1.2 | [خطوة محددة] | [الفريق] | [المخرج] |
| التنفيذ | 2.1 | [خطوة محددة] | [الفريق] | [المخرج] |
| التنفيذ | 2.2 | [خطوة محددة] | [الفريق] | [المخرج] |
| التحقق | 3.1 | [خطوة محددة] | [الفريق] | [المخرج] |
**الأدلة المطلوبة:** ☐ [دليل 1] ☐ [دليل 2] ☐ [دليل 3]

---
#### دليل معالجة الفجوة رقم 2: [اسم الفجوة الثانية]
| المرحلة | الخطوة | الوصف | المسؤول | المخرج |
|---------|--------|-------|---------|--------|
(5 صفوف مختلفة — لا تنسخ من الفجوة الأولى)
**الأدلة المطلوبة:** ☐ [دليل 1] ☐ [دليل 2] ☐ [دليل 3]

---
#### دليل معالجة الفجوة رقم 3: [اسم الفجوة الثالثة]
| المرحلة | الخطوة | الوصف | المسؤول | المخرج |
|---------|--------|-------|---------|--------|
(5 صفوف فريدة)
**الأدلة المطلوبة:** ☐ [دليل 1] ☐ [دليل 2] ☐ [دليل 3]

---
#### دليل معالجة الفجوة رقم 4: [اسم الفجوة الرابعة]
| المرحلة | الخطوة | الوصف | المسؤول | المخرج |
|---------|--------|-------|---------|--------|
(5 صفوف فريدة)
**الأدلة المطلوبة:** ☐ [دليل 1] ☐ [دليل 2] ☐ [دليل 3]

---
#### دليل معالجة الفجوة رقم 5: [اسم الفجوة الخامسة]
| المرحلة | الخطوة | الوصف | المسؤول | المخرج |
|---------|--------|-------|---------|--------|
(5 صفوف فريدة)
**الأدلة المطلوبة:** ☐ [دليل 1] ☐ [دليل 2] ☐ [دليل 3]

[SECTION]

## 4. خارطة الطريق التنفيذية

### المرحلة 1: التأسيس (0-6 أشهر)
| # | النشاط | المسؤول | الجدول الزمني | المخرجات |
|---|--------|---------|---------------|----------|
(4 أنشطة خاصة بمجال {domain})

### المرحلة 2: البناء (6-12 شهر)
| # | النشاط | المسؤول | الجدول الزمني | المخرجات |
|---|--------|---------|---------------|----------|
(4 أنشطة)

### المرحلة 3: التحسين (12-24 شهر)
| # | النشاط | المسؤول | الجدول الزمني | المخرجات |
|---|--------|---------|---------------|----------|
(4 أنشطة)

[SECTION]

## 5. مؤشرات الأداء الرئيسية

### مؤشرات الأداء:
| # | المؤشر | الوصف | القيمة الحالية | القيمة المستهدفة | المبرر (لماذا هذا المستهدف مهم) | الإطار الزمني | مصدر البيانات |
|---|--------|-------|----------------|------------------|-------------------------------------|---------------|---------------|
(8 مؤشرات KPI - عمود "القيمة الحالية" يجب أن يكون "يُحدد بعد التقييم" لجميع المؤشرات بدون استثناء)

قاعدة حاسمة للمؤشرات:
- لا تضع ضابطاً واحداً (مثل MFA) كمؤشر مستقل - بل استخدم مؤشرات على مستوى نطاق الضوابط (مثلاً: "نسبة تطبيق ضوابط إدارة الهوية والوصول" بدلاً من "تغطية MFA")
- المؤشرات يجب أن تقيس نتائج قابلة للقياس على مستوى المنظمة

### أدلة تقييم المؤشرات

تنبيه: كل مؤشر يحتاج دليل تقييم فريد بصيغة جدول. لا تكرر نفس الخطوات لكل مؤشر.
- مثال سيء (عام): "تحديد النطاق → تحديد المصادر → جمع البيانات → التحقق → الاحتساب" لكل مؤشر
- مثال جيد (محدد): لوقت الاستجابة → "استخراج أوقات من SIEM → حساب MTTD من وقت التنبيه → حساب MTTR من التنبيه للاحتواء"

كل دليل يجب أن يكون بصيغة جدول بالأعمدة التالية:

---
#### دليل تقييم المؤشر رقم 1: [اسم المؤشر الأول]
| الخطوة | الإجراء | الأداة/النظام | المسؤول | المخرج |
|--------|---------|---------------|---------|--------|
| 1 | [خطوة جمع بيانات محددة] | [اسم الأداة بالتحديد] | [الدور المحدد] | [مخرج ملموس] |
| 2 | [خطوة التحليل/الاحتساب] | [الأداة] | [الدور] | [المخرج] |
| 3 | [خطوة التحقق] | [الأداة] | [الدور] | [المخرج] |
| 4 | [خطوة إعداد التقرير] | [الأداة] | [الدور] | [المخرج] |

**صيغة الاحتساب:** [صيغة الاحتساب الدقيقة لهذا المؤشر]

---
#### دليل تقييم المؤشر رقم 2: [اسم المؤشر الثاني]
| الخطوة | الإجراء | الأداة/النظام | المسؤول | المخرج |
|--------|---------|---------------|---------|--------|
(4-5 صفوف مختلفة - لا تنسخ من المؤشر الأول)

**صيغة الاحتساب:** [صيغة مختلفة]

---
#### دليل تقييم المؤشر رقم 3: [اسم المؤشر الثالث]
| الخطوة | الإجراء | الأداة/النظام | المسؤول | المخرج |
|--------|---------|---------------|---------|--------|
(4-5 صفوف فريدة)

**صيغة الاحتساب:** [صيغة مختلفة]

---
#### دليل تقييم المؤشر رقم 4: [اسم المؤشر الرابع]
| الخطوة | الإجراء | الأداة/النظام | المسؤول | المخرج |
|--------|---------|---------------|---------|--------|
(4-5 صفوف فريدة)

**صيغة الاحتساب:** [صيغة مختلفة]

---
#### دليل تقييم المؤشر رقم 5: [اسم المؤشر الخامس]
| الخطوة | الإجراء | الأداة/النظام | المسؤول | المخرج |
|--------|---------|---------------|---------|--------|
(4-5 صفوف فريدة)

**صيغة الاحتساب:** [صيغة مختلفة]

---
#### دليل تقييم المؤشر رقم 6: [اسم المؤشر السادس]
| الخطوة | الإجراء | الأداة/النظام | المسؤول | المخرج |
|--------|---------|---------------|---------|--------|
(4-5 صفوف فريدة)

**صيغة الاحتساب:** [صيغة مختلفة]

---
#### دليل تقييم المؤشر رقم 7: [اسم المؤشر السابع]
| الخطوة | الإجراء | الأداة/النظام | المسؤول | المخرج |
|--------|---------|---------------|---------|--------|
(4-5 صفوف فريدة)

**صيغة الاحتساب:** [صيغة مختلفة]

---
#### دليل تقييم المؤشر رقم 8: [اسم المؤشر الثامن]
| الخطوة | الإجراء | الأداة/النظام | المسؤول | المخرج |
|--------|---------|---------------|---------|--------|
(4-5 صفوف فريدة)

**صيغة الاحتساب:** [صيغة مختلفة]

[SECTION]

## 6. تقييم الجاهزية والمخاطر

**درجة الثقة:** [X]%

**مبررات التقييم:**
[فقرة تفصيلية توضح أسس تقييم الدرجة مع ذكر العوامل المحددة وتأثير كل عامل على الدرجة]

### عوامل النجاح الحرجة:
| # | العامل | الوصف | الأهمية |
|---|--------|-------|---------|
(5 عوامل نجاح مرتبطة بمتطلبات {frameworks_ar})

### المخاطر الاستراتيجية:
| # | الخطر | الاحتمالية | الأثر | خطة التخفيف |
|---|-------|-----------|-------|-------------|
(4 مخاطر خاصة بمجال {domain})

---
**تاريخ الإعداد:** يُحدد لاحقاً

تنبيه هام — كتلة المزامنة المخفية:
بعد انتهاء الوثيقة بالكامل، أضف كتلة مخفية لتكامل النظام.
الصيغة: <json_sync>{{"initiatives":[{{"title":"...","pillar":"..."}}],"gaps":[{{"name":"...","priority":"high","description":"..."}}],"kpis":[{{"name":"...","target":"...","timeframe":"..."}}],"risks":[{{"name":"...","likelihood":"...","impact":"..."}}]}}</json_sync>
املأها بالبيانات الفعلية من استراتيجيتك أعلاه. هذه الكتلة تُعالج آلياً وتُزال من العرض."""
        else:
            # Map domain to context description
            domain_context = {
                'Cyber Security': 'Cyber Security - including protection of systems, networks, and data from cyber threats',
                'Data Management': 'Data Management - including data governance, quality, protection, and lifecycle management',
                'Artificial Intelligence': 'Artificial Intelligence - including AI governance, ethics, risks, and responsible AI practices',
                'Digital Transformation': 'Digital Transformation - including digitization strategy, emerging technologies, and change management',
                'Global Standards': 'Global Standards - including compliance with international standards like ISO, NIST, and COBIT',
                'Enterprise Risk Management': 'Enterprise Risk Management - including risk identification, assessment, treatment, monitoring, risk appetite, and organizational resilience per COSO ERM and ISO 31000',
            }
            domain_desc = domain_context.get(domain, domain)
            
            # Shorten framework names for clean prompt injection
            # "NCA ECC (Essential Cybersecurity Controls)" → "NCA ECC"
            import re as re_fw
            fw_short = re_fw.sub(r'\s*\([^)]+\)', '', frameworks_list).strip()
            if not fw_short:
                fw_short = frameworks_list
            
            # Build domain-specific English gap table (like Arabic prompt has)
            domain_gaps_en = {
                'Cyber Security': [
                    ('Policy & Governance Gap', 'Outdated or missing cybersecurity policies and procedures'),
                    ('Technical Controls Gap', f'Incomplete implementation of mandatory technical controls per {fw_short}'),
                    ('Training & Awareness Gap', 'Insufficient security awareness programs for employees'),
                    ('Incident Response Gap', f'Incomplete incident response plan per {fw_short} requirements'),
                    ('Asset Management Gap', 'No comprehensive and updated cybersecurity asset register'),
                ],
                'Data Management': [
                    ('Data Governance Gap', 'Absence of a formal data governance framework and roles'),
                    ('Data Quality Gap', 'No data quality standards or measurement processes'),
                    ('Data Classification Gap', 'Data not classified by sensitivity or criticality levels'),
                    ('Data Lifecycle Gap', 'Missing data retention and disposal policies'),
                    ('Data Protection Gap', 'Insufficient data encryption and access controls'),
                ],
                'Artificial Intelligence': [
                    ('AI Governance Gap', 'No formal AI governance framework or oversight committee'),
                    ('AI Ethics Gap', 'Missing ethical guidelines and bias assessment processes'),
                    ('AI Risk Gap', 'No AI-specific risk assessment methodology'),
                    ('AI Transparency Gap', 'Lack of model documentation and explainability standards'),
                    ('AI Data Gap', 'Insufficient controls for AI training data quality and bias'),
                ],
                'Digital Transformation': [
                    ('Digital Strategy Gap', 'No comprehensive digital transformation roadmap'),
                    ('Change Management Gap', 'Limited organizational change management programs'),
                    ('Digital Skills Gap', 'Workforce digital competency shortfall'),
                    ('Systems Integration Gap', 'Lack of interoperability between digital systems'),
                    ('ROI Measurement Gap', 'No metrics framework for digital transformation outcomes'),
                ],
                'Enterprise Risk Management': [
                    ('ERM Framework Gap', 'No integrated enterprise risk management framework'),
                    ('Risk Identification Gap', 'Incomplete risk identification process'),
                    ('Risk Assessment Gap', 'No standardized risk assessment methodology'),
                    ('Risk Monitoring Gap', 'Limited risk monitoring and reporting mechanisms'),
                    ('Strategic Alignment Gap', 'Risk management not linked to strategic objectives'),
                ],
            }
            
            gaps = domain_gaps_en.get(domain, [
                ('Policy Gap', 'Missing or outdated documented policies'),
                ('Awareness Gap', 'Insufficient training and awareness programs'),
                ('Risk Management Gap', 'No clear risk management methodology'),
                ('Monitoring Gap', 'Limited monitoring and oversight capabilities'),
                ('Continuity Gap', 'Incomplete business continuity plans'),
            ])
            
            # Build pre-populated gap table
            gaps_table_en = ""
            if org_structure and ('no' in org_structure.lower() or 'none' in org_structure.lower()):
                domain_structure_gap_en = {
                    'Cyber Security': ('Cybersecurity Organizational Structure Gap', 'No dedicated cybersecurity department - this is the foundation for implementing all other controls'),
                    'Data Management': ('Data Management Organizational Structure Gap', 'No dedicated data management office or data governance roles - this is the foundation for all data controls'),
                    'Artificial Intelligence': ('AI Governance Organizational Structure Gap', 'No dedicated AI governance committee or office - this is the foundation for AI oversight'),
                    'Digital Transformation': ('Digital Transformation Organizational Structure Gap', 'No dedicated digital transformation office or team - this is the foundation for driving transformation'),
                    'Enterprise Risk Management': ('Risk Management Organizational Structure Gap', 'No dedicated enterprise risk management department - this is the foundation for ERM framework'),
                    'Global Standards': ('Compliance Organizational Structure Gap', 'No dedicated compliance and standards department - this is the foundation for all standards implementation'),
                }
                gap_name, gap_desc = domain_structure_gap_en.get(domain, ('Organizational Structure Gap', 'No dedicated department - this is the foundation for implementing all controls'))
                gaps_table_en += f"| 1 | {gap_name} | {gap_desc} | Critical | Open - Confirmed |\n"
                start_idx = 2
            else:
                start_idx = 1
            for i, (gap_name, gap_desc) in enumerate(gaps, start_idx):
                priority = "High" if i <= 2 else ("Medium" if i <= 4 else "Low")
                gaps_table_en += f"| {i} | {gap_name} | {gap_desc} | {priority} | Open |\n"
            
            # Pre-compute domain-specific text for prompt injection
            _domain_isolation_en = {
                'Cyber Security': 'data governance committees, AI ethics, or digital transformation offices',
                'Data Management': 'CISO, Cybersecurity Department, SOC, CSIRT, or cybersecurity organizational structures. Use Data Protection Officer (DPO), Data Governance Committee, Chief Data Officer (CDO), and Data Privacy Office instead',
                'Artificial Intelligence': 'CISO, Cybersecurity Department, SOC, or cybersecurity gaps as fixes. Use AI Ethics Board, AI Governance Committee, Model Risk Officer instead',
                'Enterprise Risk Management': 'CISO, Cybersecurity Department, or SOC. Use Chief Risk Officer (CRO), Risk Owners, Risk Committee instead',
                'Digital Transformation': 'CISO, Cybersecurity Department, or SOC. Use Chief Digital Officer, Digital Transformation Office, Innovation Team instead',
                'Global Standards': 'domain-specific roles from other areas unless relevant to standards compliance',
            }.get(domain, 'structures from other domains')
            domain_isolation_text_en = f"DOMAIN ISOLATION — CRITICAL: This is a {domain} strategy. Do NOT mention {_domain_isolation_en}. Stay strictly within {domain} domain."
            
            _pillar_names_en = {
                'Cyber Security': ['Cybersecurity Governance & Risk Management', 'Technology & Security Architecture', 'People & Security Culture', 'Compliance & Assurance'],
                'Data Management': ['Data Governance & Stewardship', 'Data Quality & Lifecycle Management', 'Data Literacy & Culture', 'Data Privacy & Compliance'],
                'Artificial Intelligence': ['AI Governance & Ethics Framework', 'Model Transparency & Bias Management', 'AI Skills & Awareness', 'AI Compliance & Monitoring'],
                'Digital Transformation': ['Digital Strategy & Innovation', 'Technology Modernization & Integration', 'Digital Skills & Change Management', 'Digital Governance & Measurement'],
                'Enterprise Risk Management': ['Risk Governance & Framework', 'Risk Assessment & Treatment', 'Risk Culture & Awareness', 'Monitoring & Reporting'],
                'Global Standards': ['Standards Adoption & Certification', 'Process Alignment & Documentation', 'Training & Competency', 'Audit & Continuous Improvement'],
            }.get(domain, ['Governance & Framework', 'Technology & Implementation', 'People & Culture', 'Compliance & Assurance'])
            
            _gap1_title_en = {
                'Cyber Security': 'Cybersecurity Organizational Structure Gap',
                'Data Management': 'Data Management Office Structure Gap',
                'Artificial Intelligence': 'AI Governance Structure Gap',
                'Digital Transformation': 'Digital Transformation Office Gap',
                'Enterprise Risk Management': 'ERM Department Structure Gap',
                'Global Standards': 'Compliance Department Gap',
            }.get(domain, 'Organizational Structure Gap')
            
            _org_dept_en = {
                'Cyber Security': 'cybersecurity department',
                'Data Management': 'data management office',
                'Artificial Intelligence': 'AI governance committee',
                'Digital Transformation': 'digital transformation office',
                'Enterprise Risk Management': 'ERM department',
                'Global Standards': 'compliance department',
            }.get(domain, 'organizational structure')
            
            _org_roles_en = {
                'Cyber Security': 'CISO, SOC team, Governance team, CSIRT',
                'Data Management': 'CDO, Data Stewards, Data Quality team, Data Governance committee',
                'Artificial Intelligence': 'AI Ethics Officer, Model Risk team, AI Compliance team',
                'Digital Transformation': 'CDO/CTO, Innovation team, Change Management team',
                'Enterprise Risk Management': 'CRO, Risk Analysts, Risk Owners, Risk Committee',
                'Global Standards': 'CCO, Standards Liaisons, Audit Coordinators',
            }.get(domain, 'leadership, specialists, support team')
            
            _org_struct_req_en = {
                'Cyber Security': 'cybersecurity org structure requirements',
                'Data Management': 'data management office structure requirements',
                'Artificial Intelligence': 'AI governance structure requirements',
                'Digital Transformation': 'digital transformation office structure requirements',
                'Enterprise Risk Management': 'ERM department structure requirements',
                'Global Standards': 'compliance department structure requirements',
            }.get(domain, 'organizational structure requirements')
            
            _has_no_org = org_structure and ('no' in org_structure.lower() or 'none' in org_structure.lower())
            
            prompt = f"""You are a GRC expert specializing in **{domain_desc}**.

Generate a professional strategy document in Markdown. Domain: **{domain}**.

FRAMEWORK: {fw_short}
- Use "{fw_short}" naturally throughout — in objectives, gaps, KPIs, and pillars
- Do NOT mention any other framework unless it is listed above
- Do NOT echo raw instruction text, labels, or meta-phrases from this prompt into the output
- NEVER include phrases like "The user has selected", "selected frameworks", "FRAMEWORK:", or any other prompt text in the document

OUTPUT RULES:
1. Use relative timeframes: "Within 6 months", "Within 12 months", "Year 1", "Year 2". No person names.
2. Current values = "To be assessed" (never fabricate). Target values may include percentages with justification from {fw_short}.
3. Gap statuses: "Open - Confirmed" for gaps already identified, "Open" for gaps requiring assessment.
4. JUSTIFICATION COLUMN: Each justification must explain WHY the objective matters and HOW it connects to a specific {fw_short} requirement.
   - GOOD: "Human error causes majority of incidents; {fw_short} mandates regular security training"
   - GOOD: "Without a dedicated team, no control can be sustainably implemented or monitored"
   - BAD: "{fw_short} detection and response controls" (this is a label, not a justification)
   - BAD: "{fw_short} awareness requirements" (this restates the objective, not why it matters)
{"5. CRITICAL: The organization has NO dedicated " + _org_dept_en + ". Objective #1 = Establish a dedicated " + _org_dept_en + " (" + _org_roles_en + ") with reporting lines and authority. This is the FOUNDATION before any other objective can succeed." if _has_no_org else ""}

Organization: {data.get('org_name', 'Organization')} | Sector: {data.get('sector', 'General')} | Size: {data.get('size', 'Medium')} | Budget: {data.get('budget', '1M-5M')} | Maturity: {maturity} | Technologies: {tech_list} | Challenges: {data.get('challenges', 'Not specified')}

Write 6 sections separated by [SECTION].

DOCUMENT STYLE (Big-4 Consulting Standard):
- Write in the style of a Big-4 consulting firm (Deloitte, PwC, EY, KPMG) strategy document
- Use formal, authoritative executive language — no informal tone
- Tables must be comprehensive with clear headers and actionable data
- Every section must be substantive with detailed analysis, not high-level summaries
- Include specific control IDs, clause references, and measurable targets
- Use "The Organization" or "the entity" — never "you" or "your company"
- Every recommendation must be specific and actionable with clear ownership

FORMATTING RULES (STRICT — output must be valid Markdown):
- Use ## for section headings, ### for sub-section headings
- ALL tables MUST use standard Markdown pipe table syntax with EXACTLY this 3-line structure:
  1. Header row:    | Column 1 | Column 2 | Column 3 |
  2. Separator row: |----------|----------|----------|
  3. Data rows:     | Data     | Data     | Data     |
- The separator row (line 2 with dashes) is MANDATORY — without it, the table will NOT render
- Each table row MUST be on its own line — NEVER merge multiple rows onto one line
- Use • (bullet) for pillar initiatives ONLY
- Every table MUST have a blank line before AND after it
- Do NOT use conversational filler (no "Here is your strategy", "I have created", "As requested")
- Start each section directly with content — no preamble
- Use **bold** for emphasis, not ALL CAPS

{domain_isolation_text_en}
## 1. Vision & Objectives
(This section is ASPIRATIONAL — it answers WHY and WHERE: where do we want to be? Do NOT include any implementation details, workstreams, or initiatives here — those belong in Section 3.)

**Vision:**
[1-2 paragraphs: the organization's desired future state for {fw_short} compliance. Focus on OUTCOMES and aspirations — e.g., "become a trusted, resilient organization that…". This is a DESTINATION, not a roadmap. Do NOT list initiatives, projects, or workstreams here.]

### Strategic Objectives:
| # | Objective | Target Metric | Justification (why it matters + link to {fw_short}) | Timeframe |
|---|-----------|---------------|-----------------------------------------------------|-----------|
{"| 1 | Establish dedicated " + _org_dept_en + " with defined structure and authority | Approved org structure and full team | Without a dedicated team, no " + fw_short + " control can be sustainably implemented or monitored | Within 6 months |" if _has_no_org else "| 1 | [First objective] | [Metric] | [Why this matters + specific " + fw_short + " requirement] | Within X months |"}
(5-7 objectives total, each with a meaningful justification that explains the business need)

[SECTION]

## 2. Strategic Pillars
(This section is OPERATIONAL — it answers HOW: what workstreams deliver the vision? Each pillar is a named program of work with 3-4 concrete initiatives. Do NOT repeat the objectives table from Section 1. Do NOT restate the vision. Jump straight into pillar names and bullet-point initiatives.)

CRITICAL: Every pillar MUST contain 3-4 bullet-point initiatives starting with "• Initiative:". These are the CORE of the strategy — without initiatives, the strategy has no actionable value. Each initiative must describe a specific, measurable action with a clear deliverable.

PILLAR ISOLATION — CRITICAL (no overlapping content between pillars):
- Governance pillar: policies, committees, organizational structures, risk management ONLY
- Technology pillar: tools, systems, infrastructure, architecture ONLY — do NOT mention policies or training
- People pillar: training, awareness, hiring, culture ONLY — do NOT mention tools or policies
- Compliance pillar: auditing, assessment, assurance ONLY — do NOT mention technologies or hiring

### Pillar 1: [Name — e.g., {_pillar_names_en[0]}]
• Initiative: [concrete action with a deliverable]
• Initiative: [concrete action]
• Initiative: [concrete action]

### Pillar 2: [Name — e.g., {_pillar_names_en[1]}]
• Initiative: [concrete action]
• Initiative: [concrete action]
• Initiative: [concrete action]

### Pillar 3: [Name — e.g., {_pillar_names_en[2]}]
• Initiative: [concrete action]
• Initiative: [concrete action]
• Initiative: [concrete action]

### Pillar 4: [Name — e.g., {_pillar_names_en[3]}]
• Initiative: [concrete action]
• Initiative: [concrete action]
• Initiative: [concrete action]

[SECTION]

## 3. Gap Analysis

### Identified Gaps:
| # | Gap | Description | Priority | Status |
|---|-----|-------------|----------|--------|
{gaps_table_en}

### Detailed Implementation Guidelines:

Provide a SEPARATE, COMPLETE implementation guide for EACH gap above. Do NOT write "repeat" or "same as above".

---
#### Gap #1 Implementation Guide: {_gap1_title_en if _has_no_org else "[First Gap Name]"}
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
{"| Planning | 1.1 | Define " + _org_struct_req_en + " per " + fw_short + " | Executive Management | Requirements document |" if _has_no_org else "| Planning | 1.1 | [Specific step] | [Team] | [Output] |"}
{"| Planning | 1.2 | Design proposed structure (" + _org_roles_en + ") | Executive Management | Proposed org chart |" if _has_no_org else "| Planning | 1.2 | [Specific step] | [Team] | [Output] |"}
| Execution | 2.1 | [Specific step] | [Team] | [Output] |
| Execution | 2.2 | [Specific step] | [Team] | [Output] |
| Verification | 3.1 | [Specific step] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

---
#### Gap #2 Implementation Guide: [Second Gap Name]
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Planning | 1.1 | [Specific step for Gap 2] | [Team] | [Output] |
| Planning | 1.2 | [Specific step for Gap 2] | [Team] | [Output] |
| Execution | 2.1 | [Specific step for Gap 2] | [Team] | [Output] |
| Execution | 2.2 | [Specific step for Gap 2] | [Team] | [Output] |
| Verification | 3.1 | [Specific step for Gap 2] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

---
#### Gap #3 Implementation Guide: [Third Gap Name]
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Planning | 1.1 | [Specific step for Gap 3] | [Team] | [Output] |
| Planning | 1.2 | [Specific step for Gap 3] | [Team] | [Output] |
| Execution | 2.1 | [Specific step for Gap 3] | [Team] | [Output] |
| Execution | 2.2 | [Specific step for Gap 3] | [Team] | [Output] |
| Verification | 3.1 | [Specific step for Gap 3] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

---
#### Gap #4 Implementation Guide: [Fourth Gap Name]
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Planning | 1.1 | [Specific step for Gap 4] | [Team] | [Output] |
| Planning | 1.2 | [Specific step for Gap 4] | [Team] | [Output] |
| Execution | 2.1 | [Specific step for Gap 4] | [Team] | [Output] |
| Execution | 2.2 | [Specific step for Gap 4] | [Team] | [Output] |
| Verification | 3.1 | [Specific step for Gap 4] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

---
#### Gap #5 Implementation Guide: [Fifth Gap Name]
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Planning | 1.1 | [Specific step for Gap 5] | [Team] | [Output] |
| Planning | 1.2 | [Specific step for Gap 5] | [Team] | [Output] |
| Execution | 2.1 | [Specific step for Gap 5] | [Team] | [Output] |
| Execution | 2.2 | [Specific step for Gap 5] | [Team] | [Output] |
| Verification | 3.1 | [Specific step for Gap 5] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

[SECTION]

## 4. Implementation Roadmap

### Phase 1: Foundation (0-6 months)
| # | Activity | Owner | Timeline | Deliverable |
|---|----------|-------|----------|-------------|
| 1 | [Activity 1] | [Owner] | Month 1-2 | [Deliverable] |
| 2 | [Activity 2] | [Owner] | Month 2-3 | [Deliverable] |
| 3 | [Activity 3] | [Owner] | Month 3-4 | [Deliverable] |
| 4 | [Activity 4] | [Owner] | Month 4-6 | [Deliverable] |

### Phase 2: Build (6-12 months)
| # | Activity | Owner | Timeline | Deliverable |
|---|----------|-------|----------|-------------|
| 1 | [Activity 1] | [Owner] | Month 6-7 | [Deliverable] |
| 2 | [Activity 2] | [Owner] | Month 7-9 | [Deliverable] |
| 3 | [Activity 3] | [Owner] | Month 9-10 | [Deliverable] |
| 4 | [Activity 4] | [Owner] | Month 10-12 | [Deliverable] |

### Phase 3: Optimize (12-24 months)
| # | Activity | Owner | Timeline | Deliverable |
|---|----------|-------|----------|-------------|
| 1 | [Activity 1] | [Owner] | Month 12-15 | [Deliverable] |
| 2 | [Activity 2] | [Owner] | Month 15-18 | [Deliverable] |
| 3 | [Activity 3] | [Owner] | Month 18-21 | [Deliverable] |
| 4 | [Activity 4] | [Owner] | Month 21-24 | [Deliverable] |

[SECTION]

## 5. Key Performance Indicators

### KPIs:
| # | KPI | Description | Current Value | Target Value | Justification | Timeframe | Data Source |
|---|-----|-------------|---------------|--------------|---------------|-----------|-------------|
| 1 | [KPI Name] | [Brief description] | To be assessed | [Justified value] | [Why this target] | Within X months | [Source] |
(8 KPIs total)

CRITICAL KPI RULES:
- Do NOT list individual controls as standalone KPIs (e.g. "MFA coverage" is wrong — MFA is one of many controls under Identity & Access Management)
- KPIs must measure at the CONTROL DOMAIN level (e.g. "Identity & Access Management implementation rate" which covers MFA, PAM, RBAC, access reviews together)
- Each KPI must be measurable with a clear formula

### KPI Assessment Guidelines

CRITICAL: Each KPI below requires a UNIQUE measurement guide in TABLE format. Do NOT repeat the same steps for different KPIs.
Bad example (generic): "Define scope → Identify sources → Collect data → Validate → Calculate" for every KPI
Good example (specific): For incident response time → "Configure SIEM timestamps → Extract incident ticket data → Calculate MTTD from alert-to-triage → Calculate MTTR from triage-to-containment"

Each guide MUST be formatted as a table with these columns:

---
#### KPI #1 Assessment Guide: [First KPI Name]
| Step | Action | Tool/System | Owner | Output |
|------|--------|-------------|-------|--------|
| 1 | [Specific data collection step] | [Exact tool name] | [Specific role] | [Concrete deliverable] |
| 2 | [Analysis/calculation step] | [Tool] | [Role] | [Deliverable] |
| 3 | [Validation step] | [Tool] | [Role] | [Deliverable] |
| 4 | [Reporting step] | [Tool] | [Role] | [Deliverable] |

**Formula:** [Exact calculation formula for this KPI]

---
#### KPI #2 Assessment Guide: [Second KPI Name]
| Step | Action | Tool/System | Owner | Output |
|------|--------|-------------|-------|--------|
(4-5 DIFFERENT rows — do NOT copy from KPI #1)

**Formula:** [Different formula]

---
#### KPI #3 Assessment Guide: [Third KPI Name]
| Step | Action | Tool/System | Owner | Output |
|------|--------|-------------|-------|--------|
(4-5 rows unique to this measurement)

**Formula:** [Formula]

---
#### KPI #4 Assessment Guide: [Fourth KPI Name]
| Step | Action | Tool/System | Owner | Output |
|------|--------|-------------|-------|--------|
(4-5 rows unique to this measurement)

**Formula:** [Formula]

---
#### KPI #5 Assessment Guide: [Fifth KPI Name]
| Step | Action | Tool/System | Owner | Output |
|------|--------|-------------|-------|--------|
(4-5 rows unique to this measurement)

**Formula:** [Formula]

---
#### KPI #6 Assessment Guide: [Sixth KPI Name]
| Step | Action | Tool/System | Owner | Output |
|------|--------|-------------|-------|--------|
(4-5 rows unique to this measurement)

**Formula:** [Formula]

---
#### KPI #7 Assessment Guide: [Seventh KPI Name]
| Step | Action | Tool/System | Owner | Output |
|------|--------|-------------|-------|--------|
(4-5 rows unique to this measurement)

**Formula:** [Formula]

---
#### KPI #8 Assessment Guide: [Eighth KPI Name]
| Step | Action | Tool/System | Owner | Output |
|------|--------|-------------|-------|--------|
(4-5 rows unique to this measurement)

**Formula:** [Formula]

[SECTION]

## 6. Confidence Assessment & Risks

**Confidence Score:** [X]%

**Score Justification:**
[Detailed paragraph explaining the basis for the confidence score, citing specific factors and their individual impact on the score]

### Critical Success Factors:
| # | Factor | Description | Importance |
|---|--------|-------------|------------|
(5 success factors linked to {fw_short} requirements)

### Key Risks:
| # | Risk | Likelihood | Impact | Mitigation Plan |
|---|------|------------|--------|-----------------|
| 1 | [Risk] | High/Medium/Low | High/Medium/Low | [Action] |
| 2 | [Risk] | High/Medium/Low | High/Medium/Low | [Action] |
| 3 | [Risk] | High/Medium/Low | High/Medium/Low | [Action] |
| 4 | [Risk] | High/Medium/Low | High/Medium/Low | [Action] |

---
**Preparation Date:** To be determined

IMPORTANT — HIDDEN SYNC BLOCK:
After the entire document, append a hidden block for system integration.
Format: <json_sync>{{"initiatives":[{{"title":"...","pillar":"..."}}],"gaps":[{{"name":"...","priority":"high","description":"..."}}],"kpis":[{{"name":"...","target":"...","timeframe":"..."}}],"risks":[{{"name":"...","likelihood":"...","impact":"..."}}]}}</json_sync>
Populate with the ACTUAL data from your strategy above. This block is machine-parsed and will be stripped from display."""

        content = generate_ai_content(prompt, lang, content_type='strategy')
        
        if not content:
            content = generate_simulation_content(prompt, lang)
        
        import re  # Import at function level to ensure availability
        import json as json_mod
        
        # ── SILENT PIPE: Parse and strip <json_sync> block ──
        json_sync_data = None
        json_sync_match = re.search(r'<json_sync>\s*(\{.*?\})\s*</json_sync>', content, re.DOTALL)
        if json_sync_match:
            try:
                json_sync_data = json_mod.loads(json_sync_match.group(1))
                print(f"✅ JSON_SYNC parsed: {list(json_sync_data.keys())}", flush=True)
            except Exception as jse:
                print(f"⚠ JSON_SYNC parse error (non-fatal): {jse}", flush=True)
            # Strip the block from displayed content
            content = content[:json_sync_match.start()].rstrip() + content[json_sync_match.end():]
        
        # ── SILENT INSERT from json_sync data ──
        if json_sync_data:
            sync_uid = session.get('user_id')
            sync_domain = data.get('domain', 'General')
            sync_counts = {'initiatives': 0, 'gaps': 0, 'kpis': 0, 'risks': 0}
            try:
                conn = get_db()
                for item in json_sync_data.get('initiatives', []):
                    if item.get('title') and len(item['title']) > 3:
                        conn.execute('''INSERT INTO project_tasks 
                            (user_id, domain, title, description, status, priority, category, source_type)
                            VALUES (?, ?, ?, ?, 'todo', 'high', 'strategic_initiative', 'ai_json_sync')''',
                            (sync_uid, sync_domain, item['title'][:200], item.get('pillar', '')[:500]))
                        sync_counts['initiatives'] += 1
                for item in json_sync_data.get('gaps', []):
                    if item.get('name') and len(item['name']) > 3:
                        pri = item.get('priority', 'medium').lower()
                        if pri not in ('critical', 'high', 'medium', 'low'):
                            pri = 'medium'
                        conn.execute('''INSERT INTO project_tasks 
                            (user_id, domain, title, description, status, priority, category, source_type)
                            VALUES (?, ?, ?, ?, 'todo', ?, 'gap_mitigation', 'ai_json_sync')''',
                            (sync_uid, sync_domain, item['name'][:200], item.get('description', '')[:500], pri))
                        sync_counts['gaps'] += 1
                for item in json_sync_data.get('kpis', []):
                    if item.get('name') and len(item['name']) > 3:
                        conn.execute('''INSERT INTO kpis 
                            (user_id, domain, name, description, target_value, timeframe, source_type)
                            VALUES (?, ?, ?, ?, ?, ?, 'ai_json_sync')''',
                            (sync_uid, sync_domain, item['name'][:200],
                             item.get('description', '')[:500],
                             item.get('target', 'TBD')[:100],
                             item.get('timeframe', 'TBD')[:100]))
                        sync_counts['kpis'] += 1
                        # Also create tracking task
                        conn.execute('''INSERT INTO project_tasks 
                            (user_id, domain, title, description, status, priority, category, source_type)
                            VALUES (?, ?, ?, ?, 'todo', 'medium', 'kpi_tracking', 'ai_json_sync')''',
                            (sync_uid, sync_domain, f"Track KPI: {item['name'][:150]}",
                             f"Target: {item.get('target', 'TBD')}\nTimeframe: {item.get('timeframe', 'TBD')}"))
                for item in json_sync_data.get('risks', []):
                    if item.get('name') and len(item['name']) > 3:
                        # Map text likelihood/impact to integer 1-5
                        _li_map = {'critical': 5, 'high': 4, 'medium': 3, 'low': 2, 'very low': 1}
                        _li = _li_map.get(str(item.get('likelihood', 'medium')).lower(), 3)
                        _im = _li_map.get(str(item.get('impact', 'medium')).lower(), 3)
                        conn.execute('''INSERT INTO risk_register 
                            (user_id, name, description, category, likelihood, impact,
                             treatment, status, domain, source_type)
                            VALUES (?, ?, ?, 'Strategic', ?, ?, 'Mitigate', 'Open', ?, 'ai_json_sync')''',
                            (sync_uid, item['name'][:200],
                             item.get('description', item['name'])[:500],
                             _li, _im, sync_domain))
                        sync_counts['risks'] += 1
                conn.commit()
                if any(v > 0 for v in sync_counts.values()):
                    log_action(sync_uid, 'json_sync_insert', {'domain': sync_domain, 'counts': sync_counts})
                    print(f"✅ JSON_SYNC inserted: {sync_counts}", flush=True)
            except Exception as sync_err:
                print(f"⚠ JSON_SYNC DB error (non-fatal): {sync_err}", flush=True)
        
        # Content is already cleaned by generate_ai_content via clean_ai_response()
        
        # DEBUG: Print content length and check for section markers
        print(f"DEBUG: Generated content length: {len(content)}", flush=True)
        print(f"DEBUG: Contains [SECTION]: {'[SECTION]' in content}", flush=True)
        print(f"DEBUG: Contains ## 4.: {'## 4.' in content}", flush=True)
        print(f"DEBUG: Contains ## 5.: {'## 5.' in content}", flush=True)
        print(f"DEBUG: Contains ## 6.: {'## 6.' in content}", flush=True)
        print(f"DEBUG: Contains خارطة الطريق: {'خارطة الطريق' in content}", flush=True)
        print(f"DEBUG: Contains مؤشرات الأداء: {'مؤشرات الأداء' in content}", flush=True)
        
        # Parse sections - split by separator
        parts = []
        
        if '[SECTION]' in content:
            parts = content.split('[SECTION]')
            print(f"DEBUG: Split by [SECTION], got {len(parts)} parts", flush=True)
        else:
            # No [SECTION] markers - must split by section headers
            # First, try to split by ## X. pattern (works for both English and Arabic)
            section_split_pattern = r'(?=##\s*[1-6]\.)'
            parts = re.split(section_split_pattern, content)
            parts = [p for p in parts if p.strip()]
            print(f"DEBUG: Split by ## X. pattern, got {len(parts)} parts", flush=True)
            
            if len(parts) < 4:
                # Try splitting by Arabic section titles directly
                ar_section_titles = [
                    'الرؤية والأهداف',
                    'تحليل الفجوات', 
                    'الركائز الاستراتيجية',
                    'خارطة الطريق',
                    'مؤشرات الأداء',
                    'تقييم الثقة',
                    'تقييم الجاهزية'
                ]
                # Build pattern to match any of these titles
                ar_pattern = r'(?=##\s*\d*\.?\s*(?:' + '|'.join(ar_section_titles) + '))'
                parts = re.split(ar_pattern, content)
                parts = [p for p in parts if p.strip()]
                print(f"DEBUG: Split by Arabic titles, got {len(parts)} parts", flush=True)
            
            if len(parts) < 4:
                # Try splitting by English section titles
                en_section_titles = [
                    'Vision',
                    'Gap Analysis',
                    'Strategic Pillars',
                    'Implementation Roadmap',
                    'Key Performance',
                    'Confidence Assessment'
                ]
                en_pattern = r'(?=##\s*\d*\.?\s*(?:' + '|'.join(en_section_titles) + '))'
                parts = re.split(en_pattern, content)
                parts = [p for p in parts if p.strip()]
                print(f"DEBUG: Split by English titles, got {len(parts)} parts", flush=True)
        
        # If still not enough parts, try splitting by numbered headers without ##
        if len(parts) < 4:
            # Try matching "1. الرؤية" or "1. Vision" at start of line
            section_split_pattern = r'(?=^[1-6]\.\s*(?:الرؤية|تحليل|الركائز|خارطة|مؤشرات|تقييم|Vision|Gap|Strategic|Implementation|Key|Confidence))'
            parts = re.split(section_split_pattern, content, flags=re.MULTILINE | re.IGNORECASE)
            parts = [p for p in parts if p.strip()]
            print(f"DEBUG: Split by numbered pattern, got {len(parts)} parts", flush=True)
        
        # Last resort: split by any "## " followed by a number
        if len(parts) < 4:
            parts = re.split(r'(?=##\s+\d)', content)
            parts = [p for p in parts if p.strip()]
            print(f"DEBUG: Split by ## + number, got {len(parts)} parts", flush=True)
        
        # Clean parts
        parts = [p.strip() for p in parts if p.strip()]
        print(f"DEBUG: Final number of parts after cleaning: {len(parts)}", flush=True)
        
        # Print first 100 chars of each part for debugging
        for i, part in enumerate(parts[:6]):
            print(f"DEBUG: Part {i} starts with: {part[:100]}", flush=True)
        
        def fix_formatting(text, lang_code):
            """Fix markdown formatting - add ### before tables and ## before section headers."""
            import re
            lines = text.split('\n')
            fixed_lines = []
            
            for i, line in enumerate(lines):
                stripped = line.strip()
                
                # Skip empty lines
                if not stripped:
                    fixed_lines.append(line)
                    continue
                
                # Skip if already has ## or ###
                if stripped.startswith('##'):
                    fixed_lines.append(line)
                    continue
                
                # Check if this is a main section header (like "1. Vision & Objectives")
                if re.match(r'^[1-6]\.\s+\w', stripped):
                    # Add ## before the section number
                    fixed_lines.append('## ' + stripped)
                    continue
                
                # Check if this line ends with : and next non-empty line is a table
                if stripped.endswith(':') and not stripped.startswith('**'):
                    # Look ahead for table
                    next_line_idx = i + 1
                    while next_line_idx < len(lines) and not lines[next_line_idx].strip():
                        next_line_idx += 1
                    
                    if next_line_idx < len(lines):
                        next_line = lines[next_line_idx].strip()
                        # Check if next line is a table header (starts with |)
                        if next_line.startswith('|'):
                            # This is a table header without ###, add it
                            fixed_lines.append('### ' + stripped)
                            continue
                
                fixed_lines.append(line)
            
            return '\n'.join(fixed_lines)
        
        # Apply fix to each part
        print("=" * 60, flush=True)
        print("BEFORE fix_formatting - First part preview:", flush=True)
        if parts:
            print(parts[0][:150], flush=True)
        print("=" * 60, flush=True)
        
        parts = [fix_formatting(p, lang) for p in parts]
        
        print("AFTER fix_formatting - First part preview:", flush=True)
        if parts:
            print(parts[0][:150], flush=True)
        print("=" * 60, flush=True)
        
        # Section header patterns to identify each section
        section_patterns = {
            'vision': [
                '1. vision', '## 1.', '1.', 'vision & objective', 'vision and objective',
                '1. الرؤية', '## 1.', 'الرؤية والأهداف'
            ],
            'gaps': [
                '2. gap', '## 2.', '2.', 'gap analysis',
                '2. تحليل', '## 2.', 'تحليل الفجوات'
            ],
            'pillars': [
                '3. strategic', '## 3.', '3.', 'strategic pillar', 'pillar 1',
                '3. الركائز', '## 3.', 'الركائز الاستراتيجية'
            ],
            'roadmap': [
                '4. implementation', '## 4.', '4.', 'roadmap', 'phase 1 (0-6',
                '4. خارطة', '## 4.', 'خارطة الطريق', 'المرحلة 1'
            ],
            'kpis': [
                '5. key performance', '## 5.', '5.', 'kpi', 'key performance indicator',
                '5. مؤشرات', '## 5.', 'مؤشرات الأداء'
            ],
            'confidence': [
                '6. confidence', '## 6.', '6.', 'confidence assessment', 'confidence score',
                '6. تقييم الثقة', '## 6.', 'تقييم الثقة', 'درجة الثقة'
            ]
        }
        
        def identify_section(text, lang_code):
            """Identify which section type this text belongs to."""
            text_lower = text.lower()[:300]  # Check first 300 chars
            first_line = text.strip().split('\n')[0] if text.strip() else ''
            first_line_lower = first_line.lower()
            
            # MOST RELIABLE: Match "## N." at the start of the first line
            import re as _re
            num_match = _re.match(r'^##\s*(\d)\.', first_line_lower)
            if num_match:
                num = num_match.group(1)
                return {'1': 'vision', '2': 'pillars', '3': 'gaps', '4': 'roadmap', '5': 'kpis', '6': 'confidence'}.get(num)
            
            # SECOND: Match "N. Title" at the start (no ##)
            num_match = _re.match(r'^(\d)\.\s+\w', first_line_lower)
            if num_match:
                num = num_match.group(1)
                return {'1': 'vision', '2': 'pillars', '3': 'gaps', '4': 'roadmap', '5': 'kpis', '6': 'confidence'}.get(num)
            
            # THIRD: Arabic section titles (without numbers)
            if 'الرؤية' in first_line and 'الأهداف' in first_line:
                return 'vision'
            if 'تحليل الفجوات' in first_line or ('الفجوات' in first_line and 'تحليل' in first_line):
                return 'gaps'
            if 'الركائز الاستراتيجية' in first_line or 'الركائز' in first_line:
                return 'pillars'
            if 'خارطة الطريق' in first_line:
                return 'roadmap'
            if 'مؤشرات الأداء' in first_line:
                return 'kpis'
            if 'تقييم الثقة' in first_line or 'درجة الثقة' in first_line or 'تقييم الجاهزية' in first_line:
                return 'confidence'
            
            # FOURTH: English titles (without numbers) — must be in first line only
            if 'vision' in first_line_lower and 'objective' in first_line_lower:
                return 'vision'
            if 'gap analysis' in first_line_lower or ('gap' in first_line_lower and 'analy' in first_line_lower):
                return 'gaps'
            if 'pillar' in first_line_lower or 'strategic pillar' in first_line_lower:
                return 'pillars'
            if 'roadmap' in first_line_lower or 'implementation' in first_line_lower:
                return 'roadmap'
            if 'kpi' in first_line_lower or 'key performance' in first_line_lower:
                return 'kpis'
            if 'confidence' in first_line_lower and ('assessment' in first_line_lower or 'score' in first_line_lower):
                return 'confidence'
            
            # LAST RESORT: Keyword scoring in first 300 chars (weighted)
            keyword_scores = {
                'vision': [('vision statement', 3), ('desired future state', 3), ('strategic objective', 2), ('objective', 1), ('mission', 1), ('الرؤية التنفيذية', 3), ('الأهداف الاستراتيجية', 2)],
                'gaps': [('gap analysis', 3), ('identified gaps', 3), ('current gap', 2), ('gap #', 2), ('الفجوة', 2), ('الفجوات', 2), ('تحليل', 1)],
                'pillars': [('pillar 1', 3), ('pillar 2', 3), ('strategic pillar', 3), ('initiative', 2), ('workstream', 2), ('الركيزة', 3), ('المبادرات', 2), ('الركائز', 3)],
                'roadmap': [('phase 1', 3), ('phase 2', 2), ('roadmap', 3), ('timeline', 2), ('implementation', 1), ('المرحلة', 2), ('خارطة الطريق', 3)],
                'kpis': [('kpi', 3), ('key performance indicator', 3), ('metric', 2), ('target', 1), ('مؤشر', 2), ('مؤشرات الأداء', 3)],
                'confidence': [('confidence score', 3), ('confidence assessment', 3), ('overall confidence', 3), ('الثقة', 2), ('تقييم الثقة', 3), ('درجة الثقة', 3), ('تقييم الجاهزية', 3), ('الجاهزية', 2)]
            }
            
            scores = {}
            for section_type, keywords in keyword_scores.items():
                score = sum(weight for kw, weight in keywords if kw in text_lower)
                scores[section_type] = score
            
            if max(scores.values()) > 0:
                return max(scores, key=scores.get)
            return None
        
        # Initialize sections
        sections = {
            'vision': '',
            'gaps': '',
            'pillars': '',
            'roadmap': '',
            'kpis': '',
            'confidence': ''
        }
        
        def inject_implementation_guidelines(gaps_content, lang_code):
            """Inject implementation guidelines if AI didn't include them or if some are missing."""
            # Extract gaps from the table to create guidelines for each
            import re
            
            # Find gap names from table rows (pattern: | number | gap name | description |)
            gap_pattern = r'\|\s*(\d+)\s*\|\s*([^|]+)\s*\|'
            gaps_found = re.findall(gap_pattern, gaps_content)
            
            if not gaps_found or len(gaps_found) < 2:  # Need at least 2 (header row doesn't count)
                return gaps_content
            
            # Skip header row if it contains "Gap" or "الفجوة"
            gaps_list = []
            for num, name in gaps_found:
                name = name.strip()
                if name and name.lower() not in ['gap', 'الفجوة', '#', 'description', 'الوصف', 'priority', 'الأولوية', 'status', 'الحالة']:
                    gaps_list.append((num, name))
            
            if not gaps_list:
                return gaps_content
            
            # Gap-specific implementation templates (English)
            gap_templates_en = {
                'governance': {
                    'steps': [
                        ('Planning', '1.1', 'Conduct organizational assessment and define office mandate', 'Executive Sponsor', 'Mandate document'),
                        ('Planning', '1.2', 'Design organizational structure with roles (Director, Managers, Analysts)', 'HR/Strategy', 'Org chart draft'),
                        ('Planning', '1.3', 'Define RACI matrix for all governance functions', 'Project Lead', 'RACI matrix'),
                        ('Planning', '1.4', 'Develop job descriptions and competency requirements', 'HR', 'Job descriptions'),
                        ('Approval', '2.1', 'Obtain executive approval for structure and headcount', 'Executive', 'Approved structure'),
                        ('Approval', '2.2', 'Secure budget allocation for staffing and operations', 'Finance', 'Budget approval'),
                        ('Staffing', '3.1', 'Recruit/appoint office director and key positions', 'HR', 'Appointment letters'),
                        ('Staffing', '3.2', 'Onboard staff with training on mandate and frameworks', 'Training', 'Onboarding records'),
                        ('Operationalize', '4.1', 'Establish operating procedures, reporting lines, and meeting cadence', 'Office Director', 'Operating manual'),
                        ('Operationalize', '4.2', 'Launch first quarterly governance review cycle', 'Office Director', 'First review report'),
                    ],
                    'evidence': ['Approved org chart', 'Signed appointment letters', 'RACI matrix', 'First quarterly governance report']
                },
                'data_governance': {
                    'steps': [
                        ('Planning', '1.1', 'Assess current data management maturity (DCAM/DAMA-DMBOK)', 'Data Lead', 'Maturity assessment'),
                        ('Planning', '1.2', 'Define data governance framework scope and principles', 'Data Governance', 'Governance charter'),
                        ('Planning', '1.3', 'Identify data domains and assign data owners/stewards', 'Data Governance', 'Data ownership matrix'),
                        ('Framework', '2.1', 'Develop data standards (naming, quality, metadata)', 'Data Governance', 'Data standards doc'),
                        ('Framework', '2.2', 'Build enterprise data catalog and glossary', 'Data Mgmt', 'Data catalog'),
                        ('Framework', '2.3', 'Define data quality KPIs and measurement process', 'Data Quality', 'DQ scorecard template'),
                        ('Implementation', '3.1', 'Implement data governance tooling (catalog, lineage, quality)', 'IT/Data', 'Configured tools'),
                        ('Implementation', '3.2', 'Train data stewards and owners on responsibilities', 'Training', 'Training records'),
                        ('Operations', '4.1', 'Launch data governance council with regular meetings', 'Data Governance', 'Council minutes'),
                        ('Operations', '4.2', 'Execute first data quality assessment cycle', 'Data Quality', 'DQ assessment report'),
                    ],
                    'evidence': ['Data governance charter', 'Data ownership matrix', 'Enterprise data catalog', 'Data quality assessment']
                },
                'risk_management': {
                    'steps': [
                        ('Planning', '1.1', 'Review existing risk management capabilities and gaps', 'Risk Team', 'Gap assessment'),
                        ('Planning', '1.2', 'Define risk appetite and tolerance levels with leadership', 'Executive/Risk', 'Risk appetite statement'),
                        ('Framework', '2.1', 'Develop risk management framework aligned to ISO 31000/COSO', 'Risk Team', 'Risk framework'),
                        ('Framework', '2.2', 'Design risk assessment methodology (criteria, scales, scoring)', 'Risk Team', 'Assessment methodology'),
                        ('Framework', '2.3', 'Create risk register template and categorization scheme', 'Risk Team', 'Risk register template'),
                        ('Implementation', '3.1', 'Conduct initial enterprise-wide risk assessment', 'Risk Team', 'Risk assessment results'),
                        ('Implementation', '3.2', 'Develop risk treatment plans for critical/high risks', 'Risk Owners', 'Treatment plans'),
                        ('Implementation', '3.3', 'Implement risk monitoring tools and dashboards', 'IT/Risk', 'Risk dashboard'),
                        ('Operations', '4.1', 'Establish risk reporting cadence to board/executives', 'Risk Team', 'Risk report template'),
                        ('Operations', '4.2', 'Conduct first periodic risk review and update cycle', 'Risk Team', 'Updated risk register'),
                    ],
                    'evidence': ['Approved risk appetite statement', 'Risk framework document', 'Populated risk register', 'Board risk report']
                },
                'compliance_framework': {
                    'steps': [
                        ('Planning', '1.1', 'Identify all applicable laws, regulations, and standards', 'Compliance', 'Regulatory inventory'),
                        ('Planning', '1.2', 'Map regulatory requirements to organizational controls', 'Compliance', 'Requirements mapping'),
                        ('Development', '2.1', 'Develop compliance program charter and governance', 'Compliance', 'Program charter'),
                        ('Development', '2.2', 'Design compliance monitoring and testing plan', 'Compliance', 'Monitoring plan'),
                        ('Development', '2.3', 'Create compliance training curriculum', 'Training', 'Training materials'),
                        ('Implementation', '3.1', 'Deploy compliance management system/tool', 'IT', 'Configured system'),
                        ('Implementation', '3.2', 'Train compliance liaisons across departments', 'Training', 'Trained liaisons'),
                        ('Implementation', '3.3', 'Execute first compliance self-assessment', 'Compliance', 'Self-assessment results'),
                        ('Operations', '4.1', 'Establish regulatory change management process', 'Compliance', 'Change tracking process'),
                        ('Operations', '4.2', 'Generate first compliance status report to leadership', 'Compliance', 'Compliance dashboard'),
                    ],
                    'evidence': ['Regulatory inventory', 'Compliance program charter', 'Self-assessment results', 'Compliance status report']
                },
                'policy': {
                    'steps': [
                        ('Planning', '1.1', 'Inventory existing policies and identify gaps', 'InfoSec', 'Policy inventory'),
                        ('Planning', '1.2', 'Review regulatory requirements (NCA, ISO, NIST)', 'Compliance', 'Requirements matrix'),
                        ('Planning', '1.3', 'Define policy structure and approval workflow', 'InfoSec', 'Policy framework'),
                        ('Development', '2.1', 'Draft new/updated policies', 'InfoSec', 'Policy drafts'),
                        ('Development', '2.2', 'Legal and stakeholder review', 'Legal', 'Review comments'),
                        ('Development', '2.3', 'Incorporate feedback and finalize', 'InfoSec', 'Final policies'),
                        ('Approval', '3.1', 'Submit for executive approval', 'InfoSec', 'Approval request'),
                        ('Approval', '3.2', 'Obtain sign-off from leadership', 'Executive', 'Signed policies'),
                        ('Deployment', '4.1', 'Communicate policies to all staff', 'HR/Comms', 'Communication log'),
                        ('Deployment', '4.2', 'Conduct policy awareness training', 'Training', 'Training records'),
                    ],
                    'evidence': ['Approved policy documents', 'Review meeting minutes', 'Training completion records', 'Communication acknowledgments']
                },
                'technology': {
                    'steps': [
                        ('Planning', '1.1', 'Assess current technology landscape', 'IT', 'Technology inventory'),
                        ('Planning', '1.2', 'Define technical requirements', 'Security', 'Requirements doc'),
                        ('Planning', '1.3', 'Evaluate vendor solutions (RFP/RFI)', 'Procurement', 'Vendor comparison'),
                        ('Procurement', '2.1', 'Select vendor and negotiate contract', 'Procurement', 'Signed contract'),
                        ('Procurement', '2.2', 'Allocate budget and resources', 'Finance', 'Budget approval'),
                        ('Implementation', '3.1', 'Prepare infrastructure and environment', 'IT', 'Ready environment'),
                        ('Implementation', '3.2', 'Install and configure solution', 'IT/Vendor', 'Configured system'),
                        ('Implementation', '3.3', 'Integrate with existing systems', 'IT', 'Integration complete'),
                        ('Testing', '4.1', 'Conduct UAT and security testing', 'QA', 'Test results'),
                        ('Operations', '5.1', 'Train operations team', 'Training', 'Trained staff'),
                    ],
                    'evidence': ['Vendor contract', 'Installation report', 'Integration test results', 'Training certificates']
                },
                'training': {
                    'steps': [
                        ('Planning', '1.1', 'Assess current awareness levels (baseline survey)', 'InfoSec', 'Baseline report'),
                        ('Planning', '1.2', 'Identify target audiences and learning objectives', 'Training', 'Training plan'),
                        ('Planning', '1.3', 'Select training delivery methods', 'Training', 'Delivery strategy'),
                        ('Development', '2.1', 'Develop training content and materials', 'Training', 'Training modules'),
                        ('Development', '2.2', 'Create assessments and quizzes', 'Training', 'Assessment bank'),
                        ('Development', '2.3', 'Prepare phishing simulation scenarios', 'Security', 'Simulation plan'),
                        ('Execution', '3.1', 'Launch mandatory training program', 'HR', 'Training schedule'),
                        ('Execution', '3.2', 'Conduct phishing simulations', 'Security', 'Simulation results'),
                        ('Execution', '3.3', 'Provide remedial training for failures', 'Training', 'Remediation log'),
                        ('Monitoring', '4.1', 'Track completion rates and scores', 'Training', 'Progress dashboard'),
                    ],
                    'evidence': ['Training completion records', 'Assessment scores', 'Phishing simulation results', 'Improvement metrics']
                },
                'incident': {
                    'steps': [
                        ('Planning', '1.1', 'Review current incident response capabilities', 'Security', 'Gap assessment'),
                        ('Planning', '1.2', 'Define incident categories and severity levels', 'Security', 'Classification matrix'),
                        ('Development', '2.1', 'Develop incident response playbooks', 'Security', 'Playbook documents'),
                        ('Development', '2.2', 'Create escalation procedures and contact lists', 'Security', 'Escalation matrix'),
                        ('Development', '2.3', 'Design reporting templates', 'Security', 'Report templates'),
                        ('Team Setup', '3.1', 'Form CSIRT/IRT team', 'Management', 'Team charter'),
                        ('Team Setup', '3.2', 'Define roles and responsibilities (RACI)', 'Security', 'RACI matrix'),
                        ('Training', '4.1', 'Train incident response team', 'Training', 'Trained team'),
                        ('Testing', '5.1', 'Conduct tabletop exercises', 'Security', 'Exercise report'),
                        ('Testing', '5.2', 'Perform full simulation drill', 'Security', 'Drill results'),
                    ],
                    'evidence': ['Approved IR plan', 'CSIRT roster', 'Tabletop exercise report', 'Drill after-action report']
                },
                'data': {
                    'steps': [
                        ('Planning', '1.1', 'Inventory all data assets', 'Data Mgmt', 'Data inventory'),
                        ('Planning', '1.2', 'Classify data by sensitivity level', 'InfoSec', 'Classification scheme'),
                        ('Planning', '1.3', 'Map data flows and storage locations', 'IT', 'Data flow diagram'),
                        ('Implementation', '2.1', 'Apply data classification labels', 'Data Mgmt', 'Labeled data'),
                        ('Implementation', '2.2', 'Implement DLP controls', 'Security', 'DLP configured'),
                        ('Implementation', '2.3', 'Enable encryption for sensitive data', 'IT', 'Encryption report'),
                        ('Implementation', '2.4', 'Configure access controls (need-to-know)', 'IT', 'Access matrix'),
                        ('Monitoring', '3.1', 'Deploy data monitoring tools', 'Security', 'Monitoring active'),
                        ('Monitoring', '3.2', 'Establish data breach response procedures', 'Security', 'Breach procedures'),
                        ('Compliance', '4.1', 'Verify PDPL/GDPR compliance', 'Compliance', 'Compliance report'),
                    ],
                    'evidence': ['Data classification register', 'DLP policy configuration', 'Encryption certificates', 'Access control matrix']
                },
                'default': {
                    'steps': [
                        ('Planning', '1.1', 'Assess current state and identify requirements', 'Project Lead', 'Assessment report'),
                        ('Planning', '1.2', 'Define scope and success criteria', 'Project Lead', 'Project charter'),
                        ('Planning', '1.3', 'Develop implementation timeline', 'Project Lead', 'Project plan'),
                        ('Execution', '2.1', 'Allocate resources and budget', 'Management', 'Resource allocation'),
                        ('Execution', '2.2', 'Execute implementation activities', 'Implementation Team', 'Progress reports'),
                        ('Execution', '2.3', 'Train relevant personnel', 'Training', 'Training records'),
                        ('Verification', '3.1', 'Test and validate implementation', 'QA', 'Test results'),
                        ('Verification', '3.2', 'Conduct management review', 'Management', 'Review minutes'),
                        ('Closure', '4.1', 'Document lessons learned', 'Project Lead', 'Lessons learned'),
                        ('Closure', '4.2', 'Obtain formal sign-off', 'Management', 'Sign-off document'),
                    ],
                    'evidence': ['Assessment report', 'Implementation records', 'Test results', 'Sign-off document']
                }
            }
            
            # Gap-specific implementation templates (Arabic)
            gap_templates_ar = {
                'governance': {
                    'steps': [
                        ('التخطيط', '1.1', 'إجراء تقييم تنظيمي وتحديد صلاحيات المكتب/الإدارة', 'الراعي التنفيذي', 'وثيقة الصلاحيات'),
                        ('التخطيط', '1.2', 'تصميم الهيكل التنظيمي مع الأدوار (مدير، مشرفون، محللون)', 'الموارد البشرية/الاستراتيجية', 'مسودة الهيكل التنظيمي'),
                        ('التخطيط', '1.3', 'تحديد مصفوفة المسؤوليات (RACI) لجميع وظائف الحوكمة', 'قائد المشروع', 'مصفوفة RACI'),
                        ('التخطيط', '1.4', 'إعداد الأوصاف الوظيفية ومتطلبات الكفاءة', 'الموارد البشرية', 'الأوصاف الوظيفية'),
                        ('الاعتماد', '2.1', 'الحصول على اعتماد الإدارة التنفيذية للهيكل والملاك الوظيفي', 'الإدارة التنفيذية', 'الهيكل المعتمد'),
                        ('الاعتماد', '2.2', 'تأمين تخصيص الميزانية للتوظيف والعمليات', 'المالية', 'اعتماد الميزانية'),
                        ('التوظيف', '3.1', 'استقطاب/تعيين مدير المكتب والمناصب الرئيسية', 'الموارد البشرية', 'خطابات التعيين'),
                        ('التوظيف', '3.2', 'تهيئة الموظفين وتدريبهم على الصلاحيات والأطر', 'التدريب', 'سجلات التهيئة'),
                        ('التشغيل', '4.1', 'وضع إجراءات التشغيل وخطوط التقارير والاجتماعات الدورية', 'مدير المكتب', 'دليل التشغيل'),
                        ('التشغيل', '4.2', 'إطلاق أول دورة مراجعة حوكمة ربع سنوية', 'مدير المكتب', 'تقرير المراجعة الأول'),
                    ],
                    'evidence': ['الهيكل التنظيمي المعتمد', 'خطابات التعيين الموقعة', 'مصفوفة RACI', 'تقرير المراجعة الربعية الأول']
                },
                'data_governance': {
                    'steps': [
                        ('التخطيط', '1.1', 'تقييم نضج إدارة البيانات الحالي (DCAM/DAMA-DMBOK)', 'قائد البيانات', 'تقييم النضج'),
                        ('التخطيط', '1.2', 'تحديد نطاق ومبادئ إطار حوكمة البيانات', 'حوكمة البيانات', 'ميثاق الحوكمة'),
                        ('التخطيط', '1.3', 'تحديد نطاقات البيانات وتعيين ملاك وأمناء البيانات', 'حوكمة البيانات', 'مصفوفة ملكية البيانات'),
                        ('الإطار', '2.1', 'تطوير معايير البيانات (التسمية، الجودة، البيانات الوصفية)', 'حوكمة البيانات', 'وثيقة معايير البيانات'),
                        ('الإطار', '2.2', 'بناء كتالوج البيانات المؤسسي والمسرد', 'إدارة البيانات', 'كتالوج البيانات'),
                        ('الإطار', '2.3', 'تحديد مؤشرات جودة البيانات وآلية القياس', 'جودة البيانات', 'نموذج بطاقة الأداء'),
                        ('التنفيذ', '3.1', 'تطبيق أدوات حوكمة البيانات (كتالوج، تتبع، جودة)', 'تقنية المعلومات/البيانات', 'الأدوات المُكوّنة'),
                        ('التنفيذ', '3.2', 'تدريب أمناء وملاك البيانات على مسؤولياتهم', 'التدريب', 'سجلات التدريب'),
                        ('التشغيل', '4.1', 'إطلاق مجلس حوكمة البيانات باجتماعات دورية', 'حوكمة البيانات', 'محاضر المجلس'),
                        ('التشغيل', '4.2', 'تنفيذ أول دورة تقييم لجودة البيانات', 'جودة البيانات', 'تقرير تقييم الجودة'),
                    ],
                    'evidence': ['ميثاق حوكمة البيانات', 'مصفوفة ملكية البيانات', 'كتالوج البيانات المؤسسي', 'تقرير تقييم جودة البيانات']
                },
                'risk_management': {
                    'steps': [
                        ('التخطيط', '1.1', 'مراجعة قدرات إدارة المخاطر الحالية والفجوات', 'فريق المخاطر', 'تقييم الفجوات'),
                        ('التخطيط', '1.2', 'تحديد شهية المخاطر ومستويات التحمل مع القيادة', 'الإدارة التنفيذية/المخاطر', 'بيان شهية المخاطر'),
                        ('الإطار', '2.1', 'تطوير إطار إدارة المخاطر المتوافق مع ISO 31000/COSO', 'فريق المخاطر', 'إطار المخاطر'),
                        ('الإطار', '2.2', 'تصميم منهجية تقييم المخاطر (المعايير، المقاييس، التسجيل)', 'فريق المخاطر', 'منهجية التقييم'),
                        ('الإطار', '2.3', 'إنشاء نموذج سجل المخاطر وآلية التصنيف', 'فريق المخاطر', 'نموذج سجل المخاطر'),
                        ('التنفيذ', '3.1', 'إجراء أول تقييم شامل للمخاطر على مستوى المنظمة', 'فريق المخاطر', 'نتائج التقييم'),
                        ('التنفيذ', '3.2', 'وضع خطط معالجة المخاطر الحرجة والعالية', 'ملاك المخاطر', 'خطط المعالجة'),
                        ('التنفيذ', '3.3', 'تطبيق أدوات مراقبة المخاطر ولوحات المعلومات', 'تقنية المعلومات/المخاطر', 'لوحة معلومات المخاطر'),
                        ('التشغيل', '4.1', 'تأسيس دورية تقارير المخاطر للمجلس/الإدارة', 'فريق المخاطر', 'نموذج تقرير المخاطر'),
                        ('التشغيل', '4.2', 'تنفيذ أول دورة مراجعة وتحديث دورية للمخاطر', 'فريق المخاطر', 'سجل المخاطر المحدث'),
                    ],
                    'evidence': ['بيان شهية المخاطر المعتمد', 'وثيقة إطار المخاطر', 'سجل المخاطر المُعبأ', 'تقرير المخاطر للمجلس']
                },
                'compliance_framework': {
                    'steps': [
                        ('التخطيط', '1.1', 'تحديد جميع القوانين واللوائح والمعايير المنطبقة', 'الامتثال', 'جرد التنظيمات'),
                        ('التخطيط', '1.2', 'ربط المتطلبات التنظيمية بضوابط المنظمة', 'الامتثال', 'خريطة المتطلبات'),
                        ('التطوير', '2.1', 'وضع ميثاق وحوكمة برنامج الامتثال', 'الامتثال', 'ميثاق البرنامج'),
                        ('التطوير', '2.2', 'تصميم خطة المراقبة والاختبار للامتثال', 'الامتثال', 'خطة المراقبة'),
                        ('التطوير', '2.3', 'إعداد المنهج التدريبي للامتثال', 'التدريب', 'المواد التدريبية'),
                        ('التنفيذ', '3.1', 'نشر نظام/أداة إدارة الامتثال', 'تقنية المعلومات', 'النظام المُكوّن'),
                        ('التنفيذ', '3.2', 'تدريب منسقي الامتثال في الإدارات', 'التدريب', 'المنسقون المدربون'),
                        ('التنفيذ', '3.3', 'تنفيذ أول تقييم ذاتي للامتثال', 'الامتثال', 'نتائج التقييم الذاتي'),
                        ('التشغيل', '4.1', 'وضع آلية إدارة التغييرات التنظيمية', 'الامتثال', 'آلية تتبع التغييرات'),
                        ('التشغيل', '4.2', 'إصدار أول تقرير حالة الامتثال للقيادة', 'الامتثال', 'لوحة معلومات الامتثال'),
                    ],
                    'evidence': ['جرد التنظيمات', 'ميثاق برنامج الامتثال', 'نتائج التقييم الذاتي', 'تقرير حالة الامتثال']
                },
                'policy': {
                    'steps': [
                        ('التخطيط', '1.1', 'جرد السياسات الحالية وتحديد الفجوات', 'أمن المعلومات', 'قائمة السياسات'),
                        ('التخطيط', '1.2', 'مراجعة المتطلبات التنظيمية (NCA, ISO, NIST)', 'الامتثال', 'مصفوفة المتطلبات'),
                        ('التخطيط', '1.3', 'تحديد هيكل السياسات وسير الاعتماد', 'أمن المعلومات', 'إطار السياسات'),
                        ('التطوير', '2.1', 'صياغة السياسات الجديدة/المحدثة', 'أمن المعلومات', 'مسودات السياسات'),
                        ('التطوير', '2.2', 'مراجعة قانونية ومن أصحاب المصلحة', 'الشؤون القانونية', 'تعليقات المراجعة'),
                        ('التطوير', '2.3', 'دمج الملاحظات والصياغة النهائية', 'أمن المعلومات', 'السياسات النهائية'),
                        ('الاعتماد', '3.1', 'رفع للاعتماد التنفيذي', 'أمن المعلومات', 'طلب الاعتماد'),
                        ('الاعتماد', '3.2', 'الحصول على توقيع القيادة', 'الإدارة التنفيذية', 'السياسات الموقعة'),
                        ('النشر', '4.1', 'تعميم السياسات على جميع الموظفين', 'الموارد البشرية', 'سجل التعميم'),
                        ('النشر', '4.2', 'تنفيذ تدريب التوعية بالسياسات', 'التدريب', 'سجلات التدريب'),
                    ],
                    'evidence': ['السياسات المعتمدة', 'محاضر اجتماعات المراجعة', 'سجلات إتمام التدريب', 'إقرارات الاستلام']
                },
                'technology': {
                    'steps': [
                        ('التخطيط', '1.1', 'تقييم المشهد التقني الحالي', 'تقنية المعلومات', 'جرد التقنيات'),
                        ('التخطيط', '1.2', 'تحديد المتطلبات التقنية', 'فريق الأمن', 'وثيقة المتطلبات'),
                        ('التخطيط', '1.3', 'تقييم حلول الموردين (RFP/RFI)', 'المشتريات', 'مقارنة الموردين'),
                        ('الشراء', '2.1', 'اختيار المورد والتفاوض على العقد', 'المشتريات', 'عقد موقع'),
                        ('الشراء', '2.2', 'تخصيص الميزانية والموارد', 'المالية', 'اعتماد الميزانية'),
                        ('التنفيذ', '3.1', 'تجهيز البنية التحتية والبيئة', 'تقنية المعلومات', 'بيئة جاهزة'),
                        ('التنفيذ', '3.2', 'تثبيت وتكوين الحل', 'تقنية المعلومات/المورد', 'نظام مُكوّن'),
                        ('التنفيذ', '3.3', 'الدمج مع الأنظمة الحالية', 'تقنية المعلومات', 'اكتمال التكامل'),
                        ('الاختبار', '4.1', 'إجراء اختبار القبول والأمان', 'ضمان الجودة', 'نتائج الاختبار'),
                        ('التشغيل', '5.1', 'تدريب فريق التشغيل', 'التدريب', 'فريق مدرب'),
                    ],
                    'evidence': ['عقد المورد', 'تقرير التثبيت', 'نتائج اختبار التكامل', 'شهادات التدريب']
                },
                'training': {
                    'steps': [
                        ('التخطيط', '1.1', 'تقييم مستوى الوعي الحالي (استبيان أساسي)', 'أمن المعلومات', 'تقرير الأساس'),
                        ('التخطيط', '1.2', 'تحديد الفئات المستهدفة وأهداف التعلم', 'التدريب', 'خطة التدريب'),
                        ('التخطيط', '1.3', 'اختيار طرق تقديم التدريب', 'التدريب', 'استراتيجية التقديم'),
                        ('التطوير', '2.1', 'تطوير محتوى ومواد التدريب', 'التدريب', 'وحدات التدريب'),
                        ('التطوير', '2.2', 'إنشاء التقييمات والاختبارات', 'التدريب', 'بنك الأسئلة'),
                        ('التطوير', '2.3', 'إعداد سيناريوهات محاكاة التصيد', 'فريق الأمن', 'خطة المحاكاة'),
                        ('التنفيذ', '3.1', 'إطلاق برنامج التدريب الإلزامي', 'الموارد البشرية', 'جدول التدريب'),
                        ('التنفيذ', '3.2', 'تنفيذ محاكاة التصيد', 'فريق الأمن', 'نتائج المحاكاة'),
                        ('التنفيذ', '3.3', 'تقديم تدريب علاجي للفاشلين', 'التدريب', 'سجل المعالجة'),
                        ('المراقبة', '4.1', 'تتبع معدلات الإكمال والدرجات', 'التدريب', 'لوحة التقدم'),
                    ],
                    'evidence': ['سجلات إتمام التدريب', 'درجات التقييم', 'نتائج محاكاة التصيد', 'مقاييس التحسن']
                },
                'incident': {
                    'steps': [
                        ('التخطيط', '1.1', 'مراجعة قدرات الاستجابة الحالية', 'فريق الأمن', 'تقييم الفجوات'),
                        ('التخطيط', '1.2', 'تحديد فئات ومستويات خطورة الحوادث', 'فريق الأمن', 'مصفوفة التصنيف'),
                        ('التطوير', '2.1', 'تطوير أدلة الاستجابة للحوادث', 'فريق الأمن', 'وثائق الأدلة'),
                        ('التطوير', '2.2', 'إنشاء إجراءات التصعيد وقوائم الاتصال', 'فريق الأمن', 'مصفوفة التصعيد'),
                        ('التطوير', '2.3', 'تصميم قوالب التقارير', 'فريق الأمن', 'قوالب التقارير'),
                        ('بناء الفريق', '3.1', 'تشكيل فريق CSIRT/IRT', 'الإدارة', 'ميثاق الفريق'),
                        ('بناء الفريق', '3.2', 'تحديد الأدوار والمسؤوليات (RACI)', 'فريق الأمن', 'مصفوفة RACI'),
                        ('التدريب', '4.1', 'تدريب فريق الاستجابة للحوادث', 'التدريب', 'فريق مدرب'),
                        ('الاختبار', '5.1', 'إجراء تمارين الطاولة', 'فريق الأمن', 'تقرير التمرين'),
                        ('الاختبار', '5.2', 'تنفيذ محاكاة كاملة', 'فريق الأمن', 'نتائج المحاكاة'),
                    ],
                    'evidence': ['خطة IR المعتمدة', 'قائمة فريق CSIRT', 'تقرير تمرين الطاولة', 'تقرير ما بعد المحاكاة']
                },
                'data': {
                    'steps': [
                        ('التخطيط', '1.1', 'جرد جميع أصول البيانات', 'إدارة البيانات', 'جرد البيانات'),
                        ('التخطيط', '1.2', 'تصنيف البيانات حسب مستوى الحساسية', 'أمن المعلومات', 'نظام التصنيف'),
                        ('التخطيط', '1.3', 'رسم خرائط تدفق البيانات ومواقع التخزين', 'تقنية المعلومات', 'مخطط تدفق البيانات'),
                        ('التنفيذ', '2.1', 'تطبيق تسميات تصنيف البيانات', 'إدارة البيانات', 'بيانات مصنفة'),
                        ('التنفيذ', '2.2', 'تنفيذ ضوابط DLP', 'فريق الأمن', 'تكوين DLP'),
                        ('التنفيذ', '2.3', 'تفعيل التشفير للبيانات الحساسة', 'تقنية المعلومات', 'تقرير التشفير'),
                        ('التنفيذ', '2.4', 'تكوين ضوابط الوصول (الحاجة للمعرفة)', 'تقنية المعلومات', 'مصفوفة الوصول'),
                        ('المراقبة', '3.1', 'نشر أدوات مراقبة البيانات', 'فريق الأمن', 'المراقبة فعالة'),
                        ('المراقبة', '3.2', 'وضع إجراءات الاستجابة لاختراق البيانات', 'فريق الأمن', 'إجراءات الاختراق'),
                        ('الامتثال', '4.1', 'التحقق من امتثال PDPL/GDPR', 'الامتثال', 'تقرير الامتثال'),
                    ],
                    'evidence': ['سجل تصنيف البيانات', 'تكوين سياسة DLP', 'شهادات التشفير', 'مصفوفة التحكم بالوصول']
                },
                'default': {
                    'steps': [
                        ('التخطيط', '1.1', 'تقييم الوضع الحالي وتحديد المتطلبات', 'قائد المشروع', 'تقرير التقييم'),
                        ('التخطيط', '1.2', 'تحديد النطاق ومعايير النجاح', 'قائد المشروع', 'ميثاق المشروع'),
                        ('التخطيط', '1.3', 'وضع الجدول الزمني للتنفيذ', 'قائد المشروع', 'خطة المشروع'),
                        ('التنفيذ', '2.1', 'تخصيص الموارد والميزانية', 'الإدارة', 'تخصيص الموارد'),
                        ('التنفيذ', '2.2', 'تنفيذ أنشطة التطبيق', 'فريق التنفيذ', 'تقارير التقدم'),
                        ('التنفيذ', '2.3', 'تدريب الموظفين المعنيين', 'التدريب', 'سجلات التدريب'),
                        ('التحقق', '3.1', 'اختبار والتحقق من التنفيذ', 'ضمان الجودة', 'نتائج الاختبار'),
                        ('التحقق', '3.2', 'إجراء مراجعة الإدارة', 'الإدارة', 'محضر المراجعة'),
                        ('الإغلاق', '4.1', 'توثيق الدروس المستفادة', 'قائد المشروع', 'الدروس المستفادة'),
                        ('الإغلاق', '4.2', 'الحصول على الإقرار الرسمي', 'الإدارة', 'وثيقة الإقرار'),
                    ],
                    'evidence': ['تقرير التقييم', 'سجلات التنفيذ', 'نتائج الاختبار', 'وثيقة الإقرار']
                }
            }
            
            def get_gap_type(gap_name):
                """Determine gap type from name for template selection.
                Order matters: check SPECIFIC patterns first, then broader ones."""
                gap_lower = gap_name.lower()
                
                # 1. MOST SPECIFIC: Organizational structure / governance office / team setup
                if any(kw in gap_lower for kw in [
                    'organizational structure', 'org structure', 'team', 'office', 'department', 'unit',
                    'هيكل تنظيمي', 'هيكل', 'مكتب', 'إدارة البيانات', 'فريق', 'وحدة', 'قسم',
                    'governance structure', 'dmo', 'data management office', 'ciso office',
                    'إنشاء مكتب', 'تأسيس', 'حوكمة مؤسسية', 'هيكلة'
                ]):
                    return 'governance'
                
                # 2. Data governance (stewardship, quality, lifecycle — NOT protection/DLP)
                if any(kw in gap_lower for kw in [
                    'data governance', 'data quality', 'data steward', 'data lifecycle', 'data catalog',
                    'metadata', 'data lineage', 'master data', 'data standard',
                    'حوكمة البيانات', 'جودة البيانات', 'أمناء البيانات', 'دورة حياة البيانات',
                    'كتالوج البيانات', 'البيانات الوصفية', 'معايير البيانات', 'إدارة البيانات الرئيسية'
                ]):
                    return 'data_governance'
                
                # 3. Risk management framework
                if any(kw in gap_lower for kw in [
                    'risk management', 'risk framework', 'risk appetite', 'risk register',
                    'إدارة المخاطر', 'إطار المخاطر', 'شهية المخاطر', 'سجل المخاطر',
                    'risk assessment', 'تقييم المخاطر'
                ]):
                    return 'risk_management'
                
                # 4. Compliance framework / regulatory
                if any(kw in gap_lower for kw in [
                    'compliance framework', 'compliance program', 'regulatory',
                    'إطار الامتثال', 'برنامج الامتثال', 'تنظيمي'
                ]):
                    return 'compliance_framework'
                
                # 5. Policy / procedure / documentation
                if any(kw in gap_lower for kw in ['policy', 'سياس', 'procedure', 'إجراء', 'documentation', 'توثيق']):
                    return 'policy'
                
                # 6. Technology / tools
                if any(kw in gap_lower for kw in ['technology', 'تقني', 'siem', 'edr', 'tool', 'أداة', 'software', 'برنامج']):
                    return 'technology'
                
                # 7. Training / awareness
                if any(kw in gap_lower for kw in ['training', 'تدريب', 'awareness', 'توعية', 'skill', 'مهار']):
                    return 'training'
                
                # 8. Incident response
                if any(kw in gap_lower for kw in ['incident', 'حادث', 'response', 'استجابة', 'csirt', 'soc']):
                    return 'incident'
                
                # 9. Data protection (DLP, encryption, classification)
                if any(kw in gap_lower for kw in ['data', 'بيان', 'protection', 'حماية', 'privacy', 'خصوصية', 'encryption', 'تشفير', 'dlp']):
                    return 'data'
                
                return 'default'
            
            # Generate guidelines
            if lang_code == 'ar':
                # Check how many guides already exist
                existing_guides = gaps_content.count('دليل تنفيذ الفجوة رقم')
                if existing_guides >= len(gaps_list[:5]):
                    return gaps_content  # All guides present
                
                guidelines = "\n\n---\n\n### دليل التنفيذ لكل فجوة:\n\n"
                for num, gap_name in gaps_list[:5]:
                    # Check if this specific guide already exists
                    if f'دليل تنفيذ الفجوة رقم {num}' in gaps_content:
                        continue
                    
                    gap_type = get_gap_type(gap_name)
                    template = gap_templates_ar.get(gap_type, gap_templates_ar['default'])
                    
                    guidelines += f"---\n#### دليل تنفيذ الفجوة رقم {num}: {gap_name}\n\n"
                    guidelines += "**الخطوات التفصيلية:**\n"
                    guidelines += "| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |\n"
                    guidelines += "|---------|--------|-------|---------|----------|\n"
                    for step in template['steps']:
                        guidelines += f"| {step[0]} | {step[1]} | {step[2]} | {step[3]} | {step[4]} |\n"
                    guidelines += f"\n**الأدلة المطلوبة للإغلاق:** ☐ {' ☐ '.join(template['evidence'])}\n\n"
            else:
                # Check how many guides already exist
                existing_guides = gaps_content.count('Implementation Guide:')
                if existing_guides >= len(gaps_list[:5]):
                    return gaps_content  # All guides present
                
                guidelines = "\n\n---\n\n### Implementation Guidelines for Each Gap:\n\n"
                for num, gap_name in gaps_list[:5]:
                    # Check if this specific guide already exists
                    if f'Gap #{num} Implementation Guide' in gaps_content:
                        continue
                    
                    gap_type = get_gap_type(gap_name)
                    template = gap_templates_en.get(gap_type, gap_templates_en['default'])
                    
                    guidelines += f"---\n#### Gap #{num} Implementation Guide: {gap_name}\n\n"
                    guidelines += "**Step-by-Step Implementation:**\n"
                    guidelines += "| Phase | Step | Description | Owner | Deliverable |\n"
                    guidelines += "|-------|------|-------------|-------|-------------|\n"
                    for step in template['steps']:
                        guidelines += f"| {step[0]} | {step[1]} | {step[2]} | {step[3]} | {step[4]} |\n"
                    guidelines += f"\n**Evidence Required for Closure:** ☐ {' ☐ '.join(template['evidence'])}\n\n"
            
            return gaps_content + guidelines
        
        # Assign parts to sections based on content detection
        assigned = set()
        part_to_section = {}  # Track: part_index → section_name
        print(f"DEBUG: Number of parts to assign: {len(parts)}", flush=True)
        for idx, part in enumerate(parts):
            section_type = identify_section(part, lang)
            print(f"DEBUG: Part {idx} identified as: {section_type} (first 100 chars: {part[:100]})", flush=True)
            if section_type and section_type not in assigned:
                sections[section_type] = part.strip()
                assigned.add(section_type)
                part_to_section[idx] = section_type
        
        print(f"DEBUG: Assigned sections: {assigned}", flush=True)
        
        # CRITICAL FIX: If we only have 1-2 parts, the content might all be in one block
        # Try to extract sections from within the single block
        if len(parts) <= 2 and len(assigned) < 4:
            print("DEBUG: Few parts detected - trying to extract sections from content block", flush=True)
            full_content = '\n'.join(parts)
            
            # Define section markers for Arabic
            ar_section_markers = [
                (r'(##\s*1\.?\s*الرؤية[^\n]*|##\s*الرؤية والأهداف)', 'vision'),
                (r'(##\s*2\.?\s*الركائز[^\n]*|##\s*الركائز الاستراتيجية)', 'pillars'),
                (r'(##\s*3\.?\s*تحليل[^\n]*|##\s*تحليل الفجوات)', 'gaps'),
                (r'(##\s*4\.?\s*خارطة[^\n]*|##\s*خارطة الطريق)', 'roadmap'),
                (r'(##\s*5\.?\s*مؤشرات[^\n]*|##\s*مؤشرات الأداء)', 'kpis'),
                (r'(##\s*6\.?\s*تقييم[^\n]*|##\s*تقييم الثقة)', 'confidence'),
            ]
            
            # Define section markers for English
            en_section_markers = [
                (r'(##\s*1\.?\s*Vision[^\n]*)', 'vision'),
                (r'(##\s*2\.?\s*Strategic[^\n]*|##\s*2\.?\s*Pillars[^\n]*)', 'pillars'),
                (r'(##\s*3\.?\s*Gap[^\n]*)', 'gaps'),
                (r'(##\s*4\.?\s*Implementation[^\n]*|##\s*4\.?\s*Roadmap[^\n]*)', 'roadmap'),
                (r'(##\s*5\.?\s*Key Performance[^\n]*|##\s*5\.?\s*KPI[^\n]*)', 'kpis'),
                (r'(##\s*6\.?\s*Confidence[^\n]*)', 'confidence'),
            ]
            
            markers = ar_section_markers if lang == 'ar' else en_section_markers
            
            # Find positions of each section
            section_positions = []
            for pattern, section_name in markers:
                match = re.search(pattern, full_content, re.IGNORECASE)
                if match:
                    section_positions.append((match.start(), section_name, match.group()))
                    print(f"DEBUG: Found {section_name} at position {match.start()}", flush=True)
            
            # Sort by position
            section_positions.sort(key=lambda x: x[0])
            
            # Extract content for each section
            for i, (pos, section_name, header) in enumerate(section_positions):
                if section_name not in assigned or not sections.get(section_name):
                    # Find end position (start of next section or end of content)
                    end_pos = section_positions[i + 1][0] if i + 1 < len(section_positions) else len(full_content)
                    section_content = full_content[pos:end_pos].strip()
                    if section_content:
                        sections[section_name] = section_content
                        assigned.add(section_name)
                        print(f"DEBUG: Extracted {section_name} ({len(section_content)} chars)", flush=True)
        
        # Ensure all 6 sections are filled
        section_order = ['vision', 'pillars', 'gaps', 'roadmap', 'kpis', 'confidence']
        
        # Order-based assignment: fill EMPTY sections with UNASSIGNED parts only
        # This prevents the bug where part 0 (vision) gets duplicated into pillars slot
        used_parts = set(part_to_section.keys())
        for i, part in enumerate(parts[:6]):
            if i < len(section_order) and i not in used_parts:
                section_name = section_order[i]
                if not sections.get(section_name) or sections[section_name].strip() == '':
                    sections[section_name] = part.strip()
                    used_parts.add(i)
                    print(f"DEBUG: Assigned part {i} to {section_name} (order-based)", flush=True)
        
        # If we still have empty sections and more parts available, fill them
        empty_sections = [s for s in section_order if not sections.get(s) or sections[s].strip() == '']
        if empty_sections and len(parts) > 6:
            for i, extra_part in enumerate(parts[6:]):
                if i < len(empty_sections):
                    sections[empty_sections[i]] = extra_part.strip()
                    print(f"DEBUG: Assigned extra part to {empty_sections[i]}", flush=True)
        
        # Log final section status
        for section_name in section_order:
            content_len = len(sections.get(section_name, ''))
            print(f"DEBUG: Section '{section_name}' has {content_len} chars", flush=True)
        
        # INJECT IMPLEMENTATION GUIDELINES if missing from gaps section
        if sections.get('gaps'):
            print("DEBUG: Checking gaps section for implementation guidelines...", flush=True)
            original_len = len(sections['gaps'])
            sections['gaps'] = inject_implementation_guidelines(sections['gaps'], lang)
            new_len = len(sections['gaps'])
            if new_len > original_len:
                print(f"DEBUG: Injected implementation guidelines (+{new_len - original_len} chars)", flush=True)
            else:
                print("DEBUG: Guidelines already present, no injection needed", flush=True)
        
        # INJECT KPI ASSESSMENT GUIDES if missing from KPIs section
        def inject_kpi_assessment_guides(kpi_content, lang_code):
            """Inject KPI assessment guides if AI didn't include them."""
            if not kpi_content:
                return kpi_content
            
            # Check if guides already exist IN TABLE FORMAT
            has_table_guides = (
                ('Assessment Guide' in kpi_content or 'دليل تقييم المؤشر' in kpi_content) and
                ('| Step |' in kpi_content or '| Action |' in kpi_content or 
                 '| الخطوة |' in kpi_content or '| الإجراء |' in kpi_content)
            )
            
            if has_table_guides:
                return kpi_content
            
            # Strip any existing bullet-point (non-table) assessment guides before re-injecting as tables
            import re
            
            # Remove bullet-point guide sections that lack table formatting
            if ('Assessment Guide' in kpi_content or 'دليل تقييم المؤشر' in kpi_content):
                # Remove existing non-table guide sections (#### headers followed by bullets/text, not tables)
                kpi_content = re.sub(
                    r'(?:---\s*\n)?####\s*(?:KPI\s*#?\d+\s*Assessment Guide|دليل تقييم المؤشر رقم\s*\d+)[^\n]*\n(?:(?!\n####|\n---\n####|\n##\s|\n\[SECTION\]|\Z)(?!\|.*\|.*\|).)*',
                    '', kpi_content, flags=re.DOTALL
                )
                # Remove orphaned "KPI Assessment Guidelines" / "أدلة تقييم" headers
                kpi_content = re.sub(r'###\s*(?:KPI Assessment Guidelines|أدلة تقييم (?:المؤشرات|مؤشرات الأداء))\s*\n', '', kpi_content)
                # Clean up leftover whitespace
                kpi_content = re.sub(r'\n{3,}', '\n\n', kpi_content).strip()
            
            # Extract KPI names from table
            kpi_pattern = r'\|\s*(\d+)\s*\|\s*([^|]+)\s*\|'
            kpis_found = re.findall(kpi_pattern, kpi_content)
            
            kpi_list = []
            for num, name in kpis_found:
                name = name.strip()
                if name and name.lower() not in ['kpi', 'المؤشر', '#', 'description', 'الوصف', 'current', 'الحالية', 'target', 'المستهدفة']:
                    kpi_list.append((num, name))
            
            if not kpi_list or len(kpi_list) < 2:
                return kpi_content
            
            # Generate assessment guides - KPI-specific, not generic
            # Map KPI keywords to specific measurement steps
            KPI_GUIDES_AR = {
                'امتثال': [
                    ('حصر جميع الضوابط المنطبقة على المنظمة', 'فريق الامتثال', 'مصفوفة انطباق الضوابط'),
                    ('تقييم حالة التطبيق لكل ضابط: مطبق / جزئي / غير مطبق', 'فريق أمن المعلومات', 'ورقة تقييم الضوابط'),
                    ('احتساب: (مطبق بالكامل + 0.5×جزئي) / إجمالي الضوابط', 'فريق الامتثال', 'نسبة الامتثال'),
                    ('تحديد الفجوات وترتيب أولويات المعالجة حسب المخاطر', 'فريق المخاطر', 'خطة معالجة مرتبة'),
                ],
                'استجابة|حوادث|كشف|MTTD|MTTR': [
                    ('تكوين SIEM لتسجيل وقت توليد التنبيه تلقائياً', 'فريق SOC', 'قواعد الارتباط'),
                    ('استخراج الأوقات من تذاكر الحوادث آخر 90 يوم', 'فريق SOC', 'بيانات الحوادث'),
                    ('احتساب متوسط الوقت بين الاختراق والكشف (MTTD)', 'فريق SOC', 'قيمة MTTD'),
                    ('احتساب متوسط الوقت بين الكشف والاحتواء الكامل (MTTR)', 'فريق SOC', 'قيمة MTTR'),
                ],
                'تدريب|توعية|تدريب.*إكمال': [
                    ('الحصول على إجمالي عدد الموظفين من الموارد البشرية', 'الموارد البشرية', 'قائمة الموظفين'),
                    ('استخراج سجلات الإكمال من نظام إدارة التعلم (LMS)', 'فريق التدريب', 'سجلات الإكمال'),
                    ('مطابقة القائمة وتحديد غير المكملين', 'فريق التدريب', 'تقرير الفجوات'),
                    ('قياس معدل النقر في محاكاة التصيد كمقياس تكميلي', 'أمن المعلومات', 'نتائج المحاكاة'),
                ],
                'ثغرات|معالجة|SLA': [
                    ('استخراج جميع الثغرات من أدوات الفحص', 'إدارة الثغرات', 'جرد الثغرات'),
                    ('تصنيف حسب الخطورة: حرجة (7 أيام)، عالية (14 يوم)، متوسطة (30 يوم)', 'إدارة الثغرات', 'قائمة مصنفة'),
                    ('تتبع تواريخ المعالجة من نظام التغييرات', 'عمليات تقنية المعلومات', 'جدول المعالجة'),
                    ('احتساب: المعالجة ضمن الوقت / إجمالي الثغرات لكل مستوى', 'إدارة الثغرات', 'نسبة الالتزام'),
                ],
                'هوية|وصول|IAM|صلاحيات': [
                    ('حصر جميع ضوابط IAM المطلوبة (MFA، PAM، RBAC، مراجعة الصلاحيات)', 'فريق IAM', 'قائمة ضوابط IAM'),
                    ('تقييم حالة التطبيق لكل ضابط عبر جميع الأنظمة', 'فريق IAM', 'حالة IAM لكل نظام'),
                    ('التحقق: نسبة MFA، تغطية الحسابات المميزة، دورية المراجعة', 'فريق IAM', 'القياسات الفرعية'),
                    ('احتساب النتيجة المركبة: الضوابط المطبقة / إجمالي ضوابط IAM', 'فريق IAM', 'نسبة تطبيق IAM'),
                ],
                'تقنية.*ضوابط|ضوابط.*تقني|technical': [
                    ('استخراج قائمة الضوابط التقنية المطلوبة من الإطار', 'فريق الامتثال', 'سجل الضوابط التقنية'),
                    ('ربط كل ضابط بالتقنية/الأداة/التكوين المطلوب', 'فريق تقنية المعلومات', 'خريطة الضوابط'),
                    ('التحقق من التطبيق عبر تدقيق التكوينات وجمع الأدلة', 'أمن المعلومات', 'مستودع الأدلة'),
                    ('احتساب النسبة الإجمالية وتحديد الفجوات الحرجة', 'فريق الامتثال', 'لوحة معلومات التطبيق'),
                ],
                'استمرارية|BCP|تعافي': [
                    ('مراجعة وثائق خطط الاستمرارية وتحديد جميع الخطط المطلوب اختبارها', 'فريق BCM', 'جرد الخطط'),
                    ('تصميم سيناريوهات الاختبار: تمارين طاولة وتمارين وظيفية', 'فريق BCM', 'سيناريوهات الاختبار'),
                    ('تنفيذ الاختبارات وقياس: تحقيق RTO/RPO أثناء الاختبار', 'فريق BCM', 'قياس RTO/RPO'),
                    ('توثيق الدروس المستفادة وتحديث الخطط', 'فريق BCM', 'خطة محدثة'),
                ],
                'طرف ثالث|مورد|أطراف|vendor': [
                    ('حصر جميع الموردين الذين لديهم وصول للبيانات أو الأنظمة', 'المشتريات', 'سجل الموردين'),
                    ('تصنيف الموردين حسب مستوى الخطورة: حرج / عالي / متوسط / منخفض', 'فريق المخاطر', 'قائمة مصنفة'),
                    ('إرسال استبيانات التقييم الأمني للموردين الحرجين والعاليين', 'إدارة مخاطر الأطراف الثالثة', 'استبيانات مكتملة'),
                    ('احتساب: الموردون المستوفون / إجمالي الموردين المقيّمين', 'إدارة مخاطر الأطراف الثالثة', 'نسبة الامتثال'),
                ],
            }
            
            KPI_GUIDES_EN = {
                'third.party|vendor|supplier': [
                    ('Inventory all vendors with access to data or systems', 'Procurement', 'Vendor register'),
                    ('Classify vendors by risk tier: Critical / High / Medium / Low', 'Risk Team', 'Tiered vendor list'),
                    ('Send security questionnaires to Critical/High vendors', 'Third-Party Risk', 'Completed assessments'),
                    ('Calculate: Vendors meeting requirements / Total assessed', 'Third-Party Risk', 'Compliance rate'),
                ],
                'identity|access.*control|access.*manage|IAM|privileged': [
                    ('Inventory all IAM controls required (MFA, PAM, RBAC, access reviews)', 'IAM Team', 'IAM controls checklist'),
                    ('Assess implementation across all systems', 'IAM Team', 'System-by-system IAM status'),
                    ('Verify: MFA enrollment, privileged account coverage, review cadence', 'IAM Team', 'Sub-metric measurements'),
                    ('Calculate composite: controls implemented / total IAM controls', 'IAM Team', 'IAM implementation rate'),
                ],
                'detect|response|incident|MTTD|MTTR': [
                    ('Configure SIEM to timestamp alert generation automatically', 'SOC Team', 'SIEM correlation rules'),
                    ('Extract timestamps from last 90 days of incident tickets', 'SOC Team', 'Incident timeline dataset'),
                    ('Calculate average time between compromise and detection (MTTD)', 'SOC Team', 'MTTD metric'),
                    ('Calculate average time between detection and containment (MTTR)', 'SOC Team', 'MTTR metric'),
                ],
                'training|awareness|training.*completion': [
                    ('Obtain total headcount including contractors from HR', 'HR Department', 'Employee roster'),
                    ('Pull completion records from LMS', 'Training Team', 'Completion logs'),
                    ('Cross-reference roster vs completions, flag non-completions', 'Training Team', 'Gap report'),
                    ('Measure phishing simulation click rate as supplementary metric', 'InfoSec Team', 'Phishing test results'),
                ],
                'vulnerabilit|remediat|patch|SLA': [
                    ('Extract all vulnerabilities from scanning tools (Nessus, Qualys)', 'Vulnerability Mgmt', 'Full vulnerability inventory'),
                    ('Classify by severity: Critical (7d), High (14d), Medium (30d)', 'Vulnerability Mgmt', 'Classified list with SLA deadlines'),
                    ('Track remediation dates from patching/change management system', 'IT Operations', 'Remediation timeline'),
                    ('Calculate: Remediated within SLA / Total per severity', 'Vulnerability Mgmt', 'SLA compliance rate'),
                ],
                'technical.*control|control.*implement(?!.*access)': [
                    ('Extract mandatory technical controls from framework', 'Compliance Team', 'Technical controls register'),
                    ('Map each control to specific technology/tool/configuration', 'IT Team', 'Control-to-technology map'),
                    ('Verify via configuration audits and evidence collection', 'InfoSec Team', 'Evidence repository'),
                    ('Score per control and calculate aggregate rate', 'Compliance Team', 'Implementation dashboard'),
                ],
                'continuity|BCP|disaster|recovery': [
                    ('Review BCP documentation and identify plans requiring testing', 'BCM Team', 'Plans inventory'),
                    ('Design test scenarios: tabletop and functional drills', 'BCM Team', 'Test scenarios'),
                    ('Execute tests and measure RTO/RPO achievement', 'BCM Team', 'RTO/RPO measurement'),
                    ('Document lessons learned and update plans', 'BCM Team', 'Updated BCP'),
                ],
                'compliance|audit.*pass': [
                    ('Map all applicable framework controls to organizational scope', 'Compliance Team', 'Controls applicability matrix'),
                    ('Assess implementation status per control: Implemented / Partial / Not', 'InfoSec Team', 'Control-by-control status'),
                    ('Score: (Fully Implemented + 0.5×Partial) / Total Controls', 'Compliance Team', 'Compliance percentage'),
                    ('Identify gaps and prioritize remediation by risk', 'Risk Team', 'Prioritized remediation plan'),
                ],
            }
            
            def find_kpi_guide(kpi_name, guides_dict):
                """Find the BEST matching guide for a KPI name - most specific match wins."""
                matches = []
                for pattern, steps in guides_dict.items():
                    m = re.search(pattern, kpi_name, re.IGNORECASE)
                    if m:
                        # Score by: number of pattern alternatives matched (more = more specific)
                        # and length of the matched text (longer = more specific)
                        match_len = len(m.group())
                        # Count how many alternatives in the pattern actually match
                        alt_matches = sum(1 for alt in pattern.split('|') if re.search(alt, kpi_name, re.IGNORECASE))
                        matches.append((alt_matches, match_len, steps))
                if matches:
                    # Return the most specific match (most alternatives matched, then longest match)
                    matches.sort(reverse=True)
                    return matches[0][2]
                return None
            
            if lang_code == 'ar':
                guides = "\n\n### أدلة تقييم مؤشرات الأداء\n\n"
                for num, kpi_name in kpi_list[:8]:
                    steps = find_kpi_guide(kpi_name, KPI_GUIDES_AR)
                    guides += f"---\n#### دليل تقييم المؤشر رقم {num}: {kpi_name[:40]}\n"
                    guides += "| الخطوة | الإجراء | الأداة/النظام | المسؤول | المخرج |\n"
                    guides += "|--------|---------|---------------|---------|--------|\n"
                    if steps:
                        for i, (action, owner, output) in enumerate(steps, 1):
                            guides += f"| {i} | {action} | — | {owner} | {output} |\n"
                    else:
                        guides += f"| 1 | تحديد متطلبات القياس الخاصة بمؤشر {kpi_name[:30]} | — | فريق الامتثال | وثيقة المتطلبات |\n"
                        guides += f"| 2 | جمع البيانات من الأنظمة والأدوات ذات الصلة | — | الفريق التقني | البيانات المجمعة |\n"
                        guides += f"| 3 | تطبيق صيغة الاحتساب المحددة لهذا المؤشر | — | فريق القياس | القيمة المحتسبة |\n"
                        guides += f"| 4 | مراجعة النتائج وتحديد إجراءات التحسين | — | الإدارة المعنية | تقرير التحسين |\n"
                    guides += "\n"
            else:
                guides = "\n\n### KPI Assessment Guidelines\n\n"
                for num, kpi_name in kpi_list[:8]:
                    steps = find_kpi_guide(kpi_name, KPI_GUIDES_EN)
                    guides += f"---\n#### KPI #{num} Assessment Guide: {kpi_name[:40]}\n"
                    guides += "| Step | Action | Tool/System | Owner | Output |\n"
                    guides += "|------|--------|-------------|-------|--------|\n"
                    if steps:
                        for i, (action, owner, output) in enumerate(steps, 1):
                            guides += f"| {i} | {action} | — | {owner} | {output} |\n"
                    else:
                        guides += f"| 1 | Identify specific data sources for {kpi_name[:30]} | — | Compliance Team | Data source map |\n"
                        guides += f"| 2 | Collect measurements from relevant systems and tools | — | Technical Team | Raw measurements |\n"
                        guides += f"| 3 | Apply KPI-specific calculation formula | — | Measurement Team | Calculated value |\n"
                        guides += f"| 4 | Review results and identify improvement actions | — | Management | Improvement report |\n"
                    guides += "\n"
            
            return kpi_content + guides
        
        if sections.get('kpis'):
            print("DEBUG: Checking KPIs section for assessment guides...", flush=True)
            original_len = len(sections['kpis'])
            sections['kpis'] = inject_kpi_assessment_guides(sections['kpis'], lang)
            new_len = len(sections['kpis'])
            if new_len > original_len:
                print(f"DEBUG: Injected KPI assessment guides (+{new_len - original_len} chars)", flush=True)
            else:
                print("DEBUG: KPI guides already present, no injection needed", flush=True)
        
        # INJECT CONFIDENCE SECTION if missing (common when AI truncates long strategies)
        if not sections.get('confidence') or sections['confidence'].strip() == '':
            print("DEBUG: Confidence section EMPTY — injecting fallback", flush=True)
            domain_val = data.get('domain', 'Cyber Security')
            fw_val = data.get('frameworks', [])
            fw_str = ', '.join(fw_val) if isinstance(fw_val, list) else str(fw_val)
            
            # Domain-specific risks and success factors
            domain_risks_ar = {
                'Cyber Security': [
                    ('نقص الكوادر المؤهلة في الأمن السيبراني', 'عالية', 'عالي', 'برنامج تدريب وتأهيل مكثف واستقطاب خبرات خارجية'),
                    ('مقاومة التغيير التنظيمي', 'متوسطة', 'عالي', 'خطة إدارة تغيير شاملة مع دعم القيادة التنفيذية'),
                    ('تطور التهديدات السيبرانية بشكل أسرع من التطبيق', 'عالية', 'عالي', 'تبني نهج أمني تكيفي مع مراجعات دورية'),
                    ('قيود الميزانية مقابل متطلبات التنفيذ', 'متوسطة', 'متوسط', 'تحديد الأولويات وفق المخاطر والتنفيذ المرحلي'),
                ],
                'Data Management': [
                    ('غياب ثقافة إدارة البيانات في المنظمة', 'عالية', 'عالي', 'برنامج توعية وتدريب مع دعم القيادة'),
                    ('تعقيد دمج أنظمة البيانات المتعددة', 'متوسطة', 'عالي', 'نهج تكاملي مرحلي مع إطار موحد للبيانات'),
                    ('تغيرات التشريعات والأنظمة المتعلقة بالبيانات', 'متوسطة', 'متوسط', 'آلية رصد تنظيمي وتحديث دوري للسياسات'),
                    ('مقاومة ملاك البيانات لمشاركة الصلاحيات', 'عالية', 'متوسط', 'حوكمة واضحة مع تحديد الأدوار والمسؤوليات'),
                ],
                'Artificial Intelligence': [
                    ('غياب الأطر التنظيمية الناضجة للذكاء الاصطناعي', 'عالية', 'عالي', 'تبني أفضل الممارسات الدولية مع مراجعة دورية'),
                    ('مخاطر التحيز في نماذج الذكاء الاصطناعي', 'متوسطة', 'عالي', 'إطار تقييم عدالة النماذج مع مراجعة دورية'),
                    ('نقص الكفاءات المتخصصة في حوكمة الذكاء الاصطناعي', 'عالية', 'متوسط', 'برنامج تطوير الكفاءات واستقطاب الخبرات'),
                    ('صعوبة تفسير قرارات النماذج (الصندوق الأسود)', 'متوسطة', 'عالي', 'اعتماد نماذج قابلة للتفسير ووثائق الشفافية'),
                ],
                'Digital Transformation': [
                    ('مقاومة التغيير على المستويات التشغيلية', 'عالية', 'عالي', 'خطة إدارة تغيير مع أبطال التحول في كل إدارة'),
                    ('تعقيد تكامل الأنظمة القديمة مع الحديثة', 'متوسطة', 'عالي', 'نهج تحديث تدريجي مع طبقة تكامل وسيطة'),
                    ('فجوة المهارات الرقمية لدى الموظفين', 'عالية', 'متوسط', 'برنامج تدريب رقمي شامل مع مسارات تعلم'),
                    ('تجاوز الميزانية والجداول الزمنية', 'متوسطة', 'متوسط', 'حوكمة المشاريع الرشيقة مع مراجعات دورية'),
                ],
                'Enterprise Risk Management': [
                    ('ضعف ثقافة إدارة المخاطر في المنظمة', 'عالية', 'عالي', 'برنامج توعية شامل مع دعم مجلس الإدارة'),
                    ('صعوبة قياس المخاطر غير المالية', 'متوسطة', 'عالي', 'تطوير مقاييس كمية ونوعية مع منهجيات معتمدة'),
                    ('عدم كفاية البيانات لتقييم المخاطر', 'متوسطة', 'متوسط', 'بناء مستودع بيانات المخاطر مع مصادر موثوقة'),
                    ('تغير بيئة المخاطر بشكل متسارع', 'عالية', 'عالي', 'نظام إنذار مبكر مع مراجعات ربع سنوية'),
                ],
                'Global Standards': [
                    ('تعقيد التوفيق بين المعايير المتعددة', 'عالية', 'عالي', 'بناء إطار موحد يربط الضوابط المشتركة'),
                    ('تكلفة وجهد الحصول على الشهادات', 'متوسطة', 'متوسط', 'نهج مرحلي مع تحديد الأولويات حسب الأثر'),
                    ('صعوبة الحفاظ على الامتثال المستمر', 'عالية', 'متوسط', 'أتمتة المراقبة مع تقييمات دورية'),
                    ('نقص الكفاءات المتخصصة في المعايير الدولية', 'متوسطة', 'عالي', 'برنامج تأهيل مع شهادات مهنية'),
                ],
            }
            
            domain_risks_en = {
                'Cyber Security': [
                    ('Shortage of qualified cybersecurity professionals', 'High', 'High', 'Intensive training program and external expertise recruitment'),
                    ('Organizational resistance to change', 'Medium', 'High', 'Comprehensive change management with executive sponsorship'),
                    ('Evolving threats outpacing implementation', 'High', 'High', 'Adaptive security approach with periodic reviews'),
                    ('Budget constraints vs implementation requirements', 'Medium', 'Medium', 'Risk-based prioritization with phased implementation'),
                ],
                'Data Management': [
                    ('Absence of data management culture', 'High', 'High', 'Awareness and training program with leadership support'),
                    ('Complexity of integrating multiple data systems', 'Medium', 'High', 'Phased integration approach with unified data framework'),
                    ('Evolving data regulations and legislation', 'Medium', 'Medium', 'Regulatory monitoring mechanism with periodic policy updates'),
                    ('Data owners resistance to governance controls', 'High', 'Medium', 'Clear governance with defined roles and responsibilities'),
                ],
                'Artificial Intelligence': [
                    ('Lack of mature AI regulatory frameworks', 'High', 'High', 'Adopt international best practices with periodic review'),
                    ('Bias risks in AI models', 'Medium', 'High', 'Model fairness assessment framework with periodic audits'),
                    ('Shortage of AI governance expertise', 'High', 'Medium', 'Competency development program and expert recruitment'),
                    ('Model interpretability challenges (black box)', 'Medium', 'High', 'Adopt explainable models and transparency documentation'),
                ],
                'Digital Transformation': [
                    ('Operational resistance to change', 'High', 'High', 'Change management plan with transformation champions'),
                    ('Legacy system integration complexity', 'Medium', 'High', 'Gradual modernization with integration middleware'),
                    ('Digital skills gap among employees', 'High', 'Medium', 'Comprehensive digital training with learning paths'),
                    ('Budget and timeline overruns', 'Medium', 'Medium', 'Agile project governance with periodic reviews'),
                ],
                'Enterprise Risk Management': [
                    ('Weak risk management culture', 'High', 'High', 'Comprehensive awareness program with board support'),
                    ('Difficulty measuring non-financial risks', 'Medium', 'High', 'Develop quantitative and qualitative metrics'),
                    ('Insufficient data for risk assessment', 'Medium', 'Medium', 'Build risk data repository with reliable sources'),
                    ('Rapidly changing risk landscape', 'High', 'High', 'Early warning system with quarterly reviews'),
                ],
                'Global Standards': [
                    ('Complexity of harmonizing multiple standards', 'High', 'High', 'Build unified framework mapping common controls'),
                    ('Cost and effort of certification', 'Medium', 'Medium', 'Phased approach with impact-based prioritization'),
                    ('Difficulty maintaining continuous compliance', 'High', 'Medium', 'Automate monitoring with periodic assessments'),
                    ('Shortage of standards expertise', 'Medium', 'High', 'Qualification program with professional certifications'),
                ],
            }
            
            domain_factors_ar = {
                'Cyber Security': [
                    ('دعم القيادة التنفيذية', 'توفير الدعم المستمر من الإدارة العليا لمبادرات الأمن السيبراني', 'حرجة'),
                    ('توفر الكفاءات المتخصصة', 'استقطاب وتطوير كوادر مؤهلة في الأمن السيبراني', 'حرجة'),
                    ('كفاية الميزانية المخصصة', 'تخصيص ميزانية كافية لتنفيذ الضوابط والتقنيات المطلوبة', 'عالية'),
                    ('وعي الموظفين الأمني', 'رفع مستوى الوعي الأمني لجميع منسوبي المنظمة', 'عالية'),
                    ('نضج العمليات التقنية', 'وجود عمليات تقنية ناضجة تدعم التطبيق الفعال', 'متوسطة'),
                ],
                'Data Management': [
                    ('التزام القيادة بحوكمة البيانات', 'دعم الإدارة العليا لمبادرات إدارة البيانات', 'حرجة'),
                    ('تعيين ملاك وأمناء البيانات', 'تحديد المسؤوليات الواضحة لكل نطاق بيانات', 'حرجة'),
                    ('توفر أدوات إدارة البيانات', 'تأمين الأدوات التقنية اللازمة لحوكمة البيانات', 'عالية'),
                    ('جودة البيانات الأساسية', 'ضمان دقة واكتمال البيانات المؤسسية', 'عالية'),
                    ('التكامل بين الأنظمة', 'تحقيق تكامل فعال بين أنظمة البيانات المختلفة', 'متوسطة'),
                ],
            }
            
            domain_factors_en = {
                'Cyber Security': [
                    ('Executive Leadership Support', 'Continuous support from senior management for cybersecurity initiatives', 'Critical'),
                    ('Qualified Workforce Availability', 'Recruiting and developing qualified cybersecurity professionals', 'Critical'),
                    ('Adequate Budget Allocation', 'Sufficient budget for implementing required controls and technologies', 'High'),
                    ('Employee Security Awareness', 'Raising security awareness across all organization members', 'High'),
                    ('Technical Process Maturity', 'Mature technical processes supporting effective implementation', 'Medium'),
                ],
                'Data Management': [
                    ('Leadership Commitment to Data Governance', 'Senior management support for data management initiatives', 'Critical'),
                    ('Data Owners and Stewards Assignment', 'Clear responsibilities defined for each data domain', 'Critical'),
                    ('Data Management Tool Availability', 'Securing necessary technical tools for data governance', 'High'),
                    ('Baseline Data Quality', 'Ensuring accuracy and completeness of enterprise data', 'High'),
                    ('System Integration', 'Achieving effective integration between data systems', 'Medium'),
                ],
            }
            
            # Build default factors — use domain-specific if available, else generic
            default_factors_ar = domain_factors_ar.get(domain_val, [
                ('دعم القيادة التنفيذية', 'توفير الدعم المستمر من الإدارة العليا', 'حرجة'),
                ('توفر الكفاءات والموارد', 'استقطاب وتطوير الكوادر المؤهلة', 'حرجة'),
                ('كفاية الميزانية', 'تخصيص ميزانية كافية للتنفيذ', 'عالية'),
                ('وعي وثقافة المنظمة', 'رفع مستوى الوعي والالتزام المؤسسي', 'عالية'),
                ('نضج العمليات والتقنيات', 'وجود بنية تحتية تقنية وعمليات ناضجة', 'متوسطة'),
            ])
            
            default_factors_en = domain_factors_en.get(domain_val, [
                ('Executive Leadership Support', 'Continuous support from senior management', 'Critical'),
                ('Qualified Workforce and Resources', 'Recruiting and developing qualified professionals', 'Critical'),
                ('Adequate Budget Allocation', 'Sufficient budget for implementation', 'High'),
                ('Organizational Awareness and Culture', 'Raising awareness and institutional commitment', 'High'),
                ('Process and Technology Maturity', 'Mature infrastructure and operational processes', 'Medium'),
            ])
            
            risks_ar = domain_risks_ar.get(domain_val, domain_risks_ar['Cyber Security'])
            risks_en = domain_risks_en.get(domain_val, domain_risks_en['Cyber Security'])
            
            if lang == 'ar':
                confidence_section = f"""## 6. تقييم الجاهزية والمخاطر

**درجة الثقة:** 60%

**مبررات التقييم:**
تستند درجة الثقة إلى تقييم شامل للوضع الحالي للمنظمة في مجال {domain_val}، مع الأخذ بعين الاعتبار مستوى نضج العمليات الحالية، وتوفر الكوادر المؤهلة، ومدى جاهزية البنية التحتية التقنية. تعكس هذه الدرجة وجود فرص تحسين جوهرية مع الاعتراف بالتحديات التي تتطلب معالجة منهجية وفق متطلبات {fw_str if fw_str else 'الأطر التنظيمية المحددة'}.

### عوامل النجاح الحرجة:
| # | العامل | الوصف | الأهمية |
|---|--------|-------|---------|
"""
                for i, (factor, desc, importance) in enumerate(default_factors_ar, 1):
                    confidence_section += f"| {i} | {factor} | {desc} | {importance} |\n"
                
                confidence_section += f"""
### المخاطر الاستراتيجية:
| # | الخطر | الاحتمالية | الأثر | خطة التخفيف |
|---|-------|-----------|-------|-------------|
"""
                for i, (risk, likelihood, impact, mitigation) in enumerate(risks_ar, 1):
                    confidence_section += f"| {i} | {risk} | {likelihood} | {impact} | {mitigation} |\n"
                
                confidence_section += "\n---\n**تاريخ الإعداد:** يُحدد لاحقاً"
                
            else:
                confidence_section = f"""## 6. Confidence Assessment & Risks

**Confidence Score:** 60%

**Score Justification:**
The confidence score is based on a comprehensive assessment of the organization's current state in {domain_val}, considering the maturity level of existing processes, availability of qualified personnel, and technical infrastructure readiness. This score reflects substantial improvement opportunities while acknowledging challenges that require systematic remediation per {fw_str if fw_str else 'selected regulatory frameworks'}.

### Critical Success Factors:
| # | Factor | Description | Importance |
|---|--------|-------------|------------|
"""
                for i, (factor, desc, importance) in enumerate(default_factors_en, 1):
                    confidence_section += f"| {i} | {factor} | {desc} | {importance} |\n"
                
                confidence_section += f"""
### Key Risks:
| # | Risk | Likelihood | Impact | Mitigation Plan |
|---|------|------------|--------|-----------------|
"""
                for i, (risk, likelihood, impact, mitigation) in enumerate(risks_en, 1):
                    confidence_section += f"| {i} | {risk} | {likelihood} | {impact} | {mitigation} |\n"
                
                confidence_section += "\n---\n**Preparation Date:** To be determined"
            
            sections['confidence'] = confidence_section
            print(f"DEBUG: Injected confidence section ({len(confidence_section)} chars)", flush=True)
        
        # FINAL CLEANUP: Remove any remaining instruction artifacts from all sections
        import re
        def final_cleanup(text):
            if not text:
                return text
            # Remove [SECTION] markers
            text = re.sub(r'\[SECTION\]', '', text)
            # Remove warning/instruction blocks
            text = re.sub(r'⚠️⚠️⚠️[^⚠]*?⚠️⚠️⚠️', '', text, flags=re.DOTALL)
            text = re.sub(r'⚠️\s*CRITICAL[^\n]*(?:\n-[^\n]*)*', '', text)
            text = re.sub(r'⚠️\s*تعليمات[^\n]*(?:\n-[^\n]*)*', '', text)
            text = re.sub(r'⚠️\s*مهم جداً[^\n]*(?:\n-[^\n]*)*', '', text)
            # Remove instruction lines starting with warning emojis
            text = re.sub(r'\n⚠️[^\n]+', '', text)
            # Clean up multiple newlines
            text = re.sub(r'\n{4,}', '\n\n\n', text)
            return text.strip()
        
        for section_key in sections:
            sections[section_key] = final_cleanup(sections[section_key])
        
        # FRAMEWORK LEAKAGE CLEANUP: Remove unauthorized framework references
        all_known_frameworks = {
            'NCA ECC': ['NCA ECC', 'الضوابط الأساسية للأمن السيبراني'],
            'NCA CSCC': ['NCA CSCC', 'ضوابط الأمن السيبراني للأنظمة الحساسة'],
            'NCA DCC': ['NCA DCC', 'ضوابط الأمن السيبراني للبيانات'],
            'SAMA CSF': ['SAMA CSF', 'إطار الأمن السيبراني لمؤسسة النقد'],
            'ISO 27001': ['ISO 27001', 'آيزو 27001'],
            'ISO 27701': ['ISO 27701'],
            'ISO 31000': ['ISO 31000'],
            'ISO 22301': ['ISO 22301'],
            'NIST CSF': ['NIST CSF', 'NIST Cybersecurity Framework'],
            'NIST AI RMF': ['NIST AI RMF'],
            'NIST SP 800-53': ['NIST SP 800-53', 'NIST 800-53'],
            'CIS Controls': ['CIS Controls', 'CIS'],
            'COBIT': ['COBIT', 'COBIT 2019'],
            'PCI DSS': ['PCI DSS'],
            'COSO ERM': ['COSO ERM', 'COSO'],
            'NDMO': ['NDMO'],
            'PDPL': ['PDPL', 'نظام حماية البيانات الشخصية'],
            'SDAIA': ['SDAIA', 'سدايا'],
            'DAMA-DMBOK': ['DAMA-DMBOK', 'DAMA', 'DMBOK'],
        }
        
        # Determine which frameworks are selected
        selected_fw = data.get('frameworks', [])
        if isinstance(selected_fw, str):
            selected_fw = [selected_fw]
        selected_fw_str = ' '.join(selected_fw).upper()
        
        # Build list of authorized framework names (all variants)
        authorized_names = set()
        for fw_key, variants in all_known_frameworks.items():
            if any(fw_key.upper() in s.upper() for s in selected_fw):
                for v in variants:
                    authorized_names.add(v)
        
        # Build list of unauthorized framework names to clean
        unauthorized_names = []
        for fw_key, variants in all_known_frameworks.items():
            if not any(fw_key.upper() in s.upper() for s in selected_fw):
                for v in variants:
                    unauthorized_names.append(v)
        
        if unauthorized_names:
            def clean_framework_leaks(text):
                if not text:
                    return text
                import re
                for fw_name in sorted(unauthorized_names, key=len, reverse=True):
                    # Remove "per/وفق/حسب FRAMEWORK" references
                    text = re.sub(r'(?:per|وفق|حسب|وفقاً ل|according to|aligned with|متوافق مع)\s+' + re.escape(fw_name) + r'(?:\s+\d+)?', '', text, flags=re.IGNORECASE)
                    # Remove "FRAMEWORK compliance/requirements/controls" references
                    text = re.sub(re.escape(fw_name) + r'(?:\s+\d+)?\s+(?:compliance|requirements|controls|ضوابط|متطلبات|امتثال)', '', text, flags=re.IGNORECASE)
                    # Remove standalone framework mentions (not inside "ISO 27001, NIST CSF" lists)
                    text = re.sub(r'(?<!\w)' + re.escape(fw_name) + r'(?:\s+\d+)?(?:/| و| and |\s*,\s*)?' + r'(?!\w)', '', text, flags=re.IGNORECASE)
                # Clean up artifacts: double spaces, empty parentheses, orphaned commas
                text = re.sub(r'\(\s*\)', '', text)
                text = re.sub(r',\s*,', ',', text)
                text = re.sub(r'\s{2,}', ' ', text)
                text = re.sub(r'\|\s*\|', '| |', text)  # Fix empty table cells
                return text
            
            for section_key in sections:
                original = sections[section_key]
                sections[section_key] = clean_framework_leaks(sections[section_key])
                if sections[section_key] != original:
                    print(f"DEBUG: Cleaned framework leaks from {section_key}", flush=True)
        
        # MARKDOWN FORMATTING FIX: Ensure proper structure for rendering and DOCX export
        for section_key in sections:
            if sections[section_key]:
                sections[section_key] = ensure_markdown_formatting(sections[section_key])
        
        # DOMAIN ISOLATION POST-PROCESSING: Remove cross-domain terminology leaks
        import re as _re_domain
        _domain_term_replacements = {
            'Data Management': {
                # English replacements
                r'\bCISO\b': 'CDO',
                r'\bChief Information Security Officer\b': 'Chief Data Officer',
                r'\bcybersecurity department\b': 'data management office',
                r'\bCybersecurity Department\b': 'Data Management Office',
                r'\bSOC team\b': 'Data Quality team',
                r'\bSOC\b(?!\s*[2-9])': 'DGC',  # SOC → Data Governance Committee (not SOC 2)
                r'\bCSIRT\b': 'Data Stewardship team',
                r'\bSecurity Operations Center\b': 'Data Governance Center',
                r'\bcybersecurity governance\b': 'data governance',
                r'\bsecurity awareness\b': 'data literacy',
                r'\bthreat detection\b': 'data quality monitoring',
                # Arabic replacements
                r'مدير الأمن السيبراني': 'مدير البيانات (CDO)',
                r'إدارة الأمن السيبراني': 'مكتب إدارة البيانات',
                r'قسم الأمن السيبراني': 'قسم إدارة البيانات',
                r'فريق العمليات الأمنية': 'فريق جودة البيانات',
                r'فريق الاستجابة للحوادث': 'فريق حوكمة البيانات',
                r'مركز العمليات الأمنية': 'مركز حوكمة البيانات',
            },
            'Artificial Intelligence': {
                r'\bCISO\b': 'AI Ethics Officer',
                r'\bChief Information Security Officer\b': 'AI Governance Lead',
                r'\bcybersecurity department\b': 'AI governance office',
                r'\bCybersecurity Department\b': 'AI Governance Office',
                r'\bSOC team\b': 'AI Monitoring team',
                r'\bSOC\b(?!\s*[2-9])': 'AI Ethics Board',
                r'\bCSIRT\b': 'AI Risk Committee',
                r'\bSecurity Operations Center\b': 'AI Governance Center',
                r'مدير الأمن السيبراني': 'مسؤول أخلاقيات الذكاء الاصطناعي',
                r'إدارة الأمن السيبراني': 'مكتب حوكمة الذكاء الاصطناعي',
                r'قسم الأمن السيبراني': 'قسم حوكمة الذكاء الاصطناعي',
                r'فريق العمليات الأمنية': 'فريق مراقبة نماذج AI',
                r'فريق الاستجابة للحوادث': 'فريق مخاطر النماذج',
            },
            'Enterprise Risk Management': {
                r'\bCISO\b': 'CRO',
                r'\bChief Information Security Officer\b': 'Chief Risk Officer',
                r'\bcybersecurity department\b': 'risk management department',
                r'\bCybersecurity Department\b': 'Risk Management Department',
                r'\bSOC\b(?!\s*[2-9])': 'Risk Committee',
                r'\bCSIRT\b': 'Risk Response Team',
                r'مدير الأمن السيبراني': 'مدير المخاطر (CRO)',
                r'إدارة الأمن السيبراني': 'إدارة المخاطر المؤسسية',
                r'قسم الأمن السيبراني': 'قسم إدارة المخاطر',
            },
            'Digital Transformation': {
                r'\bCISO\b': 'CDO',
                r'\bChief Information Security Officer\b': 'Chief Digital Officer',
                r'\bcybersecurity department\b': 'digital transformation office',
                r'\bCybersecurity Department\b': 'Digital Transformation Office',
                r'\bSOC\b(?!\s*[2-9])': 'Innovation Lab',
                r'\bCSIRT\b': 'Digital Change Team',
                r'مدير الأمن السيبراني': 'مدير التحول الرقمي',
                r'إدارة الأمن السيبراني': 'مكتب التحول الرقمي',
                r'قسم الأمن السيبراني': 'فريق التحول الرقمي',
            },
        }
        
        replacements = _domain_term_replacements.get(domain, {})
        if replacements:
            for section_key in sections:
                if sections[section_key]:
                    original = sections[section_key]
                    for pattern, replacement in replacements.items():
                        sections[section_key] = _re_domain.sub(pattern, replacement, sections[section_key])
                    if sections[section_key] != original:
                        print(f"DEBUG: Sanitized cross-domain terms in {section_key}", flush=True)
        
        # Save to database
        try:
            conn = get_db()
            conn.execute('''INSERT INTO strategies (user_id, domain, org_name, sector, content, language)
                            VALUES (?, ?, ?, ?, ?, ?)''',
                        (session['user_id'], data.get('domain'), data.get('org_name'), 
                         data.get('sector'), content, lang))
            conn.commit()
        except Exception as db_error:
            print(f"Database error: {db_error}")
        
        log_action(session['user_id'], 'generate_strategy', {
            'domain': data.get('domain'), 'org_name': data.get('org_name'),
            'sector': data.get('sector'), 'language': lang
        })
        
        # ── MIZAN PIPE: Auto-extract tasks from strategy sections ──
        pipe_domain = data.get('domain', 'General')
        pipe_uid = session['user_id']
        pipe_counts = {'initiatives': 0, 'gaps': 0, 'roadmap': 0, 'kpis': 0}
        try:
            init_ids = pipe_initiatives_to_tasks(sections.get('pillars', ''), pipe_uid, pipe_domain)
            pipe_counts['initiatives'] = len(init_ids)
        except Exception as pe:
            print(f"Pipe initiatives error (non-fatal): {pe}", flush=True)
        try:
            gap_ids = pipe_gaps_to_tasks(sections.get('gaps', ''), pipe_uid, pipe_domain)
            pipe_counts['gaps'] = len(gap_ids)
        except Exception as pe:
            print(f"Pipe gaps error (non-fatal): {pe}", flush=True)
        try:
            road_ids = pipe_roadmap_to_tasks(sections.get('roadmap', ''), pipe_uid, pipe_domain)
            pipe_counts['roadmap'] = len(road_ids)
        except Exception as pe:
            print(f"Pipe roadmap error (non-fatal): {pe}", flush=True)
        try:
            kpi_ids = pipe_kpis_to_tasks(sections.get('kpis', ''), pipe_uid, pipe_domain)
            pipe_counts['kpis'] = len(kpi_ids)
        except Exception as pe:
            print(f"Pipe KPIs error (non-fatal): {pe}", flush=True)
        
        return jsonify({
            'success': True,
            'sections': sections,
            'pipe': pipe_counts,
            'debug_vision_preview': sections.get('vision', '')[:200] if sections.get('vision') else 'EMPTY'
        })
        
    except Exception as e:
        print(f"Strategy generation error: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ============================================================================
# BILINGUAL GENERATION API (Feature 4)
# ============================================================================

@app.route('/api/generate-bilingual', methods=['POST'])
@login_required
def api_generate_bilingual():
    """Generate document in both English and Arabic."""
    try:
        data = request.json
        doc_type = data.get('type', 'strategy')  # strategy, policy, audit, risk
        domain = data.get('domain', 'Cyber Security')
        
        # Check usage limits for both languages
        table_map = {'strategy': 'strategies', 'policy': 'policies', 'audit': 'audits', 'risk': 'risks'}
        table = table_map.get(doc_type, 'strategies')
        
        can_generate, used, limit = check_usage_limit(session['user_id'], table, domain)
        if not can_generate:
            return jsonify({
                'success': False,
                'error': f'Usage limit reached for {domain}.',
                'limit_reached': True
            }), 429
        
        results = {}
        
        # Generate English version
        data['language'] = 'en'
        en_content = generate_document_content(doc_type, data, 'en')
        results['en'] = en_content
        
        # Generate Arabic version
        data['language'] = 'ar'
        ar_content = generate_document_content(doc_type, data, 'ar')
        results['ar'] = ar_content
        
        return jsonify({
            'success': True,
            'bilingual': True,
            'content': results
        })
        
    except Exception as e:
        print(f"Bilingual generation error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

def generate_document_content(doc_type, data, lang):
    """Helper function to generate document content."""
    domain = data.get('domain', 'Cyber Security')
    
    if doc_type == 'strategy':
        # Simplified strategy prompt for bilingual
        if lang == 'ar':
            prompt = f"""أنشئ ملخص استراتيجية موجز في مجال {domain} يتضمن:
1. الرؤية والأهداف
2. الركائز الاستراتيجية
3. خارطة الطريق
اجعلها مختصرة ومهنية."""
        else:
            prompt = f"""Create a brief strategy summary for {domain} including:
1. Vision and Objectives
2. Strategic Pillars
3. Implementation Roadmap
Keep it concise and professional."""
    elif doc_type == 'policy':
        policy_name = data.get('policy_name', 'Security Policy')
        if lang == 'ar':
            prompt = f"""أنشئ ملخص سياسة {policy_name} في مجال {domain} يتضمن:
1. الغرض والنطاق
2. بنود السياسة الرئيسية
3. الأدوار والمسؤوليات"""
        else:
            prompt = f"""Create a {policy_name} summary for {domain} including:
1. Purpose and Scope
2. Key Policy Statements
3. Roles and Responsibilities"""
    else:
        prompt = f"Generate a brief {doc_type} document for {domain} in {lang}."
    
    return generate_ai_content(prompt, lang, content_type=doc_type)

# ============================================================================
# CHAT WITH DOCUMENT API (Feature 1)
# ============================================================================

@app.route('/api/chat-document', methods=['POST'])
@login_required
def api_chat_document():
    """Chat with a generated document using AI."""
    try:
        data = request.json
        document_content = data.get('content', '')
        user_question = data.get('question', '')
        lang = data.get('language', 'en')
        
        if not document_content or not user_question:
            return jsonify({'success': False, 'error': 'Missing content or question'}), 400
        
        # Truncate content if too long
        max_content_length = 8000
        if len(document_content) > max_content_length:
            document_content = document_content[:max_content_length] + "..."
        
        if lang == 'ar':
            prompt = f"""أنت مساعد ذكي متخصص في الحوكمة والمخاطر والامتثال.

لديك الوثيقة التالية:
---
{document_content}
---

سؤال المستخدم: {user_question}

أجب على السؤال بناءً على محتوى الوثيقة. إذا لم تجد الإجابة في الوثيقة، أخبر المستخدم بذلك واقترح معلومات مفيدة.
اجعل إجابتك مختصرة ومفيدة."""
        else:
            prompt = f"""You are an intelligent GRC assistant.

You have the following document:
---
{document_content}
---

User question: {user_question}

Answer the question based on the document content. If the answer is not in the document, let the user know and suggest helpful information.
Keep your response concise and helpful."""
        
        response = generate_ai_content(prompt, lang, content_type='chat')
        
        return jsonify({
            'success': True,
            'response': response
        })
        
    except Exception as e:
        print(f"Chat document error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================================================
# POLICY REVIEW API (Feature 2)
# ============================================================================

@app.route('/api/review-policy', methods=['POST'])
@login_required
def api_review_policy():
    """Review and analyze an existing policy document."""
    try:
        data = request.json
        policy_content = data.get('content', '')
        review_type = data.get('review_type', 'comprehensive')  # comprehensive, compliance, gaps
        framework = data.get('framework', 'ISO 27001')
        lang = data.get('language', 'en')
        
        if not policy_content:
            return jsonify({'success': False, 'error': 'No policy content provided'}), 400
        
        # Truncate if too long
        max_length = 10000
        if len(policy_content) > max_length:
            policy_content = policy_content[:max_length] + "..."
        
        if lang == 'ar':
            if review_type == 'compliance':
                prompt = f"""راجع السياسة التالية من حيث الامتثال لإطار {framework}:

{policy_content}

قدم تقرير مراجعة يتضمن:
## نتيجة المراجعة
### نقاط الامتثال ✅
- (قائمة بالنقاط المتوافقة)

### نقاط عدم الامتثال ❌
- (قائمة بالنقاط غير المتوافقة)

### التوصيات
| # | التوصية | الأولوية |
|---|---------|----------|
| 1 | ... | عالية/متوسطة/منخفضة |

### درجة الامتثال: X%"""
            elif review_type == 'gaps':
                prompt = f"""حلل الفجوات في السياسة التالية:

{policy_content}

قدم تقرير تحليل الفجوات:
## تحليل الفجوات
### الفجوات المحددة
| # | الفجوة | الوصف | الأثر |
|---|--------|-------|-------|
| 1 | ... | ... | عالي/متوسط/منخفض |

### خطة المعالجة
| # | الفجوة | الإجراء المطلوب | الجدول الزمني |
|---|--------|----------------|--------------|"""
            else:
                prompt = f"""قدم مراجعة شاملة للسياسة التالية:

{policy_content}

## تقرير المراجعة الشاملة

### ملخص تنفيذي
(ملخص موجز للسياسة ونتائج المراجعة)

### نقاط القوة ✅
- ...

### نقاط الضعف ❌
- ...

### التوصيات للتحسين
| # | التوصية | الأولوية | الجدول الزمني |
|---|---------|----------|--------------|

### التقييم العام: X/10"""
        else:
            if review_type == 'compliance':
                prompt = f"""Review the following policy for compliance with {framework}:

{policy_content}

Provide a compliance review report:
## Review Results
### Compliant Points ✅
- (list of compliant items)

### Non-Compliant Points ❌
- (list of non-compliant items)

### Recommendations
| # | Recommendation | Priority |
|---|----------------|----------|
| 1 | ... | High/Medium/Low |

### Compliance Score: X%"""
            elif review_type == 'gaps':
                prompt = f"""Analyze gaps in the following policy:

{policy_content}

Provide a gap analysis report:
## Gap Analysis
### Identified Gaps
| # | Gap | Description | Impact |
|---|-----|-------------|--------|
| 1 | ... | ... | High/Medium/Low |

### Remediation Plan
| # | Gap | Required Action | Timeline |
|---|-----|-----------------|----------|"""
            else:
                prompt = f"""Provide a comprehensive review of the following policy:

{policy_content}

IMPORTANT: Do NOT use specific dates or years (like Q1 2024, 2025). Use relative timeframes like "Within 30 days", "Within 3 months", "Short-term", "Medium-term".

## Comprehensive Review Report

### Executive Summary
(brief summary of policy and review findings)

### Strengths ✅
- ...

### Weaknesses ❌
- ...

### Improvement Recommendations
| # | Recommendation | Priority | Timeline |
|---|----------------|----------|----------|
| 1 | ... | High/Medium/Low | Within X days/months |

### Overall Rating: X/10"""
        
        review_result = generate_ai_content(prompt, lang, content_type='review')
        
        return jsonify({
            'success': True,
            'review': review_result,
            'review_type': review_type
        })
        
    except Exception as e:
        print(f"Policy review error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================================================
# MODIFY POLICY BASED ON REVIEW (Feature Enhancement)
# ============================================================================

@app.route('/api/modify-policy', methods=['POST'])
@login_required
def api_modify_policy():
    """Start policy modification as a background task (avoids Render 30s timeout)."""
    try:
        data = request.json
        policy_content = data.get('policy_content', '')
        review_content = data.get('review_content', '')
        framework = data.get('framework', 'ISO 27001')
        domain = data.get('domain', 'Cyber Security')
        lang = data.get('language', 'en')
        
        if not policy_content or not review_content:
            return jsonify({'success': False, 'error': 'Missing policy or review content'}), 400
        
        # Truncate to keep prompt manageable
        if len(policy_content) > 6000:
            policy_content = policy_content[:6000] + "\n..."
        if len(review_content) > 3000:
            review_content = review_content[:3000] + "\n..."
        
        if lang == 'ar':
            prompt = f"""أنت خبير في الحوكمة والمخاطر والامتثال. لديك سياسة وتقرير مراجعة لها.

## السياسة الحالية:
{policy_content}

## نتائج المراجعة:
{review_content}

المطلوب: أعد كتابة السياسة الكاملة مع تطبيق جميع التوصيات والتحسينات المذكورة في تقرير المراجعة. يجب أن:
- تحافظ على الهيكل العام للسياسة
- تضيف الأقسام المفقودة التي ذكرتها المراجعة
- تصحح نقاط عدم الامتثال
- تعزز نقاط الضعف المحددة
- تضيف التفاصيل والضوابط المطلوبة
- تتوافق مع إطار {framework}

اكتب السياسة المعدلة بالكامل (وليس فقط التغييرات):"""
        else:
            prompt = f"""You are a GRC expert. You have a policy and its review report.

## Current Policy:
{policy_content}

## Review Findings:
{review_content}

Task: Rewrite the COMPLETE policy incorporating all recommendations and improvements from the review. You must:
- Maintain the overall policy structure
- Add missing sections identified in the review
- Fix non-compliance points
- Strengthen identified weaknesses
- Add required details and controls
- Ensure alignment with {framework}

Write the complete modified policy (not just the changes):"""
        
        # Create background task in database
        task_id = str(uuid.uuid4())
        create_background_task(task_id, session['user_id'], domain)
        
        # Submit to bounded thread pool (max 5 concurrent AI tasks)
        ai_executor.submit(run_ai_task, task_id, prompt, lang, 'modify_policy')
        
        # Return immediately with task_id — client will poll
        return jsonify({
            'success': True,
            'task_id': task_id,
            'status': 'pending'
        })
        
    except Exception as e:
        print(f"Policy modification error: {e}", flush=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/task-status/<task_id>')
@login_required
def api_task_status(task_id):
    """Poll for background task result (reads from database, works across workers)."""
    task = get_background_task(task_id)
    
    if not task:
        return jsonify({'status': 'not_found', 'error': 'Task not found'}), 404
    
    status = task['status']
    
    if status == 'pending':
        return jsonify({'status': 'pending'})
    
    elif status == 'done':
        result = task['result']
        
        # Save modified policy to policies table
        try:
            user_id = task['user_id']
            domain = task['callback_domain']
            if user_id and domain:
                conn = get_db()
                latest_policy = conn.execute(
                    'SELECT id FROM policies WHERE user_id = ? AND domain = ? ORDER BY created_at DESC LIMIT 1',
                    (user_id, domain)
                ).fetchone()
                if latest_policy:
                    conn.execute('UPDATE policies SET content = ? WHERE id = ?',
                                (result, latest_policy['id']))
                    conn.commit()
        except Exception as db_err:
            print(f"DB update after modify: {db_err}", flush=True)
        
        # Cleanup
        delete_background_task(task_id)
        
        return jsonify({
            'status': 'done',
            'modified_policy': result
        })
    
    elif status == 'error':
        error = task['error'] or 'Unknown error'
        delete_background_task(task_id)
        return jsonify({'status': 'error', 'error': error})

# ============================================================================
# ERM RISK REGISTER API
# ============================================================================

@app.route('/api/risk-register', methods=['GET'])
@login_required
def api_risk_register_list():
    """List risks in user's risk register. Optional ?domain= filter."""
    try:
        conn = get_db()
        domain_filter = request.args.get('domain', '')
        
        if domain_filter:
            risks = conn.execute(
                'SELECT * FROM risk_register WHERE user_id = ? AND domain = ? ORDER BY (likelihood * impact) DESC, created_at DESC',
                (session['user_id'], domain_filter)
            ).fetchall()
        else:
            risks = conn.execute(
                'SELECT * FROM risk_register WHERE user_id = ? ORDER BY (likelihood * impact) DESC, created_at DESC',
                (session['user_id'],)
            ).fetchall()
        
        return jsonify({
            'success': True,
            'risks': [dict(r) for r in risks]
        })
    except Exception as e:
        print(f"Risk register list error: {e}", flush=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/risk-register', methods=['POST'])
@login_required
def api_risk_register_create():
    """Add a new risk to the register."""
    try:
        data = request.json
        name = data.get('name', '').strip()
        if not name:
            return jsonify({'success': False, 'error': 'Risk name is required'}), 400
        
        conn = get_db()
        conn.execute('''
            INSERT INTO risk_register (user_id, name, description, category, likelihood, impact, owner, treatment, treatment_plan, status, domain)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            session['user_id'],
            name,
            data.get('description', ''),
            data.get('category', 'Operational'),
            int(data.get('likelihood', 3)),
            int(data.get('impact', 3)),
            data.get('owner', ''),
            data.get('treatment', 'Mitigate'),
            data.get('treatment_plan', ''),
            data.get('status', 'Open'),
            data.get('domain', 'General')
        ))
        conn.commit()
        
        # Pillar 2: Auto-create mitigation task for Critical/High risks
        risk_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        likelihood = int(data.get('likelihood', 3))
        impact = int(data.get('impact', 3))
        risk_score = likelihood * impact
        auto_create_mitigation_task(session['user_id'], risk_id, name, risk_score, data.get('domain', 'General'))
        
        log_action(session['user_id'], 'create_risk', {'name': name, 'score': risk_score})
        return jsonify({'success': True, 'risk_id': risk_id, 'auto_task': risk_score >= 12})
    except Exception as e:
        print(f"Risk register create error: {e}", flush=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/risk-register/<int:risk_id>', methods=['PUT'])
@login_required
def api_risk_register_update(risk_id):
    """Update an existing risk."""
    try:
        data = request.json
        conn = get_db()
        
        # Verify ownership
        risk = conn.execute('SELECT id FROM risk_register WHERE id = ? AND user_id = ?', (risk_id, session['user_id'])).fetchone()
        if not risk:
            return jsonify({'success': False, 'error': 'Risk not found'}), 404
        
        conn.execute('''
            UPDATE risk_register SET
                name = ?, description = ?, category = ?, likelihood = ?, impact = ?,
                owner = ?, treatment = ?, treatment_plan = ?, status = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ? AND user_id = ?
        ''', (
            data.get('name', ''),
            data.get('description', ''),
            data.get('category', 'Operational'),
            int(data.get('likelihood', 3)),
            int(data.get('impact', 3)),
            data.get('owner', ''),
            data.get('treatment', 'Mitigate'),
            data.get('treatment_plan', ''),
            data.get('status', 'Open'),
            risk_id,
            session['user_id']
        ))
        conn.commit()
        
        log_action(session['user_id'], 'update_risk', {'risk_id': risk_id})
        return jsonify({'success': True})
    except Exception as e:
        print(f"Risk register update error: {e}", flush=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/risk-register/<int:risk_id>', methods=['DELETE'])
@login_required
def api_risk_register_delete(risk_id):
    """Delete a risk from the register."""
    try:
        conn = get_db()
        conn.execute('DELETE FROM risk_register WHERE id = ? AND user_id = ?', (risk_id, session['user_id']))
        conn.commit()
        log_action(session['user_id'], 'delete_risk', {'risk_id': risk_id})
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/risk-appetite', methods=['POST'])
@login_required
def api_risk_appetite():
    """Generate Risk Appetite Statement based on risk register data."""
    try:
        data = request.json
        risks = data.get('risks', [])
        lang = data.get('language', 'en')
        
        if not risks:
            return jsonify({'success': False, 'error': 'No risks provided'}), 400
        
        # Summarize risk register
        total = len(risks)
        critical = sum(1 for r in risks if r['likelihood'] * r['impact'] >= 20)
        high = sum(1 for r in risks if 12 <= r['likelihood'] * r['impact'] < 20)
        medium = sum(1 for r in risks if 6 <= r['likelihood'] * r['impact'] < 12)
        low = sum(1 for r in risks if r['likelihood'] * r['impact'] < 6)
        categories = list(set(r.get('category', '') for r in risks))
        
        risk_summary = f"Total: {total}, Critical: {critical}, High: {high}, Medium: {medium}, Low: {low}. Categories: {', '.join(categories)}"
        
        if lang == 'ar':
            prompt = f"""أنت خبير في إدارة المخاطر المؤسسية ومتخصص في إطار COSO ERM وISO 31000.

بناءً على سجل المخاطر التالي:
- إجمالي المخاطر: {total}
- حرجة: {critical}، عالية: {high}، متوسطة: {medium}، منخفضة: {low}
- الفئات: {', '.join(categories)}

أنشئ بيان شهية المخاطر المؤسسية الشامل الذي يتضمن:

## 1. بيان شهية المخاطر العام
[فقرة تحدد المستوى العام لشهية المخاطر للمنظمة]

## 2. شهية المخاطر حسب الفئة
[لكل فئة من الفئات المذكورة، حدد مستوى الشهية: عالية، متوسطة، منخفضة، صفر]

### جدول شهية المخاطر:
| الفئة | مستوى الشهية | الحد الأقصى المقبول | المبرر |
|-------|-------------|-------------------|--------|

## 3. حدود التحمل ومعايير التصعيد
[جدول يوضح حدود التحمل لكل مستوى من المخاطر ومتى يتم التصعيد]

### مصفوفة التصعيد:
| مستوى الخطر | الإجراء المطلوب | الجهة المسؤولة | الإطار الزمني |
|------------|----------------|---------------|--------------|

## 4. أدوار ومسؤوليات إدارة المخاطر
[وفق نموذج الخطوط الثلاثة]

## 5. آلية المراجعة والتحديث
[كيفية مراجعة وتحديث بيان شهية المخاطر]

اكتب بيان شهية المخاطر بشكل احترافي ومتوافق مع COSO ERM وISO 31000:"""
        else:
            prompt = f"""You are an ERM expert specializing in COSO ERM Framework and ISO 31000.

Based on this risk register summary:
- Total risks: {total}
- Critical: {critical}, High: {high}, Medium: {medium}, Low: {low}
- Categories: {', '.join(categories)}

Generate a comprehensive Risk Appetite Statement that includes:

## 1. Overall Risk Appetite Statement
[Paragraph defining the organization's general risk appetite level]

## 2. Risk Appetite by Category
[For each category, specify appetite level: High, Moderate, Low, Zero]

### Risk Appetite Table:
| Category | Appetite Level | Maximum Acceptable Threshold | Justification |
|----------|---------------|------------------------------|---------------|

## 3. Tolerance Thresholds & Escalation Criteria
[Table showing tolerance limits for each risk level and escalation triggers]

### Escalation Matrix:
| Risk Level | Required Action | Responsible Party | Timeframe |
|-----------|----------------|-------------------|-----------|

## 4. Risk Management Roles & Responsibilities
[Based on Three Lines Model]

## 5. Review & Update Mechanism
[How the risk appetite statement is reviewed and updated]

Write a professional risk appetite statement aligned with COSO ERM and ISO 31000:"""
        
        content = generate_ai_content(prompt, lang, content_type='risk_appetite')
        
        return jsonify({
            'success': True,
            'content': content
        })
    except Exception as e:
        print(f"Risk appetite error: {e}", flush=True)
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# GAP REMEDIATION API (Feature 3)
# ============================================================================

@app.route('/api/gap-remediation', methods=['POST'])
@login_required
def api_gap_remediation():
    """Generate remediation plan for identified gaps."""
    try:
        data = request.json
        gaps = data.get('gaps', [])  # List of gap descriptions
        domain = data.get('domain', 'Cyber Security')
        framework = data.get('framework', 'ISO 27001')
        priority = data.get('priority', 'all')  # all, high, medium, low
        lang = data.get('language', 'en')
        
        if not gaps:
            return jsonify({'success': False, 'error': 'No gaps provided'}), 400
        
        gaps_text = '\n'.join([f"- {gap}" for gap in gaps])
        
        if lang == 'ar':
            prompt = f"""أنت مستشار خبير في الحوكمة والمخاطر والامتثال.

الفجوات المحددة في مجال {domain} وفقاً لإطار {framework}:
{gaps_text}

أنشئ خطة معالجة شاملة:

## خطة المعالجة

### ملخص تنفيذي
(ملخص موجز للفجوات وخطة المعالجة)

### خطة العمل التفصيلية
| # | الفجوة | الإجراء | المسؤول | الموارد | الجدول الزمني | مؤشر النجاح |
|---|--------|---------|---------|---------|--------------|-------------|
| 1 | ... | ... | ... | ... | ... | ... |

### الأولويات
#### عالية (فورية)
- ...

#### متوسطة (3-6 أشهر)
- ...

#### منخفضة (6-12 شهر)
- ...

### الميزانية التقديرية
| البند | التكلفة التقديرية |
|-------|------------------|

### مؤشرات الأداء
| المؤشر | القيمة الحالية | القيمة المستهدفة |
|--------|---------------|-----------------|

### المخاطر والتحديات
| المخاطرة | الأثر | خطة التخفيف |
|----------|-------|------------|"""
        else:
            prompt = f"""You are an expert GRC consultant.

Identified gaps in {domain} according to {framework}:
{gaps_text}

Create a comprehensive remediation plan:

## Remediation Plan

### Executive Summary
(brief summary of gaps and remediation approach)

### Detailed Action Plan
| # | Gap | Action | Owner | Resources | Timeline | Success Metric |
|---|-----|--------|-------|-----------|----------|----------------|
| 1 | ... | ... | ... | ... | ... | ... |

### Priorities
#### High (Immediate)
- ...

#### Medium (3-6 months)
- ...

#### Low (6-12 months)
- ...

### Estimated Budget
| Item | Estimated Cost |
|------|----------------|

### KPIs
| KPI | Current Value | Target Value |
|-----|---------------|--------------|

### Risks & Challenges
| Risk | Impact | Mitigation Plan |
|------|--------|-----------------|"""
        
        remediation_plan = generate_ai_content(prompt, lang, content_type='gap_remediation')
        
        return jsonify({
            'success': True,
            'remediation_plan': remediation_plan
        })
        
    except Exception as e:
        print(f"Gap remediation error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500
    
@app.route('/api/generate-policy', methods=['POST'])
@login_required
def api_generate_policy():
    """Generate policy document."""
    try:
        data = request.json
        domain = data.get('domain', 'Cyber Security')
        is_procedure = data.get('is_procedure', False)
        
        # Check usage limit for this domain
        can_generate, used, limit = check_usage_limit(session['user_id'], 'policies', domain)
        if not can_generate:
            return jsonify({
                'success': False,
                'error': f'Usage limit reached for {domain}. You have used {used}/{limit} policies in this domain.',
                'limit_reached': True
            }), 429
        
        lang = data.get('language', 'en')
        
        if lang == 'ar':
            policy_name = data.get('policy_name', 'أمن المعلومات')
            framework = data.get('framework', 'NCA ECC')
            framework_ar = translate_frameworks_list_ar(framework) if framework else framework
            prompt = f"""أنشئ وثيقة سياسة **{policy_name}** احترافية ومفصلة بتنسيق Markdown.

تقييد حاسم - عنوان السياسة:
- عنوان السياسة هو: "{policy_name}" وليس أي شيء آخر
- لا تغير العنوان إلى "سياسة أمن المعلومات" أو أي عنوان عام
- جميع البنود والإجراءات يجب أن تتعلق مباشرة وحصراً بموضوع "{policy_name}"
- لا تكتب بنوداً عامة عن أمن المعلومات إذا كان الموضوع المحدد مختلفاً

أمثلة للتوضيح:
- إذا كان العنوان "سياسة الاستخدام المقبول" → اكتب بنوداً عن استخدام الأجهزة والإنترنت والبريد الإلكتروني
- إذا كان العنوان "سياسة أمن الشبكات" → اكتب بنوداً عن جدران الحماية والتجزئة وVPN والمراقبة
- إذا كان العنوان "سياسة البريد الإلكتروني" → اكتب بنوداً عن استخدام البريد والتشفير والمرفقات والتصيد

تقييد الإطار التنظيمي:
- الإطار المرجعي: {framework_ar}
- لا تذكر أي إطار أو معيار آخر (مثل ISO 27001 أو NIST) إلا إذا كان هو المختار أعلاه

تعليمات:
1. لا تستخدم أي تواريخ محددة - استخدم عبارات نسبية مثل "سنوياً" و"كل 90 يوم"
2. لا تستخدم أي أسماء أشخاص
3. للتواريخ استخدم: [سيتم إضافته عند الاعتماد]
4. اكتب 4 أقسام فرعية على الأقل في بنود السياسة، كل منها خاص بجانب من جوانب {policy_name}
5. اكتب بأسلوب شركات الاستشارات الكبرى (مثل ديلويت و PwC و EY و KPMG) — لغة رسمية وموثوقة
6. استخدم عبارات مثل "يجب على المنظمة"، "يلتزم الموظفون"، "تُطبق جميع الأنظمة"
7. ارجع إلى رموز ضوابط محددة من {framework_ar} حيثما أمكن

# سياسة {policy_name}

## 1. الغرض
(وصف الغرض من هذه السياسة بشكل خاص بموضوع {policy_name} - لا تكتب غرضاً عاماً)

## 2. النطاق
تنطبق هذه السياسة على:
- (قائمة بالأطراف المعنية بموضوع {policy_name} تحديداً)

## 3. بنود السياسة
### 3.1 (عنوان فرعي خاص بجانب من {policy_name})
- (بنود مفصلة ومحددة لهذا الجانب)
### 3.2 (عنوان فرعي لجانب آخر من {policy_name})
- (بنود مفصلة)
### 3.3 (عنوان فرعي إضافي خاص بـ {policy_name})
- (بنود مفصلة)
### 3.4 (عنوان فرعي رابع خاص بـ {policy_name})
- (بنود مفصلة)

## 4. الأدوار والمسؤوليات
| الدور | المسؤوليات المتعلقة بـ {policy_name} |
|-------|--------------------------------------|
| المسمى | الوصف |

تنبيه حاسم: لا تضف أقسام "المهارات المطلوبة" أو "الشهادات" أو "المؤهلات" ضمن الأدوار والمسؤوليات. اكتف بالمسمى الوظيفي ومسؤولياته فقط.

## 5. متطلبات الامتثال وفق {framework_ar}
- (متطلبات الامتثال المرتبطة بإطار {framework_ar} فقط - لا تذكر أي إطار آخر)

## 6. المراجعة والتحديث
- إجراءات المراجعة الدورية

## 7. العقوبات
- العقوبات على عدم الالتزام

---
**تاريخ الإصدار:** [سيتم إضافته عند الاعتماد]
**رقم الإصدار:** 1.0
**المالك:** [القسم المسؤول]"""
        else:
            policy_name = data.get('policy_name', 'Information Security')
            framework = data.get('framework', 'NCA ECC')
            prompt = f"""Generate a professional **{policy_name}** Policy document in Markdown format.

CRITICAL - POLICY TITLE ENFORCEMENT:
- The policy title is: "{policy_name}" - do NOT change it to "Information Security Policy" or any other generic title
- ALL sections and statements must be DIRECTLY and SPECIFICALLY about "{policy_name}"
- Do NOT write generic information security content if the topic is different

Examples:
- If title is "Acceptable Use Policy" → write about device usage, internet access, email conduct, social media rules
- If title is "Network Security Policy" → write about firewalls, segmentation, VPN, monitoring, DMZ
- If title is "Email Security Policy" → write about email encryption, attachments, phishing, retention, forwarding rules
- If title is "Password Policy" → write about complexity, rotation, storage, MFA, service accounts

FRAMEWORK RESTRICTION:
- Reference framework: {framework}
- Do NOT mention any other framework (like ISO 27001, NIST, COBIT) unless it IS the selected framework above

INSTRUCTIONS:
1. Do NOT use any specific dates - use relative timeframes like "Annually", "Every 90 days"
2. Do NOT use any person names
3. For dates use: [To be added upon approval]
4. Write at least 4 subsections in Policy Statements, each addressing a specific aspect of {policy_name}
5. Write in formal, professional Big-4 consulting style (Deloitte, PwC, EY, KPMG)
6. Use authoritative language — "The Organization shall", "Personnel must", "All systems shall"
7. Reference specific {framework} control clauses where applicable

# {policy_name} Policy

## 1. Purpose
(Description specific to {policy_name} - NOT a generic purpose)

## 2. Scope
This policy applies to:
- (List of stakeholders specifically relevant to {policy_name})

## 3. Policy Statements
### 3.1 (Subheading about a specific aspect of {policy_name})
- (Detailed statements specific to this aspect)
### 3.2 (Another aspect of {policy_name})
- (Detailed statements)
### 3.3 (Additional aspect of {policy_name})
- (Detailed statements)
### 3.4 (Fourth aspect of {policy_name})
- (Detailed statements)

## 4. Roles & Responsibilities
| Role | Responsibilities related to {policy_name} |
|------|-------------------------------------------|
| Title | Description |

CRITICAL: Do NOT include "Required Skills", "Certifications", "Qualifications", or similar subsections in Roles & Responsibilities. Only list the role title and its specific responsibilities.

## 5. Compliance Requirements per {framework}
- (Requirements linked to {framework} only - do NOT mention other frameworks)

## 6. Review & Update
- Review procedures

## 7. Enforcement
- Penalties for non-compliance

---
**Issue Date:** [To be added upon approval]
**Version:** 1.0
**Owner:** [Responsible Department]"""

        if is_procedure:
            fw_name = data.get('framework', '')
            proc_name = data.get('policy_name', '')
            if lang == 'ar':
                framework_ar = translate_frameworks_list_ar(fw_name) if fw_name else fw_name
                prompt = f"""أنشئ وثيقة إجراء تشغيلي احترافية ومفصلة بعنوان "{proc_name}" بتنسيق Markdown.

تحذير حاسم: هذا إجراء تشغيلي وليس سياسة. لا تكتب بنود سياسة أو عبارات مثل "يجب على المنظمة". اكتب خطوات تنفيذية عملية يومية.

الإطار التنظيمي: {framework_ar}
لا تذكر أي إطار آخر غير {framework_ar}.

# إجراء {proc_name}

## 1. الغرض والنطاق
(وصف الغرض التشغيلي من هذا الإجراء وعلاقته بمتطلبات {framework_ar})

## 2. الضوابط المطبقة من {framework_ar}
| رقم الضابط | وصف الضابط | متطلب الامتثال |
|------------|-----------|---------------|
| (رقم) | (وصف) | (متطلب) |

## 3. الأدوار والمسؤوليات (مصفوفة RACI)
| النشاط | المسؤول (R) | المحاسب (A) | المستشار (C) | المُبلَّغ (I) |
|--------|------------|------------|-------------|-------------|
| (نشاط) | (دور) | (دور) | (دور) | (دور) |

## 4. المتطلبات المسبقة والأدوات
| المتطلب | الوصف | الحالة |
|---------|-------|-------|
| (أداة/نظام) | (وصف) | (مطلوب/اختياري) |

## 5. خطوات التنفيذ التفصيلية
### 5.1 (عنوان المرحلة الأولى)
**الخطوة 1:** (وصف تفصيلي مع إشارة لرقم ضابط {framework_ar})
**الخطوة 2:** (وصف)
...

### 5.2 (عنوان المرحلة الثانية)
...

## 6. نقاط القرار
لكل نقطة قرار في الإجراء، استخدم جدول قرارات بهذا التنسيق:

| نقطة القرار | الشرط | إذا نعم → الإجراء | إذا لا → الإجراء |
|-------------|-------|-------------------|------------------|
| (وصف نقطة القرار) | (الشرط المطلوب التحقق منه) | (الخطوات عند التحقق) | (الخطوات البديلة أو التصعيد) |

أضف عدة جداول قرار حسب الحاجة لتغطية جميع المسارات التشغيلية.

## 7. التحقق والمراجعة
(خطوات التحقق من صحة التنفيذ)

## 8. معالجة الاستثناءات
(إجراءات التعامل مع الحالات غير المعتادة)

## 9. متطلبات الأدلة والتوثيق
| نوع الدليل | الوصف | طريقة الحفظ | مدة الاحتفاظ |
|-----------|-------|------------|-------------|
| (دليل) | (وصف) | (طريقة) | (مدة) |

## 10. النماذج والقوالب
لكل نموذج أو قالب مطلوب، استخدم جدول بهذا التنسيق:

**نموذج: (عنوان النموذج)**

| الحقل | القيمة |
|-------|--------|
| معرف النموذج | [AUTO-GENERATED] |
| التاريخ | [DATE] |
| (اسم الحقل) | [VALUE] |
| (اسم الحقل) | [VALUE] |
| التوقيع | [SIGNATURE] |
| تاريخ المراجعة القادمة | [REVIEW_DATE] |

## 11. متطلبات حفظ السجلات
(متطلبات الأرشفة والتوثيق)

## 12. جدول المراجعة والتحديث
| التكرار | النوع | المسؤول |
|---------|-------|--------|
| (فترة) | (نوع المراجعة) | (الدور) |

---
**تاريخ الإصدار:** [سيتم إضافته عند الاعتماد]
**رقم الإصدار:** 1.0
**المالك:** [القسم المسؤول]

تعليمات إضافية:
- كل خطوة يجب أن تشير إلى رقم ضابط محدد من {framework_ar}
- استخدم أفعال إجرائية: "نفذ"، "تحقق"، "وثّق"، "أرسل"، "راجع"
- لا تستخدم تواريخ محددة
- لا تستخدم أسماء أشخاص
- اكتب بأسلوب تشغيلي عملي وليس سياسي
- تحذير حاسم: لا تستخدم أبداً أقواس الكود الثلاثية - استخدم الجداول فقط لعرض مخططات القرار والنماذج والقوالب"""
            else:
                prompt = f"""Generate a professional operational PROCEDURE document titled "{proc_name}" in Markdown format.

CRITICAL WARNING: This is an OPERATIONAL PROCEDURE, NOT a policy. Do NOT write policy statements like "The Organization shall". Write practical, day-to-day executable steps.

Reference Framework: {fw_name}
Do NOT mention any other framework.

# {proc_name} Procedure

## 1. Purpose & Scope
(Operational purpose of this procedure and its relation to {fw_name} requirements)

## 2. Applicable {fw_name} Controls
| Control ID | Control Description | Compliance Requirement |
|-----------|-------------------|----------------------|
| (ID) | (Description) | (Requirement) |

## 3. Roles & Responsibilities (RACI Matrix)
| Activity | Responsible (R) | Accountable (A) | Consulted (C) | Informed (I) |
|----------|-----------------|-----------------|---------------|-------------|
| (Activity) | (Role) | (Role) | (Role) | (Role) |

## 4. Prerequisites & Tools
| Requirement | Description | Status |
|------------|-------------|--------|
| (Tool/System) | (Description) | (Required/Optional) |

## 5. Detailed Step-by-Step Instructions
### 5.1 (Phase 1 Title)
**Step 1:** (Detailed description referencing specific {fw_name} control ID)
**Step 2:** (Description)
...

### 5.2 (Phase 2 Title)
...

## 6. Decision Points
For each decision point in the procedure, use a decision table in this format:

| Decision Point | Condition | If Yes → Action | If No → Action |
|---------------|-----------|-----------------|----------------|
| (Description of decision point) | (Condition to evaluate) | (Steps when condition is met) | (Alternative steps or escalation) |

Add multiple decision tables as needed to cover all operational paths.

## 7. Verification & Validation
(Steps to verify correct execution)

## 8. Exception Handling
(Procedures for unusual cases)

## 9. Evidence & Documentation Requirements
| Evidence Type | Description | Storage Method | Retention Period |
|--------------|-------------|----------------|-----------------|
| (Evidence) | (Description) | (Method) | (Period) |

## 10. Forms & Templates
For each required form or template, use a table in this format:

**Form: (Form Title)**

| Field | Value |
|-------|-------|
| Form ID | [AUTO-GENERATED] |
| Date | [DATE] |
| (Field Name) | [VALUE] |
| (Field Name) | [VALUE] |
| Signature | [SIGNATURE] |
| Next Review Date | [REVIEW_DATE] |

## 11. Record Keeping Requirements
(Archival and documentation requirements)

## 12. Review & Update Schedule
| Frequency | Review Type | Responsible |
|-----------|------------|-------------|
| (Period) | (Type) | (Role) |

---
**Issue Date:** [To be added upon approval]
**Version:** 1.0
**Owner:** [Responsible Department]

Additional Instructions:
- Every step must reference a specific {fw_name} control ID
- Use action verbs: "Execute", "Verify", "Document", "Submit", "Review"
- Do NOT use specific dates
- Do NOT use person names
- Write in practical operational style, NOT policy style
- CRITICAL: NEVER use code blocks (triple backticks) - use ONLY markdown tables for decision points, forms, and templates"""
        content = generate_ai_content(prompt, lang, content_type='policy')
        
        # Save to database
        try:
            conn = get_db()
            conn.execute('''INSERT INTO policies (user_id, domain, policy_name, framework, content, language)
                            VALUES (?, ?, ?, ?, ?, ?)''',
                        (session['user_id'], data.get('domain'), data.get('policy_name'),
                         data.get('framework'), content, lang))
            conn.commit()
        except Exception as db_error:
            print(f"Policy DB save error: {db_error}")
        
        log_action(session['user_id'], 'generate_policy', {'domain': data.get('domain'), 'policy_name': data.get('policy_name'), 'language': lang})
        return jsonify({'success': True, 'content': content})
        
    except Exception as e:
        print(f"Policy generation error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/analyze-risk', methods=['POST'])
@login_required
def api_analyze_risk():
    """Analyze risk scenario."""
    try:
        data = request.json
        domain = data.get('domain', 'Cyber Security')
        
        # Check usage limit for this domain
        can_generate, used, limit = check_usage_limit(session['user_id'], 'risks', domain)
        if not can_generate:
            return jsonify({
                'success': False,
                'error': f'Usage limit reached for {domain}. You have used {used}/{limit} risk analyses in this domain.',
                'limit_reached': True
            }), 429
        
        lang = data.get('language', 'en')
        
        if lang == 'ar':
            # Determine expected controls for Arabic
            threat_text_ar = data.get('threat', '').lower()
            scenario_hint_ar = ''
            if any(k in threat_text_ar for k in ['insider', 'داخلي', 'موظف', 'تهديد داخلي']):
                scenario_hint_ar = f"""
لهذا السيناريو (تهديد داخلي يستهدف {data.get('asset', 'الأصل')}):
1. تحليلات سلوك المستخدم والكيان (UEBA) - للكشف عن السلوك غير الطبيعي
2. مراقبة نشاط قاعدة البيانات (DAM) - لمراقبة الاستعلامات المشبوهة
3. منع تسريب البيانات (DLP) - خاصة لمسارات تصدير البيانات
لا تقترح "تفعيل MFA" أو "تدريب التوعية" كضوابط أساسية للتهديدات الداخلية."""
            elif any(k in threat_text_ar for k in ['فدية', 'ransomware', 'تشفير.*خبيث']):
                scenario_hint_ar = f"""
لهذا السيناريو (فدية تستهدف {data.get('asset', 'الأصل')}):
1. EDR مع كشف سلوكي وعزل تلقائي
2. نسخ احتياطية معزولة وغير قابلة للتعديل مع إجراءات استعادة مختبرة
3. بوابة بريد مع حماية رمل للمرفقات وتفجير URL"""
            elif any(k in threat_text_ar for k in ['تصيد', 'phishing', 'هندسة اجتماعية']):
                scenario_hint_ar = f"""
لهذا السيناريو (تصيد/هندسة اجتماعية):
1. بوابة أمن بريد مع تطبيق DMARC/DKIM/SPF
2. برنامج محاكاة تصيد دوري مع تتبع المقاييس
3. قواعد كشف اختراق البريد التجاري (BEC)"""

            prompt = f"""حلل سيناريو الخطر التالي بتنسيق Markdown احترافي:
الفئة: {data.get('category', 'عام')}
الأصل: {data.get('asset', 'النظام')}
التهديد: {data.get('threat', 'وصول غير مصرح')}

⚠⚠⚠ أهم قاعدة — ضوابط خاصة بالسيناريو فقط:
أنت تحلل سيناريو محدد: "{data.get('threat', '')}" يستهدف "{data.get('asset', '')}".
الضوابط الموصى بها يجب أن تكون مخصصة لهذا التهديد والأصل بالتحديد.

ضوابط محظورة (لا تستخدمها):
❌ "تفعيل المصادقة متعددة العوامل" — عامة جداً لتحليل تهديد محدد
❌ "تحديث أنظمة الكشف عن التهديدات" — غامضة وغير مستهدفة
❌ "تدريب التوعية الأمنية للموظفين" — ممارسة عامة وليست ضابطاً خاصاً بالتهديد

{scenario_hint_ar}

كل ضابط يجب أن يسمي تقنية أو أداة أو تقنية محددة تعالج مباشرة "{data.get('threat', '')}".
الضوابط يجب أن تعالج مباشرة التهديد "{data.get('threat', '')}" الذي يستهدف الأصل "{data.get('asset', '')}".

تعليمات صارمة ومهمة جداً:
1. لا تستخدم أي تواريخ محددة مطلقاً (مثل 2024، 2025، يناير، فبراير، إلخ)
2. لا تستخدم أي أسماء أشخاص
3. استخدم فقط عبارات نسبية مثل: "خلال 30 يوم"، "خلال 60 يوم"، "خلال 90 يوم"
4. للتواريخ استخدم: [سيتم إضافته]

استخدم التنسيق التالي:

# تحليل المخاطر

## ملخص تقييم الخطر
| العنصر | القيمة |
|--------|-------|
| فئة الخطر | {data.get('category', 'عام')} |
| الأصل المتأثر | {data.get('asset', 'النظام')} |
| مستوى الخطر | [عالي/متوسط/منخفض] |
| درجة الخطر | [X/10] |

## تحليل التهديد
وصف التهديد ومصادره المحتملة

## تحليل الأثر
| نوع الأثر | الوصف | المستوى |
|----------|-------|---------|
| مالي | الوصف | المستوى |
| تشغيلي | الوصف | المستوى |

## تقييم الاحتمالية
| العامل | التقييم |
|--------|---------|
| العامل | القيمة |

## الضوابط الموصى بها

### جدول ملخص الضوابط:
| # | الضابط | الأولوية | الجدول الزمني | التكلفة المقدرة | الحالة |
|---|--------|----------|---------------|-----------------|--------|
| 1 | [اسم الضابط] | عالية | فوري | يحتاج تقييم |
| 2 | [اسم الضابط] | عالية | خلال 30 يوم | يحتاج تقييم |
| 3 | [اسم الضابط] | متوسطة | خلال 60 يوم | يحتاج تقييم |

### دليل تنفيذ الضوابط:

يجب كتابة دليل تنفيذ منفصل لكل ضابط موصى به (3 أدلة على الأقل)

---
#### دليل تنفيذ الضابط رقم 1: [اسم الضابط]

**الخطوات التفصيلية:**
| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |
|---------|--------|-------|---------|----------|
| التحضير | 1.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحضير | 1.2 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.2 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحقق | 3.1 | (خطوة محددة) | (الفريق) | (المخرج) |

**الأدلة المطلوبة للإغلاق:**
- [ ] (دليل 1)
- [ ] (دليل 2)
- [ ] (دليل 3)

---
#### دليل تنفيذ الضابط رقم 2: [اسم الضابط]

**الخطوات التفصيلية:**
| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |
|---------|--------|-------|---------|----------|
| التحضير | 1.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحضير | 1.2 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.2 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحقق | 3.1 | (خطوة محددة) | (الفريق) | (المخرج) |

**الأدلة المطلوبة للإغلاق:**
- [ ] (دليل 1)
- [ ] (دليل 2)
- [ ] (دليل 3)

---
#### دليل تنفيذ الضابط رقم 3: [اسم الضابط]

**الخطوات التفصيلية:**
| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |
|---------|--------|-------|---------|----------|
| التحضير | 1.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحقق | 3.1 | (خطوة محددة) | (الفريق) | (المخرج) |

**الأدلة المطلوبة للإغلاق:**
- [ ] (دليل 1)
- [ ] (دليل 2)
- [ ] (دليل 3)

---

### ضوابط كاشفة
- القائمة

### ضوابط تصحيحية
- القائمة

## الخطر المتبقي
| السيناريو | قبل الضوابط | بعد الضوابط |
|----------|------------|------------|

---
**تاريخ التقييم:** [سيتم إضافته]
**المراجعة القادمة:** خلال 6 أشهر"""
        else:
            # Determine expected controls based on threat keywords
            threat_text = data.get('threat', '').lower()
            asset_text = data.get('asset', '').lower()
            scenario_hint = ''
            if any(k in threat_text for k in ['insider', 'internal', 'employee', 'disgruntled']):
                scenario_hint = f"""
For this INSIDER THREAT scenario targeting {data.get('asset', 'the asset')}, the controls MUST include:
1. User and Entity Behavior Analytics (UEBA) - to detect anomalous user behavior
2. Database Activity Monitoring (DAM) - to monitor and alert on suspicious database queries
3. Data Loss Prevention (DLP) - specifically for database export paths and bulk data transfers
Do NOT suggest "Enable MFA" or "Security awareness training" as primary controls for insider threats."""
            elif any(k in threat_text for k in ['ransomware', 'ransom', 'encrypt']):
                scenario_hint = f"""
For this RANSOMWARE scenario targeting {data.get('asset', 'the asset')}, the controls MUST include:
1. EDR with behavioral detection and automated isolation
2. Immutable air-gapped backup strategy with tested restore procedures
3. Email gateway with attachment sandboxing and URL detonation"""
            elif any(k in threat_text for k in ['phishing', 'spear', 'social engineer', 'bec']):
                scenario_hint = f"""
For this PHISHING/SOCIAL ENGINEERING scenario, the controls MUST include:
1. Email security gateway with DMARC/DKIM/SPF enforcement
2. Phishing simulation campaign program with metrics tracking
3. Business Email Compromise (BEC) detection rules"""
            elif any(k in threat_text for k in ['ddos', 'denial', 'flood']):
                scenario_hint = f"""
For this DDoS scenario, the controls MUST include:
1. WAF with rate limiting and bot detection
2. CDN-based DDoS mitigation service
3. Traffic scrubbing and clean-pipe service"""
            elif any(k in threat_text for k in ['breach', 'exfiltrat', 'leak', 'exposure']):
                scenario_hint = f"""
For this DATA BREACH scenario, the controls MUST include:
1. Data classification system with automated labeling
2. DLP across all channels (email, web, USB, cloud)
3. Database activity monitoring with access anomaly detection"""
            elif any(k in threat_text for k in ['cloud', 'misconfig', 'saas']):
                scenario_hint = f"""
For this CLOUD SECURITY scenario, the controls MUST include:
1. Cloud Security Posture Management (CSPM) with continuous scanning
2. Cloud IAM hardening with least-privilege enforcement
3. Cloud workload protection platform (CWPP)"""

            prompt = f"""Analyze this risk scenario in professional Markdown format:
Category: {data.get('category', 'General')}
Asset: {data.get('asset', 'System')}
Threat: {data.get('threat', 'Unauthorized Access')}

⚠⚠⚠ MOST CRITICAL RULE — SCENARIO-SPECIFIC CONTROLS ONLY:
You are analyzing a SPECIFIC scenario: "{data.get('threat', '')}" targeting "{data.get('asset', '')}".
The recommended controls MUST be tailored to THIS EXACT threat and asset combination.

BANNED GENERIC CONTROLS (do NOT use these):
❌ "Enable Multi-Factor Authentication" — too generic for a specific threat analysis
❌ "Update Threat Detection Systems" — vague and not targeted
❌ "Employee Security Awareness Training" — this is a general practice, not a threat-specific control

{scenario_hint}

If you cannot identify threat-specific controls, you are doing this wrong. Every control must NAME a specific technology, tool, or technique that directly mitigates "{data.get('threat', '')}".

The controls MUST directly address the specific threat "{data.get('threat', '')}" targeting the specific asset "{data.get('asset', '')}".
Each control implementation guide must include tools/technologies specific to this threat scenario.

STRICT AND IMPORTANT INSTRUCTIONS:
1. Do NOT use any specific dates (like 2024, 2025, January, February, etc.)
2. Do NOT use any person names
3. Use ONLY relative timeframes like: "Within 30 days", "Within 60 days", "Within 90 days"
4. For dates use: [To be added]

Use the following format:

# Risk Analysis

## Risk Assessment Summary
| Element | Value |
|---------|-------|
| Risk Category | {data.get('category', 'General')} |
| Affected Asset | {data.get('asset', 'System')} |
| Risk Level | [High/Medium/Low] |
| Risk Score | [X/10] |

## Threat Analysis
Description of threat and potential sources

## Impact Analysis
| Impact Type | Description | Level |
|-------------|-------------|-------|
| Financial | Description | Level |
| Operational | Description | Level |

## Likelihood Assessment
| Factor | Assessment |
|--------|------------|
| Factor | Value |

## Recommended Controls

### Control Summary Table:
| # | Control | Priority | Timeline | Status |
|---|---------|----------|----------|--------|
| 1 | [Control name] | High | Within 30 days | To Be Assessed |
| 2 | [Control name] | Medium | Within 60 days | To Be Assessed |
| 3 | [Control name] | Low | Within 90 days | To Be Assessed |

### Control Implementation Guidelines:

---
#### Control #1 Implementation Guide: [Control Name]

**Step-by-Step Implementation:**
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Preparation | 1.1 | [Step description] | [Team] | [Output] |
| Preparation | 1.2 | [Step description] | [Team] | [Output] |
| Implementation | 2.1 | [Step description] | [Team] | [Output] |
| Implementation | 2.2 | [Step description] | [Team] | [Output] |
| Verification | 3.1 | [Step description] | [Team] | [Output] |

**Evidence Required:**
- [ ] [Evidence item 1]
- [ ] [Evidence item 2]
- [ ] [Evidence item 3]

---
#### Control #2 Implementation Guide: [Control Name]

**Step-by-Step Implementation:**
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Preparation | 1.1 | [Step description] | [Team] | [Output] |
| Preparation | 1.2 | [Step description] | [Team] | [Output] |
| Implementation | 2.1 | [Step description] | [Team] | [Output] |
| Implementation | 2.2 | [Step description] | [Team] | [Output] |
| Verification | 3.1 | [Step description] | [Team] | [Output] |

**Evidence Required:**
- [ ] [Evidence item 1]
- [ ] [Evidence item 2]
- [ ] [Evidence item 3]

---
#### Control #3 Implementation Guide: [Control Name]

**Step-by-Step Implementation:**
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Preparation | 1.1 | [Step description] | [Team] | [Output] |
| Implementation | 2.1 | [Step description] | [Team] | [Output] |
| Verification | 3.1 | [Step description] | [Team] | [Output] |

**Evidence Required:**
- [ ] [Evidence item 1]
- [ ] [Evidence item 2]
- [ ] [Evidence item 3]

---

### Detective Controls
- List

### Corrective Controls
- List

## Residual Risk
| Scenario | Before Controls | After Controls |
|----------|-----------------|----------------|

---
**Assessment Date:** [To be added]
**Next Review:** Within 6 months"""

        content = generate_ai_content(prompt, lang, content_type='risk')
        
        # POST-PROCESSING: Check if AI produced generic controls despite instructions
        # If generic, replace with scenario-specific simulation content
        generic_markers = [
            'enable multi-factor authentication', 'enable mfa', 
            'update threat detection', 'security awareness training',
            'تفعيل المصادقة متعددة العوامل', 'تحديث أنظمة الكشف', 'تدريب التوعية الأمنية',
        ]
        content_lower = content.lower() if content else ''
        generic_count = sum(1 for m in generic_markers if m in content_lower)
        if generic_count >= 2:
            # AI produced generic garbage — use scenario-specific simulation instead
            print(f"DEBUG: Risk AI output had {generic_count} generic controls — using simulation", flush=True)
            sim_content = generate_simulation_content(prompt, lang)
            if sim_content:
                content = sim_content
        
        # INJECT IMPLEMENTATION GUIDELINES if missing from risk content
        def inject_risk_guidelines(risk_content, lang_code):
            """Inject implementation guidelines if AI didn't include them."""
            has_guidelines = (
                'Implementation Guide' in risk_content or 
                'Step-by-Step' in risk_content or
                'دليل تنفيذ' in risk_content or
                'الخطوات التفصيلية' in risk_content
            )
            
            if has_guidelines:
                return risk_content
            
            # Extract controls from table
            import re
            control_pattern = r'\|\s*(\d+)\s*\|\s*([^|]+)\s*\|\s*(High|Medium|Low|عالية|متوسطة|منخفضة)'
            controls = re.findall(control_pattern, risk_content)
            
            controls_list = []
            for num, control, priority in controls:
                control = control.strip()
                if control and control.lower() not in ['control', 'الضابط', '#']:
                    controls_list.append((num, control, priority))
            
            if not controls_list:
                return risk_content
            
            if lang_code == 'ar':
                guidelines = "\n\n---\n\n### دليل تنفيذ الضوابط:\n\n"
                for num, control, priority in controls_list:
                    guidelines += f"""---
#### دليل تنفيذ الضابط رقم {num}: {control[:40]}

**الخطوات التفصيلية:**
| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |
|---------|--------|-------|---------|----------|
| التحضير | 1.1 | تقييم المتطلبات والموارد | فريق الأمن | تقرير المتطلبات |
| التحضير | 1.2 | إعداد خطة التنفيذ | مدير المشروع | خطة معتمدة |
| التنفيذ | 2.1 | تطبيق الضابط | الفريق التقني | ضابط مُفعّل |
| التنفيذ | 2.2 | تدريب المستخدمين | التدريب | فريق مدرب |
| التحقق | 3.1 | اختبار الفعالية | ضمان الجودة | تقرير الاختبار |

**الأدلة المطلوبة للإغلاق:** ☐ تقرير المتطلبات ☐ خطة التنفيذ ☐ سجلات التدريب ☐ تقرير الاختبار

"""
            else:
                guidelines = "\n\n---\n\n### Control Implementation Guidelines:\n\n"
                for num, control, priority in controls_list:
                    guidelines += f"""---
#### Control #{num} Implementation Guide: {control[:40]}

**Step-by-Step Implementation:**
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Preparation | 1.1 | Assess requirements and resources | Security Team | Requirements report |
| Preparation | 1.2 | Develop implementation plan | Project Manager | Approved plan |
| Implementation | 2.1 | Apply the control | Technical Team | Activated control |
| Implementation | 2.2 | Train users | Training | Trained team |
| Verification | 3.1 | Test effectiveness | QA | Test report |

**Evidence Required for Closure:** ☐ Requirements report ☐ Implementation plan ☐ Training records ☐ Test report

"""
            
            return risk_content + guidelines
        
        content = inject_risk_guidelines(content, lang)
        
        # ── Extract risk level from AI-generated content ──
        import re
        risk_level = 'HIGH'  # Default
        level_patterns = [
            (r'(?:Risk Level|مستوى الخطر)\s*\|\s*(Critical|High|Medium|Low|حرج|عالي|متوسط|منخفض)', 1),
            (r'(?:Risk Score|درجة الخطر)\s*\|\s*(\d+)/10', 1),
        ]
        for pattern, group in level_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                val = match.group(group).strip().lower()
                if val in ['critical', 'حرج'] or (val.isdigit() and int(val) >= 9):
                    risk_level = 'CRITICAL'
                elif val in ['high', 'عالي'] or (val.isdigit() and int(val) >= 7):
                    risk_level = 'HIGH'
                elif val in ['medium', 'متوسط'] or (val.isdigit() and int(val) >= 4):
                    risk_level = 'MEDIUM'
                else:
                    risk_level = 'LOW'
                break
        
        # Save to database
        risk_id = None
        try:
            conn = get_db()
            conn.execute('''INSERT INTO risks (user_id, domain, asset_name, threat, risk_level, analysis, language)
                            VALUES (?, ?, ?, ?, ?, ?, ?)''',
                        (session['user_id'], data.get('domain'), data.get('asset'),
                         data.get('threat'), risk_level, content, lang))
            conn.commit()
            risk_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        except Exception as db_error:
            print(f"Risk DB save error: {db_error}")
        
        # ── Pillar 2: Auto-create mitigation task for Critical/High AI-analyzed risks ──
        auto_task_created = False
        if risk_id and risk_level in ('CRITICAL', 'HIGH'):
            risk_score = 25 if risk_level == 'CRITICAL' else 16  # Map to numeric score for auto_create
            task_id = auto_create_mitigation_task(
                session['user_id'], risk_id,
                f"{data.get('threat', 'Risk')} → {data.get('asset', 'Asset')}",
                risk_score, data.get('domain', 'General')
            )
            auto_task_created = task_id is not None
        
        # ── PIPE: Auto-extract risk into Risk Register ──
        pipe_register_id = None
        try:
            pipe_register_id = pipe_risk_assessment_to_register(
                content, session['user_id'], data.get('domain', 'General'),
                asset=data.get('asset', ''), threat=data.get('threat', ''),
                category=data.get('category', 'Operational')
            )
        except Exception as pipe_err:
            print(f"Pipe risk-to-register error (non-fatal): {pipe_err}", flush=True)
        
        log_action(session['user_id'], 'analyze_risk', {
            'domain': data.get('domain'), 'category': data.get('category'),
            'language': lang, 'risk_level': risk_level,
            'auto_task_created': auto_task_created,
            'pipe_register_id': pipe_register_id
        })
        return jsonify({
            'success': True, 'analysis': content,
            'risk_level': risk_level, 'auto_task': auto_task_created,
            'pipe_register_id': pipe_register_id
        })
        
    except Exception as e:
        print(f"Risk analysis error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/generate-audit', methods=['POST'])
@login_required
def api_generate_audit():
    """Generate audit report."""
    try:
        domain = request.form.get('domain', 'Cyber Security')
        
        # Check usage limit for this domain
        can_generate, used, limit = check_usage_limit(session['user_id'], 'audits', domain)
        if not can_generate:
            return jsonify({
                'success': False,
                'error': f'Usage limit reached for {domain}. You have used {used}/{limit} audit reports in this domain.',
                'limit_reached': True
            }), 429
        
        lang = request.form.get('language', 'en')
        framework = request.form.get('framework', 'NCA ECC')
        audit_scope = request.form.get('audit_scope', 'full')
        audit_topic = request.form.get('audit_topic', '').strip()
        
        # Read generated policy content (sent from frontend if user generated a policy first)
        policy_content = request.form.get('policy_content', '').strip()
        if policy_content and len(policy_content) > 10000:
            policy_content = policy_content[:10000] + "\n..."
        
        # Handle file upload - read content for AI analysis
        evidence_files = request.files.getlist('evidence')
        evidence_info = []
        uploaded_content = ''
        for f in evidence_files:
            if f and f.filename:
                evidence_info.append(f.filename)
                # Read file content for AI analysis
                try:
                    if f.filename.lower().endswith('.pdf'):
                        # Try to extract text from PDF
                        import io
                        pdf_bytes = f.read()
                        try:
                            import fitz  # PyMuPDF
                            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                            pdf_text = ''
                            for page in doc:
                                pdf_text += page.get_text()
                            doc.close()
                            if pdf_text.strip():
                                uploaded_content += f"\n\n--- Content of {f.filename} ---\n{pdf_text[:8000]}\n--- End of {f.filename} ---\n"
                        except (ImportError, Exception):
                            try:
                                from PyPDF2 import PdfReader
                                reader = PdfReader(io.BytesIO(pdf_bytes))
                                pdf_text = ''
                                for page in reader.pages:
                                    pdf_text += (page.extract_text() or '')
                                if pdf_text.strip():
                                    uploaded_content += f"\n\n--- Content of {f.filename} ---\n{pdf_text[:8000]}\n--- End of {f.filename} ---\n"
                            except:
                                print(f"DEBUG: No PDF library available for {f.filename}", flush=True)
                    elif f.filename.lower().endswith(('.txt', '.md', '.csv')):
                        text_content = f.read().decode('utf-8', errors='ignore')
                        if text_content.strip():
                            uploaded_content += f"\n\n--- Content of {f.filename} ---\n{text_content[:8000]}\n--- End of {f.filename} ---\n"
                    elif f.filename.lower().endswith(('.docx',)):
                        try:
                            import docx
                            doc = docx.Document(io.BytesIO(f.read()))
                            doc_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                            if doc_text.strip():
                                uploaded_content += f"\n\n--- Content of {f.filename} ---\n{doc_text[:8000]}\n--- End of {f.filename} ---\n"
                        except:
                            pass
                except Exception as e:
                    print(f"DEBUG: Could not read file {f.filename}: {e}", flush=True)
        
        # Build topic-specific instruction for AI
        topic_instruction = ''
        if audit_topic:
            topic_instruction = f"""
AUDIT TOPIC: {audit_topic}
This audit focuses specifically on "{audit_topic}". ALL findings must be operationally specific to this topic.
Generate findings that reflect real-world gaps an auditor would find in a "{audit_topic}" context — not generic boilerplate."""
        
        # Combine all document sources
        document_to_audit = policy_content or uploaded_content or ''
        
        if document_to_audit and lang != 'ar':
            # ======================================================================
            # DOCUMENT-CENTRIC AUDIT: AI reads and audits the actual document
            # ======================================================================
            prompt = f"""You are a professional GRC auditor. You have an actual policy document to audit against {framework}.

## DOCUMENT UNDER AUDIT:
{document_to_audit}

## TASK:
Read the document above carefully. Analyze its actual content and produce an audit report identifying:
1. What the document covers well
2. What is MISSING or needs improvement against {framework} requirements
3. Specific recommendations for each finding with detailed implementation steps

⚠ CRITICAL RULES:
- ALL findings must be based on what you ACTUALLY found (or did NOT find) in the document above
- Do NOT generate generic boilerplate findings — each finding must reference a specific missing or weak area in the document
- Reference ONLY {framework} and its specific controls
- Do NOT use specific dates — use "Within 30 days", "Within 60 days"
- Do NOT use person names or auditor names
- Do NOT echo any instruction text from this prompt
- Provide a detailed implementation guide for EVERY finding (high, medium, AND low-risk)

Use the following format:

# Audit Report - {audit_topic or 'Policy'} vs {framework}

## Executive Summary
Brief overview. **Overall Result:** [Compliant / Partially Compliant / Non-Compliant] with count of findings by severity based on your analysis of the document.

## Findings & Observations

### High-Risk Findings
| # | Observation | Affected Control ({framework}) | Recommendation | Status |
|---|-------------|------------------------|----------------|--------|
(2 findings about critical gaps you found in the document)

### Medium-Risk Findings
| # | Observation | Affected Control | Recommendation | Status |
|---|-------------|-----------------|----------------|--------|
(2 findings about moderate gaps)

### Low-Risk Findings
| # | Observation | Affected Control | Recommendation | Status |
|---|-------------|-----------------|----------------|--------|
(1-2 findings about minor gaps)

## Detailed Implementation Guidelines

Provide a SEPARATE implementation guide for EVERY finding — high, medium, AND low-risk:

### Finding #1: [Title]
**Affected Control:** [Control] ({framework})
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
(5-6 specific implementation steps)
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

(Repeat for ALL findings — #2, #3, #4, #5 including low-risk findings. Every finding MUST have its own implementation guide.)

## Action Plan
| # | Action | Owner | Deadline | Priority |
|---|--------|-------|----------|----------|
(Each row must contain the actual recommendation, not "Address Finding #N")

---
**Report Date:** [To be added]
**Next Audit:** Within 6 months"""

        elif not document_to_audit and lang != 'ar':
            # ==================================================================
            # TOPIC-BASED AUDIT: No document content available
            # ==================================================================
            prompt = f"""Generate a comprehensive audit report in professional Markdown format.

Audit Parameters:
- Framework: {framework}
- Scope: {audit_scope}
- Domain: {domain}
{"- Audit Topic: " + audit_topic if audit_topic else ""}
- Evidence Files: {', '.join(evidence_info) if evidence_info else 'No evidence provided'}
{uploaded_content if uploaded_content else ""}
{"DOCUMENT UNDER AUDIT: The content above from the uploaded document(s) is the actual document being audited. Analyze it against " + framework + " requirements and produce specific findings based on what IS and IS NOT covered in the document." if uploaded_content else ""}

⚠ FRAMEWORK RULES:
- Reference ONLY {framework} and its specific control IDs
- Do NOT mention any other framework (ISO, NIST, COBIT, etc.) unless it IS {framework}
{topic_instruction}

⚠ OUTPUT RULES:
1. Do NOT use any specific dates — use relative timeframes like "Within 30 days", "Within 60 days"
2. Do NOT use any person names or auditor names
3. For dates use: [To be added] | For audit period use: [Audit Period]
4. Every finding (high, medium, AND low-risk) must include detailed implementation guidelines
5. Do NOT include "Audit Scope" or "Audit Methodology" sections — jump directly to findings
6. Do NOT echo any instruction text from this prompt into the output
7. Write in professional Big-4 consulting audit style (formal, authoritative, evidence-based)

Use the following format:

# Audit Report - {framework}{" - " + audit_topic if audit_topic else ""}

## Executive Summary
Brief overview including: **Overall Result:** [Compliant / Partially Compliant / Non-Compliant] with count of findings by severity. Do NOT write "To be determined" — provide an actual compliance assessment {"of " + audit_topic + " " if audit_topic else ""}against {framework} requirements.

## Findings & Observations

### High-Risk Findings
| # | Observation | Affected Control | Recommendation | Status |
|---|-------------|-----------------|----------------|--------|
| 1 | [First high-risk finding{"  related to " + audit_topic if audit_topic else ""}] | [Control ID] | [Action] | To Be Assessed |
| 2 | [Second high-risk finding{"  related to " + audit_topic if audit_topic else ""}] | [Control ID] | [Action] | To Be Assessed |

### Medium-Risk Findings
| # | Observation | Affected Control | Recommendation | Status |
|---|-------------|-----------------|----------------|--------|
| 1 | [First medium-risk finding{"  related to " + audit_topic if audit_topic else ""}] | [Control ID] | [Action] | To Be Assessed |
| 2 | [Second medium-risk finding{"  related to " + audit_topic if audit_topic else ""}] | [Control ID] | [Action] | To Be Assessed |

### Low-Risk Findings
| # | Observation | Affected Control | Recommendation | Status |
|---|-------------|-----------------|----------------|--------|

## Detailed Implementation Guidelines

You MUST provide SEPARATE implementation guides for EVERY finding — high, medium, AND low-risk. Do NOT write "repeat" - write actual steps for each.

---
### High-Risk Finding #1 Implementation Guide: [First Finding Title]
**Affected Control:** [Control ID]
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Planning | 1.1 | [Specific step for this finding] | [Team] | [Output] |
| Planning | 1.2 | [Specific step for this finding] | [Team] | [Output] |
| Implementation | 2.1 | [Specific step for this finding] | [Team] | [Output] |
| Implementation | 2.2 | [Specific step for this finding] | [Team] | [Output] |
| Verification | 3.1 | [Specific step for this finding] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

---
### High-Risk Finding #2 Implementation Guide: [Second Finding Title]
**Affected Control:** [Control ID]
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Planning | 1.1 | [Specific step for this finding] | [Team] | [Output] |
| Planning | 1.2 | [Specific step for this finding] | [Team] | [Output] |
| Implementation | 2.1 | [Specific step for this finding] | [Team] | [Output] |
| Implementation | 2.2 | [Specific step for this finding] | [Team] | [Output] |
| Verification | 3.1 | [Specific step for this finding] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

---
### Medium-Risk Finding #1 Implementation Guide: [First Medium Finding]
**Affected Control:** [Control ID]
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Planning | 1.1 | [Specific step for this finding] | [Team] | [Output] |
| Implementation | 2.1 | [Specific step for this finding] | [Team] | [Output] |
| Verification | 3.1 | [Specific step for this finding] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

---
### Medium-Risk Finding #2 Implementation Guide: [Second Medium Finding]
**Affected Control:** [Control ID]
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Planning | 1.1 | [Specific step for this finding] | [Team] | [Output] |
| Implementation | 2.1 | [Specific step for this finding] | [Team] | [Output] |
| Verification | 3.1 | [Specific step for this finding] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

---
### Low-Risk Finding #1 Implementation Guide: [Low Finding Title]
**Affected Control:** [Control ID]
| Phase | Step | Description | Owner | Deliverable |
|-------|------|-------------|-------|-------------|
| Planning | 1.1 | [Specific step for this finding] | [Team] | [Output] |
| Implementation | 2.1 | [Specific step for this finding] | [Team] | [Output] |
| Verification | 3.1 | [Specific step for this finding] | [Team] | [Output] |
**Evidence Required:** ☐ [Evidence 1] ☐ [Evidence 2] ☐ [Evidence 3]

## Action Plan
| # | Action | Owner | Deadline | Priority |
|---|--------|-------|----------|----------|
| 1 | Address high-risk findings | Security Team | Within 30 days | High |
| 2 | Address medium-risk findings | Compliance Team | Within 60 days | Medium |
| 3 | Address low-risk findings | Operations Team | Within 90 days | Low |

---
**Report Date:** [To be added]
**Next Audit:** Within 6 months"""

        if lang == 'ar' and document_to_audit:
            # ======================================================================
            # ARABIC DOCUMENT-CENTRIC AUDIT
            # ======================================================================
            prompt = f"""أنت مدقق GRC محترف. لديك وثيقة سياسة فعلية يجب تدقيقها مقابل إطار {framework}.

## الوثيقة الخاضعة للتدقيق:
{document_to_audit}

## المطلوب:
اقرأ الوثيقة أعلاه بعناية. حلل محتواها الفعلي وأنشئ تقرير تدقيق يحدد:
1. ما تغطيه الوثيقة بشكل جيد
2. ما ينقصها أو يحتاج تحسين مقابل متطلبات {framework}
3. توصيات محددة لكل ملاحظة مع خطوات تنفيذ مفصلة

⚠ قواعد حرجة:
- يجب أن تكون جميع الملاحظات مبنية على ما وجدته فعلياً في الوثيقة (أو ما لم تجده)
- لا تولّد ملاحظات عامة — كل ملاحظة يجب أن تشير إلى قسم محدد مفقود أو ضعيف
- ارجع فقط إلى {framework} وضوابطه
- لا تستخدم تواريخ محددة — استخدم "خلال 30 يوم" و"خلال 60 يوم"
- لا تكتب أي تعليمات من هذا النص في المخرجات
- قدم دليل تنفيذ مفصل لكل ملاحظة (عالية ومتوسطة ومنخفضة الخطورة)

# تقرير تدقيق سياسة {audit_topic or 'السياسة'} - وفق {framework}

## الملخص التنفيذي
**النتيجة العامة:** [ممتثل / ممتثل جزئياً / غير ممتثل] — مع ذكر عدد الملاحظات حسب الخطورة بناءً على تحليلك للوثيقة.

## النتائج والملاحظات

### نتائج عالية الخطورة
| # | الملاحظة | الضابط المتأثر ({framework}) | التوصية | الحالة |
|---|----------|--------------------------|---------|--------|

### نتائج متوسطة الخطورة
| # | الملاحظة | الضابط المتأثر | التوصية | الحالة |
|---|----------|---------------|---------|--------|

### نتائج منخفضة الخطورة
| # | الملاحظة | الضابط المتأثر | التوصية | الحالة |
|---|----------|---------------|---------|--------|

## أدلة التنفيذ التفصيلية
يجب كتابة دليل تنفيذ منفصل لكل ملاحظة (عالية ومتوسطة ومنخفضة الخطورة) مع جدول خطوات وأدلة مطلوبة. لا تترك أي ملاحظة بدون دليل تنفيذ.

## خطة العمل
| # | الإجراء | المسؤول | الموعد النهائي | الأولوية |
|---|--------|---------|---------------|----------|
(كل صف يحتوي التوصية الفعلية وليس "معالجة الملاحظة رقم X")

---
**تاريخ التقرير:** يُحدد لاحقاً
**التدقيق القادم:** خلال 6 أشهر"""

        elif lang == 'ar':
            topic_ar = ''
            if audit_topic:
                topic_ar = f"""
⚠ قاعدة حاسمة ضد النتائج العامة:
لا تكرر نفس الملاحظات الخمس العامة لكل سياسة. الملاحظات التالية ممنوعة إلا إذا كانت مرتبطة مباشرة بـ "{audit_topic}":
- "عدم تفعيل MFA" (فقط لتدقيق التحكم بالوصول أو المصادقة)
- "سياسات كلمات المرور ضعيفة" (فقط لتدقيق المصادقة)
- "نقص في التوثيق" (عامة جداً — حدد أي توثيق مفقود)
- "تأخر برامج التوعية" (عامة جداً — حدد أي فجوة تدريبية)
- "عدم المراجعة الدورية" (عامة جداً — كل تدقيق يقول هذا)

جميع الملاحظات يجب أن تكون تشغيلية ومحددة لموضوع "{audit_topic}".
لا تكتب ملاحظات عامة لا علاقة لها بـ "{audit_topic}"."""
            
            prompt = f"""أنشئ تقرير تدقيق احترافي {"لموضوع " + audit_topic + " وفق " if audit_topic else "لـ"}إطار {framework} في مجال {domain}.
{"الملفات المرفقة: " + ", ".join(evidence_info) if evidence_info else ""}
{uploaded_content if uploaded_content else ""}
{"الوثيقة محل التدقيق: المحتوى أعلاه من الملف المرفق هو الوثيقة التي يتم تدقيقها. حلل محتواها مقابل متطلبات " + framework + " واكتب ملاحظات محددة عما هو موجود وما هو مفقود في الوثيقة." if uploaded_content else ""}

قواعد صارمة:
- لا تكتب أي تعليمات في المخرجات
- لا تستخدم أي نسب مئوية محددة (لا 75% ولا 80%)
- استخدم "يُحدد بعد التقييم" للقيم الحالية
- استخدم شرطة (-) للقوائم
- لا تذكر أي إطار أو معيار غير {framework}
- لا تكتب قسم "نطاق التدقيق" أو "منهجية التدقيق" - هذا التدقيق يقيّم الامتثال لضوابط {framework} مباشرة، انتقل فوراً إلى النتائج والملاحظات
{topic_ar}

# تقرير التدقيق - {framework}{" - " + audit_topic if audit_topic else ""}

## الملخص التنفيذي
**النتيجة العامة:** [ممتثل / ممتثل جزئياً / غير ممتثل] — مع ذكر عدد الملاحظات حسب الخطورة. لا تكتب "سيتم تحديد" — قدم تقييم امتثال فعلي {"لموضوع " + audit_topic + " وفق " if audit_topic else "لـ"}إطار {framework}.

## النتائج والملاحظات

### نتائج عالية الخطورة
| # | الملاحظة | الضابط المتأثر | التوصية | الحالة |
|---|----------|---------------|---------|--------|
| 1 | (ملاحظة عالية الخطورة {"متعلقة بـ" + audit_topic if audit_topic else ""}) | (رمز من {framework}) | (التوصية) | يحتاج تقييم |
| 2 | (ملاحظة عالية الخطورة {"متعلقة بـ" + audit_topic if audit_topic else ""}) | (رمز من {framework}) | (التوصية) | يحتاج تقييم |

### نتائج متوسطة الخطورة
| # | الملاحظة | الضابط المتأثر | التوصية | الحالة |
|---|----------|---------------|---------|--------|
| 1 | (ملاحظة متوسطة الخطورة {"متعلقة بـ" + audit_topic if audit_topic else ""}) | (رمز من {framework}) | (التوصية) | يحتاج تقييم |
| 2 | (ملاحظة متوسطة الخطورة {"متعلقة بـ" + audit_topic if audit_topic else ""}) | (رمز من {framework}) | (التوصية) | يحتاج تقييم |

### نتائج منخفضة الخطورة
| # | الملاحظة | الضابط المتأثر | التوصية | الحالة |
|---|----------|---------------|---------|--------|
| 1 | (ملاحظة منخفضة الخطورة {"متعلقة بـ" + audit_topic if audit_topic else ""}) | (رمز من {framework}) | (التوصية) | يحتاج تقييم |

## أدلة التنفيذ التفصيلية

يجب كتابة دليل تنفيذ منفصل لكل ملاحظة (عالية ومتوسطة ومنخفضة الخطورة). لا تترك أي ملاحظة بدون دليل تنفيذ.

### دليل تنفيذ الملاحظة عالية الخطورة رقم 1: (عنوان الملاحظة)
**الضابط المتأثر:** (رمز الضابط من {framework})
| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |
|---------|--------|-------|---------|----------|
| التخطيط | 1.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التخطيط | 1.2 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.2 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحقق | 3.1 | (خطوة محددة) | (الفريق) | (المخرج) |
**الأدلة المطلوبة:** ☐ (دليل 1) ☐ (دليل 2) ☐ (دليل 3)

### دليل تنفيذ الملاحظة عالية الخطورة رقم 2: (عنوان الملاحظة)
**الضابط المتأثر:** (رمز الضابط من {framework})
| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |
|---------|--------|-------|---------|----------|
| التخطيط | 1.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التخطيط | 1.2 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.2 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحقق | 3.1 | (خطوة محددة) | (الفريق) | (المخرج) |
**الأدلة المطلوبة:** ☐ (دليل 1) ☐ (دليل 2) ☐ (دليل 3)

### دليل تنفيذ الملاحظة متوسطة الخطورة رقم 1: (عنوان الملاحظة)
**الضابط المتأثر:** (رمز الضابط من {framework})
| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |
|---------|--------|-------|---------|----------|
| التخطيط | 1.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحقق | 3.1 | (خطوة محددة) | (الفريق) | (المخرج) |
**الأدلة المطلوبة:** ☐ (دليل 1) ☐ (دليل 2) ☐ (دليل 3)

### دليل تنفيذ الملاحظة متوسطة الخطورة رقم 2: (عنوان الملاحظة)
**الضابط المتأثر:** (رمز الضابط من {framework})
| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |
|---------|--------|-------|---------|----------|
| التخطيط | 1.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحقق | 3.1 | (خطوة محددة) | (الفريق) | (المخرج) |
**الأدلة المطلوبة:** ☐ (دليل 1) ☐ (دليل 2) ☐ (دليل 3)

### دليل تنفيذ الملاحظة منخفضة الخطورة رقم 1: (عنوان الملاحظة)
**الضابط المتأثر:** (رمز الضابط من {framework})
| المرحلة | الخطوة | الوصف | المسؤول | المخرجات |
|---------|--------|-------|---------|----------|
| التخطيط | 1.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التنفيذ | 2.1 | (خطوة محددة) | (الفريق) | (المخرج) |
| التحقق | 3.1 | (خطوة محددة) | (الفريق) | (المخرج) |
**الأدلة المطلوبة:** ☐ (دليل 1) ☐ (دليل 2) ☐ (دليل 3)

## خطة العمل
| # | الإجراء | المسؤول | الموعد النهائي | الأولوية |
|---|--------|---------|---------------|----------|
| 1 | معالجة الملاحظات عالية الخطورة | فريق الأمن | خلال 30 يوم | عالية |
| 2 | معالجة الملاحظات متوسطة الخطورة | فريق الامتثال | خلال 60 يوم | متوسطة |
| 3 | معالجة الملاحظات منخفضة الخطورة | فريق التوثيق | خلال 90 يوم | منخفضة |

---
**تاريخ التقرير:** يُحدد لاحقاً
**التدقيق القادم:** خلال 6 أشهر"""

        content = generate_ai_content(prompt, lang, content_type='audit')
        
        # Save to database
        try:
            conn = get_db()
            log_action(session['user_id'], 'generate_audit', {'domain': domain, 'framework': framework, 'topic': audit_topic, 'language': lang})
            conn.execute('''INSERT INTO audits (user_id, domain, framework, scope, content, language)
                            VALUES (?, ?, ?, ?, ?, ?)''',
                        (session['user_id'], domain, framework, audit_scope, content, lang))
            conn.commit()
        except Exception as db_error:
            print(f"Database error: {db_error}")
        
        return jsonify({'success': True, 'content': content})
        
    except Exception as e:
        print(f"Audit generation error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test-docx', methods=['GET'])
@login_required
def api_test_docx():
    """Test if DOCX generation works at all on this server."""
    try:
        from io import BytesIO
        from docx import Document
        from docx.shared import Inches
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = 'Test'
        table.rows[0].cells[1].text = 'Table'
        table.rows[1].cells[0].text = '1'
        table.rows[1].cells[1].text = '2'
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        from flask import send_file
        return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True, download_name='test.docx')
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/api/generate-docx', methods=['POST'])
@login_required
def api_generate_docx():
    """Generate Word document from content."""
    from io import BytesIO
    import re
    
    try:
        data = request.json
    except Exception as je:
        print(f"DOCX ERROR: Failed to parse JSON: {je}", flush=True)
        return jsonify({'error': f'Invalid JSON: {str(je)}'}), 400
    
    if not data:
        print("DOCX ERROR: No JSON data received!", flush=True)
        return jsonify({'error': 'No data received'}), 400
    
    content = data.get('content', '')
    filename = data.get('filename', 'document')
    lang = data.get('language', 'en')
    
    try:
        log_action(session.get('user_id'), 'export_docx', {'filename': filename})
    except Exception:
        pass
    
    print("=" * 60, flush=True)
    print(f"DOCX GENERATION - lang={lang}, filename={filename}", flush=True)
    print(f"Content length: {len(content)}", flush=True)
    print(f"Content preview: {content[:200]}", flush=True)
    print("=" * 60, flush=True)
    
    if not content or not content.strip():
        print("DOCX ERROR: Empty content!", flush=True)
        return jsonify({'error': 'No content to generate document'}), 400
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls, qn
        from docx.oxml import parse_xml
        
        is_arabic = lang == 'ar'
        doc = Document()
        
        # Set page size and document-level RTL
        for section in doc.sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        
        # Set document-level RTL for Arabic
        if is_arabic:
            try:
                body = doc.element.body
                sectPr = body.find(qn('w:sectPr'))
                if sectPr is not None:
                    bidi = parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>')
                    sectPr.append(bidi)
            except Exception:
                pass  # Non-critical
        
        def set_rtl_paragraph(paragraph):
            """Set full RTL properties on a paragraph."""
            if not is_arabic:
                return
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Set paragraph-level RTL
            pPr = paragraph._p.get_or_add_pPr()
            bidi = parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>')
            pPr.append(bidi)
            # Set RTL on each run
            for run in paragraph.runs:
                rPr = run._r.get_or_add_rPr()
                rtl_elem = parse_xml(f'<w:rtl {nsdecls("w")} w:val="1"/>')
                rPr.append(rtl_elem)
        
        def set_rtl_run(run):
            """Set RTL on a single run."""
            if not is_arabic:
                return
            rPr = run._r.get_or_add_rPr()
            rtl_elem = parse_xml(f'<w:rtl {nsdecls("w")} w:val="1"/>')
            rPr.append(rtl_elem)
        
        def add_rtl_heading(text, level):
            """Add a heading with proper RTL and inline bold support."""
            # Strip ** markers from heading text but keep the text
            clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            h = doc.add_heading('', level=level)
            
            # Process inline bold within heading
            if '**' in text:
                parts = re.split(r'\*\*(.+?)\*\*', text)
                for idx_part, part in enumerate(parts):
                    if part:
                        run = h.add_run(part)
                        if idx_part % 2 == 1:
                            run.bold = True
                        if is_arabic:
                            set_rtl_run(run)
            else:
                run = h.add_run(clean_text)
                if is_arabic:
                    set_rtl_run(run)
            
            if is_arabic:
                set_rtl_paragraph(h)
            return h
        
        def add_formatted_paragraph(text, base_bold=False):
            """Add a paragraph with inline bold/formatting support."""
            p = doc.add_paragraph()
            if '**' in text:
                parts = re.split(r'\*\*(.+?)\*\*', text)
                for idx_part, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.bold = (idx_part % 2 == 1) or base_bold
                        if is_arabic:
                            set_rtl_run(run)
            else:
                run = p.add_run(text)
                if base_bold:
                    run.bold = True
                if is_arabic:
                    set_rtl_run(run)
            if is_arabic:
                set_rtl_paragraph(p)
            return p
        
        def add_rtl_paragraph(text, bold=False):
            """Add a normal paragraph with proper RTL."""
            if '**' in text:
                return add_formatted_paragraph(text, base_bold=bold)
            p = doc.add_paragraph()
            run = p.add_run(text)
            if bold:
                run.bold = True
            if is_arabic:
                set_rtl_paragraph(p)
            return p
        
        # ---- Table helpers ----
        
        def parse_markdown_table(lines, start_idx):
            """Parse markdown table starting at start_idx."""
            table_rows = []
            idx = start_idx
            while idx < len(lines):
                ln = lines[idx].strip()
                if ln.startswith('|') and ln.endswith('|'):
                    if '---' in ln or ':-' in ln or '-:' in ln:
                        idx += 1
                        continue
                    cells = [cell.strip() for cell in ln.split('|')[1:-1]]
                    if cells:
                        table_rows.append(cells)
                    idx += 1
                else:
                    break
            return table_rows, idx
        
        def add_table_to_doc(doc, table_data, lang):
            """Add a formatted table with RTL support and professional styling."""
            if not table_data or len(table_data) < 1:
                return
            
            # Normalize all rows to same column count (prevent crash)
            num_cols = max(len(row) for row in table_data)
            for row in table_data:
                while len(row) < num_cols:
                    row.append('')
            
            table = doc.add_table(rows=len(table_data), cols=num_cols)
            try:
                table.style = 'Table Grid'
            except Exception:
                pass
            
            # Auto-fit table to window width
            try:
                tbl = table._tbl
                tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
                # Set table width to 100%
                tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
                tblPr.append(tblW)
                
                # Set table RTL direction for Arabic
                if is_arabic:
                    bidi_visual = parse_xml(f'<w:bidiVisual {nsdecls("w")} w:val="1"/>')
                    tblPr.append(bidi_visual)
            except Exception:
                pass
            
            for i, row_data in enumerate(table_data):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    if j < len(row.cells):
                        cell = row.cells[j]
                        # Clear default paragraph and add formatted text
                        cell.text = ''
                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        
                        # Process inline bold within cells
                        if '**' in cell_text:
                            parts = re.split(r'\*\*(.+?)\*\*', cell_text)
                            for idx_part, part in enumerate(parts):
                                if part:
                                    run = p.add_run(part)
                                    run.bold = (idx_part % 2 == 1) or (i == 0)
                                    run.font.size = Pt(9) if i > 0 else Pt(9.5)
                                    if is_arabic:
                                        set_rtl_run(run)
                        else:
                            run = p.add_run(cell_text)
                            run.font.size = Pt(9) if i > 0 else Pt(9.5)
                            if i == 0:
                                run.bold = True
                            if is_arabic:
                                set_rtl_run(run)
                        
                        # Style header row with blue background
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for r in paragraph.runs:
                                    r.bold = True
                                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4"/>')
                            cell._tc.get_or_add_tcPr().append(shading)
                        else:
                            # Alternate row shading
                            if i % 2 == 0:
                                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F7FB"/>')
                                cell._tc.get_or_add_tcPr().append(shading)
                        
                        # Set cell padding
                        try:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcMar = parse_xml(
                                f'<w:tcMar {nsdecls("w")}>'
                                f'<w:top w:w="40" w:type="dxa"/>'
                                f'<w:start w:w="80" w:type="dxa"/>'
                                f'<w:bottom w:w="40" w:type="dxa"/>'
                                f'<w:end w:w="80" w:type="dxa"/>'
                                f'</w:tcMar>'
                            )
                            tcPr.append(tcMar)
                        except Exception:
                            pass
                        
                        # Set RTL/alignment on each cell paragraph
                        for paragraph in cell.paragraphs:
                            if is_arabic:
                                set_rtl_paragraph(paragraph)
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add spacing after table
            doc.add_paragraph()
        
        # ---- Branded logo header (Pillar 4) ----
        try:
            user_logo = None
            if session.get('user_id'):
                user_row = get_db().execute('SELECT logo_path FROM users WHERE id = ?', (session['user_id'],)).fetchone()
                if user_row and user_row['logo_path']:
                    logo_file = user_row['logo_path']
                    if os.path.isfile(logo_file):
                        user_logo = logo_file
            if user_logo:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_para.add_run().add_picture(user_logo, width=Inches(1.8))
                doc.add_paragraph()  # spacer
        except Exception as logo_err:
            print(f"Logo inject (non-fatal): {logo_err}", flush=True)
        
        # ---- Add document title ----
        if is_arabic:
            title_text = filename.replace('_', ' ')
        else:
            title_text = filename.replace('_', ' ').title()
        
        try:
            add_rtl_heading(title_text, 0)
        except Exception:
            # Fallback if Title style fails
            add_rtl_heading(title_text, 1)
        
        # ---- Process content line by line ----
        # First, ensure proper markdown formatting for reliable parsing
        try:
            content = ensure_markdown_formatting(content)
        except Exception as fmt_err:
            print(f"DOCX WARNING: ensure_markdown_formatting failed: {fmt_err}", flush=True)
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            try:
                line = lines[i].strip()
                
                # Skip empty lines and separators
                if not line:
                    i += 1
                    continue
                if line == '---' or line == '[SECTION]':
                    doc.add_paragraph('')
                    i += 1
                    continue
                
                # Table
                if line.startswith('|') and '|' in line[1:]:
                    table_data, new_idx = parse_markdown_table(lines, i)
                    if table_data:
                        add_table_to_doc(doc, table_data, lang)
                    i = new_idx
                    continue
                
                # Headings — check most specific first
                if line.startswith('#### '):
                    add_rtl_heading(line[5:], 3)
                elif line.startswith('### '):
                    add_rtl_heading(line[4:], 2)
                elif line.startswith('## '):
                    add_rtl_heading(line[3:], 1)
                elif line.startswith('# '):
                    try:
                        add_rtl_heading(line[2:], 0)
                    except Exception:
                        add_rtl_heading(line[2:], 1)
                
                # Bullet points
                elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
                    raw_text = line[2:]
                    if is_arabic:
                        p = doc.add_paragraph()
                        # Process inline bold within bullet
                        if '**' in raw_text:
                            p.add_run('• ')
                            parts = re.split(r'\*\*(.+?)\*\*', raw_text)
                            for idx_b, part in enumerate(parts):
                                if part:
                                    run = p.add_run(part)
                                    if idx_b % 2 == 1:
                                        run.bold = True
                                    set_rtl_run(run)
                        else:
                            run = p.add_run('• ' + raw_text)
                        set_rtl_paragraph(p)
                        p.paragraph_format.left_indent = Cm(1)
                    else:
                        p = doc.add_paragraph()
                        if '**' in raw_text:
                            parts = re.split(r'\*\*(.+?)\*\*', raw_text)
                            for idx_b, part in enumerate(parts):
                                if part:
                                    run = p.add_run(part)
                                    if idx_b % 2 == 1:
                                        run.bold = True
                        else:
                            p.add_run(raw_text)
                        try:
                            p.style = doc.styles['List Bullet']
                        except Exception:
                            # Fallback: manual bullet
                            first_run = p.runs[0] if p.runs else p.add_run('')
                            first_run.text = '• ' + first_run.text
                        p.paragraph_format.left_indent = Cm(1)
                
                # Numbered lists
                elif re.match(r'^(\d+)\.\s+(.+)', line):
                    m = re.match(r'^(\d+)\.\s+(.+)', line)
                    num = m.group(1)
                    raw_text = m.group(2)
                    if is_arabic:
                        p = doc.add_paragraph()
                        # Keep natural order - RTL will handle display
                        run = p.add_run(num + '. ' + raw_text)
                        set_rtl_paragraph(p)
                        p.paragraph_format.left_indent = Cm(1)
                    else:
                        try:
                            p = doc.add_paragraph(line, style='List Number')
                        except Exception:
                            p = doc.add_paragraph(line)
                
                # Full bold line
                elif line.startswith('**') and line.endswith('**'):
                    add_rtl_paragraph(line[2:-2], bold=True)
                
                # Evidence/checkbox lines (☐ markers)
                elif '☐' in line or '☑' in line:
                    p = doc.add_paragraph()
                    if '**' in line:
                        parts = re.split(r'\*\*(.+?)\*\*', line)
                        for idx_p, part in enumerate(parts):
                            if part:
                                run = p.add_run(part)
                                if idx_p % 2 == 1:
                                    run.bold = True
                                if is_arabic:
                                    set_rtl_run(run)
                    else:
                        run = p.add_run(line)
                        if is_arabic:
                            set_rtl_run(run)
                    if is_arabic:
                        set_rtl_paragraph(p)
                    p.paragraph_format.left_indent = Cm(1)
                
                # Inline bold
                elif '**' in line:
                    p = doc.add_paragraph()
                    parts = re.split(r'\*\*(.+?)\*\*', line)
                    for idx, part in enumerate(parts):
                        if part:
                            run = p.add_run(part)
                            if idx % 2 == 1:
                                run.bold = True
                            if is_arabic:
                                set_rtl_run(run)
                    if is_arabic:
                        set_rtl_paragraph(p)
                
                # Normal text
                else:
                    add_rtl_paragraph(line)
                
                i += 1
            
            except Exception as line_err:
                print(f"DOCX WARNING: Error on line {i}: '{lines[i][:80] if i < len(lines) else '?'}' - {line_err}", flush=True)
                try:
                    add_rtl_paragraph(lines[i].strip() if i < len(lines) else '')
                except Exception:
                    pass
                i += 1
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        from flask import send_file
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{filename}.docx'
        )
    except ImportError as ie:
        print(f"DOCX IMPORT ERROR: {ie}", flush=True)
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Word generation not available: {str(ie)}'}), 500
    except Exception as e:
        print(f"DOCX GENERATION ERROR: {str(e)}", flush=True)
        import traceback
        tb = traceback.format_exc()
        print(tb, flush=True)
        return jsonify({'error': f'Document generation failed: {str(e)}'}), 500

@app.route('/api/generate-pdf', methods=['POST'])
@login_required
def api_generate_pdf():
    """Generate PDF document from content with Arabic support."""
    from io import BytesIO
    import os
    import glob
    
    data = request.json
    content = data.get('content', '')
    filename = data.get('filename', 'document')
    lang = data.get('language', 'en')
    
    try:
        log_action(session.get('user_id'), 'export_pdf', {'filename': filename})
    except Exception:
        pass
    
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_RIGHT, TA_LEFT, TA_CENTER
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.units import inch, cm
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # For Arabic text
        is_arabic = lang == 'ar'
        arabic_font_name = 'Helvetica'  # Default fallback
        arabic_font_bold = 'Helvetica-Bold'
        
        if is_arabic:
            try:
                import arabic_reshaper
                from bidi.algorithm import get_display
                
                # Extended font search paths for Arabic support
                font_paths = [
                    # Noto fonts (installed via render.yaml)
                    '/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf',
                    '/usr/share/fonts/opentype/noto/NotoSansArabic-Regular.ttf',
                    '/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf',
                    # FreeFonts
                    '/usr/share/fonts/truetype/freefont/FreeSans.ttf',
                    '/usr/share/fonts/truetype/freefont/FreeSerif.ttf',
                    # DejaVu (has Arabic support)
                    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                    # Liberation fonts
                    '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                    # Local fonts
                    'static/fonts/Amiri-Regular.ttf',
                    'static/fonts/NotoSansArabic-Regular.ttf',
                    # Additional system paths
                    '/usr/share/fonts/TTF/DejaVuSans.ttf',
                    '/usr/local/share/fonts/Amiri-Regular.ttf',
                ]
                
                # Also search for any Arabic font using glob
                additional_paths = glob.glob('/usr/share/fonts/**/[Aa]rabic*.ttf', recursive=True)
                additional_paths += glob.glob('/usr/share/fonts/**/[Nn]oto*[Aa]rabic*.ttf', recursive=True)
                additional_paths += glob.glob('/usr/share/fonts/**/[Aa]miri*.ttf', recursive=True)
                font_paths = additional_paths + font_paths
                
                font_registered = False
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont('ArabicFont', font_path))
                            arabic_font_name = 'ArabicFont'
                            arabic_font_bold = 'ArabicFont'
                            font_registered = True
                            print(f"✅ Registered Arabic font: {font_path}", flush=True)
                            break
                        except Exception as font_error:
                            print(f"⚠️ Failed to register font {font_path}: {font_error}", flush=True)
                            continue
                
                if not font_registered:
                    print("⚠️ No Arabic font found, PDF will show boxes for Arabic text", flush=True)
                    # List available fonts for debugging
                    for search_path in ['/usr/share/fonts/truetype/', '/usr/share/fonts/']:
                        if os.path.exists(search_path):
                            print(f"Available fonts in {search_path}:", flush=True)
                            for item in os.listdir(search_path)[:10]:
                                print(f"  - {item}", flush=True)
                
            except ImportError as ie:
                print(f"⚠️ Arabic libraries not available: {ie}", flush=True)
                # Fallback if arabic libraries not available
                def get_display(text):
                    return text
                class _FallbackReshaper:
                    def reshape(self, text):
                        return text
                arabic_reshaper = _FallbackReshaper()
        
        # Calculate available text width for manual line wrapping
        page_width = A4[0]
        text_width = page_width - 3*cm  # 1.5cm margin each side
        # CRITICAL: Use a tighter wrap width for Arabic pre-wrapping.
        # We must wrap NARROWER than ReportLab's Paragraph frame width,
        # so ReportLab never re-wraps our visually-ordered lines.
        # If ReportLab re-wraps, it splits reversed text at wrong points
        # (e.g. "تهدف" appearing alone on a line).
        arabic_wrap_width = text_width - 12  # 12pt safety margin
        
        def process_arabic(text, font_name_for_wrap=None, font_size_for_wrap=11, extra_indent=0):
            """Process Arabic text for correct display in PDF.
            
            Strategy: reshape → manual word-wrap → get_display per line → join with <br/>
            This prevents the jumbled word order that happens when get_display is applied
            to an entire paragraph and then ReportLab re-wraps it.
            
            CRITICAL: We wrap to arabic_wrap_width (narrower than frame) so ReportLab
            never needs to re-wrap our pre-wrapped lines.
            """
            if is_arabic and text:
                try:
                    text = str(text).strip()
                    if not text:
                        return text
                    
                    from reportlab.pdfbase.pdfmetrics import stringWidth
                    wrap_font = font_name_for_wrap or arabic_font_name
                    # Use tighter width to prevent ReportLab re-wrapping
                    wrap_width = arabic_wrap_width - extra_indent
                    
                    # Handle mixed content with HTML bold tags
                    if '<b>' in text or '</b>' in text:
                        import re as re_clean
                        parts = re_clean.split(r'(<b>|</b>)', text)
                        processed_parts = []
                        for part in parts:
                            if part in ('<b>', '</b>'):
                                processed_parts.append(part)
                            elif part.strip():
                                reshaped = arabic_reshaper.reshape(part)
                                processed_parts.append(get_display(reshaped))
                            else:
                                processed_parts.append(part)
                        return ''.join(processed_parts)
                    
                    # Step 1: Reshape for letter joining
                    reshaped = arabic_reshaper.reshape(text)
                    
                    # Step 2: Check if text fits in one line (most headings, short text)
                    total_width = stringWidth(reshaped, wrap_font, font_size_for_wrap)
                    if total_width <= wrap_width:
                        # Single line — just get_display the whole thing
                        return get_display(reshaped)
                    
                    # Step 3: Manual word-wrap for multi-line text
                    words = reshaped.split(' ')
                    lines = []
                    current_line_words = []
                    current_width = 0
                    space_width = stringWidth(' ', wrap_font, font_size_for_wrap)
                    
                    for word in words:
                        word_width = stringWidth(word, wrap_font, font_size_for_wrap)
                        test_width = current_width + word_width + (space_width if current_line_words else 0)
                        
                        if test_width > wrap_width and current_line_words:
                            # Line is full — wrap to next line
                            line_text = ' '.join(current_line_words)
                            lines.append(get_display(line_text))
                            current_line_words = [word]
                            current_width = word_width
                        else:
                            current_line_words.append(word)
                            current_width = test_width
                    
                    # Last line
                    if current_line_words:
                        line_text = ' '.join(current_line_words)
                        lines.append(get_display(line_text))
                    
                    # Step 4: Join with <br/> for ReportLab Paragraph
                    if len(lines) > 1:
                        return '<br/>'.join(lines)
                    elif lines:
                        return lines[0]
                    else:
                        return get_display(reshaped)
                    
                except Exception as e:
                    print(f"Arabic processing warning: {e}", flush=True)
                    return text
            return text
        
        def process_arabic_table(text, col_width, font_size=9):
            """Process Arabic text for table cells with specific column width."""
            if is_arabic and text:
                try:
                    text = str(text).strip()
                    if not text:
                        return text
                    
                    from reportlab.pdfbase.pdfmetrics import stringWidth
                    reshaped = arabic_reshaper.reshape(text)
                    
                    # Check if fits in one line
                    total_w = stringWidth(reshaped, arabic_font_name, font_size)
                    cell_width = col_width - 16  # subtract padding + safety margin
                    if total_w <= cell_width:
                        return get_display(reshaped)
                    
                    # Wrap strictly within cell width - no overflow tolerance
                    words = reshaped.split(' ')
                    lines = []
                    current_line_words = []
                    current_width = 0
                    space_width = stringWidth(' ', arabic_font_name, font_size)
                    
                    for word in words:
                        word_width = stringWidth(word, arabic_font_name, font_size)
                        test_width = current_width + word_width + (space_width if current_line_words else 0)
                        
                        if test_width > cell_width and current_line_words:
                            line_text = ' '.join(current_line_words)
                            lines.append(get_display(line_text))
                            current_line_words = [word]
                            current_width = word_width
                        else:
                            current_line_words.append(word)
                            current_width = test_width
                    
                    if current_line_words:
                        line_text = ' '.join(current_line_words)
                        lines.append(get_display(line_text))
                    
                    return '<br/>'.join(lines) if len(lines) > 1 else (lines[0] if lines else get_display(reshaped))
                    
                except Exception as e:
                    return text
            return text
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=1.5*cm,
            leftMargin=1.5*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        if is_arabic:
            title_style = ParagraphStyle(
                'ArabicTitle',
                parent=styles['Title'],
                alignment=TA_RIGHT,
                fontSize=24,
                spaceAfter=30,
                fontName=arabic_font_bold,
            )
            heading1_style = ParagraphStyle(
                'ArabicH1',
                parent=styles['Heading1'],
                alignment=TA_RIGHT,
                fontSize=18,
                spaceAfter=12,
                spaceBefore=20,
                textColor=colors.HexColor('#1a365d'),
                fontName=arabic_font_bold,
            )
            heading2_style = ParagraphStyle(
                'ArabicH2',
                parent=styles['Heading2'],
                alignment=TA_RIGHT,
                fontSize=14,
                spaceAfter=10,
                spaceBefore=15,
                textColor=colors.HexColor('#2d3748'),
                fontName=arabic_font_bold,
            )
            normal_style = ParagraphStyle(
                'ArabicNormal',
                parent=styles['Normal'],
                alignment=TA_RIGHT,
                fontSize=11,
                spaceAfter=8,
                leading=18,
                fontName=arabic_font_name,
            )
            bullet_style = ParagraphStyle(
                'ArabicBullet',
                parent=styles['Normal'],
                alignment=TA_RIGHT,
                fontSize=11,
                spaceAfter=6,
                rightIndent=20,
                fontName=arabic_font_name,
            )
            numbered_style = ParagraphStyle(
                'ArabicNumbered',
                parent=styles['Normal'],
                alignment=TA_RIGHT,
                fontSize=11,
                spaceAfter=6,
                rightIndent=20,
                fontName=arabic_font_name,
            )
        else:
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1a365d')
            )
            heading1_style = ParagraphStyle(
                'CustomH1',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=12,
                spaceBefore=20,
                textColor=colors.HexColor('#1a365d')
            )
            heading2_style = ParagraphStyle(
                'CustomH2',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=10,
                spaceBefore=15,
                textColor=colors.HexColor('#2d3748')
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=8,
                leading=16
            )
            bullet_style = ParagraphStyle(
                'CustomBullet',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                leftIndent=20,
                bulletIndent=10
            )
            numbered_style = ParagraphStyle(
                'CustomNumbered',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                leftIndent=20
            )
        
        # Build content
        story = []
        
        # ---- Branded logo header (Pillar 4) ----
        try:
            from reportlab.platypus import Image as RLImage
            user_logo = None
            if session.get('user_id'):
                user_row = get_db().execute('SELECT logo_path FROM users WHERE id = ?', (session['user_id'],)).fetchone()
                if user_row and user_row['logo_path']:
                    logo_file = user_row['logo_path']
                    if os.path.isfile(logo_file):
                        user_logo = logo_file
            if user_logo:
                logo_img = RLImage(user_logo, width=1.5*inch, height=1.5*inch)
                logo_img.hAlign = 'CENTER'
                story.append(logo_img)
                story.append(Spacer(1, 0.2*inch))
        except Exception as logo_err:
            print(f"PDF logo inject (non-fatal): {logo_err}", flush=True)
        
        # Add title
        if is_arabic:
            title_text = process_arabic(filename.replace('_', ' '), arabic_font_bold, 24)
        else:
            title_text = filename.replace('_', ' ').title()
        story.append(Paragraph(title_text, title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Parse markdown content
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Skip separator lines
            if line == '---':
                story.append(Spacer(1, 0.2*inch))
                i += 1
                continue
            
            # Skip [SECTION] markers
            if line.strip() == '[SECTION]':
                story.append(Spacer(1, 0.15*inch))
                i += 1
                continue
            
            # Check if this is a table
            if line.startswith('|') and '|' in line[1:]:
                table_data = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_line = lines[i].strip()
                    if '---' not in row_line:  # Skip separator
                        cells = [c.strip() for c in row_line.split('|')[1:-1]]
                        if cells:
                            # Reverse column order for Arabic RTL
                            if is_arabic:
                                cells = cells[::-1]
                            table_data.append(cells)
                    i += 1
                
                if table_data:
                    # Create table with smart column widths
                    col_count = len(table_data[0]) if table_data else 1
                    
                    # Check if this is an implementation guidelines table (has Phase/Step columns)
                    header_row = [str(c).lower() for c in table_data[0]] if table_data else []
                    is_implementation_table = any(h in ['phase', 'step', 'المرحلة', 'الخطوة'] for h in header_row)
                    
                    # Calculate column widths - give more space for Arabic text
                    available_width = doc.width
                    
                    if is_implementation_table and col_count == 5:
                        # Implementation guidelines table: Phase, Step, Description, Owner, Deliverable
                        # Give more space to Phase and Description columns
                        if is_arabic:
                            # Arabic: Deliverable, Owner, Description, Step, Phase (reversed)
                            col_widths = [
                                available_width * 0.16,  # المخرجات (Deliverable)
                                available_width * 0.13,  # المسؤول (Owner)
                                available_width * 0.36,  # الوصف (Description)
                                available_width * 0.10,  # الخطوة (Step)
                                available_width * 0.25,  # المرحلة (Phase) - wider for Arabic words
                            ]
                        else:
                            col_widths = [
                                available_width * 0.22,  # Phase - wider to avoid wrapping "Planning", "Verification", etc.
                                available_width * 0.08,  # Step - just numbers like "1.1"
                                available_width * 0.38,  # Description - main content
                                available_width * 0.14,  # Owner
                                available_width * 0.18,  # Deliverable
                            ]
                    elif col_count >= 7:
                        narrow_col = 0.4 * inch
                        remaining = available_width - narrow_col
                        other_cols = remaining / (col_count - 1)
                        if is_arabic:
                            col_widths = [other_cols] * (col_count - 1) + [narrow_col]
                        else:
                            col_widths = [narrow_col] + [other_cols] * (col_count - 1)
                    elif col_count >= 4:
                        # For 4-6 column tables, distribute more evenly
                        col_widths = [available_width / col_count] * col_count
                    elif col_count == 3:
                        equal = available_width / 3
                        col_widths = [equal] * 3
                    elif col_count == 2:
                        col_widths = [available_width * 0.5, available_width * 0.5]
                    else:
                        col_widths = [available_width]
                    
                    # Wrap cell content in Paragraphs
                    cell_style = ParagraphStyle(
                        'CellStyle', 
                        fontSize=9, 
                        leading=14,
                        fontName=arabic_font_name if is_arabic else 'Helvetica',
                        alignment=TA_RIGHT if is_arabic else TA_LEFT,
                    )
                    header_cell_style = ParagraphStyle(
                        'HeaderCellStyle', 
                        fontSize=9, 
                        leading=14, 
                        textColor=colors.white,
                        fontName=arabic_font_bold if is_arabic else 'Helvetica-Bold',
                        alignment=TA_RIGHT if is_arabic else TA_LEFT,
                    )
                    
                    wrapped_data = []
                    for row_idx, row in enumerate(table_data):
                        wrapped_row = []
                        for col_idx, cell in enumerate(row):
                            cell_text = str(cell)
                            # Process Arabic: reshape and per-cell-width bidi
                            if is_arabic and cell_text:
                                cw = col_widths[col_idx] if col_idx < len(col_widths) else text_width / len(row)
                                cell_text = process_arabic_table(cell_text, cw, font_size=9)
                            if row_idx == 0:  # Header row
                                wrapped_row.append(Paragraph(cell_text, header_cell_style))
                            else:
                                wrapped_row.append(Paragraph(cell_text, cell_style))
                        wrapped_data.append(wrapped_row)
                    
                    t = Table(wrapped_data, colWidths=col_widths)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'RIGHT' if is_arabic else 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('FONTNAME', (0, 0), (-1, -1), arabic_font_name if is_arabic else 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                        ('TOPPADDING', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.15*inch))
                continue
            
            # Handle markdown headings
            # IMPORTANT: Strip markdown prefix FIRST, then process Arabic
            # Check more specific headings first (#### before ### before ##)
            if line.startswith('#### '):
                raw_text = line[5:]
                text = process_arabic(raw_text, arabic_font_bold, 14) if is_arabic else raw_text
                story.append(Paragraph(text, heading2_style))
            elif line.startswith('### '):
                raw_text = line[4:]
                text = process_arabic(raw_text, arabic_font_bold, 14) if is_arabic else raw_text
                story.append(Paragraph(text, heading2_style))
            elif line.startswith('## '):
                raw_text = line[3:]
                text = process_arabic(raw_text, arabic_font_bold, 18) if is_arabic else raw_text
                story.append(Paragraph(text, heading1_style))
            elif line.startswith('# '):
                raw_text = line[2:]
                text = process_arabic(raw_text, arabic_font_bold, 24) if is_arabic else raw_text
                story.append(Paragraph(text, title_style))
            elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
                raw_text = line[2:]
                text = process_arabic(raw_text, arabic_font_name, 11, extra_indent=20) if is_arabic else raw_text
                if is_arabic:
                    # Arabic: bullet on right, get_display already reversed the text
                    bullet_text = text + ' •'
                else:
                    bullet_text = '• ' + text
                story.append(Paragraph(bullet_text, bullet_style))
            elif len(line) > 2 and line[0].isdigit() and (line[1] == '.' or (len(line) > 2 and line[1].isdigit() and line[2] == '.')):
                # Numbered list
                import re as re_mod
                num_match = re_mod.match(r'^(\d+)\.\s*(.+)', line)
                if num_match:
                    num = num_match.group(1)
                    raw_text = num_match.group(2)
                    text = process_arabic(raw_text, arabic_font_name, 11, extra_indent=20) if is_arabic else raw_text
                    if is_arabic:
                        numbered_text = text + ' .' + num
                    else:
                        numbered_text = num + '. ' + text
                    story.append(Paragraph(numbered_text, numbered_style))
                else:
                    text = process_arabic(line, arabic_font_name, 11) if is_arabic else line
                    story.append(Paragraph(text, normal_style))
            elif line.startswith('**') and line.endswith('**'):
                raw_text = line[2:-2]
                text = process_arabic(raw_text, arabic_font_bold, 11) if is_arabic else raw_text
                story.append(Paragraph(f'<b>{text}</b>', normal_style))
            elif '**' in line:
                # Handle inline bold - strip ** markers, then process Arabic
                import re
                parts = re.split(r'\*\*(.+?)\*\*', line)
                formatted_parts = []
                for idx, part in enumerate(parts):
                    if idx % 2 == 1:  # Bold text
                        p = process_arabic(part, arabic_font_bold, 11) if is_arabic else part
                        formatted_parts.append(f'<b>{p}</b>')
                    else:
                        p = process_arabic(part, arabic_font_name, 11) if is_arabic else part
                        formatted_parts.append(p)
                formatted = ''.join(formatted_parts)
                story.append(Paragraph(formatted, normal_style))
            else:
                text = process_arabic(line, arabic_font_name, 11) if is_arabic else line
                story.append(Paragraph(text, normal_style))
            
            i += 1
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        from flask import send_file
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{filename}.pdf'
        )
        
    except ImportError as e:
        return jsonify({'error': f'PDF generation not available: {str(e)}'}), 500
    except Exception as e:
        print(f"PDF generation error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'PDF generation failed: {str(e)}'}), 500

@app.route('/api/set-language/<lang>')
def set_language(lang):
    """Set language preference."""
    session['lang'] = lang if lang in ['en', 'ar'] else 'en'
    return jsonify({'success': True, 'lang': session['lang']})

# ============================================================================
# EXCEL EXPORT
# ============================================================================

@app.route('/api/generate-excel', methods=['POST'])
@login_required
def api_generate_excel():
    """Generate Excel file from data."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io
    
    try:
    
        data = request.json
        export_type = data.get('type', 'analytics')  # analytics, risks, documents
    
        wb = Workbook()
        ws = wb.active
    
        # Styles
        header_font = Font(bold=True, color='FFFFFF', size=11)
        header_fill = PatternFill(start_color='667eea', end_color='667eea', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell_alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
    
        if export_type == 'analytics':
            ws.title = 'Analytics Summary'
        
            # Get user data
            user_id = session.get('user_id')
            if user_id:
                compliance = calculate_compliance_score(user_id)
                maturity = calculate_maturity_levels(user_id)
            
                # Compliance Score Section
                ws['A1'] = 'COMPLIANCE SCORE SUMMARY'
                ws['A1'].font = Font(bold=True, size=14, color='667eea')
                ws.merge_cells('A1:D1')
            
                headers = ['Metric', 'Value', 'Target', 'Status']
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=3, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                    cell.border = thin_border
            
                metrics = [
                    ('Overall Score', f"{compliance['score']}%", '80%', 'Good' if compliance['score'] >= 60 else 'Needs Improvement'),
                    ('Strategies', compliance['strategies'], '5', '✓' if compliance['strategies'] >= 3 else '✗'),
                    ('Policies', compliance['policies'], '10', '✓' if compliance['policies'] >= 5 else '✗'),
                    ('Audits', compliance['audits'], '5', '✓' if compliance['audits'] >= 3 else '✗'),
                    ('Risk Assessments', compliance['risks'], '10', '✓' if compliance['risks'] >= 5 else '✗'),
                    ('Domains Covered', compliance['domains_covered'], '5', '✓' if compliance['domains_covered'] >= 3 else '✗'),
                ]
            
                for row, (metric, value, target, status) in enumerate(metrics, 4):
                    ws.cell(row=row, column=1, value=metric).border = thin_border
                    ws.cell(row=row, column=2, value=value).border = thin_border
                    ws.cell(row=row, column=3, value=target).border = thin_border
                    ws.cell(row=row, column=4, value=status).border = thin_border
            
                # Maturity Section
                ws['A12'] = 'MATURITY RADAR SCORES'
                ws['A12'].font = Font(bold=True, size=14, color='667eea')
                ws.merge_cells('A12:C12')
            
                maturity_headers = ['Dimension', 'Score (1-5)', 'Level']
                for col, header in enumerate(maturity_headers, 1):
                    cell = ws.cell(row=14, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                    cell.border = thin_border
            
                def get_level(score):
                    if score < 2: return 'Initial'
                    elif score < 3: return 'Developing'
                    elif score < 4: return 'Defined'
                    elif score < 4.5: return 'Managed'
                    else: return 'Optimized'
            
                maturity_data = [
                    ('Governance', maturity['governance'], get_level(maturity['governance'])),
                    ('Risk Management', maturity['risk_mgmt'], get_level(maturity['risk_mgmt'])),
                    ('Compliance', maturity['compliance'], get_level(maturity['compliance'])),
                    ('Technology', maturity['technology'], get_level(maturity['technology'])),
                    ('Process', maturity['process'], get_level(maturity['process'])),
                    ('Overall Average', maturity['average'], get_level(maturity['average'])),
                ]
            
                for row, (dim, score, level) in enumerate(maturity_data, 15):
                    ws.cell(row=row, column=1, value=dim).border = thin_border
                    ws.cell(row=row, column=2, value=score).border = thin_border
                    ws.cell(row=row, column=3, value=level).border = thin_border
            
                # Adjust column widths
                ws.column_dimensions['A'].width = 25
                ws.column_dimensions['B'].width = 15
                ws.column_dimensions['C'].width = 15
                ws.column_dimensions['D'].width = 20
    
        elif export_type == 'risks':
            ws.title = 'Risk Register'
            user_id = session.get('user_id')
        
            if user_id:
                conn = get_db()
                risks = conn.execute('''
                    SELECT domain, asset_name, threat, risk_level, created_at 
                    FROM risks WHERE user_id = ? ORDER BY created_at DESC
                ''', (user_id,)).fetchall()
            
                headers = ['#', 'Domain', 'Asset', 'Threat', 'Risk Level', 'Date']
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                    cell.border = thin_border
            
                for row, risk in enumerate(risks, 2):
                    ws.cell(row=row, column=1, value=row-1).border = thin_border
                    ws.cell(row=row, column=2, value=risk['domain']).border = thin_border
                    ws.cell(row=row, column=3, value=risk['asset_name']).border = thin_border
                    ws.cell(row=row, column=4, value=risk['threat'][:100] if risk['threat'] else '').border = thin_border
                    ws.cell(row=row, column=5, value=risk['risk_level']).border = thin_border
                    ws.cell(row=row, column=6, value=risk['created_at'][:10] if risk['created_at'] else '').border = thin_border
            
                ws.column_dimensions['A'].width = 5
                ws.column_dimensions['B'].width = 20
                ws.column_dimensions['C'].width = 25
                ws.column_dimensions['D'].width = 50
                ws.column_dimensions['E'].width = 12
                ws.column_dimensions['F'].width = 12
    
        elif export_type == 'documents':
            ws.title = 'Document Inventory'
            user_id = session.get('user_id')
        
            if user_id:
                conn = get_db()
            
                headers = ['#', 'Type', 'Domain', 'Title/Name', 'Language', 'Created']
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                    cell.border = thin_border
            
                row = 2
            
                # Strategies
                strategies = conn.execute('SELECT domain, org_name, language, created_at FROM strategies WHERE user_id = ?', (user_id,)).fetchall()
                for s in strategies:
                    ws.cell(row=row, column=1, value=row-1).border = thin_border
                    ws.cell(row=row, column=2, value='Strategy').border = thin_border
                    ws.cell(row=row, column=3, value=s['domain']).border = thin_border
                    ws.cell(row=row, column=4, value=s['org_name']).border = thin_border
                    ws.cell(row=row, column=5, value=s['language']).border = thin_border
                    ws.cell(row=row, column=6, value=s['created_at'][:10] if s['created_at'] else '').border = thin_border
                    row += 1
            
                # Policies
                policies = conn.execute('SELECT domain, policy_name, language, created_at FROM policies WHERE user_id = ?', (user_id,)).fetchall()
                for p in policies:
                    ws.cell(row=row, column=1, value=row-1).border = thin_border
                    ws.cell(row=row, column=2, value='Policy').border = thin_border
                    ws.cell(row=row, column=3, value=p['domain']).border = thin_border
                    ws.cell(row=row, column=4, value=p['policy_name']).border = thin_border
                    ws.cell(row=row, column=5, value=p['language']).border = thin_border
                    ws.cell(row=row, column=6, value=p['created_at'][:10] if p['created_at'] else '').border = thin_border
                    row += 1
            
                # Audits
                audits = conn.execute('SELECT domain, framework, language, created_at FROM audits WHERE user_id = ?', (user_id,)).fetchall()
                for a in audits:
                    ws.cell(row=row, column=1, value=row-1).border = thin_border
                    ws.cell(row=row, column=2, value='Audit').border = thin_border
                    ws.cell(row=row, column=3, value=a['domain']).border = thin_border
                    ws.cell(row=row, column=4, value=a['framework']).border = thin_border
                    ws.cell(row=row, column=5, value=a['language']).border = thin_border
                    ws.cell(row=row, column=6, value=a['created_at'][:10] if a['created_at'] else '').border = thin_border
                    row += 1
            
            
                ws.column_dimensions['A'].width = 5
                ws.column_dimensions['B'].width = 12
                ws.column_dimensions['C'].width = 20
                ws.column_dimensions['D'].width = 35
                ws.column_dimensions['E'].width = 10
                ws.column_dimensions['F'].width = 12
    
        # Save to BytesIO
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
    
        filename = data.get('filename', f'mizan_export_{export_type}')
    
        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename={filename}.xlsx'}
        )

    except Exception as e:
        print(f"Excel generation error: {e}", flush=True)
        return jsonify({"error": str(e)}), 500

# ============================================================================
# ADMIN ROUTES
# ============================================================================

@app.route('/admin')
@admin_required
def admin_dashboard():
    """Admin dashboard with statistics."""
    lang = request.args.get('lang', session.get('lang', 'en'))
    txt = get_text(lang)
    conn = get_db()
    
    # Get statistics
    stats = {
        'total_users': conn.execute('SELECT COUNT(*) FROM users').fetchone()[0],
        'active_users': conn.execute('SELECT COUNT(*) FROM users WHERE is_active = 1').fetchone()[0],
        'total_strategies': conn.execute('SELECT COUNT(*) FROM strategies').fetchone()[0],
        'total_policies': conn.execute('SELECT COUNT(*) FROM policies').fetchone()[0],
        'total_audits': conn.execute('SELECT COUNT(*) FROM audits').fetchone()[0],
        'total_risks': conn.execute('SELECT COUNT(*) FROM risks').fetchone()[0],
        'max_users': MAX_USERS,
    }
    
    # Calculate total documents
    stats['total_documents'] = stats['total_strategies'] + stats['total_policies'] + stats['total_audits'] + stats['total_risks']
    
    # Get documents by domain
    domains_strategies = conn.execute('SELECT domain, COUNT(*) as count FROM strategies GROUP BY domain').fetchall()
    domains_policies = conn.execute('SELECT domain, COUNT(*) as count FROM policies GROUP BY domain').fetchall()
    domains_audits = conn.execute('SELECT domain, COUNT(*) as count FROM audits GROUP BY domain').fetchall()
    domains_risks = conn.execute('SELECT domain, COUNT(*) as count FROM risks GROUP BY domain').fetchall()
    
    # Get ALL users with activity counts
    recent_users = conn.execute('''
        SELECT u.id, u.username, u.email, u.role, u.is_active, u.created_at, u.last_login,
            (SELECT COUNT(*) FROM strategies WHERE user_id = u.id) as strategy_count,
            (SELECT COUNT(*) FROM policies WHERE user_id = u.id) as policy_count,
            (SELECT COUNT(*) FROM audits WHERE user_id = u.id) as audit_count,
            (SELECT COUNT(*) FROM risks WHERE user_id = u.id) as risk_count,
            (SELECT COUNT(*) FROM audit_logs WHERE user_id = u.id) as action_count
        FROM users u ORDER BY u.created_at DESC
    ''').fetchall()
    
    # Get documents per day (last 7 days)
    docs_per_day = conn.execute('''
        SELECT DATE(created_at) as date, COUNT(*) as count 
        FROM (
            SELECT created_at FROM strategies 
            UNION ALL SELECT created_at FROM policies 
            UNION ALL SELECT created_at FROM audits 
            UNION ALL SELECT created_at FROM risks
        ) 
        WHERE created_at >= DATE('now', '-7 days')
        GROUP BY DATE(created_at) 
        ORDER BY date
    ''').fetchall()
    
    # Get benchmarks
    benchmarks = conn.execute('SELECT * FROM benchmarks ORDER BY sector').fetchall()
    
    
    return render_template('admin.html',
                          txt=txt,
                          lang=lang,
                          is_rtl=(lang == 'ar'),
                          stats=stats,
                          domains_strategies=domains_strategies,
                          domains_policies=domains_policies,
                          domains_audits=domains_audits,
                          domains_risks=domains_risks,
                          recent_users=recent_users,
                          docs_per_day=docs_per_day,
                          benchmarks=benchmarks,
                          config=config)

@app.route('/admin/api/benchmark/<int:benchmark_id>', methods=['POST'])
@admin_required
def update_benchmark(benchmark_id):
    """Update benchmark data."""
    data = request.json
    conn = get_db()
    
    try:
        conn.execute('''
            UPDATE benchmarks SET 
                compliance_score_avg = ?,
                maturity_level_avg = ?,
                risk_coverage_avg = ?,
                policy_count_avg = ?,
                source = ?,
                source_year = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
        ''', (
            float(data.get('compliance_score_avg', 60)),
            float(data.get('maturity_level_avg', 2.5)),
            float(data.get('risk_coverage_avg', 50)),
            int(data.get('policy_count_avg', 5)),
            data.get('source', 'Manual Update'),
            int(data.get('source_year', 2024)),
            benchmark_id
        ))
        conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/admin/toggle-user/<int:user_id>', methods=['POST'])
@admin_required
def toggle_user(user_id):
    """Toggle user active status."""
    conn = get_db()
    user = conn.execute('SELECT is_active FROM users WHERE id = ?', (user_id,)).fetchone()
    if user:
        new_status = 0 if user['is_active'] == 1 else 1
        conn.execute('UPDATE users SET is_active = ? WHERE id = ?', (new_status, user_id))
        conn.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/delete-user/<int:user_id>', methods=['POST'])
@admin_required
def delete_user(user_id):
    """Delete user (except admin)."""
    conn = get_db()
    user = conn.execute('SELECT role FROM users WHERE id = ?', (user_id,)).fetchone()
    if user and user['role'] != 'admin':
        conn.execute('DELETE FROM users WHERE id = ?', (user_id,))
        conn.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/api/stats')
@admin_required
def admin_api_stats():
    """API endpoint for admin statistics."""
    conn = get_db()
    
    stats = {
        'users': {
            'total': conn.execute('SELECT COUNT(*) FROM users').fetchone()[0],
            'active': conn.execute('SELECT COUNT(*) FROM users WHERE is_active = 1').fetchone()[0],
            'limit': MAX_USERS
        },
        'documents': {
            'strategies': conn.execute('SELECT COUNT(*) FROM strategies').fetchone()[0],
            'policies': conn.execute('SELECT COUNT(*) FROM policies').fetchone()[0],
            'audits': conn.execute('SELECT COUNT(*) FROM audits').fetchone()[0],
            'risks': conn.execute('SELECT COUNT(*) FROM risks').fetchone()[0]
        }
    }
    stats['documents']['total'] = sum(stats['documents'].values())
    
    return jsonify(stats)

@app.route('/admin/api/ai-diagnostic')
@admin_required
def admin_ai_diagnostic():
    """Diagnose AI provider connectivity — shows which keys are set and tests each provider."""
    results = {
        'env_vars': {},
        'provider_tests': {},
        'selected_provider': None,
        'ai_provider_setting': config.AI_PROVIDER,
    }
    
    # 1. Check which env vars are set (show first/last 4 chars only)
    keys = {
        'ANTHROPIC_API_KEY': config.ANTHROPIC_API_KEY,
        'OPENAI_API_KEY': config.OPENAI_API_KEY,
        'GOOGLE_API_KEY': config.GOOGLE_API_KEY,
        'GROQ_API_KEY': config.GROQ_API_KEY,
    }
    for name, val in keys.items():
        if val and len(val) > 8:
            results['env_vars'][name] = f"{val[:4]}...{val[-4:]} ({len(val)} chars)"
        elif val:
            results['env_vars'][name] = f"SET ({len(val)} chars — may be too short)"
        else:
            results['env_vars'][name] = "NOT SET"
    
    # 2. Which provider is selected?
    results['selected_provider'] = get_ai_provider() or 'None (no keys detected)'
    
    # 3. Test each available provider with a tiny prompt
    test_prompt = "Reply with exactly: OK"
    
    if config.ANTHROPIC_API_KEY:
        try:
            r = _generate_anthropic("Reply briefly.", test_prompt, 'en')
            results['provider_tests']['anthropic'] = f"✅ OK ({len(r)} chars)"
        except Exception as e:
            results['provider_tests']['anthropic'] = f"❌ {type(e).__name__}: {str(e)[:200]}"
    
    if config.OPENAI_API_KEY:
        try:
            r = _generate_openai("Reply briefly.", test_prompt, 'en')
            results['provider_tests']['openai'] = f"✅ OK ({len(r)} chars)"
        except Exception as e:
            results['provider_tests']['openai'] = f"❌ {type(e).__name__}: {str(e)[:200]}"
    
    if config.GOOGLE_API_KEY:
        try:
            r = _generate_google("Reply briefly.", test_prompt, 'en')
            results['provider_tests']['google'] = f"✅ OK ({len(r)} chars)"
        except Exception as e:
            results['provider_tests']['google'] = f"❌ {type(e).__name__}: {str(e)[:200]}"
    
    if config.GROQ_API_KEY:
        try:
            r = _generate_groq("Reply briefly.", test_prompt, 'en')
            results['provider_tests']['groq'] = f"✅ OK ({len(r)} chars)"
        except Exception as e:
            results['provider_tests']['groq'] = f"❌ {type(e).__name__}: {str(e)[:200]}"
    
    if not results['provider_tests']:
        results['provider_tests']['none'] = "No API keys detected — check Render env vars"
    
    # 4. Check SDK availability
    sdk_check = {}
    for pkg in ['anthropic', 'openai', 'google.generativeai']:
        try:
            __import__(pkg)
            sdk_check[pkg] = "✅ installed"
        except ImportError:
            sdk_check[pkg] = "❌ NOT installed — add to requirements.txt"
    results['sdk_installed'] = sdk_check
    
    return jsonify(results)

# ============================================================================
# TASKS MANAGEMENT API
# ============================================================================

@app.route('/api/tasks', methods=['GET'])
@login_required
def api_get_tasks():
    """Get all tasks for the current user."""
    try:
        domain = request.args.get('domain', '')
        status = request.args.get('status', '')
        
        conn = get_db()
        query = 'SELECT * FROM project_tasks WHERE user_id = ?'
        params = [session['user_id']]
        
        if domain:
            query += ' AND domain = ?'
            params.append(domain)
        if status:
            query += ' AND status = ?'
            params.append(status)
        
        query += ' ORDER BY CASE priority WHEN "critical" THEN 0 WHEN "high" THEN 1 WHEN "medium" THEN 2 WHEN "low" THEN 3 END, due_date ASC'
        
        tasks = conn.execute(query, params).fetchall()
        
        return jsonify({'success': True, 'tasks': [dict(t) for t in tasks]})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/tasks', methods=['POST'])
@login_required
def api_create_task():
    """Create a new task."""
    try:
        data = request.json
        conn = get_db()
        conn.execute(
            'INSERT INTO project_tasks (user_id, domain, title, description, status, priority, owner, due_date, category) VALUES (?,?,?,?,?,?,?,?,?)',
            (session['user_id'], data.get('domain', 'General'), data.get('title', ''),
             data.get('description', ''), data.get('status', 'todo'),
             data.get('priority', 'medium'), data.get('owner', ''),
             data.get('due_date', ''), data.get('category', 'implementation'))
        )
        conn.commit()
        task_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        log_action(session['user_id'], 'create_task', {'task_id': task_id, 'title': data.get('title', '')})
        return jsonify({'success': True, 'id': task_id})
    except Exception as e:
        print(f"Create task error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/tasks/<int:task_id>', methods=['PUT'])
@login_required
def api_update_task(task_id):
    """Update a task."""
    try:
        data = request.json
        conn = get_db()
        # Verify ownership
        task = conn.execute('SELECT id FROM project_tasks WHERE id=? AND user_id=?', (task_id, session['user_id'])).fetchone()
        if not task:
            return jsonify({'success': False, 'error': 'Task not found'}), 404
        
        fields = []
        params = []
        for key in ['title', 'description', 'status', 'priority', 'owner', 'due_date', 'domain', 'category']:
            if key in data:
                fields.append(f'{key}=?')
                params.append(data[key])
        
        if fields:
            fields.append('updated_at=CURRENT_TIMESTAMP')
            params.append(task_id)
            conn.execute(f'UPDATE project_tasks SET {",".join(fields)} WHERE id=?', params)
            conn.commit()
        log_action(session['user_id'], 'update_task', {'task_id': task_id, 'fields': list(data.keys())})
        return jsonify({'success': True})
    except Exception as e:
        print(f"Update task error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/tasks/<int:task_id>', methods=['DELETE'])
@login_required
def api_delete_task(task_id):
    """Delete a task."""
    try:
        conn = get_db()
        conn.execute('DELETE FROM project_tasks WHERE id=? AND user_id=?', (task_id, session['user_id']))
        conn.commit()
        log_action(session['user_id'], 'delete_task', {'task_id': task_id})
        return jsonify({'success': True})
    except Exception as e:
        print(f"Delete task error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/tasks/stats')
@login_required
def api_task_stats():
    """Get task statistics."""
    try:
        conn = get_db()
        user_id = session['user_id']
        
        todo = conn.execute('SELECT COUNT(*) FROM project_tasks WHERE user_id=? AND status="todo"', (user_id,)).fetchone()[0]
        in_progress = conn.execute('SELECT COUNT(*) FROM project_tasks WHERE user_id=? AND status="in_progress"', (user_id,)).fetchone()[0]
        done = conn.execute('SELECT COUNT(*) FROM project_tasks WHERE user_id=? AND status="done"', (user_id,)).fetchone()[0]
        blocked = conn.execute('SELECT COUNT(*) FROM project_tasks WHERE user_id=? AND status="blocked"', (user_id,)).fetchone()[0]
        
        overdue = conn.execute(
            'SELECT COUNT(*) FROM project_tasks WHERE user_id=? AND status!="done" AND due_date < date("now") AND due_date != ""',
            (user_id,)
        ).fetchone()[0]
        
        total = todo + in_progress + done + blocked
        return jsonify({
            'success': True,
            'stats': {
                'todo': todo, 'in_progress': in_progress, 'done': done, 'blocked': blocked,
                'total': total, 'overdue': overdue,
                'completion_rate': round((done / total * 100) if total > 0 else 0)
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/tasks/by-category')
@login_required
def api_tasks_by_category():
    """Get tasks grouped by category for Tasks Management tab.
    
    Categories: strategic_initiative, implementation, kpi, gap_mitigation, risk_mitigation
    Optional ?domain=... and ?category=... filters.
    """
    try:
        conn = get_db()
        user_id = session['user_id']
        domain = request.args.get('domain', '')
        category = request.args.get('category', '')
        
        query = 'SELECT * FROM project_tasks WHERE user_id = ?'
        params = [user_id]
        
        if domain:
            query += ' AND domain = ?'
            params.append(domain)
        if category:
            query += ' AND category = ?'
            params.append(category)
        
        query += ' ORDER BY CASE priority WHEN "critical" THEN 0 WHEN "high" THEN 1 WHEN "medium" THEN 2 WHEN "low" THEN 3 END, created_at DESC'
        
        tasks = conn.execute(query, params).fetchall()
        
        # Category counts
        cq = 'SELECT COALESCE(category,"implementation") as cat, COUNT(*) as cnt FROM project_tasks WHERE user_id = ?'
        cp = [user_id]
        if domain:
            cq += ' AND domain = ?'
            cp.append(domain)
        cq += ' GROUP BY COALESCE(category,"implementation")'
        counts = {r['cat']: r['cnt'] for r in conn.execute(cq, cp).fetchall()}
        
        return jsonify({
            'success': True,
            'tasks': [dict(t) for t in tasks],
            'counts': counts,
            'total': len(tasks)
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ai/status')
@login_required
def api_ai_status():
    """Get current AI provider status."""
    try:
        return jsonify({'success': True, **get_ai_status()})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================================================
# AWARENESS MODULES API
# ============================================================================

@app.route('/api/awareness/modules')
@login_required
def api_awareness_modules():
    """Get awareness modules for a domain."""
    try:
        domain = request.args.get('domain', 'Cyber Security')
        lang = request.args.get('lang', 'en')
        domain_code = DOMAIN_CODES.get(domain, 'cyber')
        
        modules = AWARENESS_MODULES.get(domain_code, {}).get(lang, AWARENESS_MODULES.get(domain_code, {}).get('en', []))
        
        # Get user's completed modules
        conn = get_db()
        completed = conn.execute(
            'SELECT module_id, score, total, passed FROM awareness_scores WHERE user_id = ? AND domain = ?',
            (session['user_id'], domain)
        ).fetchall()
        
        completed_map = {r['module_id']: {'score': r['score'], 'total': r['total'], 'passed': r['passed']} for r in completed}
        
        # Add completion status to modules
        modules_with_status = []
        for m in modules:
            m_copy = dict(m)
            if m['id'] in completed_map:
                m_copy['completed'] = True
                m_copy['user_score'] = completed_map[m['id']]['score']
                m_copy['user_total'] = completed_map[m['id']]['total']
                m_copy['passed'] = bool(completed_map[m['id']]['passed'])
            else:
                m_copy['completed'] = False
            modules_with_status.append(m_copy)
        
        return jsonify({'success': True, 'modules': modules_with_status})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/awareness/submit', methods=['POST'])
@login_required
def api_awareness_submit():
    """Submit quiz results."""
    try:
        data = request.json
        module_id = data.get('module_id', '')
        domain = data.get('domain', '')
        score = data.get('score', 0)
        total = data.get('total', 0)
        lang = data.get('language', 'en')
        passed = 1 if total > 0 and (score / total) >= 0.7 else 0
        
        conn = get_db()
        # Update or insert score (keep best score)
        existing = conn.execute(
            'SELECT id, score FROM awareness_scores WHERE user_id = ? AND domain = ? AND module_id = ?',
            (session['user_id'], domain, module_id)
        ).fetchone()
        
        if existing:
            if score > existing['score']:
                conn.execute('UPDATE awareness_scores SET score=?, total=?, passed=?, language=?, completed_at=CURRENT_TIMESTAMP WHERE id=?',
                            (score, total, passed, lang, existing['id']))
        else:
            conn.execute('INSERT INTO awareness_scores (user_id, domain, module_id, score, total, passed, language) VALUES (?,?,?,?,?,?,?)',
                        (session['user_id'], domain, module_id, score, total, passed, lang))
        conn.commit()
        
        return jsonify({'success': True, 'score': score, 'total': total, 'passed': bool(passed), 'percentage': round((score/total)*100) if total > 0 else 0})
    except Exception as e:
        print(f"Awareness submit error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/awareness/progress')
@login_required
def api_awareness_progress():
    """Get user's overall awareness progress."""
    try:
        conn = get_db()
        scores = conn.execute(
            'SELECT domain, module_id, score, total, passed FROM awareness_scores WHERE user_id = ?',
            (session['user_id'],)
        ).fetchall()
        
        progress = {}
        for s in scores:
            d = s['domain']
            if d not in progress:
                progress[d] = {'completed': 0, 'passed': 0, 'total_score': 0, 'total_possible': 0}
            progress[d]['completed'] += 1
            progress[d]['passed'] += s['passed']
            progress[d]['total_score'] += s['score']
            progress[d]['total_possible'] += s['total']
        
        return jsonify({'success': True, 'progress': progress})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# ============================================================================
# PILLAR 2.2: COMPLIANCE SNAPSHOTS
# ============================================================================

@app.route('/api/compliance/snapshot', methods=['POST'])
@login_required
def api_compliance_snapshot():
    """Calculate current maturity levels and save to compliance_history for trend reporting."""
    try:
        user_id = session['user_id']
        conn = get_db()
        
        # Calculate current scores
        current = calculate_compliance_score(user_id)
        maturity = calculate_maturity_levels(user_id)
        
        # Save snapshot (allows multiple per day for on-demand use)
        conn.execute('''
            INSERT INTO compliance_history 
            (user_id, score, maturity_avg, strategies, policies, audits, risks, domains_covered)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (user_id, current['score'], maturity['average'], current['strategies'],
              current['policies'], current['audits'], current['risks'], current['domains_covered']))
        conn.commit()
        
        snapshot_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        log_action(user_id, 'compliance_snapshot', {'score': current['score'], 'maturity_avg': maturity['average']})
        
        # Return full snapshot data for immediate UI display
        return jsonify({
            'success': True,
            'snapshot': {
                'id': snapshot_id,
                'score': current['score'],
                'maturity': maturity,
                'document_counts': {
                    'strategies': current['strategies'],
                    'policies': current['policies'],
                    'audits': current['audits'],
                    'risks': current['risks'],
                    'domains_covered': current['domains_covered']
                },
                'recorded_at': datetime.now().isoformat()
            }
        })
    except Exception as e:
        print(f"Compliance snapshot error: {e}", flush=True)
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# PILLAR 3.2: DYNAMIC FRAMEWORKS API
# ============================================================================

@app.route('/api/frameworks', methods=['GET'])
@login_required
def api_get_frameworks():
    """Fetch frameworks from DB (replaces hardcoded Jinja injection for dynamic use)."""
    try:
        domain = request.args.get('domain', '')
        conn = get_db()
        
        if domain:
            rows = conn.execute(
                'SELECT id, domain, region, name, description FROM grc_frameworks WHERE domain = ? AND is_active = 1 ORDER BY sort_order',
                (domain,)
            ).fetchall()
        else:
            rows = conn.execute(
                'SELECT id, domain, region, name, description FROM grc_frameworks WHERE is_active = 1 ORDER BY domain, sort_order'
            ).fetchall()
        
        # Group by region for frontend consumption (matches existing frameworksByRegion format)
        by_region = {}
        for row in rows:
            region = row['region']
            if region not in by_region:
                by_region[region] = []
            by_region[region].append({
                'id': row['id'],
                'name': row['name'],
                'description': row['description'] or '',
                'domain': row['domain']
            })
        
        return jsonify({'success': True, 'frameworks': by_region})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/frameworks', methods=['POST'])
@login_required
def api_add_framework():
    """Admin: Add a new framework to the DB."""
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'Admin only'}), 403
        
        data = request.json
        conn = get_db()
        conn.execute(
            'INSERT OR IGNORE INTO grc_frameworks (domain, region, name, description, is_active, sort_order) VALUES (?,?,?,?,1,?)',
            (data['domain'], data['region'], data['name'], data.get('description', ''),
             data.get('sort_order', 999))
        )
        conn.commit()
        log_action(session['user_id'], 'add_framework', {'name': data['name'], 'domain': data['domain']})
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/frameworks/<int:fw_id>', methods=['DELETE'])
@login_required
def api_delete_framework(fw_id):
    """Admin: Soft-delete a framework."""
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'Admin only'}), 403
        conn = get_db()
        conn.execute('UPDATE grc_frameworks SET is_active = 0 WHERE id = ?', (fw_id,))
        conn.commit()
        log_action(session['user_id'], 'delete_framework', {'framework_id': fw_id})
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# PILLAR 4.1: FORM DRAFTS (SAVE / LOAD / DELETE)
# ============================================================================

@app.route('/api/drafts', methods=['GET'])
@login_required
def api_get_draft():
    """Load a saved form draft for a specific domain + form type."""
    try:
        user_id = session['user_id']
        domain = request.args.get('domain', '')
        form_type = request.args.get('form_type', 'strategy')
        
        conn = get_db()
        row = conn.execute(
            'SELECT draft_data, updated_at FROM form_drafts WHERE user_id = ? AND domain = ? AND form_type = ?',
            (user_id, domain, form_type)
        ).fetchone()
        
        if row:
            return jsonify({
                'success': True,
                'draft': json.loads(row['draft_data']),
                'updated_at': row['updated_at']
            })
        return jsonify({'success': True, 'draft': None})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/drafts', methods=['POST'])
@login_required
def api_save_draft():
    """Save partial form data (upsert by user_id + domain + form_type)."""
    try:
        user_id = session['user_id']
        data = request.json
        domain = data.get('domain', '')
        form_type = data.get('form_type', 'strategy')
        form_data = data.get('form_data', {})
        
        if not domain or not form_data:
            return jsonify({'success': False, 'error': 'Missing domain or form_data'}), 400
        
        conn = get_db()
        conn.execute('''
            INSERT INTO form_drafts (user_id, domain, form_type, draft_data, updated_at)
            VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
            ON CONFLICT(user_id, domain, form_type) 
            DO UPDATE SET draft_data = excluded.draft_data, updated_at = CURRENT_TIMESTAMP
        ''', (user_id, domain, form_type, json.dumps(form_data, ensure_ascii=False)))
        conn.commit()
        
        log_action(user_id, 'save_draft', {'domain': domain, 'form_type': form_type})
        return jsonify({'success': True, 'message': 'Draft saved'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/drafts', methods=['DELETE'])
@login_required
def api_delete_draft():
    """Delete a saved draft."""
    try:
        user_id = session['user_id']
        domain = request.args.get('domain', '')
        form_type = request.args.get('form_type', 'strategy')
        
        conn = get_db()
        conn.execute(
            'DELETE FROM form_drafts WHERE user_id = ? AND domain = ? AND form_type = ?',
            (user_id, domain, form_type)
        )
        conn.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# PILLAR 4.2: LOGO UPLOAD FOR BRANDED EXPORTS
# ============================================================================

@app.route('/api/user/logo', methods=['POST'])
@login_required
def api_upload_logo():
    """Upload a logo image for branded PDF/DOCX headers."""
    try:
        if 'logo' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['logo']
        if not file.filename:
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        # Validate file type
        allowed_ext = {'.png', '.jpg', '.jpeg', '.gif', '.svg'}
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in allowed_ext:
            return jsonify({'success': False, 'error': f'Invalid file type. Allowed: {", ".join(allowed_ext)}'}), 400
        
        # Save to uploads directory
        upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'logos')
        os.makedirs(upload_dir, exist_ok=True)
        
        safe_filename = f"user_{session['user_id']}_logo{ext}"
        logo_path = os.path.join(upload_dir, safe_filename)
        file.save(logo_path)
        
        # Update user profile
        conn = get_db()
        conn.execute('UPDATE users SET logo_path = ? WHERE id = ?', (logo_path, session['user_id']))
        conn.commit()
        
        log_action(session['user_id'], 'upload_logo', {'filename': safe_filename})
        return jsonify({'success': True, 'logo_path': f'/static/logos/{safe_filename}'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/user/logo', methods=['DELETE'])
@login_required
def api_delete_logo():
    """Remove the user's branded logo."""
    try:
        conn = get_db()
        user = conn.execute('SELECT logo_path FROM users WHERE id = ?', (session['user_id'],)).fetchone()
        if user and user['logo_path'] and os.path.isfile(user['logo_path']):
            os.remove(user['logo_path'])
        conn.execute('UPDATE users SET logo_path = NULL WHERE id = ?', (session['user_id'],))
        conn.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# PILLAR 3.1: TOKEN USAGE ADMIN ENDPOINT
# ============================================================================

@app.route('/api/user/token-usage', methods=['GET'])
@login_required
def api_token_usage():
    """Return current user's token usage and limit."""
    try:
        allowed, used, limit = check_token_quota(session['user_id'])
        return jsonify({
            'success': True,
            'token_usage': used,
            'token_limit': limit,
            'remaining': max(0, limit - used),
            'exceeded': not allowed
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/admin/reset-tokens/<int:user_id>', methods=['POST'])
@login_required
def api_admin_reset_tokens(user_id):
    """Admin: Reset a user's token usage counter."""
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'Admin only'}), 403
        conn = get_db()
        new_limit = request.json.get('new_limit', TOKEN_LIMIT_DEFAULT) if request.json else TOKEN_LIMIT_DEFAULT
        conn.execute('UPDATE users SET token_usage = 0, token_limit = ? WHERE id = ?', (new_limit, user_id))
        conn.commit()
        log_action(session['user_id'], 'reset_tokens', {'target_user': user_id, 'new_limit': new_limit})
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/admin/api/users/<int:user_id>/quota', methods=['POST'])
@login_required
def api_admin_user_quota(user_id):
    """Admin: Manage user token quota — reset usage and/or set new limit.
    
    JSON body options:
      {"reset": true}                    → Resets token_usage to 0
      {"token_limit": 1000000}           → Sets new token limit
      {"reset": true, "token_limit": X}  → Both reset and set limit
    """
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'Admin only'}), 403
        
        data = request.get_json() if request.is_json else {}
        conn = get_db()
        
        # Get current user info
        user = conn.execute('SELECT username, token_usage, token_limit FROM users WHERE id = ?', (user_id,)).fetchone()
        if not user:
            return jsonify({'success': False, 'error': f'User ID {user_id} not found'}), 404
        
        updates = []
        params = []
        action_details = {'target_user': user_id, 'username': user['username']}
        
        # Reset usage
        if data.get('reset', False):
            updates.append('token_usage = 0')
            action_details['reset_usage'] = True
            action_details['previous_usage'] = user['token_usage']
        
        # Set new limit
        new_limit = data.get('token_limit')
        if new_limit is not None:
            try:
                new_limit = int(new_limit)
                if new_limit < 0:
                    return jsonify({'success': False, 'error': 'token_limit must be >= 0'}), 400
                updates.append('token_limit = ?')
                params.append(new_limit)
                action_details['new_limit'] = new_limit
                action_details['previous_limit'] = user['token_limit']
            except (ValueError, TypeError):
                return jsonify({'success': False, 'error': 'token_limit must be an integer'}), 400
        
        if not updates:
            return jsonify({'success': False, 'error': 'No action specified. Use {"reset": true} and/or {"token_limit": N}'}), 400
        
        params.append(user_id)
        conn.execute(f"UPDATE users SET {', '.join(updates)} WHERE id = ?", params)
        conn.commit()
        
        log_action(session['user_id'], 'manage_quota', action_details)
        
        # Return updated state
        updated = conn.execute('SELECT token_usage, token_limit FROM users WHERE id = ?', (user_id,)).fetchone()
        return jsonify({
            'success': True,
            'user_id': user_id,
            'username': user['username'],
            'token_usage': updated['token_usage'],
            'token_limit': updated['token_limit']
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# PILLAR 1.2: AUDIT LOG VIEWER (ADMIN)
# ============================================================================

@app.route('/api/admin/audit-logs', methods=['GET'])
@login_required
def api_audit_logs():
    """Admin: View audit log entries."""
    try:
        if session.get('role') != 'admin':
            return jsonify({'success': False, 'error': 'Admin only'}), 403
        
        limit = request.args.get('limit', 100, type=int)
        action_filter = request.args.get('action', '')
        conn = get_db()
        
        if action_filter:
            rows = conn.execute(
                'SELECT al.*, u.username FROM audit_logs al LEFT JOIN users u ON al.user_id = u.id WHERE al.action = ? ORDER BY al.created_at DESC LIMIT ?',
                (action_filter, limit)
            ).fetchall()
        else:
            rows = conn.execute(
                'SELECT al.*, u.username FROM audit_logs al LEFT JOIN users u ON al.user_id = u.id ORDER BY al.created_at DESC LIMIT ?',
                (limit,)
            ).fetchall()
        
        logs = []
        for r in rows:
            logs.append({
                'id': r['id'],
                'user_id': r['user_id'],
                'username': r['username'] or 'system',
                'action': r['action'],
                'metadata': json.loads(r['metadata']) if r['metadata'] else None,
                'created_at': r['created_at']
            })
        
        return jsonify({'success': True, 'logs': logs})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================================================
# MAIN
# ============================================================================



# ============================================================================
# KPI API ENDPOINTS
# ============================================================================

@app.route('/api/kpis', methods=['GET'])
@login_required
def api_list_kpis():
    """List user's KPIs."""
    try:
        conn = get_db()
        domain = request.args.get('domain', '')
        if domain:
            kpis = conn.execute('SELECT * FROM kpis WHERE user_id = ? AND domain = ? ORDER BY created_at DESC',
                               (session['user_id'], domain)).fetchall()
        else:
            kpis = conn.execute('SELECT * FROM kpis WHERE user_id = ? ORDER BY created_at DESC',
                               (session['user_id'],)).fetchall()
        return jsonify({'success': True, 'kpis': [dict(k) for k in kpis]})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/kpis/<int:kpi_id>', methods=['PUT'])
@login_required
def api_update_kpi(kpi_id):
    """Update a KPI's current value or status."""
    try:
        data = request.json
        conn = get_db()
        kpi = conn.execute('SELECT id FROM kpis WHERE id = ? AND user_id = ?', (kpi_id, session['user_id'])).fetchone()
        if not kpi:
            return jsonify({'success': False, 'error': 'KPI not found'}), 404
        conn.execute("""UPDATE kpis SET current_value = ?, status = ? WHERE id = ? AND user_id = ?""",
                     (data.get('current_value', 'TBD'), data.get('status', 'pending'), kpi_id, session['user_id']))
        conn.commit()
        log_action(session['user_id'], 'update_kpi', {'kpi_id': kpi_id})
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
