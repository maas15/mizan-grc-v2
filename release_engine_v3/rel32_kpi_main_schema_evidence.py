"""REL3.2 — returned-file KPI main schema consistency evidence."""

from __future__ import annotations

import re
from typing import Any, Dict, List, Optional

from release_engine_v3.rel32_preview_table_dom import (
    evaluate_preview_dom_binding_check,
    extract_table_dom_binding,
)
from release_engine_v3.rel32_table_schema_binding import (
    REL32_KPI_MAIN_EXPECTED_SCHEMA_AR,
    emit_rel32_kpi_main_schema_consistency_diag,
    emit_rel32_kpi_owner_consistency_diag,
    evaluate_kpi_main_schema_consistency,
    evaluate_kpi_owner_consistency,
    find_kpi_main_table,
    rebind_table_spec,
)


def _cells_from_markdown_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip('|').split('|')]


def _kpi_main_section_blob(blob: str) -> str:
    try:
        from release_engine.rel27_export_checks import _kpi_section_blob
        return _kpi_section_blob(blob or '')
    except Exception:  # noqa: BLE001
        return blob or ''


def _canonical_kpi_header_cells(cells: List[str]) -> bool:
    expected = list(REL32_KPI_MAIN_EXPECTED_SCHEMA_AR)
    if len(cells) >= len(expected) and cells[:len(expected)] == expected:
        return True
    return cells == expected


def _locate_canonical_kpi_table_in_docx(doc: Any) -> tuple[List[str], List[List[str]]]:
    """Prefer the 8-column canonical KPI table over legacy KPI-like tables."""
    expected = list(REL32_KPI_MAIN_EXPECTED_SCHEMA_AR)
    best: tuple[List[str], List[List[str]]] = ([], [])
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = [(c.text or '').strip() for c in row.cells]
            if not cells or not _canonical_kpi_header_cells(cells):
                continue
            headers = cells[:len(expected)]
            rows_out: List[List[str]] = []
            for j in range(i + 1, len(table.rows)):
                rcells = [(c.text or '').strip() for c in table.rows[j].cells]
                if not rcells or not rcells[0]:
                    continue
                if rcells[0] in ('#', 'رقم') or '---' in rcells[0]:
                    continue
                if rcells[0].replace('.', '').isdigit():
                    rows_out.append(
                        rcells[:len(expected)] if len(rcells) >= len(expected) else rcells)
            if rows_out:
                return headers, rows_out
            best = (headers, rows_out)
    return best


def extract_kpi_main_header_labels_from_docx(raw: bytes) -> List[str]:
    """Extract KPI main header from structured DOCX table cells."""
    try:
        from io import BytesIO
        from docx import Document
        doc = Document(BytesIO(raw))
        headers, _rows = _locate_canonical_kpi_table_in_docx(doc)
        if headers:
            return headers
        expected = list(REL32_KPI_MAIN_EXPECTED_SCHEMA_AR)
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or '').strip() for c in row.cells]
                if not cells:
                    continue
                blob_ln = ' '.join(cells).lower()
                if not any(k in blob_ln for k in ('مؤشر', 'kpi', 'indicator', 'وصف')):
                    continue
                if 'وصف المؤشر' in blob_ln and len(cells) >= len(expected):
                    return cells[:len(expected)]
                if cells[0] in ('#', 'رقم') and len(cells) >= len(expected):
                    return cells[:len(expected)]
        return []
    except Exception:  # noqa: BLE001
        return []


def extract_kpi_main_rows_from_docx(raw: bytes) -> List[List[str]]:
    """Extract KPI main data rows from structured DOCX tables."""
    try:
        from io import BytesIO
        from docx import Document
        doc = Document(BytesIO(raw))
        _headers, rows_out = _locate_canonical_kpi_table_in_docx(doc)
        if rows_out:
            return rows_out
    except Exception:  # noqa: BLE001
        pass
    rows_out = []
    try:
        from io import BytesIO
        from docx import Document
        doc = Document(BytesIO(raw))
        expected = list(REL32_KPI_MAIN_EXPECTED_SCHEMA_AR)
        for table in doc.tables:
            header_idx = -1
            for i, row in enumerate(table.rows):
                cells = [(c.text or '').strip() for c in row.cells]
                blob_ln = ' '.join(cells).lower()
                if header_idx < 0 and 'وصف المؤشر' in blob_ln:
                    header_idx = i
                    continue
                if header_idx >= 0 and i > header_idx:
                    if not cells or not cells[0]:
                        continue
                    if cells[0] in ('#', 'رقم') or '---' in cells[0]:
                        continue
                    if cells[0].replace('.', '').isdigit():
                        rows_out.append(cells)
            if rows_out:
                break
    except Exception:  # noqa: BLE001
        return []
    return rows_out


def evaluate_kpi_main_schema_from_docx_bytes(
        raw: bytes,
        *,
        route_name: str = 'docx',
        lang: str = 'ar',
) -> Dict[str, Any]:
    headers = extract_kpi_main_header_labels_from_docx(raw)
    rows = extract_kpi_main_rows_from_docx(raw)
    if headers and rows:
        diag = evaluate_kpi_main_schema_consistency(
            route_name=route_name,
            header_labels=headers,
            rows=rows,
            lang=lang,
            repair_rows=False,
        )
        emit_rel32_kpi_main_schema_consistency_diag(diag)
        return diag
    text = ''
    try:
        from release_engine_v3.evidence.docx_text_extractor import (
            extract_docx_visible_text,
        )
        text = extract_docx_visible_text(raw)
    except Exception:  # noqa: BLE001
        pass
    return evaluate_kpi_main_schema_from_export_text(
        text, route_name=route_name, lang=lang)


def _emit_kpi_extractability(
        *,
        route_name: str,
        domain: str,
        document_type: str,
        table_render_mode: str,
        used_cards_fallback: bool,
        header_labels: List[str],
        row_count_rendered: int,
        row_count_extracted: int,
        kpi_pdf_extractable: bool,
        blocking_errors: List[str],
) -> Dict[str, Any]:
    """Emit the [REL33-PDF-KPI-MAIN-EXTRACTABILITY] diagnostic."""
    from release_engine_v3.rel33_pdf_evidence_norm import (
        emit_rel33_pdf_kpi_main_extractability,
    )
    diag = {
        'route_name': route_name,
        'domain': domain,
        'document_type': document_type,
        'table_render_mode': table_render_mode,
        'used_cards_fallback': used_cards_fallback,
        'header_labels': list(header_labels or []),
        'expected_header_labels': list(REL32_KPI_MAIN_EXPECTED_SCHEMA_AR),
        'row_count_rendered': row_count_rendered,
        'row_count_extracted': row_count_extracted,
        'kpi_pdf_extractable': kpi_pdf_extractable,
        'blocking_errors': list(blocking_errors or []),
    }
    emit_rel33_pdf_kpi_main_extractability(diag)
    return diag


def _kpi_rows_from_pdf_tables(raw: bytes) -> List[List[str]]:
    """Aggregate canonical 8-column KPI data rows across all PDF pages.

    Aggregating across pages fixes the case where a dense KPI table spans a
    page break (header + first rows on page 1, remaining rows on page 2) and
    per-page structured detection would otherwise see too few rows.
    """
    expected = list(REL32_KPI_MAIN_EXPECTED_SCHEMA_AR)
    rows: List[List[str]] = []
    try:
        import fitz
    except Exception:  # noqa: BLE001
        return rows
    try:
        doc = fitz.open(stream=raw, filetype='pdf')
    except Exception:  # noqa: BLE001
        return rows
    for page in doc:
        try:
            found = page.find_tables()
            tables = found.tables if hasattr(found, 'tables') else list(found)
        except Exception:  # noqa: BLE001
            continue
        for table in tables:
            try:
                df = table.to_pandas()
            except Exception:  # noqa: BLE001
                continue
            if df.shape[1] != len(expected):
                continue
            table_rows: List[List[str]] = []
            for _, series in df.iterrows():
                cells = [
                    str(v).strip().replace('\n', ' ')
                    for v in series.tolist()
                ]
                if cells and cells[0].replace('.', '').isdigit():
                    table_rows.append(cells)
            if not table_rows:
                continue
            blob = ' '.join(' '.join(r) for r in table_rows[:3])
            if not any(k in blob for k in ('KPI', 'SOC', 'SIEM', 'CISO', 'مدير')):
                continue
            rows.extend(table_rows)
    return rows


def evaluate_kpi_main_schema_from_pdf_bytes(
        raw: bytes,
        *,
        route_name: str = 'pdf',
        lang: str = 'ar',
        domain: str = '',
        document_type: str = 'strategy',
) -> Dict[str, Any]:
    """Evaluate KPI main schema from PDF bytes via structured table extraction.

    Uses PyMuPDF structured table detection (aggregated across pages) as the
    primary returned-file evidence — this is the canonical 8-column KPI table
    the renderer emitted. Falls back to normalized visible-text extraction so a
    KPI table that IS present in the returned PDF is not falsely reported as
    header-not-found because Arabic glyphs did not survive naive extraction.
    """
    expected = list(REL32_KPI_MAIN_EXPECTED_SCHEMA_AR)
    rows = _kpi_rows_from_pdf_tables(raw)
    if rows:
        diag = evaluate_kpi_main_schema_consistency(
            route_name=route_name,
            header_labels=list(expected),
            rows=rows,
            lang=lang,
            repair_rows=False,
        )
        if diag.get('kpi_main_schema_passed'):
            emit_rel32_kpi_main_schema_consistency_diag(diag)
            _emit_kpi_extractability(
                route_name=route_name, domain=domain,
                document_type=document_type,
                table_render_mode='structured_table',
                used_cards_fallback=False,
                header_labels=diag.get('header_labels') or expected,
                row_count_rendered=len(rows),
                row_count_extracted=diag.get('row_count') or len(rows),
                kpi_pdf_extractable=True,
                blocking_errors=[],
            )
            return diag
    from release_engine_v3.evidence.pdf_text_extractor import extract_pdf_visible_text
    text = extract_pdf_visible_text(raw)
    norm = (text or '').replace('#وصف المؤشر', '# وصف المؤشر')
    text_diag = evaluate_kpi_main_schema_from_export_text(
        norm, route_name=route_name, lang=lang)
    _emit_kpi_extractability(
        route_name=route_name, domain=domain,
        document_type=document_type,
        table_render_mode=(
            'structured_table' if rows else 'text_fallback'),
        used_cards_fallback=False,
        header_labels=text_diag.get('header_labels') or [],
        row_count_rendered=len(rows),
        row_count_extracted=text_diag.get('row_count') or 0,
        kpi_pdf_extractable=bool(text_diag.get('kpi_main_schema_passed')),
        blocking_errors=text_diag.get('blocking_errors') or [],
    )
    return text_diag


def extract_kpi_main_header_labels_from_text(blob: str) -> List[str]:
    """Extract KPI main header labels from returned DOCX/PDF/preview text."""
    section = _kpi_main_section_blob(blob)
    expected = list(REL32_KPI_MAIN_EXPECTED_SCHEMA_AR)
    for ln in section.splitlines():
        if not ln.strip().startswith('|') or '---' in ln:
            continue
        cells = _cells_from_markdown_row(ln)
        if not cells:
            continue
        blob_ln = ' '.join(cells).lower()
        if not any(k in blob_ln for k in ('مؤشر', 'kpi', 'indicator', 'وصف')):
            continue
        if cells[0] in ('#', 'رقم') or 'وصف المؤشر' in ln or 'المؤشر' in ln:
            return cells
    lines = [ln.strip() for ln in section.splitlines() if ln.strip()]
    for i, ln in enumerate(lines):
        if ln != expected[0]:
            continue
        window = lines[i:i + len(expected)]
        if window == expected:
            return list(expected)
    return []


def _extract_main_kpi_rows(blob: str) -> List[List[str]]:
    try:
        from release_engine.rel27_export_checks import _extract_kpi_main_rows
        return _extract_kpi_main_rows(blob or '')
    except Exception:  # noqa: BLE001
        return []


def evaluate_kpi_main_schema_from_model(
        model: Optional[Dict[str, Any]],
        *,
        route_name: str,
        lang: str = 'ar',
) -> Dict[str, Any]:
    tbl = find_kpi_main_table((model or {}).get('blocks') or {})
    if not tbl:
        diag = evaluate_kpi_main_schema_consistency(
            route_name=route_name,
            header_labels=[],
            rows=[],
            lang=lang,
        )
        diag['blocking_errors'] = ['rel32_kpi_main_table_missing']
        diag['kpi_main_schema_passed'] = False
        emit_rel32_kpi_main_schema_consistency_diag(diag)
        return diag
    rebound = rebind_table_spec(dict(tbl), lang=lang) or tbl
    diag = evaluate_kpi_main_schema_consistency(
        route_name=route_name,
        header_labels=rebound.get('header') or [],
        rows=rebound.get('rows') or [],
        bound_rows=rebound.get('bound_rows') or [],
        lang=lang,
    )
    emit_rel32_kpi_main_schema_consistency_diag(diag)
    return diag


def evaluate_kpi_main_schema_from_preview_html(
        html_text: str,
        *,
        route_name: str = 'preview',
) -> Dict[str, Any]:
    from release_engine_v3.rel32_kpi_owner_consistency_evidence import (
        _KpiTableRowsParser,
        _bound_rows_from_cells,
        evaluate_kpi_owner_consistency_from_preview_html,
    )
    dom_check = evaluate_preview_dom_binding_check(html_text or '', 'kpi_main')
    owner_diag = evaluate_kpi_owner_consistency_from_preview_html(
        html_text, route_name=route_name)
    parser = _KpiTableRowsParser()
    parser.feed(html_text or '')
    headers = parser.headers or list(REL32_KPI_MAIN_EXPECTED_SCHEMA_AR)
    all_br = _bound_rows_from_cells(headers, parser.rows, repair=False)
    diag = evaluate_kpi_main_schema_consistency(
        route_name=route_name,
        header_labels=headers,
        bound_rows=all_br,
        lang='ar',
    )
    if dom_check.get('blocking_errors'):
        diag['blocking_errors'] = list(dict.fromkeys(
            (diag.get('blocking_errors') or [])
            + (dom_check.get('blocking_errors') or [])))
        diag['kpi_main_schema_passed'] = not diag['blocking_errors']
    if not owner_diag.get('kpi_owner_consistency_passed'):
        diag['blocking_errors'] = list(dict.fromkeys(
            (diag.get('blocking_errors') or [])
            + (owner_diag.get('blocking_errors') or [])))
        diag['kpi_main_schema_passed'] = not diag['blocking_errors']
    diag['kpi_owner_consistency'] = owner_diag
    emit_rel32_kpi_main_schema_consistency_diag(diag)
    return diag


def evaluate_kpi_main_schema_from_export_text(
        blob: str,
        *,
        route_name: str,
        lang: str = 'ar',
) -> Dict[str, Any]:
    headers = extract_kpi_main_header_labels_from_text(blob)
    rows = _extract_main_kpi_rows(blob)
    diag = evaluate_kpi_main_schema_consistency(
        route_name=route_name,
        header_labels=headers,
        rows=rows,
        lang=lang,
        repair_rows=False,
    )
    if not headers:
        diag['blocking_errors'] = list(dict.fromkeys(
            (diag.get('blocking_errors') or [])
            + ['rel32_kpi_main_header_not_found_in_export_text']))
        diag['kpi_main_schema_passed'] = False
    emit_rel32_kpi_main_schema_consistency_diag(diag)
    return diag


def merge_kpi_main_schema_blockers(
        gate: Dict[str, Any],
        diag: Dict[str, Any],
) -> Dict[str, Any]:
    if diag.get('kpi_main_schema_passed'):
        gate['rel32_kpi_main_schema_consistency'] = diag
        return gate
    gate['blocking_errors'] = list(dict.fromkeys(
        (gate.get('blocking_errors') or [])
        + (diag.get('blocking_errors') or [])))
    gate['rel32_kpi_main_schema_consistency'] = diag
    return gate
